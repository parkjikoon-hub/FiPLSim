"""
FiPLSim 말단 수압 산출 방식 + 베르누이 MC 물리적 해석 — DOCX 문서 생성
"""
import os
from docx import Document
from docx.shared import Pt, Mm, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, "FiPLSim_Terminal_Pressure_Analysis.docx")


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    tcPr.append(shading)


def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h


def add_formula_paragraph(doc, formula_text, description=""):
    """수식 스타일 단락 추가"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(formula_text)
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    run.bold = True
    if description:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(description)
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p


def add_note_paragraph(doc, text, bold_prefix=""):
    """참고 박스 스타일 단락"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r0 = p.add_run(bold_prefix)
        r0.bold = True
        r0.font.size = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def make_table(doc, headers, rows, col_widths=None, header_color="1A3C6E"):
    """테이블 생성 헬퍼"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    # 데이터
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if i % 2 == 1:
                set_cell_shading(cell, "F0F4FA")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def build_document():
    doc = Document()

    # ── 페이지 설정 ──
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    # ── 기본 스타일 ──
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3

    # ═══════════════════════════════════════════
    #  표지
    # ═══════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("FiPLSim 말단 수압 산출 방식\n및 물리적 의미 분석")
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("Terminal Pressure Calculation Methodology\n& Bernoulli Monte Carlo Physical Interpretation")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r2.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = info_p.add_run("Fire Protection Pipe System Simulator (FiPLSim)\n2026. 02.")
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    #  목차
    # ═══════════════════════════════════════════
    add_styled_heading(doc, "목차", level=1)
    toc_items = [
        "1. 개요",
        "2. 시뮬레이션 조건",
        "3. 말단 수압 산출 방식 (Step by Step)",
        "   3.1. 1단계: 입구 압력 설정",
        "   3.2. 2단계: 교차배관 손실 계산",
        "   3.3. 3단계: 가지배관 입구 손실 (K3)",
        "   3.4. 4단계: 가지배관 8개 구간 순차 계산",
        "   3.5. 5단계: 최악 가지배관 선택",
        "4. 수리계산 공식 요약",
        "5. 베르누이 MC 시뮬레이션 결과",
        "6. 물리적 의미 해석",
        "   6.1. 평균 수압의 단조 감소",
        "   6.2. 표준편차와 이항분포",
        "   6.3. 시스템 안전 여유",
        "   6.4. 비드 위치의 중요성",
        "7. 일반 MC vs 베르누이 MC 비교",
        "8. 결론",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.size = Pt(10)
        if not item.startswith("   "):
            r.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════
    #  1. 개요
    # ═══════════════════════════════════════════
    add_styled_heading(doc, "1. 개요", level=1)
    doc.add_paragraph(
        "본 문서는 FiPLSim(Fire Protection Pipe System Simulator)이 "
        "소화배관 시스템의 말단 수압을 산출하는 정확한 방법론을 기술하고, "
        "베르누이 몬테카를로(Bernoulli Monte Carlo) 시뮬레이션 결과의 물리적 의미를 해석한다."
    )
    doc.add_paragraph(
        "FiPLSim은 입구 압력에서 시작하여 물이 실제로 흘러가는 경로를 따라 "
        "4가지 손실(마찰 + 접합부 비드 + 헤드 이음쇠 + 용접 비드)을 순차적으로 감산하는 방식으로 "
        "말단 수압을 계산한다. 베르누이 MC는 이 과정에서 '접합부 비드 손실(K1)'을 확률적으로 "
        "변동시켜 시공 품질의 불확실성이 시스템 성능에 미치는 영향을 정량화한다."
    )

    # ═══════════════════════════════════════════
    #  2. 시뮬레이션 조건
    # ═══════════════════════════════════════════
    add_styled_heading(doc, "2. 시뮬레이션 조건", level=1)
    doc.add_paragraph("베르누이 MC 시뮬레이션에 사용된 고정 조건은 다음과 같다.")

    cond_headers = ["항목", "값", "비고"]
    cond_rows = [
        ["입구 압력", "0.4 MPa", "펌프 토출 압력"],
        ["설계 유량", "2,100 LPM", "임계 유량 조건"],
        ["배관 구성", "4 가지배관 x 8 헤드 = 32 접합부", "트리(Tree) 토폴로지"],
        ["가지배관 간격", "3.5 m", "교차배관 위 배치 간격"],
        ["헤드 간격", "2.3 m", "가지배관 위 배치 간격"],
        ["비드 높이", "2.5 mm", "비드 존재 시 돌출 높이"],
        ["가지배관당 용접 비드", "8개", "직관 구간 내 균일 배치"],
        ["K_base", "0.5", "비드 0mm 기본 저항 계수"],
        ["K2 (헤드 이음쇠)", "2.5", "고정값"],
        ["K3 (분기 입구)", "1.0", "Tee-Branch 저항"],
        ["MC 반복 횟수", "10,000회", "각 확률 수준별"],
        ["비드 존재 확률 (p)", "0.1, 0.3, 0.5, 0.7, 0.9", "5개 수준"],
    ]
    make_table(doc, cond_headers, cond_rows, col_widths=[4.5, 5.0, 6.5])

    doc.add_paragraph()
    doc.add_paragraph(
        "비드 존재 확률 p는 각 접합부에 대해 독립적으로 적용된다. "
        "즉, 32개 접합부 각각에 대해 [0, 1] 균일분포 난수를 생성하고, "
        "난수 <= p이면 해당 접합부에 2.5mm 비드가 존재, 난수 > p이면 비드 없음(0mm)으로 설정한다."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    #  3. 말단 수압 산출 방식
    # ═══════════════════════════════════════════
    add_styled_heading(doc, "3. 말단 수압 산출 방식 (Step by Step)", level=1)
    doc.add_paragraph(
        "물이 입구에서 마지막 헤드(스프링클러)까지 흘러가는 경로를 따라, "
        "각 구간에서 발생하는 압력 손실을 순차적으로 감산하여 말단 수압을 산출한다. "
        "전체 과정은 5단계로 구성된다."
    )

    # ── 3.1 입구 압력 ──
    add_styled_heading(doc, "3.1. 1단계: 입구 압력 설정", level=2)
    doc.add_paragraph(
        "시뮬레이션의 시작점이다. 펌프가 배관 시스템에 공급하는 압력을 설정한다."
    )
    add_formula_paragraph(doc, "현재 압력 P = P_inlet = 0.4 MPa")
    doc.add_paragraph(
        "이 압력은 배관 입구에서 측정되는 정압(Static Pressure)으로, "
        "물이 배관을 통과하며 마찰, 방향 변환, 단면적 변화 등으로 인해 점진적으로 감소한다."
    )

    # ── 3.2 교차배관 ──
    add_styled_heading(doc, "3.2. 2단계: 교차배관 손실 계산", level=2)
    doc.add_paragraph(
        "교차배관(Cross Main, 80A)은 입구에서 들어온 물을 4개의 가지배관으로 분배하는 메인 배관이다. "
        "물이 교차배관을 따라 이동하면서 두 가지 손실이 발생한다."
    )

    doc.add_paragraph()
    p_cm1 = doc.add_paragraph()
    r_cm1 = p_cm1.add_run("(1) 마찰 손실 (주손실, Major Loss) - Darcy-Weisbach:")
    r_cm1.bold = True

    add_formula_paragraph(
        doc,
        "h_f = f x (L / D) x (V^2 / 2g)",
        "f: 마찰계수, L: 배관 길이(3.5m), D: 내경(77.92mm), V: 유속, g: 중력가속도(9.81 m/s^2)"
    )
    doc.add_paragraph(
        "배관 내벽과 물 사이의 마찰로 인한 에너지 손실이다. "
        "마찰계수 f는 Colebrook-White 방정식을 반복법으로 풀어 구한다."
    )

    doc.add_paragraph()
    p_cm2 = doc.add_paragraph()
    r_cm2 = p_cm2.add_run("(2) 티 분기 손실 (부차손실, Minor Loss):")
    r_cm2.bold = True

    add_formula_paragraph(
        doc,
        "h_m = K_TEE_RUN x (V^2 / 2g)",
        "K_TEE_RUN = 0.3 (교차배관 직진 흐름의 T자 분기 저항)"
    )
    doc.add_paragraph(
        "물이 T자 분기점을 통과할 때 유체 흐름의 방향 변화와 와류(Eddy)로 인해 발생하는 에너지 손실이다."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "교차배관 손실은 누적된다. 가지배관 1번(입구 바로 옆)은 교차배관 손실이 0이고, "
        "가지배관 4번(가장 먼 곳)은 3개 구간의 교차배관 손실이 모두 누적되어 분기점 압력이 가장 낮다."
    )

    # 교차배관 유량 감소 테이블
    cm_headers = ["교차배관 구간", "잔류 유량 (LPM)", "유속 (m/s)", "비고"]
    cm_rows = [
        ["입구 ~ 가지 1", "2,100", "7.36", "최대 유량"],
        ["가지 1 ~ 가지 2", "1,575", "5.52", "1개 분기 후"],
        ["가지 2 ~ 가지 3", "1,050", "3.68", "2개 분기 후"],
        ["가지 3 ~ 가지 4", "525", "1.84", "3개 분기 후"],
    ]
    make_table(doc, cm_headers, cm_rows, col_widths=[4, 3.5, 3, 5.5])

    # ── 3.3 K3 ──
    add_styled_heading(doc, "3.3. 3단계: 가지배관 입구 손실 (K3)", level=2)
    doc.add_paragraph(
        "교차배관에서 가지배관으로 물이 90도 방향 전환하며 들어갈 때 발생하는 분기 입구 손실이다."
    )
    add_formula_paragraph(
        doc,
        "h_K3 = K3 x (V_inlet^2 / 2g)",
        "K3 = 1.0 (Tee-Branch, T자 분기 저항 계수)"
    )
    doc.add_paragraph(
        "V_inlet은 가지배관 입구에서의 유속으로, 가지배관 전체 유량(525 LPM)을 "
        "첫 구간 배관(50A, 내경 52.51mm)의 단면적으로 나누어 계산한다."
    )

    doc.add_page_break()

    # ── 3.4 가지배관 8개 구간 ──
    add_styled_heading(doc, "3.4. 4단계: 가지배관 8개 구간 순차 계산 (핵심)", level=2)
    doc.add_paragraph(
        "가지배관에는 8개의 헤드(스프링클러)가 있고, 각 헤드 사이가 하나의 '구간(Segment)'이다. "
        "물이 구간을 하나씩 통과할 때마다 4가지 종류의 압력 손실이 발생한다."
    )

    doc.add_paragraph()
    p_key = doc.add_paragraph()
    r_key = p_key.add_run("핵심 원리: ")
    r_key.bold = True
    r_key2 = p_key.add_run(
        "헤드를 하나 지날 때마다 해당 헤드로 물이 빠져나가므로 "
        "잔류 유량이 감소한다. 유량이 줄면 유속이 낮아지고, 유속이 낮아지면 손실도 줄어든다. "
        "따라서 상류(구간 1)에서 손실이 크고, 하류(구간 8)에서 손실이 작다."
    )

    # 유량 감소 테이블
    flow_headers = ["구간", "관경", "내경 (mm)", "유량 (LPM)", "유속 (m/s)", "비고"]
    flow_rows = [
        ["1", "50A", "52.51", "525.0", "4.05", "전체 유량 (8개 헤드분)"],
        ["2", "50A", "52.51", "459.4", "3.54", "7개 헤드분"],
        ["3", "50A", "52.51", "393.8", "3.03", "6개 헤드분"],
        ["4", "40A", "40.90", "328.1", "4.16", "5개 헤드분 (관경 축소)"],
        ["5", "40A", "40.90", "262.5", "3.33", "4개 헤드분"],
        ["6", "32A", "35.04", "196.9", "3.41", "3개 헤드분 (관경 축소)"],
        ["7", "25A", "26.64", "131.3", "3.93", "2개 헤드분 (관경 축소)"],
        ["8", "25A", "26.64", "65.6", "1.97", "1개 헤드분 (최말단)"],
    ]
    make_table(doc, flow_headers, flow_rows, col_widths=[1.5, 2, 2.5, 2.5, 2.5, 5])

    doc.add_paragraph()
    doc.add_paragraph(
        "관경(Pipe Size)은 NFSC 103 기준에 따라 하류 헤드 수에 의해 자동 선정된다: "
        "8~6개 -> 50A, 5~4개 -> 40A, 3개 -> 32A, 2~1개 -> 25A."
    )

    doc.add_paragraph()
    # ── 4가지 손실 상세 설명 ──
    p_a = doc.add_paragraph()
    r_a = p_a.add_run("(a) 마찰 손실 (p_major) - Darcy-Weisbach:")
    r_a.bold = True
    r_a.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    add_formula_paragraph(
        doc,
        "p_major = f x (L / D) x (V^2 / 2g)",
        "f: Colebrook-White 마찰계수, L: 2.3m (헤드 간격), D: 해당 구간 내경, V: 해당 구간 유속"
    )
    doc.add_paragraph(
        "배관 내벽 마찰에 의한 직관 손실이다. "
        "마찰계수 f는 레이놀즈 수(Re)와 상대조도(epsilon/D)에 따라 결정되며, "
        "Colebrook-White 방정식을 Swamee-Jain 근사로 초기값을 구한 후 "
        "최대 10회 고정점 반복(Fixed-Point Iteration)으로 수렴시킨다."
    )

    doc.add_paragraph()
    p_b = doc.add_paragraph()
    r_b = p_b.add_run("(b) 접합부 비드 손실 (p_K1) - 베르누이 MC의 핵심 변수:")
    r_b.bold = True
    r_b.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    add_formula_paragraph(
        doc,
        "K1 = K_base x (D / D_eff)^4\nD_eff = D - 2 x bead_height\np_K1 = K1 x (V^2 / 2g)",
        "K_base = 0.5, D: 배관 내경(mm), D_eff: 유효 내경(mm), bead_height: 비드 높이(mm)"
    )
    doc.add_paragraph(
        "배관 이음쇠(Welded Fitting)의 용접 비드로 인한 국부 손실이다. "
        "비드가 배관 내부로 돌출되면 유효 내경이 줄어들어 유로가 좁아지고, "
        "급축소-급확대(Sudden Contraction-Expansion) 효과로 추가 와류가 발생한다."
    )

    doc.add_paragraph()
    doc.add_paragraph("비드 높이에 따른 K1값 변화 (25A 배관 기준):")

    k1_headers = ["비드 높이 (mm)", "D_eff (mm)", "내경 축소율 (%)", "K1", "K1 증가율"]
    k1_rows = [
        ["0.0", "26.64", "0.0%", "0.500", "기준값"],
        ["1.0", "24.64", "7.5%", "0.683", "+36.6%"],
        ["1.5", "23.64", "11.3%", "0.806", "+61.3%"],
        ["2.0", "22.64", "15.0%", "0.959", "+91.7%"],
        ["2.5", "21.64", "18.8%", "1.148", "+129.7%"],
        ["3.0", "20.64", "22.5%", "1.388", "+177.5%"],
    ]
    make_table(doc, k1_headers, k1_rows, col_widths=[3, 3, 3, 2.5, 3.5])

    doc.add_paragraph()
    add_note_paragraph(
        doc,
        "비드 2.5mm일 때 K1이 0.5에서 1.148로 약 2.3배 증가한다. "
        "이것은 비드 없는 배관 대비 해당 접합부에서 2.3배 더 많은 에너지가 손실됨을 의미한다.",
        "핵심: "
    )

    doc.add_paragraph()
    add_note_paragraph(
        doc,
        "베르누이 MC에서는 각 접합부마다 독립적으로 난수를 생성하여, "
        "난수 <= p이면 bead = 2.5mm (K1 = 1.148 ~ 관경별 상이), "
        "난수 > p이면 bead = 0mm (K1 = 0.5)로 설정한다.",
        "베르누이 MC: "
    )

    doc.add_paragraph()
    p_c = doc.add_paragraph()
    r_c = p_c.add_run("(c) 헤드 이음쇠 손실 (p_K2):")
    r_c.bold = True
    r_c.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    add_formula_paragraph(
        doc,
        "p_K2 = K2 x (V^2 / 2g)",
        "K2 = 2.5 (헤드 이음쇠 고유 저항 계수, 상수)"
    )
    doc.add_paragraph(
        "스프링클러 헤드가 가지배관에서 분기되는 지점의 구조적 저항이다. "
        "이 값은 비드 유무와 관계없이 항상 동일하게 적용된다."
    )

    doc.add_paragraph()
    p_d = doc.add_paragraph()
    r_d = p_d.add_run("(d) 용접 비드 손실 (p_weld_beads):")
    r_d.bold = True
    r_d.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    add_formula_paragraph(
        doc,
        "p_weld = SUM( K_bead_i x (V^2 / 2g) )  (해당 구간 내 모든 비드)",
        "가지배관당 8개의 용접 비드가 직관 구간 내에 균일 배치됨"
    )
    doc.add_paragraph(
        "배관 직관 구간 내에 존재하는 용접 비드(Weld Bead)의 국부 손실이다. "
        "이 비드들은 접합부 비드와 별개로, 배관 제작 시 용접 이음부에서 발생하는 비드를 모사한다. "
        "rng=None일 때 균일하게 배치되어 매 시행마다 동일한 위치에 존재한다."
    )

    doc.add_paragraph()
    p_sum = doc.add_paragraph()
    r_sum = p_sum.add_run("구간 합산 공식:")
    r_sum.bold = True
    r_sum.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    add_formula_paragraph(
        doc,
        "Total_seg_loss = p_major + p_K1 + p_K2 + p_weld_beads\n"
        "P_current = P_current - Total_seg_loss",
        "8개 구간을 순차적으로 반복하면, 마지막 헤드(8번)에서의 P_current가 말단 수압이 된다."
    )

    doc.add_page_break()

    # ── 3.5 최악 가지배관 ──
    add_styled_heading(doc, "3.5. 5단계: 최악 가지배관 선택", level=2)
    doc.add_paragraph(
        "4개 가지배관 모두에 대해 위의 계산을 수행한 후, "
        "가장 낮은 말단 수압을 가진 가지배관을 '최악 가지배관(Worst Branch)'으로 식별한다. "
        "이 최악 말단 수압이 MC 시뮬레이션의 기록 대상(worst_terminal_mpa)이다."
    )
    doc.add_paragraph(
        "일반적으로 교차배관 입구에서 가장 먼 가지배관(4번)이 최악이 된다. "
        "이는 교차배관 누적 손실이 가장 크기 때문이다. "
        "단, 비드 배치에 따라 다른 가지배관이 최악이 될 수도 있다."
    )

    # ═══════════════════════════════════════════
    #  4. 수리계산 공식 요약
    # ═══════════════════════════════════════════
    add_styled_heading(doc, "4. 수리계산 공식 요약", level=1)

    formula_headers = ["공식명", "수식", "변수 설명"]
    formula_rows = [
        ["레이놀즈 수", "Re = V x D / nu", "nu = 1.004e-6 m^2/s (물, 20'C)"],
        ["유속 변환", "V = Q / A\nA = pi x (D/2)^2", "Q: m^3/s, A: 단면적, D: 내경(m)"],
        ["Colebrook-White", "1/sqrt(f) = -2 log10(\nepsilon/3.7D + 2.51/Re*sqrt(f))", "epsilon = 0.045mm (탄소강)"],
        ["주손실 (Darcy)", "h_f = f x (L/D) x V^2/2g", "f: 마찰계수, L: 배관 길이"],
        ["부차손실 (K-factor)", "h_m = K x V^2/2g", "K: 손실 계수"],
        ["비드 K-factor", "K = K_base x (D/D_eff)^4\nD_eff = D - 2*bead", "K_base=0.5, bead: 비드 높이"],
        ["수두->압력 변환", "P = rho x g x h / 10^6", "rho=998 kg/m^3, g=9.81 m/s^2"],
    ]
    make_table(doc, formula_headers, formula_rows, col_widths=[3.5, 5.5, 5.5])

    doc.add_page_break()

    # ═══════════════════════════════════════════
    #  5. 베르누이 MC 결과
    # ═══════════════════════════════════════════
    add_styled_heading(doc, "5. 베르누이 MC 시뮬레이션 결과", level=1)
    doc.add_paragraph(
        "고정 조건(입구 0.4 MPa, 유량 2,100 LPM, 비드 2.5mm, 4x8 배관)에서 "
        "비드 존재 확률 p를 5개 수준으로 변화시킨 MC 10,000회 시뮬레이션 결과이다."
    )

    mc_headers = [
        "p\n(비드 확률)", "의미", "기대\n비드 수",
        "평균 수압\n(MPa)", "표준편차\n(MPa)", "최솟값\n(MPa)",
        "최댓값\n(MPa)", "Pf\n(%)", "판정"
    ]
    mc_rows = [
        ["0.1", "우수한 시공 품질", "3.2", "0.1205", "0.0022", "0.1125", "0.1279", "0.0", "PASS"],
        ["0.3", "양호한 시공 품질", "9.6", "0.1162", "0.0030", "0.1059", "0.1254", "0.0", "PASS"],
        ["0.5", "기본 설정", "16.0", "0.1125", "0.0030", "0.1009", "0.1226", "0.0", "PASS"],
        ["0.7", "열악한 시공 품질", "22.4", "0.1090", "0.0025", "0.1001", "0.1183", "0.0", "PASS"],
        ["0.9", "매우 열악한 시공 품질", "28.8", "0.1059", "0.0015", "0.1007", "0.1117", "0.0", "PASS"],
    ]
    make_table(doc, mc_headers, mc_rows, col_widths=[1.8, 3.0, 1.5, 2.0, 2.0, 1.8, 1.8, 1.0, 1.2])

    doc.add_paragraph()
    doc.add_paragraph(
        "기대 비드 수 = 전체 접합부(32개) x p. "
        "예: p=0.5이면 평균 16개의 접합부에 2.5mm 비드가 존재한다."
    )

    # ═══════════════════════════════════════════
    #  6. 물리적 의미 해석
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_styled_heading(doc, "6. 물리적 의미 해석", level=1)

    # 6.1
    add_styled_heading(doc, "6.1. 평균 수압의 단조 감소: p 증가 -> 수압 감소", level=2)
    doc.add_paragraph(
        "비드 존재 확률 p가 0.1에서 0.9로 증가하면, 비드가 존재하는 접합부가 "
        "평균 3.2개에서 28.8개로 늘어난다. 비드가 있는 접합부마다 K1이 0.5에서 "
        "약 1.15(25A 기준)로 상승하므로, 추가 국부 손실이 누적되어 말단 수압이 낮아진다."
    )

    dp_headers = ["비교 구간", "수압 차이 (MPa)", "수두 환산 (m)", "물리적 의미"]
    dp_rows = [
        ["p=0.1 -> p=0.9", "0.0146", "약 1.49 m", "비드 25.6개 추가 효과"],
        ["비드 1개당 평균", "~0.0006", "약 0.06 m", "접합부 1개의 비드 영향"],
    ]
    make_table(doc, dp_headers, dp_rows, col_widths=[3.5, 3, 3, 5.5])

    doc.add_paragraph()
    doc.add_paragraph(
        "이는 비드 존재 확률이 높을수록(시공 품질이 낮을수록) "
        "말단 수압이 선형적으로 감소함을 보여준다. "
        "p=0.1에서 p=0.9까지의 수압 차이 0.0146 MPa는 "
        "약 1.49m 수두에 해당하며, 이는 비드로 인한 총 추가 손실을 의미한다."
    )

    # 6.2
    add_styled_heading(doc, "6.2. 표준편차와 이항분포의 관계", level=2)
    doc.add_paragraph(
        "32개 접합부에서 비드가 존재하는 개수는 이항분포 B(n, p)를 따른다."
    )
    add_formula_paragraph(
        doc,
        "비드 개수 ~ B(32, p)\n분산 = n x p x (1-p) = 32 x p x (1-p)",
        "n = 32 (전체 접합부 수), p: 비드 존재 확률"
    )

    var_headers = ["p", "n x p x (1-p)", "분산", "수압 표준편차 (MPa)", "비고"]
    var_rows = [
        ["0.1", "32 x 0.1 x 0.9", "2.88", "0.0022", "낮은 변동"],
        ["0.3", "32 x 0.3 x 0.7", "6.72", "0.0030", "중간 변동"],
        ["0.5", "32 x 0.5 x 0.5", "8.00", "0.0030", "최대 변동"],
        ["0.7", "32 x 0.7 x 0.3", "6.72", "0.0025", "중간 변동"],
        ["0.9", "32 x 0.9 x 0.1", "2.88", "0.0015", "낮은 변동"],
    ]
    make_table(doc, var_headers, var_rows, col_widths=[1.5, 3.5, 2, 3.5, 3.5])

    doc.add_paragraph()
    doc.add_paragraph(
        "이항분포의 분산은 p=0.5에서 최대가 된다. 이는 동전 던지기와 동일한 원리이다. "
        "앞면 확률이 50%일 때 결과가 가장 예측 불가능한 것처럼, "
        "비드 존재 확률이 50%일 때 매 시행마다 비드 배치 조합이 가장 다양하게 변한다."
    )
    doc.add_paragraph(
        "반면, p=0.1일 때는 거의 비드가 없고, p=0.9일 때는 거의 모든 접합부에 비드가 있어 "
        "시행 간 변동이 작아진다. 이것이 수압 표준편차가 양 극단에서 작고 중간에서 큰 이유이다."
    )

    # 6.3
    add_styled_heading(doc, "6.3. 시스템 안전 여유", level=2)
    doc.add_paragraph(
        "NFSC 103 규정에 따른 말단 최소 방수압 기준은 0.1 MPa이다. "
        "모든 확률 수준에서 규정 미달 확률(Pf)이 0%로, "
        "입구 압력 0.4 MPa는 2.5mm 비드의 추가 손실을 충분히 감당한다."
    )

    margin_headers = ["p", "평균 수압 (MPa)", "기준 대비 여유 (MPa)", "여유율 (%)"]
    margin_rows = [
        ["0.1", "0.1205", "0.0205", "20.5%"],
        ["0.3", "0.1162", "0.0162", "16.2%"],
        ["0.5", "0.1125", "0.0125", "12.5%"],
        ["0.7", "0.1090", "0.0090", "9.0%"],
        ["0.9", "0.1059", "0.0059", "5.9%"],
    ]
    make_table(doc, margin_headers, margin_rows, col_widths=[2, 3.5, 4, 4.5])

    doc.add_paragraph()
    doc.add_paragraph(
        "p=0.9(매우 열악한 시공)에서도 여유가 5.9% 남아있으나, "
        "최솟값(0.1007 MPa)이 기준(0.1 MPa)에 매우 근접한다. "
        "유량이 증가하거나 입구 압력이 낮아지면 규정 미달이 발생할 수 있다."
    )

    # 6.4
    add_styled_heading(doc, "6.4. 비드 위치의 중요성", level=2)
    doc.add_paragraph(
        "같은 수의 비드라도 어디에 위치하느냐에 따라 손실이 달라진다. "
        "이는 각 구간의 관경과 유속이 다르기 때문이다."
    )

    loc_headers = ["구간", "관경", "유속 (m/s)", "K=0.648 비드\n1개 손실 (kPa)", "비고"]
    loc_rows = [
        ["1 (상류)", "50A", "4.05", "0.54", "높은 유속 -> 큰 손실"],
        ["4", "40A", "4.16", "0.57", "관경 축소로 유속 증가"],
        ["6", "32A", "3.41", "0.38", "중간 유속"],
        ["8 (말단)", "25A", "1.97", "0.13", "낮은 유속 -> 작은 손실"],
    ]
    make_table(doc, loc_headers, loc_rows, col_widths=[2.5, 2, 2.5, 3.5, 5.5])

    doc.add_paragraph()
    doc.add_paragraph(
        "상류(구간 1)의 비드 1개가 말단(구간 8)의 비드 1개보다 약 4배 더 큰 손실을 유발한다. "
        "이는 손실 공식 h_m = K x V^2/2g에서 유속의 제곱에 비례하기 때문이다. "
        "따라서 같은 p 값에서도 비드가 상류에 집중된 시행은 수압이 낮고, "
        "말단에 집중된 시행은 수압이 상대적으로 높다."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    #  7. 일반 MC vs 베르누이 MC 비교
    # ═══════════════════════════════════════════
    add_styled_heading(doc, "7. 일반 MC vs 베르누이 MC 비교", level=1)

    comp_headers = ["항목", "일반 MC\n(run_mc_batch)", "베르누이 MC\n(run_mc_bernoulli)"]
    comp_rows = [
        ["비드 배치 방식", "1~3개 '결함 위치'를\n무작위 선택", "32개 접합부 각각\n독립 확률 p로 판정"],
        ["비드 개수 분포", "균일분포 U(1, 3)", "이항분포 B(32, p)"],
        ["비드 높이", "고정값 (0/2.0/2.5mm)", "2.5mm (존재 시) or 0mm"],
        ["비드 개수 범위", "1~3개 (고정)", "0~32개 (확률적)"],
        ["물리적 의미", "소수 결함의 위치 영향", "시공 품질 수준별 영향"],
        ["Pf (bead 2.5, 2100LPM)", "2.43%", "0% (모든 p 수준)"],
        ["차이 원인", "비드가 특정 가지배관에\n집중 가능", "비드가 4개 가지배관에\n골고루 분산"],
    ]
    make_table(doc, comp_headers, comp_rows, col_widths=[3.5, 5.5, 5.5])

    doc.add_paragraph()
    doc.add_paragraph(
        "일반 MC에서 Pf=2.43%가 나왔던 반면 베르누이 MC에서 Pf=0%인 주된 이유는 "
        "비드 배치의 공간적 분포 차이이다."
    )
    doc.add_paragraph(
        "일반 MC: 결함 1~3개가 같은 가지배관의 같은 구간에 집중될 수 있다. "
        "예를 들어, 가지배관 4번(최악 가지배관)의 상류 구간에 비드 3개가 몰리면 "
        "해당 가지배관의 말단 수압이 크게 떨어진다."
    )
    doc.add_paragraph(
        "베르누이 MC: 각 접합부가 독립적이므로, p=0.9에서도 비드가 "
        "4개 가지배관에 대략 균등하게 분포한다 (가지배관당 약 7.2개). "
        "특정 가지배관에 비드가 극단적으로 집중될 확률이 매우 낮아 "
        "최악 말단 수압의 하한이 높다."
    )

    # ═══════════════════════════════════════════
    #  8. 결론
    # ═══════════════════════════════════════════
    add_styled_heading(doc, "8. 결론", level=1)

    conclusions = [
        (
            "산출 방법론의 물리적 타당성",
            "FiPLSim은 Darcy-Weisbach 주손실과 K-factor 부차손실을 결합한 "
            "표준 수리계산 방법론을 사용하며, Colebrook-White 마찰계수 산정, "
            "NFSC 103 기반 자동 관경 선정 등 실무 설계 기준을 충실히 반영한다."
        ),
        (
            "비드 K-factor 모델의 물리적 근거",
            "K = K_base x (D/D_eff)^4 모델은 비드로 인한 유효 내경 감소가 "
            "급축소-급확대 효과를 유발하는 물리적 메커니즘을 반영한다. "
            "이는 Idelchik의 급축소 손실 계수와 Borda-Carnot 급확대 손실 계수를 "
            "통합한 실험적/이론적 모델이다."
        ),
        (
            "베르누이 MC의 공학적 의미",
            "시공 품질을 확률 변수로 모델링하여 시스템 신뢰성을 정량화할 수 있다. "
            "p=0.9(매우 열악한 시공)에서도 Pf=0%로 시스템이 충분한 안전 여유를 가지지만, "
            "여유율 5.9%는 유량 증가나 배관 노화 시 취약해질 수 있음을 시사한다."
        ),
        (
            "비드 위치 효과",
            "동일한 비드 개수라도 상류(고유속 구간)에 위치할 때 손실이 최대 4배까지 "
            "커진다. 이는 시공 관리에서 상류 구간의 용접 품질이 특히 중요함을 의미한다."
        ),
        (
            "일반 MC와의 차이",
            "일반 MC(결함 집중 가능)와 베르누이 MC(결함 분산)는 서로 다른 물리적 시나리오를 "
            "모사한다. 실제 현장에서는 두 모델을 병행 적용하여 보수적 설계(일반 MC)와 "
            "통계적 품질 관리(베르누이 MC)를 동시에 수행하는 것이 바람직하다."
        ),
    ]

    for i, (title, content) in enumerate(conclusions, 1):
        p = doc.add_paragraph()
        r_num = p.add_run(f"{i}. {title}: ")
        r_num.bold = True
        r_num.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
        r_content = p.add_run(content)
        r_content.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # 문서 정보
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_f = footer_p.add_run("Generated by FiPLSim (Fire Protection Pipe System Simulator)")
    r_f.font.size = Pt(8)
    r_f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r_f.italic = True

    # ── 저장 ──
    doc.save(OUT_PATH)
    print(f"문서 저장 완료: {OUT_PATH}")
    return OUT_PATH


if __name__ == "__main__":
    build_document()
