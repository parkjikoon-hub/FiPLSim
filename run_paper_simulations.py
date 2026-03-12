#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FiPLSim 논문용 배치 자동화 스크립트
====================================
JESTECH 논문: 소화배관 용접비드가 말단 수압에 미치는 영향 분석

실행: PYTHONIOENCODING=utf-8 python3 run_paper_simulations.py [--stage 1|2|3]
"""

import sys
import os
import argparse
import json
import time
import importlib
import math
import warnings
from pathlib import Path
from datetime import datetime

import numpy as np
import pandas as pd

# ──────────────────────────────────────────────
# 논문용 물성치 오버라이드 (코드 기본값과 다른 값)
# ──────────────────────────────────────────────
import constants
# 먼저 상수 모듈의 값을 변경
constants.EPSILON_MM = 0.046
constants.EPSILON_M = 0.046 / 1000.0
constants.RHO = 1000.0
constants.MU = 1.002e-3
constants.NU = constants.MU / constants.RHO

# 다른 모듈을 reload하면 변경된 constants 값으로 재설정됨
# (constants는 reload하지 않음 — reload하면 원본 파일값으로 복구됨)
import hydraulics
importlib.reload(hydraulics)

import pipe_network as pn
importlib.reload(pn)

import simulation as sim
importlib.reload(sim)

# matplotlib 한글 설정
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

# 한글 폰트 설정 시도
_KOREAN_FONTS = ["Malgun Gothic", "맑은 고딕", "NanumGothic", "AppleGothic"]
_font_set = False
for _fname in _KOREAN_FONTS:
    _flist = [f for f in fm.fontManager.ttflist if _fname in f.name]
    if _flist:
        plt.rcParams["font.family"] = _fname
        plt.rcParams["axes.unicode_minus"] = False
        _font_set = True
        break
if not _font_set:
    warnings.warn("한글 폰트를 찾을 수 없습니다. 그래프 제목이 깨질 수 있습니다.")

# ──────────────────────────────────────────────
# 논문용 공통 파라미터
# ──────────────────────────────────────────────
NUM_BRANCHES = 4
HEADS_PER_BRANCH = 8
BRANCH_SPACING_M = 3.5
HEAD_SPACING_M = 2.1
DESIGN_FLOW_LPM = 2560.0        # 32 × 80 LPM (NFPC 103)
BEAD_HEIGHT_MM = 1.5
BEAD_HEIGHT_STD_MM = 0.5         # 비균일 모델 기본 σ
BEADS_PER_BRANCH = 4
BRANCH_INLET_CONFIG = "80A-65A"
SUPPLY_PIPE_SIZE = "100A"
MC_ITERATIONS = 10000

# K-factor 상수
K1_BASE = constants.K1_BASE      # 0.5
K2 = constants.K2                # 2.5
K3 = constants.K3                # 1.0

# 장비 K-factors (계획서: 6종, 총 K=6.20)
EQUIPMENT_K_FACTORS = {
    "알람밸브 (습식)":     {"K": 2.0,  "qty": 1},
    "유수검지장치":        {"K": 1.0,  "qty": 1},
    "게이트밸브 (전개)":   {"K": 0.15, "qty": 2},
    "체크밸브 (스윙형)":   {"K": 2.0,  "qty": 1},
    "90° 엘보":           {"K": 0.75, "qty": 1},
    "리듀서 (점축소)":     {"K": 0.15, "qty": 1},
}

# 배관 치수 (mm) - constants.py에서 가져옴
PIPE_DIMS = constants.PIPE_DIMENSIONS

# 합격 기준
PASS_THRESHOLD_MPA = 0.1  # 말단 수압 ≥ 0.1 MPa

# 출력 디렉토리
BASE_DIR = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "outputs"
STAGE1_DIR = OUTPUT_DIR / "stage1_기반확립"
STAGE2_DIR = OUTPUT_DIR / "stage2_파라메트릭"
STAGE3_DIR = OUTPUT_DIR / "stage3_확률론적"

# 시스템 특성화 저장 파일
CHAR_FILE = OUTPUT_DIR / "system_characterization.json"

# ──────────────────────────────────────────────
# 공통 유틸리티
# ──────────────────────────────────────────────
def ensure_dirs():
    """출력 디렉토리 생성"""
    for d in [STAGE1_DIR / "data", STAGE1_DIR / "figures",
              STAGE2_DIR / "data", STAGE2_DIR / "figures",
              STAGE3_DIR / "data", STAGE3_DIR / "figures"]:
        d.mkdir(parents=True, exist_ok=True)

def save_csv(df, path):
    """CSV 저장 (UTF-8 BOM)"""
    df.to_csv(path, index=False, encoding="utf-8-sig")
    print(f"  → CSV 저장: {path.name}")

def save_fig(fig, path, dpi=200):
    """PNG 저장"""
    fig.savefig(path, dpi=dpi, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  → 그래프 저장: {path.name}")

def save_system_characterization(char_data):
    """시스템 특성화 결과를 JSON으로 저장"""
    char_data["timestamp"] = datetime.now().isoformat()
    char_data["param_hash"] = _compute_param_hash()
    CHAR_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(CHAR_FILE, "w", encoding="utf-8") as f:
        json.dump(char_data, f, ensure_ascii=False, indent=2)
    print(f"  → P_ref = {char_data['p_ref_mpa']:.4f} MPa 저장됨")

def load_system_characterization():
    """저장된 시스템 특성화 결과 로드"""
    if not CHAR_FILE.exists():
        raise FileNotFoundError(
            "시스템 특성화 파일이 없습니다. 1단계(--stage 1)를 먼저 실행하세요."
        )
    with open(CHAR_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data

def load_p_ref():
    """P_ref 값만 빠르게 로드"""
    data = load_system_characterization()
    return data["p_ref_mpa"]

def _compute_param_hash():
    """현재 런타임 파라미터의 해시 계산 (변경 감지용)"""
    import hashlib
    params = {
        "flow": DESIGN_FLOW_LPM,
        "branches": NUM_BRANCHES,
        "heads": HEADS_PER_BRANCH,
        "h_spacing": HEAD_SPACING_M,
        "b_spacing": BRANCH_SPACING_M,
        "inlet_config": BRANCH_INLET_CONFIG,
        "supply_pipe": SUPPLY_PIPE_SIZE,
        "epsilon": constants.EPSILON_MM,
        "rho": constants.RHO,
        "nu": constants.NU,
    }
    param_str = json.dumps(params, sort_keys=True)
    return hashlib.md5(param_str.encode()).hexdigest()[:12]

def verify_system_characterization():
    """시스템 특성화가 현재 파라미터와 일치하는지 검증"""
    if not CHAR_FILE.exists():
        return False
    data = load_system_characterization()
    stored_hash = data.get("param_hash", "")
    current_hash = _compute_param_hash()
    if stored_hash != current_hash:
        print("  !! 파라미터 변경 감지. 시스템 특성화 재실행이 필요합니다.")
        return False
    return True

def common_params():
    """모든 시뮬레이션에 공통으로 전달되는 파라미터 dict"""
    return dict(
        num_branches=NUM_BRANCHES,
        heads_per_branch=HEADS_PER_BRANCH,
        branch_spacing_m=BRANCH_SPACING_M,
        head_spacing_m=HEAD_SPACING_M,
        equipment_k_factors=EQUIPMENT_K_FACTORS,
        supply_pipe_size=SUPPLY_PIPE_SIZE,
        branch_inlet_config=BRANCH_INLET_CONFIG,
    )

def common_params_mc():
    """MC 전용 공통 파라미터"""
    return dict(
        **common_params(),
        K1_base=K1_BASE,
        K2_val=K2,
        K3_val=K3,
        beads_per_branch=BEADS_PER_BRANCH,
    )

def elapsed(t0):
    """경과 시간 포맷"""
    s = time.time() - t0
    if s < 60:
        return f"{s:.1f}초"
    return f"{s/60:.1f}분"


# ══════════════════════════════════════════════
#  1단계: 기반 확립 (Part 0.5, Part 1, Part 2)
# ══════════════════════════════════════════════

def run_stage1():
    """1단계: 프로그램 검증 + 기준선 분석"""
    print("\n" + "=" * 60)
    print("  1단계: 기반 확립 (Part 0.5 + Part 1 + Part 2)")
    print("=" * 60)
    t0 = time.time()

    data_dir = STAGE1_DIR / "data"
    fig_dir = STAGE1_DIR / "figures"

    # ── Part 0.5: 시스템 특성화 (NFPC 103 역산 + 이중 검증) ──
    p_ref = run_part0_5(data_dir)

    # ── Part 1: 프로그램 검증 ──
    run_part1(data_dir, fig_dir, p_ref)

    # ── Part 2: 기준선 정적 비교 ──
    run_part2(data_dir, fig_dir, p_ref)

    # ── DOCX 보고서 생성 ──
    generate_stage1_report(data_dir, fig_dir, p_ref)

    print(f"\n✅ 1단계 완료 (소요시간: {elapsed(t0)})")
    return p_ref


# ── Part 0.5: 시스템 특성화 (NFPC 103 역산 + 이중 검증) ──
def run_part0_5(data_dir):
    """수리계산 직접 역산 + 시뮬레이션 이중 검증 → P_ref 결정"""
    print("\n── Part 0.5: 시스템 특성화 (NFPC 103 역산 + 이중 검증) ──")

    analytical_params = dict(
        total_flow_lpm=DESIGN_FLOW_LPM,
        num_branches=NUM_BRANCHES,
        heads_per_branch=HEADS_PER_BRANCH,
        branch_spacing_m=BRANCH_SPACING_M,
        head_spacing_m=HEAD_SPACING_M,
        K1_base=K1_BASE,
        K2_val=K2,
        K3_val=K3,
        equipment_k_factors=EQUIPMENT_K_FACTORS,
        supply_pipe_size=SUPPLY_PIPE_SIZE,
        branch_inlet_config=BRANCH_INLET_CONFIG,
    )

    # [Step 1] 수리계산 직접 역산 (Analytical)
    print("  Step 1: 수리계산 직접 역산 (Analytical)")
    analytical_B = pn.calculate_system_delta_p(
        bead_height_mm=0.0, beads_per_branch=0,
        **analytical_params,
    )
    analytical_A = pn.calculate_system_delta_p(
        bead_height_mm=BEAD_HEIGHT_MM, beads_per_branch=BEADS_PER_BRANCH,
        **analytical_params,
    )
    dp_analytical_B = analytical_B["delta_p_total_mpa"]
    dp_analytical_A = analytical_A["delta_p_total_mpa"]
    print(f"    ΔP_analytical_B = {dp_analytical_B:.6f} MPa")
    print(f"    ΔP_analytical_A = {dp_analytical_A:.6f} MPa")

    # [Step 2] 시뮬레이션 검증 (P_in=0.5, 1.0)
    print("  Step 2: 시뮬레이션 검증 (2회)")
    sim_pressures = [0.5, 1.0]
    dp_sim_B = []
    dp_sim_A = []
    for p_in in sim_pressures:
        comp = pn.compare_dynamic_cases(
            bead_height_existing=BEAD_HEIGHT_MM,
            bead_height_new=0.0,
            inlet_pressure_mpa=p_in,
            total_flow_lpm=DESIGN_FLOW_LPM,
            beads_per_branch=BEADS_PER_BRANCH,
            **common_params(),
        )
        dp_b = round(p_in - comp["terminal_B_mpa"], 6)
        dp_a = round(p_in - comp["terminal_A_mpa"], 6)
        dp_sim_B.append(dp_b)
        dp_sim_A.append(dp_a)
        print(f"    P_in={p_in:.1f}: ΔP_B={dp_b:.6f}, ΔP_A={dp_a:.6f}")

    # [Step 3] 교차 검증
    print("  Step 3: 교차 검증")
    dp_sim_B_avg = sum(dp_sim_B) / len(dp_sim_B)
    dp_sim_A_avg = sum(dp_sim_A) / len(dp_sim_A)

    # 입구 압력 독립성 검증
    indep_err_B = abs(dp_sim_B[0] - dp_sim_B[1]) / dp_sim_B_avg * 100 if dp_sim_B_avg else 0
    indep_err_A = abs(dp_sim_A[0] - dp_sim_A[1]) / dp_sim_A_avg * 100 if dp_sim_A_avg else 0

    # 수리계산 vs 시뮬레이션 교차 검증
    cross_err_B = abs(dp_analytical_B - dp_sim_B_avg) / dp_sim_B_avg * 100 if dp_sim_B_avg else 0
    cross_err_A = abs(dp_analytical_A - dp_sim_A_avg) / dp_sim_A_avg * 100 if dp_sim_A_avg else 0

    max_cross_err = max(cross_err_B, cross_err_A)
    max_indep_err = max(indep_err_B, indep_err_A)
    if max_cross_err < 1.0 and max_indep_err < 1.0:
        status = "PASS"
    elif max_cross_err < 5.0 and max_indep_err < 5.0:
        status = "WARNING"
    else:
        status = "FAIL"
    print(f"    교차 검증 오차: B={cross_err_B:.4f}%, A={cross_err_A:.4f}%")
    print(f"    독립성 오차: B={indep_err_B:.4f}%, A={indep_err_A:.4f}%")
    print(f"    판정: {status}")

    # 최종 ΔP는 시뮬레이션 평균값 사용
    dp_B = dp_sim_B_avg
    dp_A = dp_sim_A_avg

    # [Step 4] NFPC 103 범위 산출
    print("  Step 4: NFPC 103 범위 산출")
    p_min_B = round(0.1 + dp_B, 4)
    p_max_B = round(1.2 + dp_B, 4)
    p_min_A = round(0.1 + dp_A, 4)
    p_max_A = round(1.2 + dp_A, 4)
    p_ref = p_min_B
    print(f"    P_ref = {p_ref:.4f} MPa (= 0.1 + ΔP_B = 0.1 + {dp_B:.4f})")
    print(f"    NFPC 103 범위 B: [{p_min_B:.4f}, {p_max_B:.4f}] MPa")
    print(f"    NFPC 103 범위 A: [{p_min_A:.4f}, {p_max_A:.4f}] MPa")

    # 저장
    char_data = {
        "design_flow_lpm": DESIGN_FLOW_LPM,
        "validation": {
            "delta_p_analytical_B_mpa": round(dp_analytical_B, 6),
            "delta_p_analytical_A_mpa": round(dp_analytical_A, 6),
            "delta_p_simulation_B_mpa": dp_sim_B,
            "delta_p_simulation_A_mpa": dp_sim_A,
            "cross_check_error_B_pct": round(cross_err_B, 4),
            "cross_check_error_A_pct": round(cross_err_A, 4),
            "inlet_independence_error_B_pct": round(indep_err_B, 4),
            "inlet_independence_error_A_pct": round(indep_err_A, 4),
            "status": status,
        },
        "nfpc103_range": {
            "min_terminal_mpa": 0.1,
            "max_terminal_mpa": 1.2,
            "p_min_B_mpa": p_min_B,
            "p_max_B_mpa": p_max_B,
            "p_min_A_mpa": p_min_A,
            "p_max_A_mpa": p_max_A,
        },
        "p_ref_mpa": p_ref,
        "p_ref_basis": "NFPC 103 최소 말단 압력(0.1 MPa) + ΔP_system_B 역산",
        "analytical_detail_B": analytical_B,
        "analytical_detail_A": analytical_A,
    }
    save_system_characterization(char_data)

    # CSV 출력
    validation_rows = [
        {"항목": "ΔP_analytical_B (MPa)", "값": round(dp_analytical_B, 6)},
        {"항목": "ΔP_analytical_A (MPa)", "값": round(dp_analytical_A, 6)},
        {"항목": "ΔP_simulation_B @0.5MPa", "값": dp_sim_B[0]},
        {"항목": "ΔP_simulation_B @1.0MPa", "값": dp_sim_B[1]},
        {"항목": "ΔP_simulation_A @0.5MPa", "값": dp_sim_A[0]},
        {"항목": "ΔP_simulation_A @1.0MPa", "값": dp_sim_A[1]},
        {"항목": "교차 검증 오차 B (%)", "값": round(cross_err_B, 4)},
        {"항목": "교차 검증 오차 A (%)", "값": round(cross_err_A, 4)},
        {"항목": "독립성 오차 B (%)", "값": round(indep_err_B, 4)},
        {"항목": "독립성 오차 A (%)", "값": round(indep_err_A, 4)},
        {"항목": "판정", "값": status},
        {"항목": "P_min_B (MPa)", "값": p_min_B},
        {"항목": "P_max_B (MPa)", "값": p_max_B},
        {"항목": "P_min_A (MPa)", "값": p_min_A},
        {"항목": "P_max_A (MPa)", "값": p_max_A},
        {"항목": "P_ref (MPa)", "값": p_ref},
    ]
    save_csv(pd.DataFrame(validation_rows), data_dir / "part0_5_system_characterization.csv")

    return p_ref


# ── Part 1: 프로그램 검증 ──
def run_part1(data_dir, fig_dir, p_ref):
    """V1~V6 검증 계산"""
    print("\n── Part 1: 프로그램 검증 ──")

    # --- V1: 단일 배관 마찰손실 ---
    print("  V1: 단일 배관 마찰손실")
    branch_flow = DESIGN_FLOW_LPM / NUM_BRANCHES  # 640 LPM
    pipe_id_m = PIPE_DIMS["50A"]["id_mm"] / 1000.0  # 52.51mm → 0.05251m
    V = hydraulics.velocity_from_flow(branch_flow, pipe_id_m)
    Re = hydraulics.reynolds_number(V, pipe_id_m, constants.NU)
    f = hydraulics.friction_factor(Re, constants.EPSILON_M, pipe_id_m)
    h_f = hydraulics.major_loss(f, HEAD_SPACING_M, pipe_id_m, V)
    h_f_mpa = hydraulics.head_to_mpa(h_f, constants.RHO)

    v1_data = [{
        "parameter": "유량 Q (LPM)", "value": branch_flow,
    }, {
        "parameter": "관경 D (mm)", "value": PIPE_DIMS["50A"]["id_mm"],
    }, {
        "parameter": "유속 V (m/s)", "value": round(V, 4),
    }, {
        "parameter": "레이놀즈수 Re", "value": round(Re, 0),
    }, {
        "parameter": "마찰계수 f", "value": round(f, 6),
    }, {
        "parameter": "관 길이 L (m)", "value": HEAD_SPACING_M,
    }, {
        "parameter": "마찰손실 h_f (m)", "value": round(h_f, 6),
    }, {
        "parameter": "마찰손실 (MPa)", "value": round(h_f_mpa, 6),
    }]
    save_csv(pd.DataFrame(v1_data), data_dir / "V1_single_pipe_friction.csv")

    # --- V2: K-factor 부차손실 ---
    print("  V2: K-factor 부차손실")
    k_values = {"K1_BASE (이음쇠)": K1_BASE, "K2 (헤드)": K2,
                "K_equip (장비 총합)": 6.20}
    v2_rows = []
    for name, k in k_values.items():
        h_m = hydraulics.minor_loss(k, V)
        v2_rows.append({
            "항목": name, "K": k,
            "유속 V (m/s)": round(V, 4),
            "부차손실 h_m (m)": round(h_m, 6),
            "부차손실 (MPa)": round(hydraulics.head_to_mpa(h_m, constants.RHO), 6),
        })
    save_csv(pd.DataFrame(v2_rows), data_dir / "V2_minor_loss.csv")

    # --- V3: 비드 K_eff 공식 ---
    print("  V3: 비드 K_eff 공식")
    bead_heights = [0.0, 0.5, 1.0, 1.5, 2.0, 2.5, 3.0]
    pipe_sizes = ["25A", "32A", "40A", "50A"]
    v3_rows = []
    for h_b in bead_heights:
        row = {"bead_height_mm": h_b}
        for ps in pipe_sizes:
            pipe_id_mm = PIPE_DIMS[ps]["id_mm"]
            k_eff = hydraulics.k_welded_fitting(h_b, pipe_id_mm, K1_BASE)
            row[f"K_eff_{ps}"] = round(k_eff, 4)
        v3_rows.append(row)
    df_v3 = pd.DataFrame(v3_rows)
    save_csv(df_v3, data_dir / "V3_K_eff_table.csv")

    # V3 그래프: K_eff 곡선
    fig, ax = plt.subplots(figsize=(8, 5))
    for ps in pipe_sizes:
        ax.plot(bead_heights, df_v3[f"K_eff_{ps}"], "o-", label=ps, markersize=4)
    ax.set_xlabel("비드 높이 (mm)")
    ax.set_ylabel("K_eff")
    ax.set_title("관경별 K_eff vs 비드 높이")
    ax.legend()
    ax.grid(True, alpha=0.3)
    save_fig(fig, fig_dir / "fig_V3_K_eff_curves.png")

    # --- V4: CRANE TP-410 비교 ---
    print("  V4: CRANE TP-410 비교")
    crane_ref = {
        "90° 엘보 (50A)": {"K_calc": 0.75, "K_crane": 0.81, "source": "CRANE TP-410"},
        "티 분기 (50A)":  {"K_calc": 1.0,  "K_crane": 1.01, "source": "CRANE TP-410"},
        "게이트밸브 (전개)": {"K_calc": 0.15, "K_crane": 0.16, "source": "CRANE TP-410"},
        "스윙 체크밸브": {"K_calc": 2.0,  "K_crane": 2.00, "source": "CRANE TP-410"},
    }
    v4_rows = []
    for name, vals in crane_ref.items():
        err = abs(vals["K_calc"] - vals["K_crane"]) / vals["K_crane"] * 100
        v4_rows.append({
            "부속": name,
            "FiPLSim K": vals["K_calc"],
            "CRANE K": vals["K_crane"],
            "오차 (%)": round(err, 2),
            "출처": vals["source"],
        })
    save_csv(pd.DataFrame(v4_rows), data_dir / "V4_crane_comparison.csv")

    # --- V6: MC 수렴성 테스트 ---
    print("  V6: MC 수렴성 테스트")
    mc_ns = [10, 50, 100, 200, 500, 1000, 2000, 5000, 10000]
    v6_rows = []
    for n in mc_ns:
        print(f"    N={n} ...", end=" ", flush=True)
        r = sim.run_dynamic_monte_carlo(
            n_iterations=n,
            min_defects=1, max_defects=3,
            bead_height_mm=BEAD_HEIGHT_MM,
            bead_height_std_mm=BEAD_HEIGHT_STD_MM,
            total_flow_lpm=DESIGN_FLOW_LPM,
            inlet_pressure_mpa=p_ref,
            **common_params_mc(),
        )
        v6_rows.append({
            "N": n,
            "mean_mpa": round(r["mean_pressure"], 6),
            "std_mpa": round(r["std_pressure"], 6),
            "min_mpa": round(r["min_pressure"], 6),
            "max_mpa": round(r["max_pressure"], 6),
            "Pf (%)": round(r["p_below_threshold"], 2),
        })
        print(f"mean={r['mean_pressure']:.4f}")
    df_v6 = pd.DataFrame(v6_rows)
    save_csv(df_v6, data_dir / "V6_mc_convergence.csv")


# ── Part 2: 기준선 정적 비교 ──
def run_part2(data_dir, fig_dir, p_ref):
    """B1~B3 기준선 분석"""
    print("\n── Part 2: 기준선 정적 비교 ──")

    # --- B1: 설계점 비교 (Case A vs B) ---
    print("  B1: Case A/B 설계점 비교")
    comp = pn.compare_dynamic_cases(
        bead_height_existing=BEAD_HEIGHT_MM,
        bead_height_new=0.0,
        inlet_pressure_mpa=p_ref,
        total_flow_lpm=DESIGN_FLOW_LPM,
        beads_per_branch=BEADS_PER_BRANCH,
        **common_params(),
    )

    b1_summary = pd.DataFrame([{
        "항목": "Case A (기존 용접)",
        "말단압력 (MPa)": round(comp["terminal_A_mpa"], 4),
        "배관손실 (MPa)": round(comp["system_A"]["loss_pipe_mpa"], 4),
        "이음쇠손실 (MPa)": round(comp["system_A"]["loss_fitting_mpa"], 4),
        "비드손실 (MPa)": round(comp["system_A"]["loss_bead_mpa"], 4),
        "장비손실 (MPa)": round(comp["system_A"]["equipment_loss_mpa"], 4),
        "PASS/FAIL": "PASS" if comp["pass_fail_A"] else "FAIL",
    }, {
        "항목": "Case B (신공법)",
        "말단압력 (MPa)": round(comp["terminal_B_mpa"], 4),
        "배관손실 (MPa)": round(comp["system_B"]["loss_pipe_mpa"], 4),
        "이음쇠손실 (MPa)": round(comp["system_B"]["loss_fitting_mpa"], 4),
        "비드손실 (MPa)": round(comp["system_B"]["loss_bead_mpa"], 4),
        "장비손실 (MPa)": round(comp["system_B"]["equipment_loss_mpa"], 4),
        "PASS/FAIL": "PASS" if comp["pass_fail_B"] else "FAIL",
    }])
    save_csv(b1_summary, data_dir / "B1_case_AB_comparison.csv")

    # B1 압력 프로파일
    worst_A = comp["case_A"]
    worst_B = comp["case_B"]
    profile_rows = []
    n_pts = min(len(worst_A["positions"]), len(worst_B["positions"]))
    for i in range(n_pts):
        profile_rows.append({
            "position_m": worst_A["positions"][i],
            "pressure_A_mpa": round(worst_A["pressures_mpa"][i], 4),
            "pressure_B_mpa": round(worst_B["pressures_mpa"][i], 4),
        })
    df_profile = pd.DataFrame(profile_rows)
    save_csv(df_profile, data_dir / "B1_pressure_profile.csv")

    # Fig.1: 압력 프로파일 비교
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(worst_A["positions"], worst_A["pressures_mpa"], "r-o",
            label=f"Case A (기존): {comp['terminal_A_mpa']:.3f} MPa", markersize=4)
    ax.plot(worst_B["positions"], worst_B["pressures_mpa"], "b-s",
            label=f"Case B (신공법): {comp['terminal_B_mpa']:.3f} MPa", markersize=4)
    ax.axhline(y=PASS_THRESHOLD_MPA, color="green", linestyle="--", alpha=0.7,
               label=f"기준: {PASS_THRESHOLD_MPA} MPa")
    ax.set_xlabel("배관 위치 (m)")
    ax.set_ylabel("압력 (MPa)")
    ax.set_title(f"최악 경로 압력 프로파일 (Q={DESIGN_FLOW_LPM} LPM, P_ref={p_ref:.3f} MPa)")
    ax.legend()
    ax.grid(True, alpha=0.3)
    save_fig(fig, fig_dir / "fig01_pressure_profile.png")

    # Fig.2: 3항 손실 분리 (적층 막대)
    fig, ax = plt.subplots(figsize=(8, 5))
    cases = ["Case A\n(기존)", "Case B\n(신공법)"]
    pipe_losses = [comp["system_A"]["loss_pipe_mpa"], comp["system_B"]["loss_pipe_mpa"]]
    fit_losses = [comp["system_A"]["loss_fitting_mpa"], comp["system_B"]["loss_fitting_mpa"]]
    bead_losses = [comp["system_A"]["loss_bead_mpa"], comp["system_B"]["loss_bead_mpa"]]
    x = range(len(cases))
    ax.bar(x, pipe_losses, label="배관 마찰", color="#2196F3")
    ax.bar(x, fit_losses, bottom=pipe_losses, label="이음쇠", color="#FF9800")
    b_bottom = [p + f for p, f in zip(pipe_losses, fit_losses)]
    ax.bar(x, bead_losses, bottom=b_bottom, label="비드", color="#F44336")
    ax.set_xticks(x)
    ax.set_xticklabels(cases)
    ax.set_ylabel("압력손실 (MPa)")
    ax.set_title("3항 압력손실 분리")
    ax.legend()
    ax.grid(True, alpha=0.3, axis="y")
    save_fig(fig, fig_dir / "fig02_loss_decomposition.png")

    # --- B2: 유량 스윕 ---
    print("  B2: 유량 스윕 (1200~2600 LPM)")
    b2 = sim.run_variable_sweep(
        sweep_variable="design_flow",
        start_val=1200, end_val=2600, step_val=40,
        inlet_pressure_mpa=p_ref,
        bead_height_mm=BEAD_HEIGHT_MM,
        beads_per_branch=BEADS_PER_BRANCH,
        total_flow_lpm=DESIGN_FLOW_LPM,
        **common_params(),
    )
    b2_rows = []
    for i, q in enumerate(b2["sweep_values"]):
        b2_rows.append({
            "flow_lpm": q,
            "terminal_A_mpa": round(b2["terminal_A"][i], 4),
            "terminal_B_mpa": round(b2["terminal_B"][i], 4),
            "improvement_pct": round(b2["improvement_pct"][i], 2),
            "pass_A": b2["pass_fail_A"][i],
            "pass_B": b2["pass_fail_B"][i],
        })
    df_b2 = pd.DataFrame(b2_rows)
    save_csv(df_b2, data_dir / "B2_flow_sweep.csv")

    # Fig.3: 유량 스윕
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(b2["sweep_values"], b2["terminal_A"], "r-", label="Case A (기존)", linewidth=1.5)
    ax.plot(b2["sweep_values"], b2["terminal_B"], "b-", label="Case B (신공법)", linewidth=1.5)
    ax.axhline(y=PASS_THRESHOLD_MPA, color="green", linestyle="--", alpha=0.7,
               label=f"기준: {PASS_THRESHOLD_MPA} MPa")
    ax.axvline(x=DESIGN_FLOW_LPM, color="gray", linestyle=":", alpha=0.5,
               label=f"설계유량: {DESIGN_FLOW_LPM} LPM")
    ax.set_xlabel("설계 유량 (LPM)")
    ax.set_ylabel("말단 압력 (MPa)")
    ax.set_title(f"유량별 말단 압력 (P_ref={p_ref:.3f} MPa)")
    ax.legend()
    ax.grid(True, alpha=0.3)
    save_fig(fig, fig_dir / "fig03_flow_sweep.png")

    # --- B3: 입구 압력 스윕 ---
    print("  B3: 입구 압력 스윕 (0.2~2.0 MPa)")
    b3 = sim.run_variable_sweep(
        sweep_variable="inlet_pressure",
        start_val=0.2, end_val=2.0, step_val=0.05,
        total_flow_lpm=DESIGN_FLOW_LPM,
        bead_height_mm=BEAD_HEIGHT_MM,
        beads_per_branch=BEADS_PER_BRANCH,
        **common_params(),
    )
    b3_rows = []
    for i, p_in in enumerate(b3["sweep_values"]):
        b3_rows.append({
            "inlet_pressure_mpa": p_in,
            "terminal_A_mpa": round(b3["terminal_A"][i], 4),
            "terminal_B_mpa": round(b3["terminal_B"][i], 4),
            "improvement_pct": round(b3["improvement_pct"][i], 2),
            "pass_A": b3["pass_fail_A"][i],
            "pass_B": b3["pass_fail_B"][i],
        })
    df_b3 = pd.DataFrame(b3_rows)
    save_csv(df_b3, data_dir / "B3_pressure_sweep.csv")

    # Fig.4: 입구 압력 스윕
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(b3["sweep_values"], b3["terminal_A"], "r-", label="Case A (기존)", linewidth=1.5)
    ax.plot(b3["sweep_values"], b3["terminal_B"], "b-", label="Case B (신공법)", linewidth=1.5)
    ax.axhline(y=PASS_THRESHOLD_MPA, color="green", linestyle="--", alpha=0.7)
    ax.axvline(x=p_ref, color="gray", linestyle=":", alpha=0.5, label=f"P_ref={p_ref:.3f}")
    ax.set_xlabel("입구 압력 (MPa)")
    ax.set_ylabel("말단 압력 (MPa)")
    ax.set_title(f"입구 압력별 말단 압력 (Q={DESIGN_FLOW_LPM} LPM)")
    ax.legend()
    ax.grid(True, alpha=0.3)
    save_fig(fig, fig_dir / "fig04_pressure_sweep.png")

    return comp


# ══════════════════════════════════════════════
#  2단계: 파라메트릭 스터디 (Part 3, 4, 5)
# ══════════════════════════════════════════════

def run_stage2(p_ref):
    """2단계: 비드 높이/개수 파라메트릭 + 민감도"""
    print("\n" + "=" * 60)
    print("  2단계: 파라메트릭 스터디 (Part 3 + Part 4 + Part 5)")
    print("=" * 60)
    t0 = time.time()

    data_dir = STAGE2_DIR / "data"
    fig_dir = STAGE2_DIR / "figures"

    run_part3(data_dir, fig_dir, p_ref)
    run_part4(data_dir, fig_dir, p_ref)
    run_part5(data_dir, fig_dir, p_ref)

    generate_stage2_report(data_dir, fig_dir, p_ref)

    print(f"\n✅ 2단계 완료 (소요시간: {elapsed(t0)})")


# ── Part 3: 비드 높이 파라메트릭 ──
def run_part3(data_dir, fig_dir, p_ref):
    """H1~H3"""
    print("\n── Part 3: 비드 높이 파라메트릭 스터디 ──")

    # --- H1: 비드 높이 스윕 (설계점) ---
    print("  H1: 비드 높이 스윕 (0.0~3.0 mm)")
    h1 = sim.run_variable_sweep(
        sweep_variable="bead_height",
        start_val=0.0, end_val=3.0, step_val=0.25,
        total_flow_lpm=DESIGN_FLOW_LPM,
        inlet_pressure_mpa=p_ref,
        beads_per_branch=BEADS_PER_BRANCH,
        **common_params(),
    )
    h1_rows = []
    for i, hb in enumerate(h1["sweep_values"]):
        h1_rows.append({
            "bead_height_mm": hb,
            "terminal_A_mpa": round(h1["terminal_A"][i], 4),
            "terminal_B_mpa": round(h1["terminal_B"][i], 4),
            "improvement_pct": round(h1["improvement_pct"][i], 2),
            "pass_A": h1["pass_fail_A"][i],
            "pass_B": h1["pass_fail_B"][i],
        })
    df_h1 = pd.DataFrame(h1_rows)
    save_csv(df_h1, data_dir / "H1_height_sweep.csv")

    # Fig.5: 비드 높이 스윕 곡선
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(h1["sweep_values"], h1["terminal_A"], "r-o", label="Case A", markersize=4)
    ax.plot(h1["sweep_values"], h1["terminal_B"], "b-s", label="Case B", markersize=4)
    ax.axhline(y=PASS_THRESHOLD_MPA, color="green", linestyle="--", alpha=0.7)
    ax.set_xlabel("비드 높이 (mm)")
    ax.set_ylabel("말단 압력 (MPa)")
    ax.set_title(f"비드 높이별 말단 압력 (Q={DESIGN_FLOW_LPM}, P_ref={p_ref:.3f})")
    ax.legend()
    ax.grid(True, alpha=0.3)
    save_fig(fig, fig_dir / "fig05_height_sweep.png")

    # --- H2: 다중 유량 × 높이 매트릭스 ---
    print("  H2: 다중 유량 × 높이 매트릭스 (5유량 × 13높이)")
    flow_levels = [1200, 1600, 2000, 2560, 3000]
    h2_all_rows = []
    h2_plot_data = {}

    for flow in flow_levels:
        print(f"    Q={flow} LPM ...", end=" ", flush=True)
        r = sim.run_variable_sweep(
            sweep_variable="bead_height",
            start_val=0.0, end_val=3.0, step_val=0.25,
            total_flow_lpm=flow,
            inlet_pressure_mpa=p_ref,
            beads_per_branch=BEADS_PER_BRANCH,
            **common_params(),
        )
        h2_plot_data[flow] = (r["sweep_values"], r["terminal_A"])
        for i, hb in enumerate(r["sweep_values"]):
            h2_all_rows.append({
                "flow_lpm": flow,
                "bead_height_mm": hb,
                "terminal_A_mpa": round(r["terminal_A"][i], 4),
                "terminal_B_mpa": round(r["terminal_B"][i], 4),
                "improvement_pct": round(r["improvement_pct"][i], 2),
            })
        print("완료")

    save_csv(pd.DataFrame(h2_all_rows), data_dir / "H2_multiflow_height.csv")

    # Fig.6: 곡선 가족
    fig, ax = plt.subplots(figsize=(10, 6))
    colors_map = {1200: "blue", 1600: "cyan", 2000: "green", 2560: "orange", 3000: "red"}
    for flow, (heights, terminals) in h2_plot_data.items():
        ax.plot(heights, terminals, "-o", color=colors_map[flow],
                label=f"Q={flow} LPM", markersize=3)
    ax.axhline(y=PASS_THRESHOLD_MPA, color="green", linestyle="--", alpha=0.7)
    ax.set_xlabel("비드 높이 (mm)")
    ax.set_ylabel("말단 압력 (MPa)")
    ax.set_title("다중 유량 × 비드 높이 매트릭스 (Case A)")
    ax.legend()
    ax.grid(True, alpha=0.3)
    save_fig(fig, fig_dir / "fig06_multiflow_height.png")

    # --- H3: 관경별 K_eff (순수 계산) ---
    print("  H3: 관경별 K_eff 테이블")
    bead_heights = [0.0, 0.5, 1.0, 1.5, 2.0, 2.5, 3.0]
    all_pipes = ["25A", "32A", "40A", "50A", "65A", "80A"]
    h3_rows = []
    for h_b in bead_heights:
        row = {"bead_height_mm": h_b}
        for ps in all_pipes:
            pid = PIPE_DIMS[ps]["id_mm"]
            k_eff = hydraulics.k_welded_fitting(h_b, pid, K1_BASE)
            row[f"K_eff_{ps}"] = round(k_eff, 4)
        h3_rows.append(row)
    save_csv(pd.DataFrame(h3_rows), data_dir / "H3_K_eff_all_pipes.csv")

    # Fig.7: K_eff by pipe size
    fig, ax = plt.subplots(figsize=(10, 6))
    df_h3 = pd.DataFrame(h3_rows)
    for ps in all_pipes:
        ax.plot(bead_heights, df_h3[f"K_eff_{ps}"], "o-", label=ps, markersize=4)
    ax.set_xlabel("비드 높이 (mm)")
    ax.set_ylabel("K_eff")
    ax.set_title("관경별 유효 K-factor (K_eff)")
    ax.legend()
    ax.grid(True, alpha=0.3)
    ax.set_yscale("log")
    save_fig(fig, fig_dir / "fig07_K_eff_by_pipe.png")


# ── Part 4: 비드 개수 파라메트릭 ──
def run_part4(data_dir, fig_dir, p_ref):
    """N1~N2"""
    print("\n── Part 4: 비드 개수 파라메트릭 스터디 ──")

    # --- N1: 비드 개수 스윕 ---
    print("  N1: 비드 개수 스윕 (0~20개)")
    n_beads_list = [0, 1, 2, 3, 4, 6, 8, 12, 16, 20]
    n1_rows = []
    for nb in n_beads_list:
        print(f"    n_beads={nb} ...", end=" ", flush=True)
        system = pn.generate_dynamic_system(
            bead_heights_2d=None,
            bead_height_for_weld_mm=BEAD_HEIGHT_MM,
            beads_per_branch=nb,
            total_flow_lpm=DESIGN_FLOW_LPM,
            inlet_pressure_mpa=p_ref,
            branch_inlet_config=BRANCH_INLET_CONFIG,
            **{k: v for k, v in common_params().items()
               if k not in ("equipment_k_factors", "supply_pipe_size",
                            "branch_inlet_config")},
        )
        result = pn.calculate_dynamic_system(
            system,
            equipment_k_factors=EQUIPMENT_K_FACTORS,
            supply_pipe_size=SUPPLY_PIPE_SIZE,
        )
        n1_rows.append({
            "beads_per_branch": nb,
            "total_beads": nb * NUM_BRANCHES,
            "terminal_mpa": round(result["worst_terminal_mpa"], 4),
            "loss_pipe_mpa": round(result["loss_pipe_mpa"], 4),
            "loss_fitting_mpa": round(result["loss_fitting_mpa"], 4),
            "loss_bead_mpa": round(result["loss_bead_mpa"], 4),
            "pass": result["worst_terminal_mpa"] >= PASS_THRESHOLD_MPA,
        })
        print(f"P_term={result['worst_terminal_mpa']:.4f}")
    df_n1 = pd.DataFrame(n1_rows)
    save_csv(df_n1, data_dir / "N1_count_sweep.csv")

    # Fig.8: 비드 개수 막대그래프
    fig, ax = plt.subplots(figsize=(10, 6))
    x = range(len(n_beads_list))
    colors = ["green" if p else "red" for p in df_n1["pass"]]
    ax.bar(x, df_n1["terminal_mpa"], color=colors, alpha=0.7)
    ax.axhline(y=PASS_THRESHOLD_MPA, color="orange", linestyle="--", linewidth=2)
    ax.set_xticks(x)
    ax.set_xticklabels(n_beads_list)
    ax.set_xlabel("가지배관당 비드 개수")
    ax.set_ylabel("말단 압력 (MPa)")
    ax.set_title(f"비드 개수별 말단 압력 (Q={DESIGN_FLOW_LPM}, h_b={BEAD_HEIGHT_MM}mm)")
    save_fig(fig, fig_dir / "fig08_bead_count.png")

    # --- N2: 이음쇠 vs 직관 비드 분리 ---
    print("  N2: 이음쇠 vs 직관 비드 분리 (3조건)")
    n2_cases = [
        {"name": "(a) 이음쇠만", "bead_h": BEAD_HEIGHT_MM, "n_beads": 0},
        {"name": "(b) 직관만",   "bead_h": 0.0,            "n_beads": BEADS_PER_BRANCH},
        {"name": "(c) 복합",     "bead_h": BEAD_HEIGHT_MM, "n_beads": BEADS_PER_BRANCH},
    ]
    n2_rows = []
    for case in n2_cases:
        system = pn.generate_dynamic_system(
            bead_height_for_weld_mm=case["bead_h"],
            beads_per_branch=case["n_beads"],
            total_flow_lpm=DESIGN_FLOW_LPM,
            inlet_pressure_mpa=p_ref,
            branch_inlet_config=BRANCH_INLET_CONFIG,
            **{k: v for k, v in common_params().items()
               if k not in ("equipment_k_factors", "supply_pipe_size",
                            "branch_inlet_config")},
        )
        result = pn.calculate_dynamic_system(
            system, equipment_k_factors=EQUIPMENT_K_FACTORS,
            supply_pipe_size=SUPPLY_PIPE_SIZE,
        )
        n2_rows.append({
            "조건": case["name"],
            "bead_height_mm": case["bead_h"],
            "beads_per_branch": case["n_beads"],
            "terminal_mpa": round(result["worst_terminal_mpa"], 4),
            "loss_pipe_mpa": round(result["loss_pipe_mpa"], 4),
            "loss_fitting_mpa": round(result["loss_fitting_mpa"], 4),
            "loss_bead_mpa": round(result["loss_bead_mpa"], 4),
        })
    save_csv(pd.DataFrame(n2_rows), data_dir / "N2_contribution_split.csv")


# ── Part 5: 민감도 분석 ──
def run_part5(data_dir, fig_dir, p_ref):
    """S1~S2"""
    print("\n── Part 5: 민감도 분석 ──")

    # --- S1: 위치별 단일 비드 민감도 ---
    print("  S1: 위치별 민감도")
    s1 = sim.run_dynamic_sensitivity(
        bead_height_mm=BEAD_HEIGHT_MM,
        total_flow_lpm=DESIGN_FLOW_LPM,
        inlet_pressure_mpa=p_ref,
        **common_params_mc(),
    )

    s1_rows = []
    worst_branch = s1["worst_branch"]
    for i, pos in enumerate(s1["ranking"]):
        delta = s1["deltas"][pos]
        s1_rows.append({
            "rank": i + 1,
            "branch": worst_branch,
            "position": pos,
            "pipe_size": s1["pipe_sizes"][pos] if pos < len(s1["pipe_sizes"]) else "N/A",
            "delta_mpa": round(delta, 6),
            "pressure_mpa": round(s1["single_bead_pressures"][pos], 4),
        })
    df_s1 = pd.DataFrame(s1_rows)
    save_csv(df_s1, data_dir / "S1_sensitivity_rank.csv")

    # Fig.9: 민감도 순위 색상 막대
    fig, ax = plt.subplots(figsize=(12, 6))
    top_n = min(20, len(s1_rows))
    labels = [f"B{r['branch']+1}-H{r['position']+1}" for r in s1_rows[:top_n]]
    deltas = [r["delta_mpa"] for r in s1_rows[:top_n]]
    colors_s1 = plt.cm.Reds(np.linspace(0.3, 1.0, top_n))
    ax.barh(range(top_n), deltas, color=colors_s1)
    ax.set_yticks(range(top_n))
    ax.set_yticklabels(labels)
    ax.invert_yaxis()
    ax.set_xlabel("압력 감소량 ΔP (MPa)")
    ax.set_title(f"위치별 민감도 순위 (상위 {top_n}개)")
    ax.grid(True, alpha=0.3, axis="x")
    save_fig(fig, fig_dir / "fig09_sensitivity_rank.png")

    # --- S2: 다중 높이 민감도 ---
    print("  S2: 다중 높이 민감도")
    height_levels = [0.5, 1.0, 1.5, 2.0, 3.0]
    s2_rows = []
    s2_worst_deltas = {}

    for h_b in height_levels:
        print(f"    h_b={h_b}mm ...", end=" ", flush=True)
        r = sim.run_dynamic_sensitivity(
            bead_height_mm=h_b,
            total_flow_lpm=DESIGN_FLOW_LPM,
            inlet_pressure_mpa=p_ref,
            **common_params_mc(),
        )
        # 최악 가지배관에서 각 위치별 delta
        wb = r["worst_branch"]
        all_deltas = []
        for pos_idx in range(len(r["deltas"])):
            d = r["deltas"][pos_idx]
            all_deltas.append(d)
            s2_rows.append({
                "bead_height_mm": h_b,
                "branch": wb,
                "position": pos_idx,
                "delta_mpa": round(d, 6),
            })
        s2_worst_deltas[h_b] = max(all_deltas) if all_deltas else 0
        print(f"worst_delta={s2_worst_deltas[h_b]:.6f}")

    save_csv(pd.DataFrame(s2_rows), data_dir / "S2_multi_height_sensitivity.csv")

    # Fig.10: 다중 높이 히트맵
    fig, ax = plt.subplots(figsize=(10, 6))
    heatmap_data = np.zeros((len(height_levels), HEADS_PER_BRANCH))
    for h_idx, h_b in enumerate(height_levels):
        subset = [r for r in s2_rows if r["bead_height_mm"] == h_b]
        for r in subset:
            if r["position"] < HEADS_PER_BRANCH:
                heatmap_data[h_idx, r["position"]] = r["delta_mpa"]

    im = ax.imshow(heatmap_data, aspect="auto", cmap="Reds")
    ax.set_xticks(range(HEADS_PER_BRANCH))
    ax.set_xticklabels([f"H{i+1}" for i in range(HEADS_PER_BRANCH)])
    ax.set_yticks(range(len(height_levels)))
    ax.set_yticklabels([f"{h}mm" for h in height_levels])
    ax.set_xlabel("헤드 위치")
    ax.set_ylabel("비드 높이")
    ax.set_title("비드 높이 × 위치 민감도 히트맵 (최악 가지배관)")
    fig.colorbar(im, ax=ax, label="ΔP (MPa)")
    save_fig(fig, fig_dir / "fig10_multi_height_heatmap.png")


# ══════════════════════════════════════════════
#  3단계: 확률론적 분석 (Part 6, 7, 8)
# ══════════════════════════════════════════════

def run_stage3(p_ref):
    """3단계: MC + 2인자 + 설계 시사점"""
    print("\n" + "=" * 60)
    print("  3단계: 확률론적 분석 (Part 6 + Part 7 + Part 8)")
    print("=" * 60)
    t0 = time.time()

    data_dir = STAGE3_DIR / "data"
    fig_dir = STAGE3_DIR / "figures"

    run_part6(data_dir, fig_dir, p_ref)
    run_part7(data_dir, fig_dir, p_ref)
    run_part8(data_dir, fig_dir, p_ref)

    generate_stage3_report(data_dir, fig_dir, p_ref)

    print(f"\n✅ 3단계 완료 (소요시간: {elapsed(t0)})")


# ── Part 6: 몬테카를로 확률론적 분석 ──
def run_part6(data_dir, fig_dir, p_ref):
    """MC1~MC5"""
    print("\n── Part 6: 몬테카를로 확률론적 분석 ──")

    # --- MC1: 전통 MC 기준선 ---
    print("  MC1: 전통 MC 기준선 (N=10,000)")
    mc1 = sim.run_dynamic_monte_carlo(
        n_iterations=MC_ITERATIONS,
        min_defects=1, max_defects=3,
        bead_height_mm=BEAD_HEIGHT_MM,
        bead_height_std_mm=BEAD_HEIGHT_STD_MM,
        total_flow_lpm=DESIGN_FLOW_LPM,
        inlet_pressure_mpa=p_ref,
        **common_params_mc(),
    )
    mc1_summary = pd.DataFrame([{
        "항목": "전통 MC",
        "N": MC_ITERATIONS,
        "min_defects": 1, "max_defects": 3,
        "mean_mpa": round(mc1["mean_pressure"], 4),
        "std_mpa": round(mc1["std_pressure"], 4),
        "min_mpa": round(mc1["min_pressure"], 4),
        "max_mpa": round(mc1["max_pressure"], 4),
        "Pf (%)": round(mc1["p_below_threshold"], 2),
    }])
    save_csv(mc1_summary, data_dir / "MC1_traditional.csv")

    # Fig.11: MC 히스토그램
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.hist(mc1["terminal_pressures"], bins=50, color="#2196F3", alpha=0.7,
            edgecolor="white", density=True)
    ax.axvline(x=PASS_THRESHOLD_MPA, color="red", linestyle="--", linewidth=2,
               label=f"기준: {PASS_THRESHOLD_MPA} MPa")
    ax.axvline(x=mc1["mean_pressure"], color="orange", linestyle="-", linewidth=2,
               label=f"평균: {mc1['mean_pressure']:.4f} MPa")
    ax.set_xlabel("말단 압력 (MPa)")
    ax.set_ylabel("확률 밀도")
    ax.set_title(f"전통 MC 분포 (N={MC_ITERATIONS}, Pf={mc1['p_below_threshold']:.1f}%)")
    ax.legend()
    ax.grid(True, alpha=0.3)
    save_fig(fig, fig_dir / "fig11_mc_histogram.png")

    # --- MC2: 결함 범위 연구 ---
    print("  MC2: 결함 범위 연구 (4 시나리오)")
    defect_scenarios = [(0, 1), (1, 3), (3, 5), (5, 10)]
    mc2_rows = []
    mc2_pressures = {}

    for min_d, max_d in defect_scenarios:
        print(f"    defects=[{min_d},{max_d}] ...", end=" ", flush=True)
        r = sim.run_dynamic_monte_carlo(
            n_iterations=MC_ITERATIONS,
            min_defects=min_d, max_defects=max_d,
            bead_height_mm=BEAD_HEIGHT_MM,
            bead_height_std_mm=BEAD_HEIGHT_STD_MM,
            total_flow_lpm=DESIGN_FLOW_LPM,
            inlet_pressure_mpa=p_ref,
            **common_params_mc(),
        )
        mc2_rows.append({
            "min_defects": min_d, "max_defects": max_d,
            "mean_mpa": round(r["mean_pressure"], 4),
            "std_mpa": round(r["std_pressure"], 4),
            "min_mpa": round(r["min_pressure"], 4),
            "max_mpa": round(r["max_pressure"], 4),
            "Pf (%)": round(r["p_below_threshold"], 2),
        })
        mc2_pressures[f"{min_d}-{max_d}"] = r["terminal_pressures"]
        print(f"Pf={r['p_below_threshold']:.1f}%")

    save_csv(pd.DataFrame(mc2_rows), data_dir / "MC2_defect_range.csv")

    # Fig.12: 결함 범위 히트맵 (간이 박스플롯)
    fig, ax = plt.subplots(figsize=(10, 6))
    bp_data = [mc2_pressures[f"{m}-{x}"] for m, x in defect_scenarios]
    bp_labels = [f"[{m},{x}]" for m, x in defect_scenarios]
    bp = ax.boxplot(bp_data, tick_labels=bp_labels, patch_artist=True)
    colors_bp = ["#4CAF50", "#FF9800", "#F44336", "#9C27B0"]
    for patch, color in zip(bp["boxes"], colors_bp):
        patch.set_facecolor(color)
        patch.set_alpha(0.6)
    ax.axhline(y=PASS_THRESHOLD_MPA, color="red", linestyle="--", alpha=0.7)
    ax.set_xlabel("결함 범위 [min, max]")
    ax.set_ylabel("말단 압력 (MPa)")
    ax.set_title("결함 범위별 말단 압력 분포")
    save_fig(fig, fig_dir / "fig13_defect_range_boxplot.png")

    # --- MC3: 베르누이 p 스윕 ---
    print("  MC3: 베르누이 p 스윕 (0.05~0.90)")
    p_values = [round(0.05 + i * 0.05, 2) for i in range(18)]  # 0.05~0.90
    mc3 = sim.run_bernoulli_sweep(
        p_values=p_values,
        n_iterations=MC_ITERATIONS,
        bead_height_mm=BEAD_HEIGHT_MM,
        bead_height_std_mm=BEAD_HEIGHT_STD_MM,
        total_flow_lpm=DESIGN_FLOW_LPM,
        inlet_pressure_mpa=p_ref,
        **common_params_mc(),
    )

    mc3_rows = []
    summary = mc3["summary"]
    for i, p in enumerate(summary["p_values"]):
        mc3_rows.append({
            "p_bead": p,
            "mean_mpa": round(summary["mean_pressures"][i], 4),
            "std_mpa": round(summary["std_pressures"][i], 4),
            "min_mpa": round(summary["min_pressures"][i], 4),
            "max_mpa": round(summary["max_pressures"][i], 4),
            "Pf (%)": round(summary["pf_percents"][i], 2),
            "expected_beads": round(summary["expected_bead_counts"][i], 1),
            "mean_beads": round(summary["mean_bead_counts"][i], 1),
        })
    save_csv(pd.DataFrame(mc3_rows), data_dir / "MC3_bernoulli_sweep.csv")

    # Fig.14: 베르누이 p 스윕 (이중축)
    fig, ax1 = plt.subplots(figsize=(10, 6))
    ax2 = ax1.twinx()
    ax1.plot(summary["p_values"], summary["mean_pressures"], "b-o",
             label="평균 압력", markersize=4)
    ax2.plot(summary["p_values"], summary["pf_percents"], "r-s",
             label="파괴확률 Pf", markersize=4)
    ax1.axhline(y=PASS_THRESHOLD_MPA, color="green", linestyle="--", alpha=0.5)
    ax1.set_xlabel("비드 발생 확률 p")
    ax1.set_ylabel("평균 말단 압력 (MPa)", color="blue")
    ax2.set_ylabel("파괴확률 Pf (%)", color="red")
    ax1.set_title("베르누이 p 스윕")
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper right")
    ax1.grid(True, alpha=0.3)
    save_fig(fig, fig_dir / "fig14_bernoulli_p_sweep.png")

    # --- MC4: 비균일 sigma 스윕 ---
    print("  MC4: 비균일 σ 스윕 (7개)")
    sigma_values = [0.0, 0.1, 0.2, 0.3, 0.5, 0.7, 1.0]
    mc4_rows = []

    for sigma in sigma_values:
        print(f"    σ={sigma}mm ...", end=" ", flush=True)
        r = sim.run_bernoulli_monte_carlo(
            p_bead=0.5,
            n_iterations=MC_ITERATIONS,
            bead_height_mm=BEAD_HEIGHT_MM,
            bead_height_std_mm=sigma,
            total_flow_lpm=DESIGN_FLOW_LPM,
            inlet_pressure_mpa=p_ref,
            **common_params_mc(),
        )
        mc4_rows.append({
            "sigma_mm": sigma,
            "mean_mpa": round(r["mean_pressure"], 4),
            "std_mpa": round(r["std_pressure"], 4),
            "min_mpa": round(r["min_pressure"], 4),
            "max_mpa": round(r["max_pressure"], 4),
            "Pf (%)": round(r["p_below_threshold"], 2),
        })
        print(f"mean={r['mean_pressure']:.4f}, Pf={r['p_below_threshold']:.1f}%")

    df_mc4 = pd.DataFrame(mc4_rows)
    save_csv(df_mc4, data_dir / "MC4_nonuniform_sigma.csv")

    # Fig.15: Jensen 부등식 (σ vs 평균 압력)
    fig, ax1 = plt.subplots(figsize=(10, 6))
    ax2 = ax1.twinx()
    ax1.plot(df_mc4["sigma_mm"], df_mc4["mean_mpa"], "b-o",
             label="평균 압력", markersize=5)
    ax2.plot(df_mc4["sigma_mm"], df_mc4["Pf (%)"], "r-s",
             label="Pf (%)", markersize=5)
    ax1.set_xlabel("비드 높이 표준편차 σ (mm)")
    ax1.set_ylabel("평균 말단 압력 (MPa)", color="blue")
    ax2.set_ylabel("파괴확률 Pf (%)", color="red")
    ax1.set_title("비균일 비드 효과 (Jensen 부등식)")
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2)
    ax1.grid(True, alpha=0.3)
    save_fig(fig, fig_dir / "fig15_jensen_sigma.png")

    # --- MC5: 베르누이 MC 수렴성 ---
    print("  MC5: 베르누이 MC 수렴성")
    mc5_ns = [100, 500, 1000, 2000, 5000, 10000]
    mc5_rows = []
    for n in mc5_ns:
        print(f"    N={n} ...", end=" ", flush=True)
        r = sim.run_bernoulli_monte_carlo(
            p_bead=0.5,
            n_iterations=n,
            bead_height_mm=BEAD_HEIGHT_MM,
            bead_height_std_mm=BEAD_HEIGHT_STD_MM,
            total_flow_lpm=DESIGN_FLOW_LPM,
            inlet_pressure_mpa=p_ref,
            **common_params_mc(),
        )
        mc5_rows.append({
            "N": n,
            "mean_mpa": round(r["mean_pressure"], 6),
            "std_mpa": round(r["std_pressure"], 6),
            "Pf (%)": round(r["p_below_threshold"], 2),
        })
        print(f"mean={r['mean_pressure']:.4f}")
    save_csv(pd.DataFrame(mc5_rows), data_dir / "MC5_convergence.csv")


# ── Part 7: 2인자 실험계획법 ──
def run_part7(data_dir, fig_dir, p_ref):
    """TF1: p_bead × h_bead 히트맵"""
    print("\n── Part 7: 2인자 실험계획법 ──")

    p_bead_vals = [0.05, 0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.8, 0.9]
    bead_h_vals = [0.5, 1.0, 1.5, 2.0, 2.5, 3.0]

    print(f"  TF1: {len(p_bead_vals)}×{len(bead_h_vals)} = {len(p_bead_vals)*len(bead_h_vals)} 셀 × {MC_ITERATIONS} MC")
    tf1 = sim.run_two_factor_sweep(
        p_bead_values=p_bead_vals,
        bead_height_values=bead_h_vals,
        n_iterations=MC_ITERATIONS,
        bead_height_std_mm=BEAD_HEIGHT_STD_MM,
        total_flow_lpm=DESIGN_FLOW_LPM,
        inlet_pressure_mpa=p_ref,
        **common_params_mc(),
    )

    # Pf 매트릭스 저장 — 원본은 (n_h, n_p), 전치하여 (n_p, n_h)로 변환
    pf_matrix = np.array(tf1["pf_matrix"]).T
    mean_matrix = np.array(tf1["mean_pressure_matrix"]).T

    pf_rows = []
    for i, p in enumerate(p_bead_vals):
        row = {"p_bead": p}
        for j, h in enumerate(bead_h_vals):
            row[f"h_{h}mm_Pf"] = round(pf_matrix[i][j], 2)
        pf_rows.append(row)
    save_csv(pd.DataFrame(pf_rows), data_dir / "TF1_pf_matrix.csv")

    mp_rows = []
    for i, p in enumerate(p_bead_vals):
        row = {"p_bead": p}
        for j, h in enumerate(bead_h_vals):
            row[f"h_{h}mm_mean_mpa"] = round(mean_matrix[i][j], 4)
        mp_rows.append(row)
    save_csv(pd.DataFrame(mp_rows), data_dir / "TF1_mean_pressure.csv")

    # Fig.16: Pf 히트맵
    fig, ax = plt.subplots(figsize=(10, 8))
    im = ax.imshow(pf_matrix, aspect="auto", cmap="RdYlGn_r",
                   vmin=0, vmax=100)
    ax.set_xticks(range(len(bead_h_vals)))
    ax.set_xticklabels([f"{h}" for h in bead_h_vals])
    ax.set_yticks(range(len(p_bead_vals)))
    ax.set_yticklabels([f"{p}" for p in p_bead_vals])
    ax.set_xlabel("비드 높이 (mm)")
    ax.set_ylabel("비드 발생 확률 p")
    ax.set_title("2인자 실험: 파괴확률 Pf (%) 히트맵")
    # 셀 값 표시
    for i in range(len(p_bead_vals)):
        for j in range(len(bead_h_vals)):
            val = pf_matrix[i][j]
            color = "white" if val > 50 else "black"
            ax.text(j, i, f"{val:.1f}", ha="center", va="center",
                    color=color, fontsize=8)
    fig.colorbar(im, ax=ax, label="Pf (%)")
    save_fig(fig, fig_dir / "fig16_twofactor_pf_heatmap.png")

    # Fig.17: 평균 압력 히트맵
    fig, ax = plt.subplots(figsize=(10, 8))
    im = ax.imshow(mean_matrix, aspect="auto", cmap="RdYlGn")
    ax.set_xticks(range(len(bead_h_vals)))
    ax.set_xticklabels([f"{h}" for h in bead_h_vals])
    ax.set_yticks(range(len(p_bead_vals)))
    ax.set_yticklabels([f"{p}" for p in p_bead_vals])
    ax.set_xlabel("비드 높이 (mm)")
    ax.set_ylabel("비드 발생 확률 p")
    ax.set_title("2인자 실험: 평균 말단 압력 (MPa) 히트맵")
    for i in range(len(p_bead_vals)):
        for j in range(len(bead_h_vals)):
            val = mean_matrix[i][j]
            color = "white" if val < 0.15 else "black"
            ax.text(j, i, f"{val:.3f}", ha="center", va="center",
                    color=color, fontsize=7)
    fig.colorbar(im, ax=ax, label="평균 압력 (MPa)")
    save_fig(fig, fig_dir / "fig17_twofactor_pressure_heatmap.png")


# ── Part 8: 설계 시사점 ──
def run_part8(data_dir, fig_dir, p_ref):
    """D1~D3"""
    print("\n── Part 8: 설계 시사점 ──")

    # --- D1: 임계 유량 추출 ---
    print("  D1: 임계 유량 탐색")
    # Part 2 B2와 동일 스윕 결과에서 추출
    b2 = sim.run_variable_sweep(
        sweep_variable="design_flow",
        start_val=1200, end_val=2600, step_val=40,
        inlet_pressure_mpa=p_ref,
        bead_height_mm=BEAD_HEIGHT_MM,
        beads_per_branch=BEADS_PER_BRANCH,
        total_flow_lpm=DESIGN_FLOW_LPM,
        **common_params(),
    )

    # Case A에서 FAIL 전환점 찾기
    critical_A = None
    critical_B = None
    for i, q in enumerate(b2["sweep_values"]):
        if not b2["pass_fail_A"][i] and critical_A is None:
            critical_A = q
        if not b2["pass_fail_B"][i] and critical_B is None:
            critical_B = q

    d1_df = pd.DataFrame([{
        "항목": "Case A 임계유량 (LPM)", "값": critical_A if critical_A else "> 2600",
    }, {
        "항목": "Case B 임계유량 (LPM)", "값": critical_B if critical_B else "> 2600",
    }, {
        "항목": "P_ref (MPa)", "값": p_ref,
    }, {
        "항목": "설계유량 (LPM)", "값": DESIGN_FLOW_LPM,
    }])
    save_csv(d1_df, data_dir / "D1_critical_flow.csv")

    # --- D2: NFPC 103 범위 내 안전 마진 ---
    print("  D2: NFPC 103 범위 내 안전 마진 (6 절대 압력 × 유량 스윕)")
    char_data = load_system_characterization()
    p_min_A = char_data["nfpc103_range"]["p_min_A_mpa"]
    p_min_B = char_data["nfpc103_range"]["p_min_B_mpa"]
    pressure_points = [
        ("P_min_B (P_ref)", p_min_B),
        ("P_min_A", p_min_A),
        ("0.6 MPa", 0.6),
        ("0.8 MPa", 0.8),
        ("1.0 MPa", 1.0),
        ("1.2 MPa", 1.2),
    ]
    d2_rows = []

    for label, p_in in pressure_points:
        p_in = round(p_in, 4)
        print(f"    P={p_in:.4f} MPa ({label}) ...", end=" ", flush=True)
        r = sim.run_variable_sweep(
            sweep_variable="design_flow",
            start_val=1200, end_val=2600, step_val=40,
            inlet_pressure_mpa=p_in,
            bead_height_mm=BEAD_HEIGHT_MM,
            beads_per_branch=BEADS_PER_BRANCH,
            total_flow_lpm=DESIGN_FLOW_LPM,
            **common_params(),
        )
        for i, q in enumerate(r["sweep_values"]):
            d2_rows.append({
                "pressure_label": label,
                "inlet_pressure_mpa": p_in,
                "flow_lpm": q,
                "terminal_A_mpa": round(r["terminal_A"][i], 4),
                "terminal_B_mpa": round(r["terminal_B"][i], 4),
                "pass_A": r["pass_fail_A"][i],
                "pass_B": r["pass_fail_B"][i],
            })
        print("완료")

    save_csv(pd.DataFrame(d2_rows), data_dir / "D2_safety_margin.csv")

    # --- D3: NFPC/NFTC 적합성 판정 ---
    print("  D3: NFPC/NFTC 적합성 판정")
    comp = pn.compare_dynamic_cases(
        bead_height_existing=BEAD_HEIGHT_MM,
        bead_height_new=0.0,
        inlet_pressure_mpa=p_ref,
        total_flow_lpm=DESIGN_FLOW_LPM,
        beads_per_branch=BEADS_PER_BRANCH,
        **common_params(),
    )

    # 유속 기준 확인 (NFTC 103: 가지배관 6 m/s 이하)
    branch_flow = DESIGN_FLOW_LPM / NUM_BRANCHES
    pipe_id_50A = PIPE_DIMS["50A"]["id_mm"] / 1000.0
    V_branch = hydraulics.velocity_from_flow(branch_flow, pipe_id_50A)

    d3_rows = [{
        "기준": "NFPC 103: 말단 수압 ≥ 0.1 MPa",
        "Case A": f"{comp['terminal_A_mpa']:.4f} MPa",
        "판정 A": "적합" if comp["pass_fail_A"] else "부적합",
        "Case B": f"{comp['terminal_B_mpa']:.4f} MPa",
        "판정 B": "적합" if comp["pass_fail_B"] else "부적합",
    }, {
        "기준": "NFTC 103: 가지배관 유속 ≤ 6 m/s",
        "Case A": f"{V_branch:.2f} m/s",
        "판정 A": "적합" if V_branch <= 6.0 else "부적합",
        "Case B": f"{V_branch:.2f} m/s",
        "판정 B": "적합" if V_branch <= 6.0 else "부적합",
    }, {
        "기준": f"설계 유량 = {DESIGN_FLOW_LPM} LPM (32×80)",
        "Case A": f"Q = {DESIGN_FLOW_LPM} LPM",
        "판정 A": "적합",
        "Case B": f"Q = {DESIGN_FLOW_LPM} LPM",
        "판정 B": "적합",
    }]
    save_csv(pd.DataFrame(d3_rows), data_dir / "D3_nfpc_compliance.csv")


# ══════════════════════════════════════════════
#  DOCX 보고서 생성
# ══════════════════════════════════════════════

def _create_docx_report(title, stage_dir, data_dir, fig_dir, sections):
    """python-docx로 DOCX 보고서 생성"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  ⚠ python-docx가 설치되지 않아 DOCX 보고서를 생성할 수 없습니다.")
        print("    설치: pip install python-docx")
        return

    doc = Document()

    # 제목
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"FiPLSim v2.0 | 설계유량: {DESIGN_FLOW_LPM} LPM")
    doc.add_paragraph("")

    # 시뮬레이션 조건 테이블
    doc.add_heading("시뮬레이션 조건", level=1)
    params_table = doc.add_table(rows=1, cols=2)
    params_table.style = "Table Grid"
    hdr = params_table.rows[0].cells
    hdr[0].text = "파라미터"
    hdr[1].text = "값"
    conditions = [
        ("가지배관 수", str(NUM_BRANCHES)),
        ("가지당 헤드 수", str(HEADS_PER_BRANCH)),
        ("헤드 간격", f"{HEAD_SPACING_M} m"),
        ("분기 간격", f"{BRANCH_SPACING_M} m"),
        ("설계 유량", f"{DESIGN_FLOW_LPM} LPM"),
        ("비드 높이", f"{BEAD_HEIGHT_MM} mm"),
        ("비드 높이 σ", f"{BEAD_HEIGHT_STD_MM} mm"),
        ("가지당 비드 수", str(BEADS_PER_BRANCH)),
        ("분기 구성", BRANCH_INLET_CONFIG),
        ("공급배관", SUPPLY_PIPE_SIZE),
        ("관 조도", f"{constants.EPSILON_MM} mm"),
        ("물 밀도", f"{constants.RHO} kg/m³"),
        ("동점성계수", f"{constants.NU:.6e} m²/s"),
        ("MC 반복수", str(MC_ITERATIONS)),
    ]
    for name, val in conditions:
        row = params_table.add_row().cells
        row[0].text = name
        row[1].text = val

    doc.add_paragraph("")

    # 각 섹션 추가
    for section in sections:
        doc.add_heading(section["title"], level=1)

        if "description" in section:
            doc.add_paragraph(section["description"])

        # CSV 데이터를 테이블로
        if "csv_file" in section:
            csv_path = data_dir / section["csv_file"]
            if csv_path.exists():
                df = pd.read_csv(csv_path, encoding="utf-8-sig")
                if len(df) > 0:
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.style = "Table Grid"
                    for j, col in enumerate(df.columns):
                        table.rows[0].cells[j].text = str(col)
                    for _, row_data in df.iterrows():
                        row_cells = table.add_row().cells
                        for j, val in enumerate(row_data):
                            row_cells[j].text = str(val)
                    doc.add_paragraph("")

        # 그래프 삽입
        if "figures" in section:
            for fig_name in section["figures"]:
                fig_path = fig_dir / fig_name
                if fig_path.exists():
                    doc.add_picture(str(fig_path), width=Inches(6.0))
                    doc.add_paragraph("")

    # 저장
    report_path = stage_dir / f"{stage_dir.name}_분석보고서.docx"
    doc.save(str(report_path))
    print(f"  → DOCX 보고서 저장: {report_path.name}")


def generate_stage1_report(data_dir, fig_dir, p_ref):
    """1단계 보고서"""
    print("\n  1단계 DOCX 보고서 생성 중...")
    sections = [
        {
            "title": "Part 0.5: 시스템 특성화 (NFPC 103 역산 + 이중 검증)",
            "description": f"NFPC 103 역산 + 이중 검증으로 시스템 총 손실(ΔP)을 산출하고, P_ref = {p_ref:.4f} MPa를 확정했습니다.",
            "csv_file": "part0_5_system_characterization.csv",
        },
        {
            "title": "Part 1-V1: 단일 배관 마찰손실",
            "description": "Darcy-Weisbach 방정식을 이용한 단일 배관 구간 마찰손실을 검증합니다.",
            "csv_file": "V1_single_pipe_friction.csv",
        },
        {
            "title": "Part 1-V2: K-factor 부차손실",
            "csv_file": "V2_minor_loss.csv",
        },
        {
            "title": "Part 1-V3: 비드 K_eff 공식",
            "description": "K_eff = K_base × (D/D_eff)⁴ 공식의 관경별 계산 결과입니다.",
            "csv_file": "V3_K_eff_table.csv",
            "figures": ["fig_V3_K_eff_curves.png"],
        },
        {
            "title": "Part 1-V4: CRANE TP-410 비교",
            "csv_file": "V4_crane_comparison.csv",
        },
        {
            "title": "Part 1-V6: MC 수렴성 테스트",
            "description": "몬테카를로 시뮬레이션의 반복 횟수에 따른 수렴성을 확인합니다.",
            "csv_file": "V6_mc_convergence.csv",
        },
        {
            "title": "Part 2-B1: Case A/B 설계점 비교",
            "description": f"설계 유량 {DESIGN_FLOW_LPM} LPM, P_ref = {p_ref:.4f} MPa (NFPC 103 역산)에서의 기존 용접(Case A) vs 신공법(Case B) 비교입니다.",
            "csv_file": "B1_case_AB_comparison.csv",
            "figures": ["fig01_pressure_profile.png", "fig02_loss_decomposition.png"],
        },
        {
            "title": "Part 2-B2: 유량 스윕",
            "csv_file": "B2_flow_sweep.csv",
            "figures": ["fig03_flow_sweep.png"],
        },
        {
            "title": "Part 2-B3: 입구 압력 스윕",
            "csv_file": "B3_pressure_sweep.csv",
            "figures": ["fig04_pressure_sweep.png"],
        },
    ]
    _create_docx_report("1단계: 기반 확립 분석 보고서", STAGE1_DIR, data_dir, fig_dir, sections)


def generate_stage2_report(data_dir, fig_dir, p_ref):
    """2단계 보고서"""
    print("\n  2단계 DOCX 보고서 생성 중...")
    sections = [
        {
            "title": "Part 3-H1: 비드 높이 스윕",
            "description": "비드 높이 0.0~3.0mm 범위에서의 말단 압력 변화를 분석합니다.",
            "csv_file": "H1_height_sweep.csv",
            "figures": ["fig05_height_sweep.png"],
        },
        {
            "title": "Part 3-H2: 다중 유량 × 높이 매트릭스",
            "description": "5개 유량 × 13개 높이 = 65개 시뮬레이션 결과입니다.",
            "csv_file": "H2_multiflow_height.csv",
            "figures": ["fig06_multiflow_height.png"],
        },
        {
            "title": "Part 3-H3: 관경별 K_eff",
            "csv_file": "H3_K_eff_all_pipes.csv",
            "figures": ["fig07_K_eff_by_pipe.png"],
        },
        {
            "title": "Part 4-N1: 비드 개수 스윕",
            "description": "가지배관당 비드 개수에 따른 말단 압력 변화를 분석합니다.",
            "csv_file": "N1_count_sweep.csv",
            "figures": ["fig08_bead_count.png"],
        },
        {
            "title": "Part 4-N2: 이음쇠 vs 직관 비드 분리",
            "csv_file": "N2_contribution_split.csv",
        },
        {
            "title": "Part 5-S1: 위치별 민감도 분석",
            "description": "각 접합부 위치에 단일 비드를 배치했을 때의 압력 영향을 순위화합니다.",
            "csv_file": "S1_sensitivity_rank.csv",
            "figures": ["fig09_sensitivity_rank.png"],
        },
        {
            "title": "Part 5-S2: 다중 높이 민감도",
            "csv_file": "S2_multi_height_sensitivity.csv",
            "figures": ["fig10_multi_height_heatmap.png"],
        },
    ]
    _create_docx_report("2단계: 파라메트릭 스터디 분석 보고서", STAGE2_DIR, data_dir, fig_dir, sections)


def generate_stage3_report(data_dir, fig_dir, p_ref):
    """3단계 보고서"""
    print("\n  3단계 DOCX 보고서 생성 중...")
    sections = [
        {
            "title": "Part 6-MC1: 전통 MC 기준선",
            "description": f"결함 1~3개, N={MC_ITERATIONS}회 몬테카를로 시뮬레이션 결과입니다.",
            "csv_file": "MC1_traditional.csv",
            "figures": ["fig11_mc_histogram.png"],
        },
        {
            "title": "Part 6-MC2: 결함 범위 연구",
            "csv_file": "MC2_defect_range.csv",
            "figures": ["fig13_defect_range_boxplot.png"],
        },
        {
            "title": "Part 6-MC3: 베르누이 p 스윕",
            "csv_file": "MC3_bernoulli_sweep.csv",
            "figures": ["fig14_bernoulli_p_sweep.png"],
        },
        {
            "title": "Part 6-MC4: 비균일 σ 스윕",
            "description": "비드 높이 표준편차(σ)에 따른 파괴확률 변화 — Jensen 부등식 효과를 확인합니다.",
            "csv_file": "MC4_nonuniform_sigma.csv",
            "figures": ["fig15_jensen_sigma.png"],
        },
        {
            "title": "Part 6-MC5: 베르누이 MC 수렴성",
            "csv_file": "MC5_convergence.csv",
        },
        {
            "title": "Part 7-TF1: 2인자 실험계획법",
            "description": "p_bead × h_bead 조합에 따른 파괴확률 및 평균 압력 히트맵입니다.",
            "csv_file": "TF1_pf_matrix.csv",
            "figures": ["fig16_twofactor_pf_heatmap.png", "fig17_twofactor_pressure_heatmap.png"],
        },
        {
            "title": "Part 8-D1: 임계 유량",
            "csv_file": "D1_critical_flow.csv",
        },
        {
            "title": "Part 8-D2: NFPC 103 범위 내 안전 마진",
            "csv_file": "D2_safety_margin.csv",
        },
        {
            "title": "Part 8-D3: NFPC/NFTC 적합성 판정",
            "csv_file": "D3_nfpc_compliance.csv",
        },
    ]
    _create_docx_report("3단계: 확률론적 분석 보고서", STAGE3_DIR, data_dir, fig_dir, sections)


# ══════════════════════════════════════════════
#  추가1: 비드 높이 × 입구 압력 2D 스윕
# ══════════════════════════════════════════════

ADD1_DIR = STAGE1_DIR / "추가1"

def run_additional1():
    """비드 높이 × 입구 압력 2D 파라메트릭 스윕 (Case A/B, 전 구간 상세)"""
    print("\n" + "=" * 60)
    print("  추가1: 비드 높이 × 입구 압력 2D 스윕")
    print("  (NFTC 103 가지배관 유속 ≤ 6 m/s 기준 데이터)")
    print("=" * 60)
    t0 = time.time()

    ADD1_DIR.mkdir(parents=True, exist_ok=True)

    # 스윕 범위
    bead_heights = [round(h * 0.25, 2) for h in range(13)]   # 0.0 ~ 3.0 (13개)
    inlet_pressures = [round(0.1 + i * 0.05, 2) for i in range(19)]  # 0.1 ~ 1.0 (19개)
    total = len(bead_heights) * len(inlet_pressures)

    print(f"  비드 높이: {bead_heights[0]}~{bead_heights[-1]} mm ({len(bead_heights)}개)")
    print(f"  입구 압력: {inlet_pressures[0]}~{inlet_pressures[-1]} MPa ({len(inlet_pressures)}개)")
    print(f"  총 조합: {total} 시뮬레이션\n")

    summary_rows = []
    detail_rows = []
    count = 0

    # 히트맵용 2D 배열
    n_h = len(bead_heights)
    n_p = len(inlet_pressures)
    terminal_A_map = np.zeros((n_h, n_p))
    terminal_B_map = np.zeros((n_h, n_p))
    velocity_A_map = np.zeros((n_h, n_p))
    velocity_B_map = np.zeros((n_h, n_p))

    for h_idx, h_b in enumerate(bead_heights):
        for p_idx, p_in in enumerate(inlet_pressures):
            count += 1
            if count % 20 == 1 or count == total:
                print(f"  [{count}/{total}] h_b={h_b}mm, P_in={p_in} MPa ...", flush=True)

            comp = pn.compare_dynamic_cases(
                bead_height_existing=h_b,
                bead_height_new=0.0,
                inlet_pressure_mpa=p_in,
                total_flow_lpm=DESIGN_FLOW_LPM,
                beads_per_branch=BEADS_PER_BRANCH,
                **common_params(),
            )

            # 최악 가지배관의 세그먼트 상세
            worst_A = comp["case_A"]
            worst_B = comp["case_B"]
            segs_A = worst_A["segment_details"]
            segs_B = worst_B["segment_details"]

            # 최대 유속 (첫 구간 = 유량 최대)
            max_v_A = max(s["velocity_ms"] for s in segs_A) if segs_A else 0
            max_v_B = max(s["velocity_ms"] for s in segs_B) if segs_B else 0

            # 히트맵 데이터
            terminal_A_map[h_idx, p_idx] = comp["terminal_A_mpa"]
            terminal_B_map[h_idx, p_idx] = comp["terminal_B_mpa"]
            velocity_A_map[h_idx, p_idx] = max_v_A
            velocity_B_map[h_idx, p_idx] = max_v_B

            # ── 요약 행 ──
            summary_rows.append({
                "bead_height_mm": h_b,
                "inlet_pressure_mpa": p_in,
                "terminal_A_mpa": round(comp["terminal_A_mpa"], 4),
                "terminal_B_mpa": round(comp["terminal_B_mpa"], 4),
                "improvement_pct": round(comp["improvement_pct"], 2),
                "pass_A": comp["pass_fail_A"],
                "pass_B": comp["pass_fail_B"],
                "max_velocity_A_ms": round(max_v_A, 4),
                "max_velocity_B_ms": round(max_v_B, 4),
                "velocity_pass_A": max_v_A <= 6.0,
                "velocity_pass_B": max_v_B <= 6.0,
                "loss_pipe_A_mpa": round(comp["system_A"]["loss_pipe_mpa"], 4),
                "loss_fitting_A_mpa": round(comp["system_A"]["loss_fitting_mpa"], 4),
                "loss_bead_A_mpa": round(comp["system_A"]["loss_bead_mpa"], 4),
                "loss_pipe_B_mpa": round(comp["system_B"]["loss_pipe_mpa"], 4),
                "loss_fitting_B_mpa": round(comp["system_B"]["loss_fitting_mpa"], 4),
                "loss_bead_B_mpa": round(comp["system_B"]["loss_bead_mpa"], 4),
                "equipment_loss_A_mpa": round(comp["system_A"]["equipment_loss_mpa"], 4),
                "equipment_loss_B_mpa": round(comp["system_B"]["equipment_loss_mpa"], 4),
                "cross_main_loss_A_mpa": round(comp["system_A"]["cross_main_cumulative"], 4),
                "cross_main_loss_B_mpa": round(comp["system_B"]["cross_main_cumulative"], 4),
                "worst_branch_A": comp["worst_branch_A"],
                "worst_branch_B": comp["worst_branch_B"],
            })

            # ── 세그먼트 상세 행 (Case A/B × 8 헤드) ──
            for case_label, segs, wb_idx in [
                ("A", segs_A, comp["worst_branch_A"]),
                ("B", segs_B, comp["worst_branch_B"]),
            ]:
                for seg in segs:
                    detail_rows.append({
                        "bead_height_mm": h_b,
                        "inlet_pressure_mpa": p_in,
                        "case": case_label,
                        "branch_idx": wb_idx,
                        "head_number": seg["head_number"],
                        "pipe_size": seg["pipe_size"],
                        "inner_diameter_mm": seg["inner_diameter_mm"],
                        "flow_lpm": round(seg["flow_lpm"], 2),
                        "velocity_ms": round(seg["velocity_ms"], 4),
                        "reynolds": round(seg["reynolds"], 0),
                        "friction_factor": round(seg["friction_factor"], 6),
                        "major_loss_mpa": round(seg["major_loss_mpa"], 6),
                        "K1_value": round(seg["K1_value"], 4),
                        "K1_loss_mpa": round(seg["K1_loss_mpa"], 6),
                        "K2_loss_mpa": round(seg["K2_loss_mpa"], 6),
                        "weld_bead_loss_mpa": round(seg["weld_bead_loss_mpa"], 6),
                        "total_seg_loss_mpa": round(seg["total_seg_loss_mpa"], 6),
                        "pressure_after_mpa": round(seg["pressure_after_mpa"], 4),
                        "bead_height_at_joint_mm": round(seg["bead_height_mm"], 2),
                    })

    # ── CSV 저장 ──
    save_csv(pd.DataFrame(summary_rows), ADD1_DIR / "ADD1_summary.csv")
    save_csv(pd.DataFrame(detail_rows), ADD1_DIR / "ADD1_segment_details.csv")

    # ── 그래프 생성 ──
    _plot_add1_heatmaps(bead_heights, inlet_pressures,
                        terminal_A_map, terminal_B_map,
                        velocity_A_map, velocity_B_map)

    print(f"\n✅ 추가1 완료 (소요시간: {elapsed(t0)})")
    print(f"  결과 위치: {ADD1_DIR}")


def _plot_add1_heatmaps(bead_heights, inlet_pressures,
                        terminal_A, terminal_B,
                        velocity_A, velocity_B):
    """추가1 히트맵 4개 생성"""

    p_labels = [f"{p:.2f}" for p in inlet_pressures]
    h_labels = [f"{h:.2f}" for h in bead_heights]

    # --- Fig: Case A 말단압력 히트맵 ---
    fig, ax = plt.subplots(figsize=(12, 7))
    im = ax.imshow(terminal_A, aspect="auto", cmap="RdYlGn",
                   origin="lower", vmin=0)
    ax.set_xticks(range(len(inlet_pressures)))
    ax.set_xticklabels(p_labels, rotation=45, ha="right", fontsize=7)
    ax.set_yticks(range(len(bead_heights)))
    ax.set_yticklabels(h_labels, fontsize=8)
    ax.set_xlabel("입구 압력 (MPa)")
    ax.set_ylabel("비드 높이 (mm)")
    ax.set_title(f"Case A 말단 압력 (MPa) — Q={DESIGN_FLOW_LPM} LPM")
    # 0.1 MPa 등고선
    cs = ax.contour(terminal_A, levels=[PASS_THRESHOLD_MPA],
                    colors="red", linewidths=2, origin="lower")
    ax.clabel(cs, fmt="%.1f", fontsize=9)
    fig.colorbar(im, ax=ax, label="말단 압력 (MPa)")
    save_fig(fig, ADD1_DIR / "fig_ADD1_terminal_heatmap_A.png")

    # --- Fig: Case B 말단압력 히트맵 ---
    fig, ax = plt.subplots(figsize=(12, 7))
    im = ax.imshow(terminal_B, aspect="auto", cmap="RdYlGn",
                   origin="lower", vmin=0)
    ax.set_xticks(range(len(inlet_pressures)))
    ax.set_xticklabels(p_labels, rotation=45, ha="right", fontsize=7)
    ax.set_yticks(range(len(bead_heights)))
    ax.set_yticklabels(h_labels, fontsize=8)
    ax.set_xlabel("입구 압력 (MPa)")
    ax.set_ylabel("비드 높이 (mm)")
    ax.set_title(f"Case B 말단 압력 (MPa) — Q={DESIGN_FLOW_LPM} LPM")
    cs = ax.contour(terminal_B, levels=[PASS_THRESHOLD_MPA],
                    colors="red", linewidths=2, origin="lower")
    ax.clabel(cs, fmt="%.1f", fontsize=9)
    fig.colorbar(im, ax=ax, label="말단 압력 (MPa)")
    save_fig(fig, ADD1_DIR / "fig_ADD1_terminal_heatmap_B.png")

    # --- Fig: Case A 최대유속 히트맵 ---
    fig, ax = plt.subplots(figsize=(12, 7))
    im = ax.imshow(velocity_A, aspect="auto", cmap="YlOrRd",
                   origin="lower")
    ax.set_xticks(range(len(inlet_pressures)))
    ax.set_xticklabels(p_labels, rotation=45, ha="right", fontsize=7)
    ax.set_yticks(range(len(bead_heights)))
    ax.set_yticklabels(h_labels, fontsize=8)
    ax.set_xlabel("입구 압력 (MPa)")
    ax.set_ylabel("비드 높이 (mm)")
    ax.set_title(f"Case A 최대 유속 (m/s) — NFTC 103 기준 6 m/s")
    # 6 m/s 등고선
    cs = ax.contour(velocity_A, levels=[6.0],
                    colors="blue", linewidths=2, linestyles="--", origin="lower")
    ax.clabel(cs, fmt="%.0f m/s", fontsize=9)
    fig.colorbar(im, ax=ax, label="최대 유속 (m/s)")
    save_fig(fig, ADD1_DIR / "fig_ADD1_velocity_heatmap_A.png")

    # --- Fig: Case B 최대유속 히트맵 ---
    fig, ax = plt.subplots(figsize=(12, 7))
    im = ax.imshow(velocity_B, aspect="auto", cmap="YlOrRd",
                   origin="lower")
    ax.set_xticks(range(len(inlet_pressures)))
    ax.set_xticklabels(p_labels, rotation=45, ha="right", fontsize=7)
    ax.set_yticks(range(len(bead_heights)))
    ax.set_yticklabels(h_labels, fontsize=8)
    ax.set_xlabel("입구 압력 (MPa)")
    ax.set_ylabel("비드 높이 (mm)")
    ax.set_title(f"Case B 최대 유속 (m/s) — NFTC 103 기준 6 m/s")
    cs = ax.contour(velocity_B, levels=[6.0],
                    colors="blue", linewidths=2, linestyles="--", origin="lower")
    ax.clabel(cs, fmt="%.0f m/s", fontsize=9)
    fig.colorbar(im, ax=ax, label="최대 유속 (m/s)")
    save_fig(fig, ADD1_DIR / "fig_ADD1_velocity_heatmap_B.png")


# ══════════════════════════════════════════════
#  메인 실행
# ══════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="FiPLSim 논문용 배치 자동화 스크립트"
    )
    parser.add_argument("--stage", type=int, choices=[1, 2, 3],
                        help="실행할 단계 (생략 시 전체 실행)")
    parser.add_argument("--additional1", action="store_true",
                        help="추가1: 비드 높이 × 입구 압력 2D 스윕")
    args = parser.parse_args()

    print("=" * 60)
    print("  FiPLSim 논문용 시뮬레이션 배치 실행")
    print(f"  설계유량: {DESIGN_FLOW_LPM} LPM (NFPC 103)")
    print(f"  물성치: ε={constants.EPSILON_MM}mm, ρ={constants.RHO}, ν={constants.NU:.3e}")
    print("=" * 60)

    ensure_dirs()
    t_total = time.time()

    # 추가 시뮬레이션 모드
    if args.additional1:
        run_additional1()
        print(f"\n  총 소요시간: {elapsed(t_total)}")
        return

    if args.stage is None or args.stage == 1:
        p_ref = run_stage1()
    else:
        p_ref = load_p_ref()
        print(f"\n  저장된 P_ref = {p_ref:.4f} MPa 로드됨")
        if not verify_system_characterization():
            print("  !! 경고: 파라미터가 변경되었습니다. --stage 1 재실행을 권장합니다.")

    if args.stage is None or args.stage == 2:
        run_stage2(p_ref)

    if args.stage is None or args.stage == 3:
        run_stage3(p_ref)

    print("\n" + "=" * 60)
    print(f"  전체 완료! 총 소요시간: {elapsed(t_total)}")
    print(f"  결과 위치: {OUTPUT_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()
