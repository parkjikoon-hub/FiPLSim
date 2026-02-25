"""
FiPLSim 논문용 '실험방법(Methodology)' 섹션 DOCX 생성 스크립트
SCI 저널 투고용 — 수치해석 이론, 수식, 검증 로직 포함
"""

import io, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def build_methods_docx() -> bytes:
    doc = Document()

    # ── 스타일 ──
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    navy = RGBColor(0x00, 0x00, 0x00)  # 논문은 흑색 기본

    def h1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.font.color.rgb = navy
        return h

    def h2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.font.color.rgb = navy
        return h

    def h3(text):
        h = doc.add_heading(text, level=3)
        for r in h.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.font.color.rgb = navy
        return h

    def para(text, bold=False, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        r.bold = bold
        r.italic = italic
        return p

    def eq(text, label=""):
        """수식 삽입 (가운데 정렬)"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Cambria Math"
        r.font.size = Pt(11)
        r.italic = True
        if label:
            r2 = p.add_run(f"    ({label})")
            r2.font.name = "Times New Roman"
            r2.font.size = Pt(11)
        return p

    def tbl(headers, rows, caption=""):
        if caption:
            pc = doc.add_paragraph()
            rc = pc.add_run(caption)
            rc.font.name = "Times New Roman"
            rc.font.size = Pt(10)
            rc.bold = True
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, hd in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = hd
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.name = "Times New Roman"
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                c = t.rows[ri + 1].cells[ci]
                c.text = str(val)
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.name = "Times New Roman"
        doc.add_paragraph()  # spacing
        return t

    # ═══════════════════════════════════════════
    #  메인 제목
    # ═══════════════════════════════════════════
    title = doc.add_heading(
        "Methodology: Numerical Simulation of Fire Protection Piping Systems", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.font.color.rgb = navy

    doc.add_paragraph()

    # ═══════════════════════════════════════════
    #  1. 시뮬레이션 개발 (Simulation Development)
    # ═══════════════════════════════════════════
    h1("1. Simulation Development")

    # ── 1.1 개발 환경 및 도구 ──
    h2("1.1 Development Environment and Tools")

    para(
        "The numerical simulation platform, designated FiPLSim "
        "(Fire Protection Pipe Let Simulator), was developed using the Python "
        "programming language (version 3.12+) with the following computational "
        "libraries: NumPy (≥1.24) for array operations and random number generation, "
        "SciPy (≥1.11) for cubic spline interpolation and root-finding algorithms, "
        "Pandas (≥2.0) for tabular data handling, and Plotly (≥5.18) for interactive "
        "visualization. The graphical user interface was implemented using Streamlit "
        "(≥1.30), a web-based dashboard framework, enabling real-time parameter "
        "adjustment and result inspection."
    )

    para(
        "All hydraulic calculations are performed in SI units with pressure expressed "
        "in MPa, flow rate in L/min (LPM), pipe dimensions in meters, and velocity "
        "in m/s. Unit conversions are applied internally at the calculation boundaries "
        "to maintain consistency throughout the computation pipeline."
    )

    # ── 1.2 수치해석 이론 ──
    h2("1.2 Governing Equations and Numerical Theories")

    # 1.2.1 Darcy-Weisbach
    h3("1.2.1 Darcy–Weisbach Equation for Major (Friction) Losses")

    para(
        "Pressure loss due to wall friction along straight pipe segments is computed "
        "using the Darcy–Weisbach equation, which is applicable to all flow regimes "
        "and pipe materials without the empirical limitations of Hazen–Williams or "
        "Manning formulas:"
    )

    eq("h_f = f · (L / D) · (V² / 2g)", "1")

    para(
        "where h_f is the friction head loss (m), f is the Darcy friction factor "
        "(dimensionless), L is the pipe segment length (m), D is the internal "
        "diameter (m), V is the mean flow velocity (m/s), and g is the gravitational "
        "acceleration (9.81 m/s²)."
    )

    # 1.2.2 Colebrook-White
    h3("1.2.2 Colebrook–White Equation for Friction Factor")

    para(
        "The Darcy friction factor f is determined by the implicit Colebrook–White "
        "equation, which is valid for the entire turbulent flow regime in commercial "
        "pipes:"
    )

    eq("1/√f = −2.0 · log₁₀ [ (ε/D)/3.7 + 2.51/(Re·√f) ]", "2")

    para(
        "where ε is the absolute pipe wall roughness (m) and Re is the Reynolds "
        "number. For carbon steel pipes used in fire protection systems, ε = 0.045 mm "
        "(4.5 × 10⁻⁵ m) is adopted in accordance with standard engineering references "
        "(Moody, 1944; Colebrook, 1939)."
    )

    para(
        "Since Eq. (2) is implicit in f, it is solved iteratively using a fixed-point "
        "iteration scheme. The initial estimate is obtained from the explicit "
        "Swamee–Jain approximation (Swamee and Jain, 1976):"
    )

    eq("f₀ = 0.25 / { log₁₀ [ (ε/D)/3.7 + 2.51/(Re·√0.02) ] }²", "3")

    para(
        "Convergence is declared when the relative change |f_{n+1} − f_n| / f_n "
        "< 10⁻⁸, typically achieved within 3–5 iterations. For laminar flow (Re < 2300), "
        "the Hagen–Poiseuille solution f = 64/Re is applied directly without iteration."
    )

    # 1.2.3 Reynolds Number
    h3("1.2.3 Reynolds Number and Flow Regime Classification")

    eq("Re = V · D / ν", "4")

    para(
        "where ν is the kinematic viscosity of water at 20°C (1.004 × 10⁻⁶ m²/s). "
        "The flow regime is classified as follows: Re < 2300 (laminar), "
        "2300 ≤ Re ≤ 4000 (transitional), and Re > 4000 (fully turbulent). "
        "Fire protection systems typically operate in the turbulent regime "
        "(Re ≈ 10⁴–10⁵)."
    )

    # 1.2.4 Minor Losses
    h3("1.2.4 Minor (Local) Loss Model")

    para(
        "Localized pressure losses at fittings, junctions, and appurtenances are "
        "computed using the K-factor method:"
    )

    eq("h_m = K · (V² / 2g)", "5")

    para(
        "where K is the dimensionless loss coefficient. The simulator accounts for "
        "the following loss components at each segment of the branch pipe:"
    )

    tbl(
        ["Loss Component", "Symbol", "Value", "Description"],
        [
            ("Welded fitting (junction)", "K₁", "f(h_bead)", "Depends on weld bead protrusion height"),
            ("Sprinkler head orifice", "K₂", "2.5", "Intrinsic head fitting resistance"),
            ("Branch entry (Tee)", "K₃", "1.0", "Cross-main to branch-pipe junction"),
            ("Cross-main Tee-run", "K_tee", "0.3", "Straight-through flow in cross-main"),
        ],
        caption="Table 1. Loss coefficient (K-factor) assignments for each hydraulic component.",
    )

    # 1.2.5 Weld Bead K-factor model
    h3("1.2.5 Weld Bead Protrusion Loss Model")

    para(
        "A critical innovation of this study is the quantitative modeling of "
        "internal weld bead protrusions and their effect on hydraulic resistance. "
        "When pipes are joined by conventional welding, the internal bead "
        "protrudes into the flow cross-section, reducing the effective diameter. "
        "The resulting increase in the local loss coefficient is modeled as:"
    )

    eq("K₁ = K₁,base · (D / D_eff)⁴", "6")

    eq("D_eff = D − 2h", "7")

    para(
        "where K₁,base is the baseline loss coefficient for an ideal joint (0.5), "
        "D is the nominal internal diameter (mm), h is the bead protrusion "
        "height (mm), and D_eff is the effective diameter after accounting for "
        "bilateral bead protrusion."
    )

    para(
        "The fourth-power dependence on the diameter ratio in Eq. (6) is derived "
        "from the Hagen–Poiseuille law for viscous flow through a constriction, "
        "where the flow rate is proportional to D⁴. For turbulent flow through "
        "an orifice-like constriction, the pressure loss scales with the area "
        "ratio squared, i.e., (A/A_eff)² = (D/D_eff)⁴, which is consistent with "
        "the energy equation applied at a sudden contraction (Idelchik, 1986). "
        "This model enables direct comparison between conventional welding "
        "(h > 0, Case A) and shape-controlled welding technology (h ≈ 0, Case B)."
    )

    # 1.2.6 Branch Pressure Profile
    h3("1.2.6 Branch Pipe Pressure Profile Computation")

    para(
        "The cumulative pressure drop along each branch pipe is computed "
        "sequentially from the cross-main junction to the terminal sprinkler head. "
        "At each segment j (j = 1, 2, ..., m), the local flow rate decreases "
        "linearly due to sequential discharge at upstream heads:"
    )

    eq("Q_j = Q_total − (j − 1) · (Q_total / m)", "8")

    para(
        "where Q_total is the total branch flow rate and m is the number of "
        "sprinkler heads. The cumulative pressure at position j is then:"
    )

    eq("P_j = P_{j-1} − (ρg)⁻¹ · 10⁶ · [ h_{f,j} + h_{K1,j} + h_{K2,j} + Σ h_{weld,j} ]", "9")

    para(
        "where P₀ is the branch inlet pressure (after subtracting the K₃ branch "
        "entry loss from the cross-main pressure), and the summation Σh_weld,j "
        "accounts for all weld bead losses within segment j. The terminal "
        "pressure P_m at the last head is the critical value for compliance "
        "verification."
    )

    # 1.2.7 Pipe sizing
    h3("1.2.7 Automatic Pipe Sizing Algorithm")

    para(
        "Pipe diameters are automatically selected based on the number of "
        "downstream sprinkler heads in accordance with NFSC 103 (Korean National "
        "Fire Safety Code) pipe sizing tables. The sizing rules ensure adequate "
        "flow capacity while maintaining economically efficient pipe diameters:"
    )

    tbl(
        ["Downstream Heads", "Branch Pipe Size", "Specification"],
        [
            ("1–2", "25A", "OD 33.40 mm, ID 26.64 mm"),
            ("3", "32A", "OD 42.16 mm, ID 35.04 mm"),
            ("4–5", "40A", "OD 48.26 mm, ID 40.90 mm"),
            ("6–11", "50A", "OD 60.33 mm, ID 52.51 mm"),
            ("≥ 12", "65A", "OD 73.03 mm, ID 62.71 mm"),
        ],
        caption="Table 2. Automatic pipe sizing rules (JIS/KS Schedule 40 carbon steel).",
    )

    tbl(
        ["Total System Heads", "Cross-Main Size", "Internal Diameter (mm)"],
        [
            ("< 20", "65A", "62.71"),
            ("20–39", "80A", "77.92"),
            ("≥ 40", "100A", "102.26"),
        ],
        caption="Table 3. Cross-main pipe sizing based on total system head count.",
    )

    # 1.2.8 Hardy-Cross
    h2("1.3 Grid Network Analysis: Hardy–Cross Method")

    para(
        "For grid (looped) piping configurations with both top and bottom "
        "cross-mains, the flow distribution cannot be determined by simple "
        "sequential calculation. The Hardy–Cross method (Cross, 1936) is "
        "employed to iteratively balance the flow in each closed loop until "
        "energy conservation is satisfied."
    )

    h3("1.3.1 Formulation")

    para(
        "Consider a network with n branches forming n rectangular loops. "
        "For each loop L, the algebraic sum of head losses must equal zero "
        "(Kirchhoff's voltage law analogue):"
    )

    eq("Σ h_{f,i} · sign(Q_i) = 0    for each loop L", "10")

    para(
        "The flow correction for loop L at each iteration is computed as:"
    )

    eq("ΔQ_L = − Σ h_{f,i} / Σ (∂h_{f,i}/∂Q_i)", "11")

    para(
        "For turbulent flow, the head loss–flow relationship is approximately "
        "quadratic (h_f ∝ Q²), yielding the derivative approximation:"
    )

    eq("∂h_f/∂Q ≈ 2 · h_f / |Q|", "12")

    para(
        "Each pipe's flow rate is then updated as:"
    )

    eq("Q_i^(k+1) = Q_i^(k) + ΔQ_L · d_i · ω", "13")

    para(
        "where d_i is the direction coefficient (+1 or −1) indicating the "
        "pipe's orientation within the loop, ω is the under-relaxation factor "
        "(default 0.5, adjustable range 0.1–1.0), and k is the iteration index."
    )

    h3("1.3.2 Convergence Criteria")

    para(
        "The iterative process is terminated when both of the following "
        "dual convergence criteria are simultaneously satisfied:"
    )

    tbl(
        ["Criterion", "Tolerance", "Physical Meaning"],
        [
            ("max |Σh_f| over all loops", "< 0.001 m (≈ 0.01 kPa)", "Loop energy imbalance"),
            ("max |ΔQ| over all loops", "< 0.0001 LPM", "Flow correction magnitude"),
        ],
        caption="Table 4. Hardy–Cross dual convergence criteria.",
    )

    para(
        "The maximum number of iterations is set to 1,000. To detect "
        "divergence, the algorithm monitors whether the loop imbalance "
        "increases for three consecutive iterations; if so, the computation "
        "is terminated and a warning is issued. This dual convergence "
        "criterion ensures both energy balance accuracy (pressure) and "
        "mass conservation (flow) are satisfied to engineering precision."
    )

    h3("1.3.3 Node Pressure Computation via BFS")

    para(
        "After flow convergence, node pressures are computed using a "
        "Breadth-First Search (BFS) traversal starting from the inlet node "
        "(known pressure boundary condition). For each pipe connecting "
        "nodes i and j with known pressure P_i, the downstream pressure is:"
    )

    eq("P_j = P_i − Δp_{i→j}(Q_{pipe})", "14")

    para(
        "where Δp is the pressure drop computed from the converged flow rate "
        "using the Darcy–Weisbach equation with all applicable minor loss "
        "coefficients. This ensures consistent pressure distribution "
        "throughout the grid network."
    )

    # 1.2.9 Pump analysis
    h2("1.4 Pump Performance Analysis")

    h3("1.4.1 Pump Curve Interpolation")

    para(
        "Manufacturer-provided pump performance data (Q–H points) are "
        "interpolated using a cubic spline (SciPy interp1d, kind='cubic') "
        "to generate a continuous pump characteristic curve H_pump(Q). "
        "This approach preserves the physical shape of the pump curve, "
        "including the shut-off head and end-of-curve behavior, without "
        "requiring an assumed polynomial order."
    )

    h3("1.4.2 System Resistance Curve")

    para(
        "The system resistance curve H_sys(Q) is constructed by computing "
        "the total head loss across the piping network at multiple flow rates. "
        "For each test flow Q_test, the full network simulation is executed "
        "and the total head loss is determined as:"
    )

    eq("H_sys(Q) = [P_inlet − P_terminal(Q)] · 10⁶ / (ρg) + H_min", "15")

    para(
        "where H_min is the minimum terminal head requirement "
        "(corresponding to 0.1 MPa per design code)."
    )

    h3("1.4.3 Operating Point Determination")

    para(
        "The operating point is found at the intersection of the pump curve "
        "and the system resistance curve, i.e., the flow Q* satisfying:"
    )

    eq("H_pump(Q*) − H_sys(Q*) = 0", "16")

    para(
        "This root-finding problem is solved using the Brent method "
        "(scipy.optimize.brentq) with a tolerance of ±0.1 LPM, which "
        "guarantees convergence for any continuous, monotonically diverging "
        "pump and system curves within the bracketed interval. "
        "The corresponding pump power consumption is then computed as:"
    )

    eq("P_pump = ρ · g · Q* · H* / η", "17")

    para(
        "where η is the pump efficiency (decimal fraction) and Q* is "
        "expressed in m³/s."
    )

    # ── 1.3 신뢰성 로직 ──
    h2("1.5 Reliability and Accuracy Assurance")

    h3("1.5.1 Input Validation")

    para(
        "All user-supplied parameters undergo boundary checking before "
        "computation. Invalid configurations (e.g., zero diameter, negative "
        "pressure, or head count exceeding physical limits) raise a "
        "ValidationError exception, preventing the propagation of "
        "nonsensical results."
    )

    h3("1.5.2 Friction Factor Convergence Control")

    para(
        "The Colebrook–White iteration employs a relative convergence "
        "criterion of 10⁻⁸ with a maximum of 10 iterations. Since the "
        "Colebrook–White function is contractive in the turbulent regime, "
        "convergence is theoretically guaranteed (Brkić, 2011). The "
        "Swamee–Jain initial estimate provides an excellent starting "
        "point, typically reducing the required iterations to 3–5."
    )

    h3("1.5.3 Hardy–Cross Stability Measures")

    para(
        "Three mechanisms are implemented to ensure robust convergence "
        "of the Hardy–Cross solver:"
    )

    p1 = doc.add_paragraph(style="List Bullet")
    p1.add_run(
        "Under-relaxation (ω = 0.5 default): Prevents oscillatory "
        "divergence in strongly coupled loops by damping the flow correction."
    ).font.size = Pt(11)

    p2 = doc.add_paragraph(style="List Bullet")
    p2.add_run(
        "Divergence detection: If the maximum loop imbalance increases "
        "for three consecutive iterations, the solver terminates early "
        "and reports a divergence warning rather than producing unreliable results."
    ).font.size = Pt(11)

    p3 = doc.add_paragraph(style="List Bullet")
    p3.add_run(
        "Dual convergence criteria: Both energy balance (head) and mass "
        "conservation (flow) must be satisfied simultaneously, preventing "
        "premature termination when only one criterion is met."
    ).font.size = Pt(11)

    h3("1.5.4 NFPC Code Compliance Verification")

    para(
        "Post-simulation, an automated compliance check verifies the "
        "results against the National Fire Protection Code (NFPC) "
        "requirements:"
    )

    tbl(
        ["Criterion", "Requirement", "Verification Method"],
        [
            ("Branch pipe velocity", "≤ 6.0 m/s", "Check all segment velocities"),
            ("Cross-main velocity", "≤ 10.0 m/s", "Compute V = Q/(πD²/4) for cross-main"),
            ("Terminal pressure (min)", "≥ 0.1 MPa", "Check all branch terminal pressures"),
            ("Terminal pressure (max)", "≤ 1.2 MPa", "Check all branch terminal pressures"),
        ],
        caption="Table 5. NFPC compliance verification criteria.",
    )

    h3("1.5.5 Regression Test Suite")

    para(
        "The simulation engine is validated by a comprehensive regression "
        "test suite comprising 130 unit and integration tests: 46 tests for "
        "grid network calculations (Hardy–Cross convergence, node pressure "
        "consistency, Kirchhoff law verification), 46 tests for integrated "
        "system behavior (Case A/B comparison, improvement percentage, "
        "pipe sizing), and 38 tests for weld bead loss model accuracy "
        "(K-factor scaling, random bead placement, segment mapping). All "
        "130 tests must pass before any code modification is accepted, "
        "ensuring regression-free development."
    )

    # ═══════════════════════════════════════════
    #  2. 시뮬레이션 방법 (Simulation Methods)
    # ═══════════════════════════════════════════
    doc.add_page_break()
    h1("2. Simulation Methods")

    h2("2.1 Network Topology Configurations")

    para(
        "Two piping network topologies are supported to represent the "
        "range of configurations encountered in actual fire protection "
        "installations:"
    )

    h3("2.1.1 Tree (Branch) Topology")

    para(
        "A single cross-main pipe supplies n branch pipes, each serving "
        "m sprinkler heads. Flow proceeds unidirectionally from the "
        "cross-main inlet through each branch to the terminal head. "
        "Pressure is computed sequentially along each branch using "
        "Eqs. (1)–(9). This topology represents the most common "
        "installation pattern in small- to medium-scale buildings."
    )

    h3("2.1.2 Full Grid (Looped) Topology")

    para(
        "Both the top and bottom ends of the branch pipes are connected "
        "to cross-mains, forming n closed rectangular loops. Flow is "
        "bidirectional, entering each branch from both ends. The "
        "Hardy–Cross iterative method (Section 1.3) is employed to "
        "determine the equilibrium flow distribution. This topology "
        "provides hydraulic redundancy and is mandated for large-scale "
        "or high-risk occupancies."
    )

    h2("2.2 Comparative Case Analysis (Case A vs. Case B)")

    para(
        "Each simulation run produces a paired comparison between two "
        "welding technology scenarios under identical boundary conditions:"
    )

    tbl(
        ["Parameter", "Case A (Conventional)", "Case B (Shape-Controlled)"],
        [
            ("Weld bead height h", "User-defined (e.g., 1.5 mm)", "0 mm (ideal joint)"),
            ("Fitting loss K₁", "K₁,base · (D/D_eff)⁴", "K₁,base = 0.5"),
            ("Physical meaning", "Existing welding technology", "Advanced shape-control technology"),
        ],
        caption="Table 6. Comparative case definitions.",
    )

    para(
        "The improvement metric is computed as the percentage change in "
        "the worst-case terminal pressure:"
    )

    eq("Improvement (%) = (P_B − P_A) / |P_A| × 100", "18")

    h2("2.3 Monte Carlo Probabilistic Analysis")

    h3("2.3.1 Stochastic Defect Modeling")

    para(
        "To assess the probabilistic impact of weld defect location "
        "variability, a Monte Carlo simulation framework is implemented. "
        "In each trial i (i = 1, 2, ..., N):"
    )

    p1 = doc.add_paragraph(style="List Number")
    r1 = p1.add_run(
        "The number of defective fittings n_d is sampled uniformly from "
        "[n_min, n_max]."
    )
    r1.font.size = Pt(11)

    p2 = doc.add_paragraph(style="List Number")
    r2 = p2.add_run(
        "The n_d defect positions are selected without replacement from "
        "the set of all fitting locations {0, 1, ..., n×m − 1} using "
        "NumPy's default_rng random generator (PCG64 algorithm)."
    )
    r2.font.size = Pt(11)

    p3 = doc.add_paragraph(style="List Number")
    r3 = p3.add_run(
        "If straight-pipe weld beads are enabled (beads_per_branch > 0), "
        "the bead positions are re-randomized within each branch using "
        "a uniform distribution over the branch length, simulating "
        "construction variability."
    )
    r3.font.size = Pt(11)

    p4 = doc.add_paragraph(style="List Number")
    r4 = p4.add_run(
        "The full system hydraulic calculation is executed, and the "
        "worst-case terminal pressure P_terminal,i is recorded."
    )
    r4.font.size = Pt(11)

    h3("2.3.2 Statistical Output Metrics")

    para(
        "The following statistics are computed from the N terminal "
        "pressure samples {P₁, P₂, ..., P_N}:"
    )

    eq("μ = (1/N) · Σ P_i", "19")

    eq("σ = √[ Σ(P_i − μ)² / (N − 1) ]", "20")

    para(
        "Note that the sample standard deviation (Bessel-corrected, "
        "N − 1 denominator) is used rather than the population standard "
        "deviation to provide an unbiased estimate. The failure "
        "probability P_fail is computed as the empirical proportion "
        "of trials where the terminal pressure falls below the "
        "regulatory minimum:"
    )

    eq("P_fail = (1/N) · Σ I(P_i < 0.1 MPa)", "21")

    para(
        "where I(·) is the indicator function. For N = 1,000 iterations, "
        "this estimator has a standard error of at most √[0.5 × 0.5 / 1000] "
        "= 0.016, providing adequate precision for engineering decision-making. "
        "For higher precision, N can be increased up to 10,000."
    )

    h2("2.4 Deterministic Sensitivity Analysis")

    para(
        "A single-bead perturbation analysis is performed to identify "
        "the most hydraulically sensitive positions within the piping "
        "network. The procedure is as follows:"
    )

    p1 = doc.add_paragraph(style="List Number")
    r1 = p1.add_run(
        "Baseline computation: The system is simulated with no fitting "
        "bead protrusions (K₁ = K₁,base for all fittings), yielding "
        "the baseline terminal pressure P_base."
    )
    r1.font.size = Pt(11)

    p2 = doc.add_paragraph(style="List Number")
    r2 = p2.add_run(
        "For each head position j (j = 1, ..., m) in the worst-case "
        "branch, a single bead of height h is placed at position j "
        "while all other positions remain at h = 0."
    )
    r2.font.size = Pt(11)

    p3 = doc.add_paragraph(style="List Number")
    r3 = p3.add_run(
        "The pressure drop due to the single bead is computed as "
        "Δp_j = P_base − P_j."
    )
    r3.font.size = Pt(11)

    p4 = doc.add_paragraph(style="List Number")
    r4 = p4.add_run(
        "Positions are ranked by Δp_j in descending order. The position "
        "with the largest Δp is designated the critical point."
    )
    r4.font.size = Pt(11)

    para(
        "This analysis reveals which joint locations have the greatest "
        "leverage on system performance, providing actionable guidance "
        "for quality control prioritization during construction."
    )

    h2("2.5 Variable Sweep (Parametric Study)")

    para(
        "To identify the system's operational envelope and failure "
        "boundaries, a variable sweep analysis systematically varies "
        "a single design parameter while holding all other parameters "
        "constant. The sweep procedure executes a full Case A/B "
        "comparison at each parameter value and records:"
    )

    p1 = doc.add_paragraph(style="List Bullet")
    r1 = p1.add_run(
        "The worst-case terminal pressure for both Case A and Case B."
    )
    r1.font.size = Pt(11)

    p2 = doc.add_paragraph(style="List Bullet")
    r2 = p2.add_run(
        "The improvement percentage."
    )
    r2.font.size = Pt(11)

    p3 = doc.add_paragraph(style="List Bullet")
    r3 = p3.add_run(
        "The PASS/FAIL status against the 0.1 MPa minimum threshold."
    )
    r3.font.size = Pt(11)

    para(
        "The critical point Q_crit (or P_crit, h_crit, etc.) is defined "
        "as the parameter value at which the system first transitions "
        "from PASS to FAIL. The following design parameters are available "
        "for sweep analysis: design flow rate (LPM), inlet pressure (MPa), "
        "weld bead height (mm), and number of heads per branch."
    )

    h2("2.6 Economic Analysis (Life-Cycle Cost)")

    para(
        "The energy savings resulting from reduced hydraulic resistance "
        "(Case B vs. Case A) are quantified through a Life-Cycle Cost (LCC) "
        "analysis. At each operating point, the pump power consumption is "
        "computed using Eq. (17), and the differential power is:"
    )

    eq("ΔP = P_pump,A − P_pump,B", "22")

    eq("Annual Energy Savings = ΔP × t_op  (kWh/year)", "23")

    eq("Annual Cost Savings = ΔP × t_op × C_e  (KRW/year)", "24")

    para(
        "where t_op is the annual pump operating hours (default 2,000 h/yr) "
        "and C_e is the unit electricity cost (default 120 KRW/kWh). "
        "These default values can be adjusted by the user to reflect "
        "site-specific conditions."
    )

    # ── 물성치 요약 ──
    doc.add_page_break()
    h2("2.7 Physical Properties and Constants")

    tbl(
        ["Property", "Symbol", "Value", "Unit", "Source"],
        [
            ("Water density (20°C)", "ρ", "998.0", "kg/m³", "Engineering data"),
            ("Kinematic viscosity (20°C)", "ν", "1.004 × 10⁻⁶", "m²/s", "Engineering data"),
            ("Dynamic viscosity (20°C)", "μ", "1.002 × 10⁻³", "Pa·s", "Engineering data"),
            ("Gravitational acceleration", "g", "9.81", "m/s²", "—"),
            ("Pipe roughness (carbon steel)", "ε", "0.045", "mm", "Moody (1944)"),
            ("Base fitting loss coefficient", "K₁,base", "0.5", "—", "Idelchik (1986)"),
            ("Sprinkler head loss coefficient", "K₂", "2.5", "—", "NFSC 103"),
            ("Branch entry loss coefficient", "K₃", "1.0", "—", "NFSC 103"),
            ("Min. terminal pressure", "P_min", "0.1", "MPa", "NFPC"),
            ("Max. terminal pressure", "P_max", "1.2", "MPa", "NFPC"),
            ("Max. branch pipe velocity", "V_max,br", "6.0", "m/s", "NFPC"),
            ("Max. cross-main velocity", "V_max,cm", "10.0", "m/s", "NFPC"),
        ],
        caption="Table 7. Physical properties and design constants used in the simulation.",
    )

    # ── 참고문헌 ──
    doc.add_page_break()
    h1("References")

    refs = [
        "Brkić, D. (2011). Review of explicit approximations to the Colebrook relation for flow friction. "
        "Journal of Petroleum Science and Engineering, 77(1), 34–48.",

        "Colebrook, C. F. (1939). Turbulent flow in pipes, with particular reference to the transition "
        "region between the smooth and rough pipe laws. Journal of the Institution of Civil Engineers, "
        "11(4), 133–156.",

        "Cross, H. (1936). Analysis of flow in networks of conduits or conductors. "
        "Bulletin No. 286, University of Illinois Engineering Experiment Station.",

        "Idelchik, I. E. (1986). Handbook of Hydraulic Resistance, 2nd ed. Hemisphere Publishing.",

        "Moody, L. F. (1944). Friction factors for pipe flow. Transactions of the ASME, 66(8), 671–684.",

        "Swamee, P. K., & Jain, A. K. (1976). Explicit equations for pipe-flow problems. "
        "Journal of the Hydraulics Division, 102(5), 657–664.",

        "National Fire Protection Code (NFPC), Republic of Korea.",

        "National Fire Safety Code 103 (NFSC 103): Sprinkler system design standards, "
        "Republic of Korea.",
    ]

    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

    # ── 저장 ──
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == "__main__":
    data = build_methods_docx()
    out_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "FiPLSim_Paper_Methodology.docx",
    )
    with open(out_path, "wb") as f:
        f.write(data)
    print(f"문서 생성 완료: {out_path}")
    print(f"파일 크기: {len(data):,} bytes")
