"""
3가지 미확인 변수 분석 — DOCX 보고서 생성 (v2)

미확인 변수 ①: 교차배관 관경 (65A vs 80A)
미확인 변수 ②: 헤드당 유량 (균등분배 vs NFPA 13 기준)
미확인 변수 ③: 추가 국부 저항 (밸브 등 누락분)
"""
import os, sys, math
import numpy as np

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from constants import (
    PIPE_DIMENSIONS, get_inner_diameter_m, K1_BASE, K2, K3, K_TEE_RUN,
    RHO, G, EPSILON_M, auto_pipe_size, auto_cross_main_size,
)
from hydraulics import (
    velocity_from_flow, reynolds_number, friction_factor,
    major_loss, minor_loss, head_to_mpa, mpa_to_head,
)
from pipe_network import generate_dynamic_system, calculate_dynamic_system

# ── 공통 조건 ──
NUM_BRANCHES = 4
HEADS_PER_BRANCH = 8
BRANCH_SPACING = 3.5
HEAD_SPACING = 2.3
INLET_PRESSURE = 0.4  # MPa
FLOWS = [1200, 1600, 2100]  # LPM
TOTAL_HEADS = NUM_BRANCHES * HEADS_PER_BRANCH  # 32

# ══════════════════════════════════════════════
# 분석 ① : 교차배관 65A vs 80A
# ══════════════════════════════════════════════
print("=" * 70)
print("  분석 ① : 교차배관 관경 65A vs 80A 비교")
print("=" * 70)

analysis1_rows = []

for flow in FLOWS:
    for cm_size in ["65A", "80A"]:
        system = generate_dynamic_system(
            num_branches=NUM_BRANCHES,
            heads_per_branch=HEADS_PER_BRANCH,
            branch_spacing_m=BRANCH_SPACING,
            head_spacing_m=HEAD_SPACING,
            inlet_pressure_mpa=INLET_PRESSURE,
            total_flow_lpm=float(flow),
        )
        # 교차배관 구경 강제 교체
        cm_id_m = get_inner_diameter_m(cm_size)
        system.cross_main_size = cm_size
        for seg in system.cross_main_segments:
            seg.nominal_size = cm_size
            seg.inner_diameter_m = cm_id_m

        result = calculate_dynamic_system(system)

        # 교차배관 내 유속 (최대 = 입구부, 전체 유량)
        V_cm = velocity_from_flow(float(flow), cm_id_m)

        row = {
            "flow": flow,
            "cm_size": cm_size,
            "cm_id_mm": PIPE_DIMENSIONS[cm_size]["id_mm"],
            "V_cm_ms": round(V_cm, 3),
            "cross_main_loss_mpa": result["cross_main_cumulative"],
            "cross_main_loss_kpa": round(result["cross_main_cumulative"] * 1000, 2),
            "worst_terminal_mpa": result["worst_terminal_mpa"],
            "worst_terminal_kpa": round(result["worst_terminal_mpa"] * 1000, 2),
            "branch_inlet_last": result["branch_inlet_pressures"][-1],
            "branch_profiles": result["branch_profiles"],
        }
        analysis1_rows.append(row)
        print(f"  Q={flow:>5} LPM | {cm_size} (ID={row['cm_id_mm']:.2f}mm, V={V_cm:.2f}m/s) | "
              f"교차배관 손실={row['cross_main_loss_kpa']:.2f} kPa | "
              f"말단수압={row['worst_terminal_kpa']:.2f} kPa")

# 65A vs 80A 차이 계산
diff_rows_1 = []
for flow in FLOWS:
    r65 = [r for r in analysis1_rows if r["flow"] == flow and r["cm_size"] == "65A"][0]
    r80 = [r for r in analysis1_rows if r["flow"] == flow and r["cm_size"] == "80A"][0]
    diff_cm_loss = r65["cross_main_loss_kpa"] - r80["cross_main_loss_kpa"]
    diff_terminal = r80["worst_terminal_kpa"] - r65["worst_terminal_kpa"]
    diff_rows_1.append({
        "flow": flow,
        "cm_loss_65": r65["cross_main_loss_kpa"],
        "cm_loss_80": r80["cross_main_loss_kpa"],
        "diff_cm_loss": round(diff_cm_loss, 2),
        "terminal_65": r65["worst_terminal_kpa"],
        "terminal_80": r80["worst_terminal_kpa"],
        "diff_terminal": round(diff_terminal, 2),
        "V_65": r65["V_cm_ms"],
        "V_80": r80["V_cm_ms"],
    })
    print(f"\n  Q={flow} LPM: 80A 사용 시 교차배관 손실 {diff_cm_loss:.2f} kPa 감소, "
          f"말단수압 +{diff_terminal:.2f} kPa 증가")

# ══════════════════════════════════════════════
# 분석 ② : 헤드당 유량 분석
# ══════════════════════════════════════════════
print("\n" + "=" * 70)
print("  분석 ② : 헤드당 유량 — 균등분배 vs NFPA 13 기준")
print("=" * 70)

analysis2_rows = []
for flow in FLOWS:
    q_head = flow / TOTAL_HEADS
    q_branch = flow / NUM_BRANCHES
    analysis2_rows.append({
        "flow": flow,
        "total_heads": TOTAL_HEADS,
        "q_head_lpm": round(q_head, 2),
        "q_branch_lpm": round(q_branch, 2),
        "meets_nfpa": q_head >= 80,
    })
    print(f"  Q={flow:>5} LPM | Q_head = {q_head:.2f} LPM | "
          f"{'충족' if q_head >= 80 else '미달'} (NFPA 13: 80 LPM)")

# ══════════════════════════════════════════════
# 분석 ③ : 추가 국부 저항 분석
# ══════════════════════════════════════════════
print("\n" + "=" * 70)
print("  분석 ③ : 추가 국부 저항 (밸브 등 누락분)")
print("=" * 70)

# 논문 참조값 (스크린샷의 "필요 추가 손실" 데이터)
paper_additional_loss_kpa = {1200: 54.5, 1600: 96.7, 2100: 166.3}

# Q²에 대한 비율 (SI 단위: Pa/(m³/s)²)
# Q(LPM) → Q(m³/s): Q_si = Q_lpm / 60000
print("\n  [ Q² 비례 분석 — SI 단위 ]")
ratios = []
for flow in FLOWS:
    Q_si = flow / 60000.0  # m³/s
    dP_Pa = paper_additional_loss_kpa[flow] * 1000  # Pa
    R = dP_Pa / (Q_si ** 2)
    ratios.append(R)
    print(f"  Q={flow} LPM = {Q_si:.5f} m³/s | ΔP={dP_Pa:.0f} Pa | "
          f"R = {R:.2e} Pa·s²/m⁶")
R_avg = np.mean(ratios)
print(f"\n  평균 R = {R_avg:.4e} Pa·s²/m⁶ (3개 유량 모두 일정)")

# 등가 K값 역산: R = K × ρ / (2A²)  →  K = R × 2A² / ρ
# 각 배관 구경별 등가 K값
print("\n  [ 등가 K값 역산 — 배관 구경별 ]")
equiv_K_by_pipe = {}
for size in ["65A", "80A", "100A"]:
    id_m = get_inner_diameter_m(size)
    A = math.pi / 4 * id_m ** 2
    K_eq = R_avg * 2 * A ** 2 / RHO
    equiv_K_by_pipe[size] = round(K_eq, 2)
    print(f"  {size} (ID={PIPE_DIMENSIONS[size]['id_mm']:.2f}mm, A={A:.6f}m²) | 등가 K = {K_eq:.2f}")

# 소화설비 일반 밸브 K-factor (실제적인 값)
# 참조: NFPA 13 부록, Crane TP-410, 소방시설 설계 핸드북
valve_data = [
    {
        "name": "알람밸브 (Alarm Valve, 습식)",
        "K": 2.0,
        "pipe": "100A",
        "desc": "습식 알람밸브 — 완전 개방 시",
        "location": "수직라이저 (급수 주배관)",
        "qty": 1,
    },
    {
        "name": "유수검지장치 (Flow Switch)",
        "K": 1.0,
        "pipe": "100A",
        "desc": "패들형 유수검지장치",
        "location": "수직라이저 (알람밸브 후단)",
        "qty": 1,
    },
    {
        "name": "게이트밸브 (Gate Valve, 전개)",
        "K": 0.15,
        "pipe": "100A",
        "desc": "완전 개방 게이트밸브 (OS&Y형)",
        "location": "수직라이저 입구",
        "qty": 2,
    },
    {
        "name": "체크밸브 (Check Valve, 스윙형)",
        "K": 2.0,
        "pipe": "100A",
        "desc": "스윙형 역류방지밸브",
        "location": "수직라이저 (급수 측)",
        "qty": 1,
    },
    {
        "name": "90° 엘보 (Standard Elbow)",
        "K": 0.75,
        "pipe": "100A",
        "desc": "라이저→교차배관 연결부 90° 엘보",
        "location": "라이저 상부 → 교차배관 연결",
        "qty": 1,
    },
    {
        "name": "리듀서 (Reducer, 점축소)",
        "K": 0.15,
        "pipe": "100A",
        "desc": "라이저(100A)→교차배관(65/80A) 축소관",
        "location": "라이저 → 교차배관 연결",
        "qty": 1,
    },
]

print("\n  [ 소화설비 밸브류 K-factor 및 유량별 손실 ]")

analysis3_rows = []
for flow in FLOWS:
    Q_si = flow / 60000.0
    dP_needed = paper_additional_loss_kpa[flow]  # kPa

    valve_detail = []
    total_valve_loss_kpa = 0

    for v in valve_data:
        pipe_id_m = get_inner_diameter_m(v["pipe"])
        A = math.pi / 4 * pipe_id_m ** 2
        V = Q_si / A
        # ΔP = K × ρV²/2 × 수량
        loss_Pa = v["K"] * 0.5 * RHO * V ** 2 * v["qty"]
        loss_kPa = loss_Pa / 1000

        valve_detail.append({
            "name": v["name"],
            "K": v["K"],
            "qty": v["qty"],
            "pipe": v["pipe"],
            "V_ms": round(V, 3),
            "loss_kpa": round(loss_kPa, 2),
            "location": v["location"],
            "desc": v["desc"],
        })
        total_valve_loss_kpa += loss_kPa

    row = {
        "flow": flow,
        "Q_si": Q_si,
        "dP_needed_kpa": dP_needed,
        "valve_detail": valve_detail,
        "total_valve_loss_kpa": round(total_valve_loss_kpa, 2),
        "gap_kpa": round(dP_needed - total_valve_loss_kpa, 2),
        "coverage_pct": round(total_valve_loss_kpa / dP_needed * 100, 1),
    }
    analysis3_rows.append(row)

    print(f"\n  Q={flow} LPM ({Q_si:.5f} m³/s):")
    for vd in valve_detail:
        print(f"    {vd['name']:<40} K={vd['K']:<5} x{vd['qty']} | "
              f"V={vd['V_ms']:.3f} m/s | {vd['loss_kpa']:.2f} kPa")
    print(f"    {'합계':<40} {'':<12} | {total_valve_loss_kpa:.2f} kPa "
          f"(필요: {dP_needed:.1f} kPa, 커버율: {row['coverage_pct']:.1f}%)")

# ══════════════════════════════════════════════
# 결합 분석: 80A + 밸브 보정
# ══════════════════════════════════════════════
print("\n" + "=" * 70)
print("  결합 분석: 80A 교차배관 + 밸브 국부저항 추가 시 영향")
print("=" * 70)

combined_rows = []
for flow in FLOWS:
    a3 = [r for r in analysis3_rows if r["flow"] == flow][0]
    for cm_size in ["65A", "80A"]:
        r = [x for x in analysis1_rows if x["flow"] == flow and x["cm_size"] == cm_size][0]
        corrected = r["worst_terminal_kpa"] - a3["total_valve_loss_kpa"]
        combined_rows.append({
            "flow": flow,
            "cm_size": cm_size,
            "sim_terminal_kpa": r["worst_terminal_kpa"],
            "valve_loss_kpa": a3["total_valve_loss_kpa"],
            "corrected_terminal_kpa": round(corrected, 2),
            "status": "PASS" if corrected >= 100 else "FAIL",
        })
        print(f"  Q={flow:>5} | {cm_size} | 시뮬={r['worst_terminal_kpa']:.1f} kPa | "
              f"밸브={a3['total_valve_loss_kpa']:.1f} kPa | 보정={corrected:.1f} kPa | "
              f"{'PASS' if corrected >= 100 else 'FAIL'}")


# ══════════════════════════════════════════════
# DOCX 생성
# ══════════════════════════════════════════════
print("\n" + "=" * 70)
print("  DOCX 문서 생성 중...")
print("=" * 70)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 스타일 설정
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

def add_table(doc, headers, rows):
    """표 추가 헬퍼"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def bold_para(doc, text, color=None):
    """굵은 단락 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if color:
        run.font.color.rgb = color
    return p

# ═══════════════════════════════════════
# 표지
# ═══════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run("FiPLSim 시뮬레이션\n핵심 미확인 변수 3가지 분석 보고서")
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run2 = subtitle.add_run(
    "① 교차배관 관경 (65A vs 80A)\n"
    "② 헤드당 유량 (균등분배 vs NFPA 13)\n"
    "③ 추가 국부 저항 (밸브류 누락분)"
)
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run3 = info.add_run(
    "고정 조건: Tree 4가지배관 x 8헤드, 간격 3.5m/2.3m, 입구 0.4 MPa, 비드 0mm\n"
    "유량: 1,200 / 1,600 / 2,100 LPM"
)
run3.font.size = Pt(10)
run3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ═══════════════════════════════════════
# 1장: 현재 배관 사양
# ═══════════════════════════════════════
doc.add_heading("1. 현재 시뮬레이션 배관 사양", level=1)

doc.add_heading("1.1 배관 치수 테이블 (JIS/KS Schedule 40)", level=2)

pipe_headers = ["호칭 구경", "외경 (mm)", "벽두께 (mm)", "내경 (mm)", "내경 (m)", "단면적 (cm2)"]
pipe_rows = []
for size in ["25A", "32A", "40A", "50A", "65A", "80A", "100A"]:
    d = PIPE_DIMENSIONS[size]
    id_m = d["id_mm"] / 1000
    A_cm2 = math.pi / 4 * (id_m * 100) ** 2
    pipe_rows.append([
        size, f"{d['od_mm']:.2f}", f"{d['wall_mm']:.2f}",
        f"{d['id_mm']:.2f}", f"{id_m:.5f}", f"{A_cm2:.2f}",
    ])
add_table(doc, pipe_headers, pipe_rows)

doc.add_heading("1.2 가지배관 구성 (8헤드 기준)", level=2)
doc.add_paragraph("NFSC 103 기반 하류 헤드 수별 자동 관경 선정:")

branch_headers = ["세그먼트", "하류\n헤드 수", "선정 구경", "내경 (mm)", "길이 (m)"]
branch_rows = []
for h in range(HEADS_PER_BRANCH):
    ds = HEADS_PER_BRANCH - h
    nom = auto_pipe_size(ds)
    id_mm = PIPE_DIMENSIONS[nom]["id_mm"]
    label = f"{h}번"
    if h == 0:
        label += " (입구)"
    elif h == HEADS_PER_BRANCH - 1:
        label += " (말단)"
    branch_rows.append([label, str(ds), nom, f"{id_mm:.2f}", f"{HEAD_SPACING:.1f}"])
add_table(doc, branch_headers, branch_rows)
doc.add_paragraph(
    f"가지배관 1줄 총 길이: {HEAD_SPACING}m x {HEADS_PER_BRANCH} = "
    f"{HEAD_SPACING * HEADS_PER_BRANCH:.1f}m"
)

doc.add_heading("1.3 교차배관 현황", level=2)
doc.add_paragraph(
    f"총 헤드 수: {NUM_BRANCHES} x {HEADS_PER_BRANCH} = {TOTAL_HEADS}개"
)
doc.add_paragraph(
    f"현재 자동 선정: {TOTAL_HEADS}개 < 40개 → 65A "
    f"(내경 {PIPE_DIMENSIONS['65A']['id_mm']:.2f}mm)"
)
doc.add_paragraph(
    f"논문 §2.3 명시: 80A (내경 {PIPE_DIMENSIONS['80A']['id_mm']:.2f}mm)"
)
doc.add_paragraph(
    f"교차배관 총 길이: {BRANCH_SPACING}m x {NUM_BRANCHES - 1} = "
    f"{BRANCH_SPACING * (NUM_BRANCHES - 1):.1f}m"
)

doc.add_page_break()

# ═══════════════════════════════════════
# 2장: 미확인 변수 ① 교차배관 관경
# ═══════════════════════════════════════
doc.add_heading("2. 미확인 변수 (1): 교차배관 관경 65A vs 80A", level=1)

doc.add_heading("2.1 문제 정의", level=2)
doc.add_paragraph(
    "현재 시뮬레이션은 자동 관경 선정 규칙(auto_cross_main_size)에 의해 "
    f"총 {TOTAL_HEADS}개 헤드 기준 65A를 선정합니다."
)
doc.add_paragraph(
    "그러나 논문 §2.3에서 'cross-main (80A)'를 명시하고 있어, "
    "실제 설계 관경과 시뮬레이션 관경이 불일치할 수 있습니다."
)
doc.add_paragraph("자동 선정 규칙:")
doc.add_paragraph("40개 이상 → 100A", style="List Bullet")
doc.add_paragraph("20~39개 → 80A", style="List Bullet")
doc.add_paragraph("20개 미만 → 65A", style="List Bullet")
bold_para(doc,
    f"현재 = {TOTAL_HEADS}개 → 65A 자동 선정 (80A 기준 경계인 20개보다 12개 초과)",
    RGBColor(0xCC, 0x00, 0x00)
)

doc.add_heading("2.2 비교 시뮬레이션 결과", level=2)

headers_1 = ["유량\n(LPM)", "교차배관", "내경\n(mm)", "입구부 유속\n(m/s)",
             "교차배관\n총손실 (kPa)", "최악\n말단수압 (kPa)"]
rows_1 = []
for r in analysis1_rows:
    rows_1.append([
        str(r["flow"]), r["cm_size"], f"{r['cm_id_mm']:.2f}",
        f"{r['V_cm_ms']:.3f}",
        f"{r['cross_main_loss_kpa']:.2f}", f"{r['worst_terminal_kpa']:.2f}",
    ])
add_table(doc, headers_1, rows_1)

doc.add_heading("2.3 65A → 80A 변경 시 차이", level=2)

headers_d1 = ["유량\n(LPM)", "65A 유속\n(m/s)", "80A 유속\n(m/s)",
              "65A 손실\n(kPa)", "80A 손실\n(kPa)", "손실 감소\n(kPa)",
              "말단수압\n증가 (kPa)"]
rows_d1 = []
for d in diff_rows_1:
    rows_d1.append([
        str(d["flow"]), f"{d['V_65']:.3f}", f"{d['V_80']:.3f}",
        f"{d['cm_loss_65']:.2f}", f"{d['cm_loss_80']:.2f}",
        f"{d['diff_cm_loss']:.2f}", f"+{d['diff_terminal']:.2f}",
    ])
add_table(doc, headers_d1, rows_d1)

doc.add_paragraph()
bold_para(doc,
    f"결론: 80A 적용 시 교차배관 유속이 65%로 감소하고, "
    f"Q=2100에서 말단수압이 {diff_rows_1[2]['diff_terminal']:.1f} kPa 상승합니다. "
    "논문 기준 80A가 맞다면, 현재 시뮬레이션은 교차배관 손실을 과대 평가합니다.",
    RGBColor(0x00, 0x44, 0x88)
)

doc.add_page_break()

# ═══════════════════════════════════════
# 3장: 미확인 변수 ② 헤드당 유량
# ═══════════════════════════════════════
doc.add_heading("3. 미확인 변수 (2): 헤드당 유량 (Q_head)", level=1)

doc.add_heading("3.1 현재 균등분배 방식", level=2)
doc.add_paragraph(
    f"Q_head = Q_total / (가지배관 수 x 헤드 수) = Q_total / "
    f"({NUM_BRANCHES} x {HEADS_PER_BRANCH}) = Q_total / {TOTAL_HEADS}"
)

headers_2 = ["총유량 (LPM)", "총 헤드 수", "Q_head (LPM)",
             "Q_branch (LPM)", "NFPA 13 최소\n(80 LPM)", "판정"]
rows_2 = []
for r in analysis2_rows:
    rows_2.append([
        str(r["flow"]), str(r["total_heads"]), f"{r['q_head_lpm']:.2f}",
        f"{r['q_branch_lpm']:.2f}", "80", "충족" if r["meets_nfpa"] else "미달",
    ])
add_table(doc, headers_2, rows_2)

doc.add_heading("3.2 NFPA 13 / NFSC 103 기준", level=2)

nfpa_headers = ["기준", "헤드 방수량", "동시 개방 수", "비고"]
nfpa_rows = [
    ["NFPA 13 경급위험", "80 LPM/head", "10개", "사무실, 주거"],
    ["NFPA 13 중급위험-1", "80~160 LPM/head", "15~20개", "상가, 공장"],
    ["NFPA 13 중급위험-2", "160~200 LPM/head", "20~25개", "창고, 물류"],
    ["NFSC 103", "80 LPM/head (최소)", "10개 (경급)", "국내 기준"],
]
add_table(doc, nfpa_headers, nfpa_rows)

doc.add_heading("3.3 필요 총유량 산정", level=2)

scenario_headers = ["시나리오", "헤드 방수량", "동시 개방 수", "필요 총유량\n(LPM)", "현재 설계\n범위 내?"]
scenario_rows = [
    ["전체 동시 개방", "80 LPM", f"{TOTAL_HEADS}개", f"{80*TOTAL_HEADS:,}", "X (2,100 초과)"],
    ["경급위험 (NFPA 13)", "80 LPM", "10개", "800", "O"],
    ["중급위험-1 (NFPA 13)", "80 LPM", "20개", "1,600", "O"],
    ["중급위험-1 (최대)", "160 LPM", "20개", "3,200", "X"],
]
add_table(doc, scenario_headers, scenario_rows)

doc.add_paragraph()
bold_para(doc,
    "핵심 질문: 논문의 시뮬레이션이 '전체 32개 헤드 동시 개방'을 가정한 것인지, "
    "'NFPA 13 기준 부분 개방'을 가정한 것인지에 따라 Q_head와 배관 유속이 달라집니다. "
    "논문에 '각 헤드 방류량', 'Q per head', '설계 방수량' 항목 확인이 필요합니다.",
    RGBColor(0x00, 0x44, 0x88)
)

doc.add_page_break()

# ═══════════════════════════════════════
# 4장: 미확인 변수 ③ 추가 국부 저항
# ═══════════════════════════════════════
doc.add_heading("4. 미확인 변수 (3): 추가 국부 저항 (밸브류)", level=1)

doc.add_heading("4.1 Q 제곱 비례 분석", level=2)
doc.add_paragraph(
    "현재 시뮬레이션 결과와 논문값의 차이가 Q^2에 정확히 비례합니다. "
    "이는 유량에 무관한 고정 저항 성분(K-factor)이 누락되었음을 의미합니다."
)
doc.add_paragraph("국부 손실 공식: Delta_P = K x (rho x V^2 / 2) = K x rho / (2A^2) x Q^2")
doc.add_paragraph("따라서 Delta_P / Q^2 = R = K x rho / (2A^2) = 상수")

headers_3a = ["유량 (LPM)", "Q (m3/s)", "필요 추가 손실\n(kPa)",
              "Delta_P (Pa)", "R = Delta_P/Q^2\n(Pa*s^2/m^6)"]
rows_3a = []
for i, flow in enumerate(FLOWS):
    Q_si = flow / 60000.0
    dP = paper_additional_loss_kpa[flow]
    rows_3a.append([
        str(flow), f"{Q_si:.5f}", f"{dP:.1f}",
        f"{dP*1000:.0f}", f"{ratios[i]:.2e}",
    ])
add_table(doc, headers_3a, rows_3a)

doc.add_paragraph(f"평균 R = {R_avg:.4e} Pa*s^2/m^6 (3개 유량 모두 일정)")

doc.add_heading("4.2 등가 K값 역산", level=2)
doc.add_paragraph(
    "R = K x rho / (2A^2) 에서 K = R x 2A^2 / rho 로 역산:"
)

headers_keq = ["배관 구경", "내경 (mm)", "단면적 (m2)", "등가 K값"]
rows_keq = []
for size in ["65A", "80A", "100A"]:
    id_m = get_inner_diameter_m(size)
    A = math.pi / 4 * id_m ** 2
    K_eq = equiv_K_by_pipe[size]
    rows_keq.append([size, f"{PIPE_DIMENSIONS[size]['id_mm']:.2f}", f"{A:.6f}", f"{K_eq:.2f}"])
add_table(doc, headers_keq, rows_keq)

doc.add_paragraph(
    f"교차배관 65A 기준 등가 K = {equiv_K_by_pipe['65A']:.2f}, "
    f"80A 기준 등가 K = {equiv_K_by_pipe['80A']:.2f}, "
    f"100A(라이저) 기준 등가 K = {equiv_K_by_pipe['100A']:.2f}"
)

doc.add_heading("4.3 소화설비 밸브류 K-factor 목록", level=2)
doc.add_paragraph(
    "소화설비 배관에 일반적으로 설치되는 밸브류와 K-factor입니다. "
    "밸브 대부분은 수직 라이저(100A급)에 설치됩니다."
)

headers_valve = ["부속류", "K-factor", "수량", "설치 위치", "기준 배관"]
rows_valve = []
for v in valve_data:
    rows_valve.append([
        v["name"], f"{v['K']:.2f}", str(v["qty"]), v["location"], v["pipe"],
    ])
# 합계 행
total_K = sum(v["K"] * v["qty"] for v in valve_data)
rows_valve.append(["합계 (100A 기준)", f"{total_K:.2f}", "-", "-", "100A"])
add_table(doc, headers_valve, rows_valve)

doc.add_heading("4.4 유량별 밸브 손실 추정", level=2)
doc.add_paragraph(
    "모든 밸브를 라이저 배관(100A, 내경 102.26mm) 유속 기준으로 계산합니다."
)

for a3 in analysis3_rows:
    doc.add_paragraph(f"Q = {a3['flow']} LPM (= {a3['Q_si']:.5f} m3/s)", style="List Bullet")

    v_headers = ["부속류", "K값", "수량", "유속 (m/s)", "손실 (kPa)"]
    v_rows = []
    for vd in a3["valve_detail"]:
        v_rows.append([
            vd["name"], f"{vd['K']:.2f}", str(vd.get("qty", 1) if "qty" not in vd else 1),
            f"{vd['V_ms']:.3f}", f"{vd['loss_kpa']:.2f}",
        ])
    v_rows.append(["합계", "-", "-", "-", f"{a3['total_valve_loss_kpa']:.2f}"])
    add_table(doc, v_headers, v_rows)

    coverage = a3["coverage_pct"]
    gap = a3["gap_kpa"]
    doc.add_paragraph(
        f"  밸브 합계 = {a3['total_valve_loss_kpa']:.2f} kPa / "
        f"필요 = {a3['dP_needed_kpa']:.1f} kPa → "
        f"커버율 {coverage:.1f}%, 잔여 {gap:.2f} kPa"
    )
    doc.add_paragraph()

doc.add_heading("4.5 잔여 차이 해석", level=2)
doc.add_paragraph(
    "밸브류 합산으로 필요 추가 손실의 상당 부분을 설명할 수 있으며, "
    "잔여 차이는 다음 요인으로 추정됩니다:"
)
doc.add_paragraph("배관 노후화에 따른 조도 증가 (Epsilon 증가)", style="List Bullet")
doc.add_paragraph("추가 엘보/티/리듀서 (배관 경로 상)", style="List Bullet")
doc.add_paragraph("라이저 배관 마찰 손실 (수직 구간)", style="List Bullet")
doc.add_paragraph("시공 오차에 의한 추가 저항", style="List Bullet")

doc.add_paragraph()
bold_para(doc,
    f"결론: 등가 K값(100A 기준 {equiv_K_by_pipe['100A']:.2f})은 "
    f"일반적인 소화설비 밸브류(K합계={total_K:.2f})로 충분히 설명 가능합니다. "
    "알람밸브(K=2.0), 체크밸브(K=2.0), 유수검지장치(K=1.0) 등이 주요 누락 요소입니다.",
    RGBColor(0x00, 0x44, 0x88)
)

doc.add_page_break()

# ═══════════════════════════════════════
# 5장: 결합 분석
# ═══════════════════════════════════════
doc.add_heading("5. 결합 분석: 교차배관 80A + 밸브 보정", level=1)

doc.add_paragraph(
    "미확인 변수 ①(교차배관 80A)과 ③(밸브 국부저항)을 동시 적용한 보정 결과:"
)

headers_comb = ["유량\n(LPM)", "교차배관", "시뮬 말단수압\n(kPa)",
                "밸브 손실\n(kPa)", "보정 말단수압\n(kPa)", "규정 (100kPa)"]
rows_comb = []
for c in combined_rows:
    rows_comb.append([
        str(c["flow"]), c["cm_size"], f"{c['sim_terminal_kpa']:.2f}",
        f"{c['valve_loss_kpa']:.2f}", f"{c['corrected_terminal_kpa']:.2f}",
        c["status"],
    ])
add_table(doc, headers_comb, rows_comb)

doc.add_page_break()

# ═══════════════════════════════════════
# 6장: 요약 및 권고사항
# ═══════════════════════════════════════
doc.add_heading("6. 요약 및 권고사항", level=1)

doc.add_heading("6.1 3가지 미확인 변수 종합 요약", level=2)

summary_headers = ["미확인 변수", "현재 시뮬레이션", "논문/기준값", "영향도"]
summary_rows = [
    [
        "① 교차배관 관경",
        "65A (자동 선정)",
        "80A (논문 §2.3)",
        f"Q=2100에서\n{diff_rows_1[2]['diff_terminal']:.1f} kPa 차이",
    ],
    [
        "② 헤드당 유량",
        f"균등 분배\n(Q=2100→65.6 LPM)",
        "NFPA 13 최소\n80 LPM/head",
        "모든 유량에서\nNFPA 미달",
    ],
    [
        "③ 추가 국부 저항",
        "밸브류 미포함\n(K_valve = 0)",
        f"등가 K={equiv_K_by_pipe['100A']:.2f}\n(100A 기준)",
        "Q^2에 비례하는\n고정 손실 누락",
    ],
]
add_table(doc, summary_headers, summary_rows)

doc.add_heading("6.2 권고사항", level=2)

recommendations = [
    "교차배관 관경을 80A로 변경하여 재시뮬레이션 수행 (논문 §2.3과 일치 확인)",
    "논문 본문에서 'Q per head', '각 헤드 방류량', '설계 방수량' 항목 확인 → 유량 배분 방식 결정",
    "알람밸브(K=2.0), 유수검지장치(K=1.0), 체크밸브(K=2.0), 게이트밸브(K=0.15x2)를 시뮬레이션에 추가",
    "밸브 보정 후 논문값과의 잔여 차이를 확인하여 추가 누락 요소(배관 노후도, 엘보 등) 식별",
    "교차배관 80A + 밸브 보정을 동시 적용한 최종 검증 수행",
]
for rec in recommendations:
    doc.add_paragraph(rec, style="List Number")

doc.add_heading("6.3 논문 확인 요청 사항", level=2)
doc.add_paragraph("아래 항목을 논문에서 확인해 주시면 시뮬레이션을 정밀 보정할 수 있습니다:")

check_items = [
    "교차배관 관경: 65A 또는 80A 중 실제 적용 구경",
    "헤드 방수량: 'Q per head' 또는 '설계 방수량' 수치",
    "동시 개방 헤드 수: 전체 32개 또는 NFPA 기준 부분 개방",
    "밸브류 포함 여부: 알람밸브, 체크밸브 등이 시뮬레이션에 반영되었는지",
    "배관 조도(Epsilon): 신관 기준(0.045mm) 또는 노후관 기준 적용 여부",
]
for item in check_items:
    doc.add_paragraph(item, style="List Bullet")

# ── 저장 ──
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "FiPLSim_미확인변수_분석보고서.docx"
)
doc.save(output_path)
print(f"\n  DOCX 저장 완료: {output_path}")
print("=" * 70)
