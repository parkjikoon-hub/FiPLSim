"""
Markdown → DOCX 변환 스크립트
FiPLSim 프로젝트용 — Overview & Design Manual 변환
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def setup_styles(doc):
    """문서 기본 스타일 설정"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # 제목 스타일
    for level in range(1, 5):
        name = f'Heading {level}'
        if name in doc.styles:
            h = doc.styles[name]
            h.font.name = '맑은 고딕'
            h.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
            if level == 1:
                h.font.size = Pt(18)
                h.font.bold = True
            elif level == 2:
                h.font.size = Pt(14)
                h.font.bold = True
            elif level == 3:
                h.font.size = Pt(12)
                h.font.bold = True
            elif level == 4:
                h.font.size = Pt(11)
                h.font.bold = True


def add_rich_text(paragraph, text):
    """볼드/코드 마크다운을 처리하여 paragraph에 추가"""
    # 패턴: **bold**, `code`, 일반텍스트
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
        else:
            paragraph.add_run(part)


def add_table(doc, rows):
    """마크다운 테이블을 DOCX 테이블로 변환"""
    if len(rows) < 2:
        return

    # 구분선 행(---|---) 제거
    data_rows = []
    for row in rows:
        cells = [c.strip() for c in row.strip('|').split('|')]
        cells = [c for c in cells if c]
        if cells and all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        if cells:
            data_rows.append(cells)

    if not data_rows:
        return

    n_cols = max(len(r) for r in data_rows)
    # 모든 행의 열 수 맞추기
    for r in data_rows:
        while len(r) < n_cols:
            r.append('')

    table = doc.add_table(rows=len(data_rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_rich_text(p, cell_text.strip())
            p.style.font.size = Pt(9)

            # 헤더 행 스타일
            if i == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                # 헤더 배경색
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'D9E2F3'
                })
                shading.append(shading_elm)

    doc.add_paragraph('')  # 테이블 후 간격


def add_code_block(doc, lines):
    """코드 블록을 회색 배경 텍스트로 추가"""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def convert_md_to_docx(md_path, docx_path, title=None):
    """마크다운 파일을 DOCX로 변환"""
    md_path = Path(md_path)
    docx_path = Path(docx_path)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    setup_styles(doc)

    i = 0
    table_buffer = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 코드 블록 처리
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # 테이블 버퍼 비우기
                if table_buffer:
                    add_table(doc, table_buffer)
                    table_buffer = []
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 테이블 행 수집
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_buffer.append(line)
            i += 1
            continue
        elif table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # 구분선 (---)
        if re.match(r'^-{3,}$', line.strip()):
            i += 1
            continue

        # LaTeX 수식 ($$...$$)
        if line.strip().startswith('$$'):
            math_text = line.strip().strip('$')
            if not math_text:
                # 다음 줄에서 수식 찾기
                i += 1
                math_lines = []
                while i < len(lines) and not lines[i].strip().startswith('$$'):
                    math_lines.append(lines[i].strip())
                    i += 1
                math_text = ' '.join(math_lines)
            # 수식을 이탤릭 텍스트로 표시
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # LaTeX → 간단한 텍스트 변환
            math_text = math_text.replace('\\times', '×')
            math_text = math_text.replace('\\frac{', '(')
            math_text = math_text.replace('}{', ')/(')
            math_text = math_text.replace('}', ')')
            math_text = math_text.replace('\\left(', '(')
            math_text = math_text.replace('\\right)', ')')
            math_text = math_text.replace('\\sqrt', '√')
            math_text = math_text.replace('\\log_{10}', 'log₁₀')
            math_text = math_text.replace('\\varepsilon', 'ε')
            math_text = math_text.replace('\\textbf{', '')
            math_text = math_text.replace('\\sum', 'Σ')
            math_text = math_text.replace('\\partial', '∂')
            math_text = math_text.replace('\\Delta', 'Δ')
            math_text = math_text.replace('\\beta', 'β')
            math_text = math_text.replace('\\sin', 'sin')
            math_text = math_text.replace('\\theta', 'θ')
            math_text = re.sub(r'\^(\d)', lambda m: '⁰¹²³⁴⁵⁶⁷⁸⁹'[int(m.group(1))], math_text)
            math_text = re.sub(r'_\{([^}]+)\}', r'_\1', math_text)
            run = p.add_run(math_text)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)
            run.italic = True
            i += 1
            continue

        # 제목 (# ~ ####)
        heading_match = re.match(r'^(#{1,4})\s+(.+)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            # 볼드 마크다운 제거
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # 인용문 (>)
        if line.strip().startswith('>'):
            text = re.sub(r'^>\s*', '', line.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_rich_text(p, text)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.italic = True
            i += 1
            continue

        # 리스트 항목 (- 또는 숫자.)
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.+)', line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            bullet = list_match.group(2)
            if bullet in ['-', '*']:
                p.add_run('• ')
            else:
                p.add_run(f'{bullet} ')
            add_rich_text(p, text)
            i += 1
            continue

        # 일반 단락
        p = doc.add_paragraph()
        add_rich_text(p, line)
        i += 1

    # 남은 테이블 처리
    if table_buffer:
        add_table(doc, table_buffer)

    doc.save(str(docx_path))
    print(f"  생성 완료: {docx_path}")
    return docx_path


if __name__ == '__main__':
    base = Path(__file__).parent

    print("=" * 60)
    print("FiPLSim 문서 MD → DOCX 변환")
    print("=" * 60)

    # 1. Overview
    print("\n[1/2] FiPLSim_Overview.md → DOCX 변환 중...")
    convert_md_to_docx(
        base / 'FiPLSim_Overview.md',
        base / 'FiPLSim_Overview.docx',
        title='FiPLSim 프로그램 개요'
    )

    # 2. Design Manual
    print("[2/2] FiPLSim_Design_Manual.md → DOCX 변환 중...")
    convert_md_to_docx(
        base / 'FiPLSim_Design_Manual.md',
        base / 'FiPLSim_Design_Manual.docx',
        title='FiPLSim 설계 매뉴얼'
    )

    print("\n" + "=" * 60)
    print("변환 완료!")
    print("=" * 60)
