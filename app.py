# ! 소화배관 시뮬레이션 — Streamlit UI 대시보드 (동적 배관망 버전)
# * 사이드바: 동적 배관망 입력 + KPI 대시보드 + 5개 탭

import sys
import os
import io
import streamlit as st
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from constants import (
    PUMP_DATABASE, FITTING_SPACING_OPTIONS, PIPE_DIMENSIONS,
    DEFAULT_INLET_PRESSURE_MPA, DEFAULT_TOTAL_FLOW_LPM,
    DEFAULT_FITTING_SPACING_M, DEFAULT_BEAD_HEIGHT_MM,
    MIN_TERMINAL_PRESSURE_MPA, MAX_TERMINAL_PRESSURE_MPA,
    MAX_VELOCITY_BRANCH_MS, MAX_VELOCITY_OTHER_MS,
    DEFAULT_MC_ITERATIONS, DEFAULT_MIN_DEFECTS, DEFAULT_MAX_DEFECTS,
    DEFAULT_OPERATING_HOURS_PER_YEAR, DEFAULT_ELECTRICITY_RATE_KRW,
    DEFAULT_NUM_BRANCHES, DEFAULT_HEADS_PER_BRANCH,
    DEFAULT_BRANCH_SPACING_M, DEFAULT_HEAD_SPACING_M,
    MAX_BRANCHES, MAX_HEADS_PER_BRANCH,
    DEFAULT_BEADS_PER_BRANCH, MAX_BEADS_PER_BRANCH,
    K1_BASE, K2, K3,
    HC_RELAXATION_FACTOR, HC_RELAXATION_MIN, HC_RELAXATION_MAX,
    DEFAULT_EQUIPMENT_K_FACTORS, DEFAULT_SUPPLY_PIPE_SIZE,
)
from pipe_network import (
    compare_dynamic_cases, compare_dynamic_cases_with_topology,
    check_nfpc_compliance, ValidationError,
)
from pump import (
    DynamicSystemCurve, load_pump, find_operating_point, calculate_energy_savings,
)
from simulation import (
    run_dynamic_monte_carlo, run_dynamic_sensitivity, run_variable_sweep,
    run_bernoulli_monte_carlo, run_bernoulli_sweep,
)


# ──────────────────────────────────────────────
# ? 페이지 설정
# ──────────────────────────────────────────────
_favicon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "favicon.png")
_page_icon = _favicon_path if os.path.exists(_favicon_path) else "🔥"

st.set_page_config(
    page_title="FiPLSim",
    page_icon=_page_icon,
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── 모바일 홈화면 아이콘 메타태그 ──
_apple_icon = os.path.join(os.path.dirname(os.path.abspath(__file__)), "apple-touch-icon.png")
if os.path.exists(_apple_icon):
    import base64 as _b64
    with open(_apple_icon, "rb") as _f:
        _icon_b64 = _b64.b64encode(_f.read()).decode()
    st.markdown(
        f'<link rel="apple-touch-icon" href="data:image/png;base64,{_icon_b64}">'
        f'<link rel="icon" type="image/png" sizes="192x192" '
        f'href="data:image/png;base64,{_icon_b64}">'
        f'<meta name="apple-mobile-web-app-capable" content="yes">'
        f'<meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">'
        f'<meta name="apple-mobile-web-app-title" content="FiPLSim">',
        unsafe_allow_html=True,
    )

st.markdown(
    '<h1 style="margin-bottom:0">Fi<span style="color:#4A9EFF">PLS</span>im: '
    'Advanced Fire Protection <span style="color:#4A9EFF">P</span>ipe '
    '<span style="color:#4A9EFF">L</span>et '
    '<span style="color:#4A9EFF">S</span>imulator</h1>',
    unsafe_allow_html=True,
)
st.caption("동적 배관망 생성 및 몬테카를로 기반 유체역학 해석 엔진 (PLS)")


def tooltip(term: str, explanation: str) -> str:
    return (
        f'<span title="{explanation}" '
        f'style="border-bottom: 1px dotted #888; cursor: help; font-weight: 600;">'
        f'{term}</span>'
    )


# ══════════════════════════════════════════════
#  다크/라이트 모드 토글
# ══════════════════════════════════════════════
if "theme_mode" not in st.session_state:
    st.session_state["theme_mode"] = "dark"

_theme_col1, _theme_col2 = st.sidebar.columns(2)
with _theme_col1:
    if st.button(":material/dark_mode: Dark", use_container_width=True,
                 type="primary" if st.session_state["theme_mode"] == "dark" else "secondary"):
        st.session_state["theme_mode"] = "dark"
        st.rerun()
with _theme_col2:
    if st.button(":material/light_mode: Light", use_container_width=True,
                 type="primary" if st.session_state["theme_mode"] == "light" else "secondary"):
        st.session_state["theme_mode"] = "light"
        st.rerun()

if st.session_state["theme_mode"] == "light":
    st.markdown("""<style>
        /* ── 전역 배경 및 텍스트 ── */
        [data-testid="stAppViewContainer"], [data-testid="stApp"],
        .main, .block-container {
            background-color: #FFFFFF !important; color: #262730 !important;
        }
        [data-testid="stSidebar"], [data-testid="stSidebar"] > div {
            background-color: #F0F2F6 !important; color: #262730 !important;
        }
        header[data-testid="stHeader"] { background-color: #FFFFFF !important; }

        /* ── 모든 텍스트 요소 ── */
        h1, h2, h3, h4, h5, h6, p, span, label, div, li, td, th,
        .stMarkdown, .stCaption, [data-testid="stMetricValue"],
        [data-testid="stMetricLabel"], [class*="st-"] {
            color: #262730 !important;
        }

        /* ── 입력 필드 (number_input, text_input, selectbox 등) ── */
        [data-testid="stNumberInput"] input,
        [data-testid="stTextInput"] input,
        input[type="number"], input[type="text"],
        [data-baseweb="input"] input,
        [data-baseweb="input"],
        [data-baseweb="base-input"] {
            background-color: #FFFFFF !important;
            color: #262730 !important;
            border-color: #CCC !important;
        }

        /* ── selectbox / dropdown ── */
        [data-baseweb="select"] > div,
        [data-baseweb="select"] span,
        [data-baseweb="popover"] li {
            background-color: #FFFFFF !important;
            color: #262730 !important;
        }

        /* ── number_input 버튼 (+/-) ── */
        [data-testid="stNumberInput"] button {
            background-color: #E8E8E8 !important;
            color: #262730 !important;
            border-color: #CCC !important;
        }

        /* ── 슬라이더 ── */
        [data-testid="stSlider"] div[data-baseweb="slider"] div {
            background-color: #E0E0E0 !important;
        }

        /* ── 탭 ── */
        .stTabs [data-baseweb="tab-list"] { background-color: #FFFFFF !important; }
        .stTabs [data-baseweb="tab"] { color: #262730 !important; }

        /* ── 테이블 / 데이터프레임 ── */
        [data-testid="stDataFrame"], [data-testid="stTable"],
        table, thead, tbody, tr, td, th {
            background-color: #FFFFFF !important;
            color: #262730 !important;
        }

        /* ── 코드 블록 / 다이어그램 영역 ── */
        [data-testid="stCodeBlock"], pre, code {
            background-color: #F5F5F5 !important;
            color: #262730 !important;
        }

        /* ── 알림/정보 박스 ── */
        [data-testid="stAlert"] {
            background-color: #E8F4FD !important;
            color: #262730 !important;
        }

        /* ── expander ── */
        [data-testid="stExpander"] {
            background-color: #F8F8F8 !important;
            border-color: #DDD !important;
        }
        [data-testid="stExpander"] summary span { color: #262730 !important; }

        /* ── metric delta ── */
        [data-testid="stMetricDelta"] svg { fill: #262730 !important; }

        /* ── radio / checkbox ── */
        [data-testid="stRadio"] label span,
        [data-testid="stCheckbox"] label span { color: #262730 !important; }
    </style>""", unsafe_allow_html=True)

st.sidebar.divider()

# ══════════════════════════════════════════════
#  사이드바 입력
# ══════════════════════════════════════════════

# ── 0. 배관망 토폴로지 선택 ──
st.sidebar.header(":material/account_tree: 배관망 구조")
topology_label = st.sidebar.radio(
    "배관망 토폴로지",
    ["Tree (가지형)", "Full Grid (격자형)"],
    help="Tree: 교차배관 1개에서 일방향 분배. "
         "Full Grid: 교차배관 2개가 평행, 가지배관 양끝 연결된 격자 구조 (Hardy-Cross 수렴 계산).",
)
topology_key = "tree" if "Tree" in topology_label else "grid"

if topology_key == "grid":
    st.sidebar.caption(
        "**Full Grid**: 교차배관 TOP/BOT 2개가 평행 배치되고, "
        "가지배관 양 끝이 연결된 격자 구조입니다. "
        "Hardy-Cross 반복법으로 유량을 수렴 계산합니다."
    )
    with st.sidebar.expander("Hardy-Cross 고급 설정"):
        hc_relaxation = st.slider(
            "이완 계수 (Under-Relaxation Factor)",
            min_value=HC_RELAXATION_MIN,
            max_value=HC_RELAXATION_MAX,
            value=HC_RELAXATION_FACTOR,
            step=0.05,
            help="유량 보정값에 곱하는 감쇠 계수입니다. "
                 "값이 작을수록 수렴이 느리지만 안정적이고, "
                 "값이 클수록 빠르지만 발산 위험이 있습니다. "
                 "대규모 배관망(50개 이상)에서는 0.3~0.5를 권장합니다.",
        )
        st.caption(
            f"**현재 설정**: 이완 계수 = {hc_relaxation} | "
            f"최대 반복 = 1,000회 | 수두 허용 오차 = 0.001m | "
            f"유량 허용 오차 = 0.0001 LPM"
        )
else:
    hc_relaxation = HC_RELAXATION_FACTOR

# ── 1. 동적 배관망 구성 ──
st.sidebar.header(":material/construction: 배관망 구성 (동적 생성)")

num_branches = st.sidebar.number_input(
    "양방향 가지배관 총 개수 (n)",
    min_value=1, max_value=MAX_BRANCHES, value=DEFAULT_NUM_BRANCHES, step=1,
    help="교차배관(Cross Main) 좌우로 뻗어 나가는 가지배관의 총 개수입니다. "
         f"최대 {MAX_BRANCHES}개까지 입력 가능합니다.",
)

heads_per_branch = st.sidebar.number_input(
    "가지배관당 헤드 수 (m)",
    min_value=1, max_value=MAX_HEADS_PER_BRANCH, value=DEFAULT_HEADS_PER_BRANCH, step=1,
    help="각 가지배관에 일정 간격으로 설치되는 스프링클러 헤드 수입니다. "
         f"최대 {MAX_HEADS_PER_BRANCH}개까지 입력 가능합니다.",
)

branch_spacing = st.sidebar.number_input(
    "가지배관 사이 간격 (m)",
    min_value=1.0, max_value=10.0, value=DEFAULT_BRANCH_SPACING_M, step=0.5,
    help="교차배관 위에서 가지배관 분기점 사이의 거리입니다.",
)

head_spacing = st.sidebar.select_slider(
    "헤드 간격 (m)",
    options=FITTING_SPACING_OPTIONS,
    value=DEFAULT_FITTING_SPACING_M,
    help="가지배관 위에서 스프링클러 헤드 사이의 배관 길이입니다.",
)

total_heads = num_branches * heads_per_branch
st.sidebar.info(f"전체 헤드: **{total_heads}개** = {num_branches} 가지배관 × {heads_per_branch} 헤드")

# ── 2. 운전 조건 ──
st.sidebar.header(":material/settings: 운전 조건")

inlet_pressure = st.sidebar.slider(
    "입구 압력 (MPa)",
    min_value=0.1, max_value=2.0, value=0.4, step=0.05,
    help="교차배관(Cross Main) 입구의 설계 압력입니다.",
)

design_flow = st.sidebar.slider(
    "전체 설계 유량 (LPM)",
    min_value=100, max_value=3000, value=int(DEFAULT_TOTAL_FLOW_LPM), step=50,
    help="교차배관 입구로 유입되는 총 유량(리터/분)입니다.",
)

# ── 3. 비드 설정 ──
st.sidebar.header(":material/build: 용접 비드 설정")

bead_height = st.sidebar.slider(
    "기존 기술 비드 높이 (mm)",
    min_value=0.1, max_value=5.0, value=DEFAULT_BEAD_HEIGHT_MM, step=0.1,
    help="기존 용접 기술의 내면 비드 돌출 높이입니다.",
)

beads_per_branch = st.sidebar.number_input(
    "가지배관당 용접 비드 개수 (개)",
    min_value=0, max_value=MAX_BEADS_PER_BRANCH, value=DEFAULT_BEADS_PER_BRANCH, step=1,
    help="각 가지배관의 직관 구간(헤드 사이 배관) 내에 배치되는 용접 비드의 개수입니다. "
         "몬테카를로 시뮬레이션 시 매 반복마다 비드 위치가 무작위로 재배치되어 "
         "위치 변화에 따른 말단 압력 산포도(Variance)를 분석합니다. "
         f"범위: 0~{MAX_BEADS_PER_BRANCH}개, 0이면 직관 용접 비드 미적용.",
)
total_weld_beads = beads_per_branch * num_branches
if beads_per_branch > 0:
    st.sidebar.caption(
        f"직관 용접 비드: 가지배관당 **{beads_per_branch}개** × "
        f"{num_branches}개 = 전체 **{total_weld_beads}개**"
    )

# ── 3.5. 밸브/기기류 국부 손실 ──
st.sidebar.header(":material/valve: 추가 기기 손실 (밸브류)")
st.sidebar.caption(
    "수직 라이저(공급배관)에 설치되는 밸브/기기류의 국부 손실입니다. "
    "체크박스를 켜면 해당 밸브의 K값 손실이 시뮬레이션에 반영됩니다."
)

# 공급배관 구경 선택
supply_pipe_options = [k for k in PIPE_DIMENSIONS.keys() if int(k.replace("A","")) >= 50]
supply_pipe_size = st.sidebar.selectbox(
    "공급배관(라이저) 구경",
    supply_pipe_options,
    index=supply_pipe_options.index(DEFAULT_SUPPLY_PIPE_SIZE),
    help="밸브가 설치된 수직 라이저 배관의 구경입니다. 유속 계산에 사용됩니다.",
)

with st.sidebar.expander("밸브류 ON/OFF 및 K값 설정", expanded=False):
    equipment_k_factors = {}
    for name, info in DEFAULT_EQUIPMENT_K_FACTORS.items():
        col_chk, col_k, col_q = st.columns([2, 1.2, 0.8])
        with col_chk:
            enabled = st.checkbox(name, value=True, key=f"equip_{name}")
        with col_k:
            k_val = st.number_input(
                "K", min_value=0.0, max_value=20.0,
                value=info["K"], step=0.05, key=f"equip_K_{name}",
                label_visibility="collapsed",
            )
        with col_q:
            qty = st.number_input(
                "qty", min_value=0, max_value=10,
                value=info["qty"], step=1, key=f"equip_qty_{name}",
                label_visibility="collapsed",
            )
        if enabled and qty > 0:
            equipment_k_factors[name] = {"K": k_val, "qty": qty}

    # 합산 K값 표시
    total_equiv_K = sum(v["K"] * v["qty"] for v in equipment_k_factors.values())
    st.caption(f"**선택된 밸브 등가 K값 합계**: {total_equiv_K:.2f}")

if not equipment_k_factors:
    equipment_k_factors = None

# ── 4. 펌프 선택 ──
st.sidebar.header(":material/water_pump: 펌프 선택")

pump_model = st.sidebar.radio(
    "펌프 모델",
    list(PUMP_DATABASE.keys()),
    format_func=lambda x: f"{x} ({PUMP_DATABASE[x]['description']})",
)

# ── 5. 시뮬레이션 파라미터 ──
st.sidebar.header(":material/science: 시뮬레이션")

mc_iterations = st.sidebar.number_input(
    "몬테카를로 반복 횟수", min_value=10, max_value=10000,
    value=DEFAULT_MC_ITERATIONS, step=10,
)

col_d1, col_d2 = st.sidebar.columns(2)
with col_d1:
    min_defects = st.number_input(
        "최소 결함", min_value=1, max_value=total_heads,
        value=min(DEFAULT_MIN_DEFECTS, total_heads),
    )
with col_d2:
    max_defects = st.number_input(
        "최대 결함", min_value=1, max_value=total_heads,
        value=min(DEFAULT_MAX_DEFECTS, total_heads),
    )

# ── 6. 경제성 ──
st.sidebar.header(":material/payments: 경제성")

operating_hours = st.sidebar.number_input(
    "연간 운전시간 (hr)", min_value=500, max_value=8760,
    value=int(DEFAULT_OPERATING_HOURS_PER_YEAR), step=100,
)
electricity_rate = st.sidebar.number_input(
    "전기요금 (KRW/kWh)", min_value=50, max_value=500,
    value=int(DEFAULT_ELECTRICITY_RATE_KRW), step=10,
)

run_button = st.sidebar.button(":material/rocket_launch: 시뮬레이션 실행", type="primary", use_container_width=True)


# ══════════════════════════════════════════════
#  메인 영역
# ══════════════════════════════════════════════

if run_button or "results" in st.session_state:

    if run_button:
        try:
            with st.spinner("동적 배관망 수리계산 실행 중..."):
                case_results = compare_dynamic_cases_with_topology(
                    topology=topology_key,
                    num_branches=num_branches,
                    heads_per_branch=heads_per_branch,
                    branch_spacing_m=branch_spacing,
                    head_spacing_m=head_spacing,
                    inlet_pressure_mpa=inlet_pressure,
                    total_flow_lpm=float(design_flow),
                    bead_height_existing=bead_height,
                    bead_height_new=0.0,
                    beads_per_branch=beads_per_branch,
                    relaxation=hc_relaxation,
                    equipment_k_factors=equipment_k_factors,
                    supply_pipe_size=supply_pipe_size,
                )

                # * 안전장치 4: Grid 모드 수렴 실패 / 발산 감지 시 에러 메시지
                if topology_key == "grid" and "system_A" in case_results:
                    sys_A_res = case_results["system_A"]
                    if sys_A_res.get("hc_converged") is False:
                        if sys_A_res.get("diverged", False):
                            st.error(
                                "**연산 수렴 실패 (발산 감지)**: "
                                "배관망 규모가 너무 크거나 구조가 불안정합니다. "
                                "가지배관 개수를 줄이거나 교차배관 구경을 늘려보세요. "
                                "또는 고급 설정에서 이완 계수를 낮춰보세요 "
                                f"(현재: {hc_relaxation})."
                            )
                            st.stop()
                        else:
                            st.warning(
                                "**연산 수렴 미완료**: "
                                f"최대 반복 횟수(1,000회) 내에 수렴하지 못했습니다. "
                                f"최종 오차: {sys_A_res.get('hc_max_imbalance_m', 0):.6f}m. "
                                "고급 설정에서 이완 계수를 조정하거나, "
                                "배관망 규모를 줄여보세요."
                            )

                pump = load_pump(pump_model)
                beads_A_2d = [[bead_height] * heads_per_branch for _ in range(num_branches)]
                beads_B_2d = [[0.0] * heads_per_branch for _ in range(num_branches)]

                sys_A = DynamicSystemCurve(
                    num_branches=num_branches, heads_per_branch=heads_per_branch,
                    branch_spacing_m=branch_spacing, head_spacing_m=head_spacing,
                    bead_heights_2d=beads_A_2d,
                    beads_per_branch=beads_per_branch,
                    bead_height_for_weld_mm=bead_height,
                    topology=topology_key,
                    relaxation=hc_relaxation,
                )
                sys_B = DynamicSystemCurve(
                    num_branches=num_branches, heads_per_branch=heads_per_branch,
                    branch_spacing_m=branch_spacing, head_spacing_m=head_spacing,
                    bead_heights_2d=beads_B_2d,
                    beads_per_branch=0,
                    topology=topology_key,
                    relaxation=hc_relaxation,
                )
                op_A = find_operating_point(pump, sys_A)
                op_B = find_operating_point(pump, sys_B)

                energy = None
                if op_A and op_B:
                    energy = calculate_energy_savings(
                        op_A, op_B,
                        operating_hours_per_year=float(operating_hours),
                        electricity_rate_krw=float(electricity_rate),
                    )

            with st.spinner("몬테카를로 시뮬레이션 중..."):
                mc_results = run_dynamic_monte_carlo(
                    n_iterations=mc_iterations,
                    min_defects=min_defects, max_defects=max_defects,
                    bead_height_mm=bead_height,
                    num_branches=num_branches, heads_per_branch=heads_per_branch,
                    branch_spacing_m=branch_spacing, head_spacing_m=head_spacing,
                    inlet_pressure_mpa=inlet_pressure,
                    total_flow_lpm=float(design_flow),
                    beads_per_branch=beads_per_branch,
                    topology=topology_key,
                    relaxation=hc_relaxation,
                    equipment_k_factors=equipment_k_factors,
                    supply_pipe_size=supply_pipe_size,
                )

            with st.spinner("민감도 분석 중..."):
                sens_results = run_dynamic_sensitivity(
                    bead_height_mm=bead_height,
                    num_branches=num_branches, heads_per_branch=heads_per_branch,
                    branch_spacing_m=branch_spacing, head_spacing_m=head_spacing,
                    inlet_pressure_mpa=inlet_pressure,
                    total_flow_lpm=float(design_flow),
                    beads_per_branch=beads_per_branch,
                    topology=topology_key,
                    relaxation=hc_relaxation,
                    equipment_k_factors=equipment_k_factors,
                    supply_pipe_size=supply_pipe_size,
                )

            st.session_state["results"] = {
                "case": case_results, "pump": pump,
                "sys_A": sys_A, "sys_B": sys_B,
                "op_A": op_A, "op_B": op_B, "energy": energy,
                "mc": mc_results, "sens": sens_results,
                "params": {
                    "num_branches": num_branches,
                    "heads_per_branch": heads_per_branch,
                    "branch_spacing": branch_spacing,
                    "head_spacing": head_spacing,
                    "inlet_pressure": inlet_pressure,
                    "design_flow": design_flow,
                    "bead_height": bead_height,
                    "beads_per_branch": beads_per_branch,
                    "pump_model": pump_model,
                    "topology": topology_key,
                    "equipment_k_factors": equipment_k_factors,
                    "supply_pipe_size": supply_pipe_size,
                },
            }

        except ValidationError as e:
            st.error(f"입력 오류: {e}")
            st.stop()

    res = st.session_state["results"]
    case_results = res["case"]
    pump = res["pump"]
    op_A, op_B = res["op_A"], res["op_B"]
    energy = res["energy"]
    mc_results = res["mc"]
    sens_results = res["sens"]
    params = res["params"]
    n_b = params["num_branches"]
    n_h = params["heads_per_branch"]

    # ── KPI 대시보드 ──
    st.markdown("---")
    bpb_display = params.get("beads_per_branch", 0)
    topo_display = params.get("topology", "tree")
    topo_label = "Full Grid (격자형)" if topo_display == "grid" else "Tree (가지형)"
    sys_info = (
        f"**시스템**: {case_results['cross_main_size']} 교차배관 → "
        f"{n_b}개 가지배관 × {n_h}개 헤드 = **{case_results['total_heads']}개 헤드** | "
        f"토폴로지: **{topo_label}** | "
        f"최악 가지배관: **B#{case_results['worst_branch_A']+1}** (Case A)"
    )
    if bpb_display > 0:
        sys_info += f" | 직관 용접 비드: 가지배관당 **{bpb_display}개**"
    # * 밸브/기기류 손실 정보
    equip_loss_A = case_results.get("system_A", {}).get("equipment_loss_mpa", 0.0)
    if equip_loss_A > 0:
        sys_info += f" | 밸브류 손실: **{equip_loss_A * 1000:.1f} kPa**"
    # * Grid 모드 수렴 정보 표시
    if topo_display == "grid" and "system_A" in case_results:
        sys_A_res = case_results["system_A"]
        if "hc_iterations" in sys_A_res:
            sys_info += (
                f" | HC 수렴: **{sys_A_res['hc_iterations']}회** "
                f"(오차: {sys_A_res['hc_max_imbalance_m']:.4f}m)"
            )
    st.markdown(sys_info)

    kpi1, kpi2, kpi3, kpi4 = st.columns(4)
    with kpi1:
        st.metric(
            "말단 압력 (신기술 B)",
            f"{case_results['terminal_B_mpa']:.4f} MPa",
            delta=f"+{case_results['improvement_pct']:.1f}%" if case_results['improvement_pct'] > 0 else f"{case_results['improvement_pct']:.1f}%",
        )
    with kpi2:
        st.metric("말단 압력 (기존 A)", f"{case_results['terminal_A_mpa']:.4f} MPa")
    with kpi3:
        if energy:
            st.metric("에너지 절감", f"{energy['delta_power_kw']:.2f} kW",
                       delta=f"{energy['annual_cost_savings_krw']:,.0f} KRW/yr")
        else:
            st.metric("에너지 절감", "N/A")
    with kpi4:
        st.metric("0.1 MPa 기준",
                   "PASS ✔" if case_results["pass_fail_B"] else "FAIL ✘")

    st.markdown("---")

    # ── NFPC 규정 준수 자동 판정 ──
    compliance_A = check_nfpc_compliance(case_results["system_A"])
    compliance_B = check_nfpc_compliance(case_results["system_B"])

    if not compliance_A["is_compliant"] or not compliance_B["is_compliant"]:
        for case_label, comp in [("Case A (기존)", compliance_A), ("Case B (신기술)", compliance_B)]:
            for v in comp["velocity_violations"]:
                if v["pipe_type"] == "cross_main":
                    st.error(
                        f"**[{case_label}] NFPC 유속 위반**: "
                        f"교차배관 ({v['pipe_size']}) — "
                        f"**{v['velocity_ms']:.2f} m/s** > {v['limit_ms']} m/s 제한 초과"
                    )
                else:
                    st.error(
                        f"**[{case_label}] NFPC 유속 위반**: "
                        f"가지배관 B#{v['branch']+1} Head #{v['head']} ({v['pipe_size']}) — "
                        f"**{v['velocity_ms']:.2f} m/s** > {v['limit_ms']} m/s 제한 초과"
                    )
            for v in comp["pressure_violations"]:
                if v["type"] == "over":
                    st.error(
                        f"**[{case_label}] NFPC 수압 위반**: "
                        f"가지배관 B#{v['branch']+1} — "
                        f"말단 수압 **{v['pressure_mpa']:.4f} MPa** > {v['limit_mpa']} MPa 상한 초과"
                    )
                else:
                    st.error(
                        f"**[{case_label}] NFPC 수압 위반**: "
                        f"가지배관 B#{v['branch']+1} — "
                        f"말단 수압 **{v['pressure_mpa']:.4f} MPa** < {v['limit_mpa']} MPa 하한 미달"
                    )

    # ── 탭 ──
    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        ":material/show_chart: 압력 프로파일",
        ":material/ssid_chart: P-Q 곡선",
        ":material/casino: 몬테카를로",
        ":material/bar_chart: 민감도 분석",
        ":material/download: 데이터 추출",
        ":material/search: 변수 스캐닝",
        ":material/science: 베르누이 MC",
    ])

    # ═══ Tab 1: 압력 프로파일 ═══
    with tab1:
        st.subheader("최악 가지배관 — 전 구간 누적 압력 프로파일")

        # * 밸브/기기류 손실 상세 분해 표시
        equip_details_A = case_results.get("system_A", {}).get("equipment_loss_details", [])
        if equip_details_A:
            with st.expander("밸브/기기류 국부 손실 상세 분해", expanded=True):
                eq_rows = []
                for d in equip_details_A:
                    eq_rows.append({
                        "부속류": d["name"],
                        "K값": d["K"],
                        "수량": d["qty"],
                        "손실 (kPa)": round(d["loss_mpa"] * 1000, 2),
                    })
                eq_total = sum(r["손실 (kPa)"] for r in eq_rows)
                eq_rows.append({
                    "부속류": "합계",
                    "K값": sum(d["K"] * d["qty"] for d in equip_details_A),
                    "수량": "",
                    "손실 (kPa)": round(eq_total, 2),
                })
                st.dataframe(pd.DataFrame(eq_rows), use_container_width=True, hide_index=True)
                supply_ps = params.get("supply_pipe_size", "100A")
                st.caption(
                    f"공급배관(라이저) 구경: **{supply_ps}** | "
                    f"총 밸브 손실: **{eq_total:.2f} kPa** ({eq_total/1000:.4f} MPa)"
                )

        worst_A = case_results["case_A"]
        worst_B = case_results["case_B"]
        pipe_sizes_worst = res["sens"]["pipe_sizes"]

        labels = ["입구"] + [
            f"H#{i+1}\n({pipe_sizes_worst[i] if i < len(pipe_sizes_worst) else ''})"
            for i in range(n_h)
        ]

        fig_p = go.Figure()
        fig_p.add_trace(go.Scatter(
            x=labels, y=worst_A["pressures_mpa"],
            name=f"Case A (비드 {params['bead_height']}mm)",
            mode="lines+markers",
            line=dict(color="#EF553B", dash="dash", width=2), marker=dict(size=8),
        ))
        fig_p.add_trace(go.Scatter(
            x=labels, y=worst_B["pressures_mpa"],
            name="Case B (비드 0mm, 신기술)",
            mode="lines+markers",
            line=dict(color="#636EFA", width=3), marker=dict(size=8),
        ))
        fig_p.add_hline(y=MIN_TERMINAL_PRESSURE_MPA,
                         line_dash="dot", line_color="green", line_width=2,
                         annotation_text="최소 방수압 0.1 MPa")
        fig_p.update_layout(
            xaxis_title="위치", yaxis_title="압력 (MPa)",
            template="plotly_white", height=500,
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        )
        st.plotly_chart(fig_p, use_container_width=True)

        # * 가지배관별 말단 압력 비교
        st.subheader("전체 가지배관 말단 압력 비교")
        fig_branches = go.Figure()
        fig_branches.add_trace(go.Bar(
            x=[f"B#{i+1}" for i in range(n_b)],
            y=case_results["system_A"]["all_terminal_pressures"],
            name=f"Case A (비드 {params['bead_height']}mm)",
            marker_color="#EF553B", opacity=0.7,
        ))
        fig_branches.add_trace(go.Bar(
            x=[f"B#{i+1}" for i in range(n_b)],
            y=case_results["system_B"]["all_terminal_pressures"],
            name="Case B (비드 0mm)",
            marker_color="#636EFA", opacity=0.7,
        ))
        fig_branches.add_hline(y=MIN_TERMINAL_PRESSURE_MPA,
                                line_dash="dot", line_color="green")
        fig_branches.update_layout(
            barmode="group", xaxis_title="가지배관",
            yaxis_title="말단 압력 (MPa)", template="plotly_white", height=400,
        )
        st.plotly_chart(fig_branches, use_container_width=True)

        with st.expander("최악 가지배관 구간별 상세"):
            det_A = worst_A["segment_details"]
            det_B = worst_B["segment_details"]
            detail_dict = {
                "헤드#": [d["head_number"] for d in det_A],
                "관경": [d["pipe_size"] for d in det_A],
                "유량(LPM)": [d["flow_lpm"] for d in det_A],
                "유속(m/s)": [d["velocity_ms"] for d in det_A],
                "A K1": [d["K1_value"] for d in det_A],
                "B K1": [d["K1_value"] for d in det_B],
            }
            # * 용접 비드 정보 (있을 경우)
            if det_A[0].get("weld_beads_in_seg") is not None:
                detail_dict["A 비드수"] = [d["weld_beads_in_seg"] for d in det_A]
                detail_dict["A 비드손실"] = [d["weld_bead_loss_mpa"] for d in det_A]
            detail_dict.update({
                "A 손실(MPa)": [d["total_seg_loss_mpa"] for d in det_A],
                "B 손실(MPa)": [d["total_seg_loss_mpa"] for d in det_B],
                "A 잔여(MPa)": [d["pressure_after_mpa"] for d in det_A],
                "B 잔여(MPa)": [d["pressure_after_mpa"] for d in det_B],
            })
            df_c = pd.DataFrame(detail_dict)
            st.dataframe(df_c, use_container_width=True, hide_index=True)

        # * Grid 모드: Hardy-Cross 수렴 이력 그래프 (논문용)
        sys_A_data = case_results.get("system_A", {})
        if sys_A_data.get("topology") == "grid" and sys_A_data.get("imbalance_history"):
            st.subheader("Hardy-Cross 수렴 이력")
            hist_imb = sys_A_data["imbalance_history"]
            hist_dq = sys_A_data.get("delta_Q_history", [])

            fig_conv = make_subplots(
                rows=1, cols=2,
                subplot_titles=(
                    "루프 수두 불균형 수렴",
                    "유량 보정값 수렴",
                ),
                horizontal_spacing=0.12,
            )
            fig_conv.add_trace(go.Scatter(
                x=list(range(1, len(hist_imb) + 1)),
                y=hist_imb,
                mode="lines",
                name="Max Loop Imbalance (m)",
                line=dict(color="#636EFA", width=2),
            ), row=1, col=1)
            fig_conv.add_hline(
                y=0.001, line_dash="dash", line_color="red",
                annotation_text="수렴 기준 (0.001 m)",
                annotation_position="top left",
                row=1, col=1,
            )
            if hist_dq:
                fig_conv.add_trace(go.Scatter(
                    x=list(range(1, len(hist_dq) + 1)),
                    y=hist_dq,
                    mode="lines",
                    name="Max ΔQ (LPM)",
                    line=dict(color="#EF553B", width=2),
                ), row=1, col=2)
                fig_conv.add_hline(
                    y=0.0001, line_dash="dash", line_color="red",
                    annotation_text="수렴 기준 (0.0001 LPM)",
                    annotation_position="top left",
                    row=1, col=2,
                )
            fig_conv.update_yaxes(type="log", row=1, col=1)
            fig_conv.update_yaxes(type="log", row=1, col=2)
            fig_conv.update_xaxes(title_text="반복 횟수 (Iteration)", row=1, col=1)
            fig_conv.update_xaxes(title_text="반복 횟수 (Iteration)", row=1, col=2)
            fig_conv.update_yaxes(title_text="최대 루프 불균형 (m)", row=1, col=1)
            fig_conv.update_yaxes(title_text="최대 유량 보정 (LPM)", row=1, col=2)
            fig_conv.update_layout(
                template="plotly_white",
                height=420,
                font=dict(family="Arial", size=13),
                showlegend=False,
                margin=dict(t=50, b=50),
            )
            st.plotly_chart(fig_conv, use_container_width=True)
            st.caption(
                f"총 {sys_A_data.get('hc_iterations', '?')}회 반복 후 수렴 완료. "
                f"최종 루프 오차: {sys_A_data.get('hc_max_imbalance_m', 0):.6f} m, "
                f"최종 유량 보정: {sys_A_data.get('hc_max_delta_Q_lpm', 0):.6f} LPM"
            )

    # ═══ Tab 2: P-Q 곡선 ═══
    with tab2:
        st.subheader("펌프 P-Q 곡선 및 운전점 분석")

        fig_pq = go.Figure()
        Q_pump, H_pump = pump.get_curve_points(100)
        fig_pq.add_trace(go.Scatter(x=Q_pump, y=H_pump,
                                     name=f"펌프: {pump.name}",
                                     line=dict(color="#00CC96", width=3)))

        sys_A_curve = res["sys_A"]
        sys_B_curve = res["sys_B"]
        Q_sA, H_sA = sys_A_curve.get_curve_points(30, q_max=pump.max_flow)
        Q_sB, H_sB = sys_B_curve.get_curve_points(30, q_max=pump.max_flow)

        fig_pq.add_trace(go.Scatter(x=Q_sA, y=H_sA,
                                     name=f"시스템 A (비드 {params['bead_height']}mm)",
                                     line=dict(color="#EF553B", dash="dash", width=2)))
        fig_pq.add_trace(go.Scatter(x=Q_sB, y=H_sB,
                                     name="시스템 B (비드 0mm)",
                                     line=dict(color="#636EFA", dash="dash", width=2)))

        if op_A:
            fig_pq.add_trace(go.Scatter(
                x=[op_A["flow_lpm"]], y=[op_A["head_m"]],
                name=f"운전점 A ({op_A['flow_lpm']:.0f}LPM, {op_A['head_m']:.1f}m)",
                mode="markers", marker=dict(size=15, color="#EF553B", symbol="circle"),
            ))
        if op_B:
            fig_pq.add_trace(go.Scatter(
                x=[op_B["flow_lpm"]], y=[op_B["head_m"]],
                name=f"운전점 B ({op_B['flow_lpm']:.0f}LPM, {op_B['head_m']:.1f}m)",
                mode="markers", marker=dict(size=15, color="#636EFA", symbol="circle"),
            ))

        fig_pq.update_layout(
            xaxis_title="유량 Q (LPM)", yaxis_title="양정 H (m)",
            template="plotly_white", height=500,
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        )
        st.plotly_chart(fig_pq, use_container_width=True)

        if energy:
            st.markdown("#### 에너지 절감 요약")
            ec1, ec2, ec3, ec4 = st.columns(4)
            ec1.metric("양정 감소", f"{energy['delta_head_m']:.2f} m")
            ec2.metric("동력 절감", f"{energy['delta_power_kw']:.3f} kW")
            ec3.metric("연간 절감", f"{energy['annual_energy_kwh']:.1f} kWh")
            ec4.metric("비용 절감", f"₩{energy['annual_cost_savings_krw']:,.0f}")

    # ═══ Tab 3: 몬테카를로 ═══
    with tab3:
        st.subheader("몬테카를로 시뮬레이션 결과")
        mc_bpb = mc_results.get("beads_per_branch", 0)
        mc_desc = (
            f"전체 **{mc_results['total_fittings']}개** 이음쇠 중 "
            f"무작위 {min_defects}~{max_defects}개 결함 비드"
        )
        if mc_bpb > 0:
            mc_desc += (
                f" + 가지배관당 **{mc_bpb}개** 직관 용접 비드 "
                f"(전체 {mc_bpb * n_b}개, **매 반복 위치 무작위 재배치**)"
            )
        mc_desc += f" → {mc_iterations}회 반복"
        st.markdown(mc_desc)

        mc1, mc2, mc3, mc4 = st.columns(4)
        mc1.metric("평균 말단 압력", f"{mc_results['mean_pressure']:.4f} MPa")
        mc2.metric("표준편차", f"{mc_results['std_pressure']:.4f} MPa")
        mc3.metric("최저 말단 압력", f"{mc_results['min_pressure']:.4f} MPa")
        mc4.metric("0.1 MPa 미달", f"{mc_results['p_below_threshold']*100:.1f}%")

        # ── 논문 품질 히스토그램 + 결함 빈도 ──
        mc_tp = mc_results["terminal_pressures"]
        mean_p = float(np.mean(mc_tp))
        std_p = float(np.std(mc_tp, ddof=1)) if len(mc_tp) > 1 else 0.0

        fig_mc = make_subplots(
            rows=1, cols=2,
            subplot_titles=(
                f"최악 말단 압력 분포 (N={mc_results['n_iterations']})",
                "가지배관별 결함 빈도",
            ),
            horizontal_spacing=0.12,
        )
        fig_mc.add_trace(go.Histogram(
            x=mc_tp, nbinsx=30,
            marker_color="rgba(99,110,250,0.7)",
            marker_line=dict(color="#636EFA", width=0.5),
            name="빈도",
        ), row=1, col=1)
        fig_mc.add_vline(
            x=MIN_TERMINAL_PRESSURE_MPA,
            line_dash="dash", line_color="red",
            annotation_text=f"최소 기준 ({MIN_TERMINAL_PRESSURE_MPA} MPa)",
            annotation_position="top right",
            row=1, col=1,
        )
        fig_mc.add_vline(
            x=mean_p, line_dash="dot", line_color="#00CC96",
            annotation_text=f"μ = {mean_p:.4f}",
            annotation_position="top left",
            row=1, col=1,
        )

        fig_mc.add_trace(go.Bar(
            x=[f"B#{i+1}" for i in range(n_b)],
            y=mc_results["defect_frequency"].tolist(),
            marker_color="rgba(239,85,59,0.7)",
            marker_line=dict(color="#EF553B", width=0.5),
            name="결함 빈도",
        ), row=1, col=2)

        fig_mc.update_xaxes(title_text="말단 압력 (MPa)", row=1, col=1)
        fig_mc.update_yaxes(title_text="빈도 (Frequency)", row=1, col=1)
        fig_mc.update_xaxes(title_text="가지배관 (Branch)", row=1, col=2)
        fig_mc.update_yaxes(title_text="결함 빈도 (Count)", row=1, col=2)
        fig_mc.update_layout(
            template="plotly_white", height=500, showlegend=False,
            font=dict(family="Arial", size=13),
            margin=dict(t=60, b=60),
        )
        st.plotly_chart(fig_mc, use_container_width=True)
        st.caption(f"통계 요약: μ = {mean_p:.4f} MPa, σ = {std_p:.4f} MPa, "
                   f"Min = {mc_results['min_pressure']:.4f} MPa, "
                   f"Max = {float(np.max(mc_tp)):.4f} MPa")

        # ── 논문 품질 박스플롯 + 산포도(Jitter) ──
        fig_box = go.Figure()
        fig_box.add_trace(go.Box(
            y=mc_tp,
            name="말단 압력",
            boxpoints="all",
            jitter=0.3,
            pointpos=-1.5,
            marker=dict(color="rgba(99,110,250,0.4)", size=4),
            line=dict(color="#636EFA"),
        ))
        fig_box.add_hline(
            y=MIN_TERMINAL_PRESSURE_MPA, line_dash="dot", line_color="red",
            annotation_text=f"최소 기준 ({MIN_TERMINAL_PRESSURE_MPA} MPa)",
            annotation_position="bottom right",
        )
        fig_box.add_hline(
            y=MAX_TERMINAL_PRESSURE_MPA, line_dash="dot", line_color="orange",
            annotation_text=f"최대 기준 ({MAX_TERMINAL_PRESSURE_MPA} MPa)",
            annotation_position="top right",
        )
        fig_box.update_layout(
            yaxis_title="말단 압력 (MPa)",
            template="plotly_white", height=400,
            font=dict(family="Arial", size=13),
        )
        st.plotly_chart(fig_box, use_container_width=True)

    # ═══ Tab 4: 민감도 분석 ═══
    with tab4:
        st.subheader("민감도 분석 — 최악 가지배관 헤드 위치별 영향도")
        st.markdown(
            f"가지배관 B#{sens_results['worst_branch']+1}의 각 헤드에 "
            f"비드({params['bead_height']}mm) 단독 배치 → 영향 비교"
        )

        crit = sens_results["critical_point"]
        p_sizes = sens_results["pipe_sizes"]
        st.info(
            f"**임계점**: Head #{crit+1} ({p_sizes[crit]}) — "
            f"압력 강하 {sens_results['deltas'][crit]*1000:.2f} kPa"
        )

        colors = ["#EF553B" if i == crit else "#636EFA" for i in range(n_h)]
        fig_s = go.Figure()
        fig_s.add_trace(go.Bar(
            x=[f"H#{i+1}\n({p_sizes[i]})" for i in range(n_h)],
            y=[d * 1000 for d in sens_results["deltas"]],
            marker_color=colors,
            text=[f"{d*1000:.2f}" for d in sens_results["deltas"]],
            textposition="outside",
        ))
        fig_s.update_layout(
            xaxis_title="헤드 위치", yaxis_title="압력 강하 (kPa)",
            template="plotly_white", height=450,
        )
        st.plotly_chart(fig_s, use_container_width=True)

        rank_data = []
        for rank, idx in enumerate(sens_results["ranking"]):
            rank_data.append({
                "순위": rank + 1,
                "위치": f"Head #{idx+1}",
                "관경": p_sizes[idx],
                "말단 압력 (MPa)": f"{sens_results['single_bead_pressures'][idx]:.4f}",
                "강하량 (kPa)": f"{sens_results['deltas'][idx]*1000:.2f}",
            })
        st.dataframe(pd.DataFrame(rank_data), use_container_width=True, hide_index=True)

    # ═══ Tab 5: 데이터 추출 ═══
    with tab5:
        st.subheader("시뮬레이션 결과 다운로드")

        def gen_excel() -> bytes:
            buf = io.BytesIO()
            with pd.ExcelWriter(buf, engine="openpyxl") as w:
                # Sheet 1: 압력 프로파일 (최악 가지배관)
                worst_A = case_results["case_A"]
                worst_B = case_results["case_B"]
                pd.DataFrame({
                    "위치": ["입구"] + [f"Head #{i+1}" for i in range(n_h)],
                    "Case A (MPa)": worst_A["pressures_mpa"],
                    "Case B (MPa)": worst_B["pressures_mpa"],
                }).to_excel(w, sheet_name="압력 프로파일", index=False)

                # Sheet 2: 가지배관별 말단 압력
                pd.DataFrame({
                    "가지배관": [f"B#{i+1}" for i in range(n_b)],
                    "Case A 말단 (MPa)": case_results["system_A"]["all_terminal_pressures"],
                    "Case B 말단 (MPa)": case_results["system_B"]["all_terminal_pressures"],
                }).to_excel(w, sheet_name="가지배관 말단", index=False)

                # Sheet 3-4: Case A/B 상세 (내경·유량·유속 포함)
                pd.DataFrame(worst_A["segment_details"]).to_excel(w, sheet_name="Case A 상세", index=False)
                pd.DataFrame(worst_B["segment_details"]).to_excel(w, sheet_name="Case B 상세", index=False)

                # Sheet 5: 몬테카를로 + 누적 통계
                tp = mc_results["terminal_pressures"]
                tp_arr = np.array(tp)
                n_mc = mc_results["n_iterations"]
                # 누적 통계 계산
                cum_mean = np.cumsum(tp_arr) / np.arange(1, n_mc + 1)
                cum_std = np.array([float(np.std(tp_arr[:i+1], ddof=1)) if i > 0 else 0.0 for i in range(n_mc)])
                cum_min = np.minimum.accumulate(tp_arr)
                cum_max = np.maximum.accumulate(tp_arr)
                cum_pf = np.cumsum(tp_arr < MIN_TERMINAL_PRESSURE_MPA) / np.arange(1, n_mc + 1) * 100.0

                mc_rows = []
                for idx_mc in range(n_mc):
                    mc_rows.append({
                        "Trial": idx_mc + 1,
                        "Worst Terminal (MPa)": round(float(tp[idx_mc]), 6),
                        "Defect Positions": str(mc_results["defect_configs"][idx_mc]),
                        "누적 평균 (μ, MPa)": round(float(cum_mean[idx_mc]), 6),
                        "누적 표준편차 (σ, MPa)": round(float(cum_std[idx_mc]), 6),
                        "누적 최솟값 (Min, MPa)": round(float(cum_min[idx_mc]), 6),
                        "누적 최댓값 (Max, MPa)": round(float(cum_max[idx_mc]), 6),
                        "규정 미달 확률 (Pf, %)": round(float(cum_pf[idx_mc]), 2),
                    })
                # 최종 통계 요약 행
                mc_rows.append({k: "" for k in mc_rows[0]})
                mc_rows.append({"Trial": "최종 통계 요약", "Worst Terminal (MPa)": "", "Defect Positions": "",
                                "누적 평균 (μ, MPa)": "", "누적 표준편차 (σ, MPa)": "",
                                "누적 최솟값 (Min, MPa)": "", "누적 최댓값 (Max, MPa)": "", "규정 미달 확률 (Pf, %)": ""})
                mc_rows.append({"Trial": "평균 (Mean)", "Worst Terminal (MPa)": round(float(np.mean(tp_arr)), 6), "Defect Positions": "",
                                "누적 평균 (μ, MPa)": "", "누적 표준편차 (σ, MPa)": "",
                                "누적 최솟값 (Min, MPa)": "", "누적 최댓값 (Max, MPa)": "", "규정 미달 확률 (Pf, %)": ""})
                mc_rows.append({"Trial": "표준편차 (Std)", "Worst Terminal (MPa)": round(float(np.std(tp_arr, ddof=1)) if n_mc > 1 else 0.0, 6), "Defect Positions": "",
                                "누적 평균 (μ, MPa)": "", "누적 표준편차 (σ, MPa)": "",
                                "누적 최솟값 (Min, MPa)": "", "누적 최댓값 (Max, MPa)": "", "규정 미달 확률 (Pf, %)": ""})
                mc_rows.append({"Trial": "최솟값 (Min)", "Worst Terminal (MPa)": round(float(np.min(tp_arr)), 6), "Defect Positions": "",
                                "누적 평균 (μ, MPa)": "", "누적 표준편차 (σ, MPa)": "",
                                "누적 최솟값 (Min, MPa)": "", "누적 최댓값 (Max, MPa)": "", "규정 미달 확률 (Pf, %)": ""})
                mc_rows.append({"Trial": "최댓값 (Max)", "Worst Terminal (MPa)": round(float(np.max(tp_arr)), 6), "Defect Positions": "",
                                "누적 평균 (μ, MPa)": "", "누적 표준편차 (σ, MPa)": "",
                                "누적 최솟값 (Min, MPa)": "", "누적 최댓값 (Max, MPa)": "", "규정 미달 확률 (Pf, %)": ""})
                mc_rows.append({"Trial": "규정 미달 확률", "Worst Terminal (MPa)": f"{float(cum_pf[-1]):.2f}%", "Defect Positions": "",
                                "누적 평균 (μ, MPa)": "", "누적 표준편차 (σ, MPa)": "",
                                "누적 최솟값 (Min, MPa)": "", "누적 최댓값 (Max, MPa)": "", "규정 미달 확률 (Pf, %)": ""})
                mc_rows.append({"Trial": "시행 횟수 (N)", "Worst Terminal (MPa)": n_mc, "Defect Positions": "",
                                "누적 평균 (μ, MPa)": "", "누적 표준편차 (σ, MPa)": "",
                                "누적 최솟값 (Min, MPa)": "", "누적 최댓값 (Max, MPa)": "", "규정 미달 확률 (Pf, %)": ""})
                pd.DataFrame(mc_rows).to_excel(w, sheet_name="몬테카를로", index=False)

                # Sheet 6: 민감도
                pd.DataFrame({
                    "Head #": [i+1 for i in range(n_h)],
                    "관경": sens_results["pipe_sizes"],
                    "말단 압력 (MPa)": sens_results["single_bead_pressures"],
                    "강하량 (MPa)": sens_results["deltas"],
                }).to_excel(w, sheet_name="민감도", index=False)

                # Sheet 7: 에너지 절감 (펌프 운전점 데이터 강화)
                if energy:
                    energy_data = dict(energy)
                    if op_A:
                        energy_data["Case A 요구 양정 (m)"] = op_A["head_m"]
                        energy_data["Case A 요구 유량 (LPM)"] = op_A["flow_lpm"]
                    if op_B:
                        energy_data["Case B 요구 양정 (m)"] = op_B["head_m"]
                        energy_data["Case B 요구 유량 (LPM)"] = op_B["flow_lpm"]
                    pd.DataFrame([energy_data]).to_excel(w, sheet_name="에너지 절감", index=False)

                # Sheet 8: 입력 파라미터
                pd.DataFrame([params]).to_excel(w, sheet_name="입력 파라미터", index=False)

                # Sheet 9: Full Grid 노드 데이터 (Grid 모드 전용)
                sys_A = case_results["system_A"]
                if sys_A.get("topology") == "grid" and "node_data" in sys_A:
                    grid_rows = []
                    for nd in sys_A["node_data"]:
                        grid_rows.append({
                            "Node ID": nd["node_id"],
                            "위치": nd["position"],
                            "행 (Row)": nd["row"],
                            "열 (Col)": nd["col"],
                            "입구 노드": "Yes" if nd["is_inlet"] else "No",
                            "수요 유량 (LPM)": nd["demand_lpm"],
                            "유입 유량 (LPM)": nd["inflow_lpm"],
                            "유출 유량 (LPM)": nd["outflow_lpm"],
                            "유량 균형 (LPM)": nd["balance_lpm"],
                            "노드 수압 (MPa)": nd["pressure_mpa"],
                        })
                    # 수렴 정보 요약 행
                    grid_rows.append({k: "" for k in grid_rows[0]})
                    grid_rows.append({
                        "Node ID": "HC 수렴 정보",
                        "위치": "",
                        "행 (Row)": "",
                        "열 (Col)": "",
                        "입구 노드": "",
                        "수요 유량 (LPM)": "",
                        "유입 유량 (LPM)": "",
                        "유출 유량 (LPM)": "",
                        "유량 균형 (LPM)": "",
                        "노드 수압 (MPa)": "",
                    })
                    grid_rows.append({
                        "Node ID": "수렴 반복 횟수",
                        "위치": sys_A.get("hc_iterations", "N/A"),
                        "행 (Row)": "",
                        "열 (Col)": "",
                        "입구 노드": "",
                        "수요 유량 (LPM)": "",
                        "유입 유량 (LPM)": "",
                        "유출 유량 (LPM)": "",
                        "유량 균형 (LPM)": "",
                        "노드 수압 (MPa)": "",
                    })
                    grid_rows.append({
                        "Node ID": "최종 루프 오차 (m)",
                        "위치": sys_A.get("hc_max_imbalance_m", "N/A"),
                        "행 (Row)": "",
                        "열 (Col)": "",
                        "입구 노드": "",
                        "수요 유량 (LPM)": "",
                        "유입 유량 (LPM)": "",
                        "유출 유량 (LPM)": "",
                        "유량 균형 (LPM)": "",
                        "노드 수압 (MPa)": "",
                    })
                    grid_rows.append({
                        "Node ID": "최종 유량 보정값 (LPM)",
                        "위치": sys_A.get("hc_max_delta_Q_lpm", "N/A"),
                        "행 (Row)": "",
                        "열 (Col)": "",
                        "입구 노드": "",
                        "수요 유량 (LPM)": "",
                        "유입 유량 (LPM)": "",
                        "유출 유량 (LPM)": "",
                        "유량 균형 (LPM)": "",
                        "노드 수압 (MPa)": "",
                    })
                    grid_rows.append({
                        "Node ID": "수렴 여부",
                        "위치": "Yes" if sys_A.get("hc_converged", False) else "No",
                        "행 (Row)": "",
                        "열 (Col)": "",
                        "입구 노드": "",
                        "수요 유량 (LPM)": "",
                        "유입 유량 (LPM)": "",
                        "유출 유량 (LPM)": "",
                        "유량 균형 (LPM)": "",
                        "노드 수압 (MPa)": "",
                    })
                    pd.DataFrame(grid_rows).to_excel(w, sheet_name="Full Grid 노드 데이터", index=False)

                # Sheet 10: 베르누이 MC 요약 (실행된 경우)
                bern_doc = st.session_state.get("bernoulli_results")
                if bern_doc:
                    bern_sum = bern_doc["summary"]
                    pd.DataFrame({
                        "p (비드 확률)": bern_sum["p_values"],
                        "기대 비드 수": bern_sum["expected_bead_counts"],
                        "실측 비드 수": bern_sum["mean_bead_counts"],
                        "평균 수압 (MPa)": bern_sum["mean_pressures"],
                        "표준편차 (MPa)": bern_sum["std_pressures"],
                        "최솟값 (MPa)": bern_sum["min_pressures"],
                        "최댓값 (MPa)": bern_sum["max_pressures"],
                        "규정 미달 Pf (%)": bern_sum["pf_percents"],
                    }).to_excel(w, sheet_name="Bernoulli MC", index=False)

            return buf.getvalue()

        def gen_report_html() -> bytes:
            """논문/정부과제 제출용 상세 분석 리포트 HTML 생성"""
            from datetime import datetime

            topo_kr = "Full Grid (격자형)" if params.get("topology") == "grid" else "Tree (가지형)"
            now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
            tp = mc_results["terminal_pressures"]
            tp_arr = np.array(tp)
            mc_mean = float(np.mean(tp_arr))
            mc_std = float(np.std(tp_arr, ddof=1)) if len(tp_arr) > 1 else 0.0
            mc_min = float(np.min(tp_arr))
            mc_max = float(np.max(tp_arr))
            mc_n = len(tp_arr)
            p_below = mc_results["p_below_threshold"] * 100

            comp_A = check_nfpc_compliance(case_results["system_A"])
            comp_B = check_nfpc_compliance(case_results["system_B"])

            # ── Grid 수렴 정보 ──
            sys_A_data = case_results.get("system_A", {})
            grid_info_html = ""
            if sys_A_data.get("topology") == "grid":
                grid_info_html = f"""
                <tr><td>Hardy-Cross 수렴 반복 횟수</td><td>{sys_A_data.get('hc_iterations', 'N/A')}회</td></tr>
                <tr><td>최종 루프 오차</td><td>{sys_A_data.get('hc_max_imbalance_m', 0):.6f} m</td></tr>
                <tr><td>수렴 여부</td><td>{'Yes' if sys_A_data.get('hc_converged', False) else 'No'}</td></tr>
                """

            # ── NFPC 위반 상세 ──
            def nfpc_detail_rows(comp, label):
                rows = ""
                for v in comp["velocity_violations"]:
                    loc = f"교차배관 ({v['pipe_size']})" if v["pipe_type"] == "cross_main" \
                        else f"B#{v['branch']+1} Head #{v['head']} ({v['pipe_size']})"
                    rows += f'<tr class="fail"><td>{label}</td><td>유속</td><td>{loc}</td><td>{v["velocity_ms"]:.2f} m/s &gt; {v["limit_ms"]} m/s</td></tr>\n'
                for v in comp["pressure_violations"]:
                    kind = "상한 초과" if v["type"] == "over" else "하한 미달"
                    rows += f'<tr class="fail"><td>{label}</td><td>수압</td><td>B#{v["branch"]+1}</td><td>{v["pressure_mpa"]:.4f} MPa — {kind}</td></tr>\n'
                return rows

            violation_rows = nfpc_detail_rows(comp_A, "Case A") + nfpc_detail_rows(comp_B, "Case B")
            nfpc_overall_A = '<span class="pass">PASS</span>' if comp_A["is_compliant"] else '<span class="fail-badge">FAIL</span>'
            nfpc_overall_B = '<span class="pass">PASS</span>' if comp_B["is_compliant"] else '<span class="fail-badge">FAIL</span>'

            # ── 에너지/경제성 ──
            energy_html = ""
            if energy:
                e_head_A = f"{op_A['head_m']:.2f}" if op_A else "N/A"
                e_flow_A = f"{op_A['flow_lpm']:.1f}" if op_A else "N/A"
                e_head_B = f"{op_B['head_m']:.2f}" if op_B else "N/A"
                e_flow_B = f"{op_B['flow_lpm']:.1f}" if op_B else "N/A"
                energy_html = f"""
                <table>
                    <tr><th>항목</th><th>Case A (기존)</th><th>Case B (신기술)</th><th>절감량</th></tr>
                    <tr><td>요구 양정 (m)</td><td>{e_head_A}</td><td>{e_head_B}</td><td>{energy['delta_head_m']:.2f} m</td></tr>
                    <tr><td>요구 유량 (LPM)</td><td>{e_flow_A}</td><td>{e_flow_B}</td><td>{energy['delta_flow_lpm']:.2f} LPM</td></tr>
                    <tr><td>소비 동력 (kW)</td><td>{energy['case_A_power_kw']:.3f}</td><td>{energy['case_B_power_kw']:.3f}</td><td>{energy['delta_power_kw']:.3f} kW</td></tr>
                    <tr><td>연간 에너지 (kWh/yr)</td><td colspan="2" style="text-align:center;">—</td><td><strong>{energy['annual_energy_kwh']:.1f}</strong></td></tr>
                    <tr><td>연간 비용 절감 (KRW/yr)</td><td colspan="2" style="text-align:center;">—</td><td><strong>₩{energy['annual_cost_savings_krw']:,.0f}</strong></td></tr>
                </table>
                """
            else:
                energy_html = "<p>펌프 운전점을 찾을 수 없어 경제성 분석이 생략되었습니다.</p>"

            # ── 위반 테이블 또는 전체 통과 메시지 ──
            if violation_rows:
                violation_table = f"""
                <table>
                    <tr><th>Case</th><th>위반 유형</th><th>위치</th><th>상세</th></tr>
                    {violation_rows}
                </table>
                """
            else:
                violation_table = '<p class="pass" style="font-size:1.1em;">전 구간 규정 위반 사항 없음</p>'

            html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>FiPLSim 시뮬레이션 분석 리포트</title>
<style>
    @media print {{ @page {{ margin: 20mm; }} }}
    * {{ box-sizing: border-box; }}
    body {{ font-family: 'Malgun Gothic', Arial, sans-serif; color: #222; max-width: 900px; margin: 0 auto; padding: 30px; line-height: 1.6; }}
    h1 {{ text-align: center; border-bottom: 3px solid #1a3c6e; padding-bottom: 10px; color: #1a3c6e; }}
    h2 {{ color: #1a3c6e; border-left: 4px solid #1a3c6e; padding-left: 12px; margin-top: 40px; }}
    h3 {{ color: #333; margin-top: 25px; }}
    .subtitle {{ text-align: center; color: #666; margin-top: -10px; font-size: 0.95em; }}
    .meta {{ text-align: center; color: #999; font-size: 0.85em; margin-bottom: 30px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 0.92em; }}
    th {{ background: #1a3c6e; color: white; padding: 10px 12px; text-align: left; }}
    td {{ padding: 8px 12px; border-bottom: 1px solid #ddd; }}
    tr:nth-child(even) {{ background: #f8f9fa; }}
    tr.fail td {{ background: #fff0f0; color: #c0392b; }}
    .pass {{ color: #27ae60; font-weight: bold; }}
    .fail-badge {{ color: #c0392b; font-weight: bold; }}
    .stat-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 15px; margin: 15px 0; }}
    .stat-card {{ background: #f8f9fa; border: 1px solid #dee2e6; border-radius: 8px; padding: 15px; text-align: center; }}
    .stat-card .value {{ font-size: 1.5em; font-weight: bold; color: #1a3c6e; }}
    .stat-card .label {{ font-size: 0.85em; color: #666; }}
    .highlight {{ background: #fff3cd; border: 1px solid #ffc107; border-radius: 6px; padding: 12px 18px; margin: 15px 0; }}
    .critical {{ background: #f8d7da; border: 1px solid #f5c6cb; border-radius: 6px; padding: 12px 18px; margin: 15px 0; }}
    .success {{ background: #d4edda; border: 1px solid #c3e6cb; border-radius: 6px; padding: 12px 18px; margin: 15px 0; }}
    .footer {{ margin-top: 50px; border-top: 1px solid #ddd; padding-top: 15px; text-align: center; color: #999; font-size: 0.8em; }}
</style>
</head>
<body>

<h1>FiPLSim Simulation Analysis Report</h1>
<p class="subtitle">소화배관 시뮬레이션 상세 분석 리포트 — 동적 배관망 유체역학 해석</p>
<p class="meta">생성 일시: {now_str} | FiPLSim: Advanced Fire Protection Pipe Let Simulator</p>

<!-- ═══ Section 1: 시뮬레이션 개요 ═══ -->
<h2>1. 시뮬레이션 개요 (Simulation Overview)</h2>
<h3>1.1 경계 조건 (Boundary Conditions)</h3>
<table>
    <tr><th>항목</th><th>값</th></tr>
    <tr><td>배관망 토폴로지</td><td><strong>{topo_kr}</strong></td></tr>
    <tr><td>가지배관 수 (n)</td><td>{params['num_branches']}개</td></tr>
    <tr><td>가지배관당 헤드 수 (m)</td><td>{params['heads_per_branch']}개</td></tr>
    <tr><td>총 헤드 수</td><td><strong>{case_results['total_heads']}개</strong></td></tr>
    <tr><td>교차배관 구경</td><td>{case_results['cross_main_size']}</td></tr>
    <tr><td>가지배관 간격</td><td>{params['branch_spacing']} m</td></tr>
    <tr><td>헤드 간격</td><td>{params['head_spacing']} m</td></tr>
    <tr><td>입구 압력</td><td>{params['inlet_pressure']} MPa</td></tr>
    <tr><td>설계 유량</td><td>{params['design_flow']} LPM</td></tr>
    <tr><td>기존 비드 높이 (Case A)</td><td>{params['bead_height']} mm</td></tr>
    <tr><td>직관 용접 비드 수</td><td>가지배관당 {params.get('beads_per_branch', 0)}개</td></tr>
    <tr><td>펌프 모델</td><td>{params.get('pump_model', 'N/A')}</td></tr>
    <tr><td>몬테카를로 반복 횟수</td><td><strong>{mc_n}회</strong></td></tr>
    {grid_info_html}
</table>

<h3>1.2 해석 조건</h3>
<table>
    <tr><th>항목</th><th>Case A (기존 용접)</th><th>Case B (형상제어 신기술)</th></tr>
    <tr><td>이음쇠 비드 높이</td><td>{params['bead_height']} mm</td><td>0.0 mm</td></tr>
    <tr><td>최악 가지배관</td><td>B#{case_results['worst_branch_A']+1}</td><td>B#{case_results['worst_branch_B']+1}</td></tr>
    <tr><td>최악 말단 수압</td><td>{case_results['terminal_A_mpa']:.4f} MPa</td><td>{case_results['terminal_B_mpa']:.4f} MPa</td></tr>
</table>

<!-- ═══ Section 2: 몬테카를로 통계 분석 ═══ -->
<h2>2. 몬테카를로 통계 분석 (Statistical Analysis)</h2>

<p>기존 용접 기술(Case A) 조건에서 용접 결함 위치를 무작위로 변화시킨 <strong>{mc_n}회</strong> 시뮬레이션의 말단 수압 산포도 통계입니다.</p>

<div class="stat-grid">
    <div class="stat-card"><div class="value">{mc_mean:.4f} MPa</div><div class="label">평균 (Mean)</div></div>
    <div class="stat-card"><div class="value">{mc_std:.4f} MPa</div><div class="label">표준편차 (Std Dev)</div></div>
    <div class="stat-card"><div class="value">{mc_min:.4f} MPa</div><div class="label">최솟값 (Min)</div></div>
    <div class="stat-card"><div class="value">{mc_max:.4f} MPa</div><div class="label">최댓값 (Max)</div></div>
</div>

<div class="{'critical' if p_below > 0 else 'success'}">
    <strong>치명적 결함 확률</strong>: 시뮬레이션 {mc_n}회 중 최소 방수압(0.1 MPa) 미달 발생 확률: <strong>{p_below:.1f}%</strong>
    {'— 규정 미달 위험이 존재합니다.' if p_below > 0 else '— 전 시행에서 규정을 만족합니다.'}
</div>

<!-- ═══ Section 3: 기술 비교 및 경제성 ═══ -->
<h2>3. 기술 비교 및 경제성 (Comparative &amp; Economic Analysis)</h2>

<h3>3.1 최악 조건 압력 개선</h3>
<div class="highlight">
    형상제어 신기술(Case B) 적용 시, 최악 조건(Worst-case) 말단 수압이
    <strong>{case_results['terminal_A_mpa']:.4f} MPa → {case_results['terminal_B_mpa']:.4f} MPa</strong>로
    <strong>+{case_results['improvement_pct']:.1f}%</strong> 개선되었습니다.
</div>

<h3>3.2 펌프 운전점 및 LCC 경제성 분석</h3>
{energy_html}

<!-- ═══ Section 4: NFPC 규정 준수 판정 ═══ -->
<h2>4. NFPC 규정 준수 판정 (Code Compliance)</h2>

<table>
    <tr><th>규정 항목</th><th>기준</th><th>Case A</th><th>Case B</th></tr>
    <tr>
        <td>가지배관 유속 제한</td><td>≤ 6.0 m/s</td>
        <td>{'<span class="pass">PASS</span>' if not any(v['pipe_type']=='branch' for v in comp_A['velocity_violations']) else '<span class="fail-badge">FAIL</span>'}</td>
        <td>{'<span class="pass">PASS</span>' if not any(v['pipe_type']=='branch' for v in comp_B['velocity_violations']) else '<span class="fail-badge">FAIL</span>'}</td>
    </tr>
    <tr>
        <td>교차배관 유속 제한</td><td>≤ 10.0 m/s</td>
        <td>{'<span class="pass">PASS</span>' if not any(v['pipe_type']=='cross_main' for v in comp_A['velocity_violations']) else '<span class="fail-badge">FAIL</span>'}</td>
        <td>{'<span class="pass">PASS</span>' if not any(v['pipe_type']=='cross_main' for v in comp_B['velocity_violations']) else '<span class="fail-badge">FAIL</span>'}</td>
    </tr>
    <tr>
        <td>말단 수압 하한</td><td>≥ 0.1 MPa</td>
        <td>{'<span class="pass">PASS</span>' if not any(v['type']=='under' for v in comp_A['pressure_violations']) else '<span class="fail-badge">FAIL</span>'}</td>
        <td>{'<span class="pass">PASS</span>' if not any(v['type']=='under' for v in comp_B['pressure_violations']) else '<span class="fail-badge">FAIL</span>'}</td>
    </tr>
    <tr>
        <td>말단 수압 상한</td><td>≤ 1.2 MPa</td>
        <td>{'<span class="pass">PASS</span>' if not any(v['type']=='over' for v in comp_A['pressure_violations']) else '<span class="fail-badge">FAIL</span>'}</td>
        <td>{'<span class="pass">PASS</span>' if not any(v['type']=='over' for v in comp_B['pressure_violations']) else '<span class="fail-badge">FAIL</span>'}</td>
    </tr>
    <tr style="font-weight:bold; background:#e8f5e9;">
        <td>종합 판정</td><td>—</td>
        <td>{nfpc_overall_A}</td>
        <td>{nfpc_overall_B}</td>
    </tr>
</table>

{violation_table if violation_rows else ''}

<div class="footer">
    <p>본 리포트는 <strong>FiPLSim (Fire Protection Pipe Let Simulator)</strong>에 의해 자동 생성되었습니다.</p>
    <p>동적 배관망 생성 및 몬테카를로 기반 유체역학 해석 엔진 (PLS) | {now_str}</p>
</div>

</body>
</html>"""
            return html.encode("utf-8")

        def gen_report_docx() -> bytes:
            """논문/정부과제 제출용 상세 분석 리포트 DOCX 생성"""
            from datetime import datetime
            from docx import Document
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()

            # ── 스타일 설정 ──
            style = doc.styles["Normal"]
            style.font.name = "맑은 고딕"
            style.font.size = Pt(10)
            style.paragraph_format.space_after = Pt(4)

            navy = RGBColor(0x1A, 0x3C, 0x6E)
            red = RGBColor(0xC0, 0x39, 0x2B)
            green = RGBColor(0x27, 0xAE, 0x60)

            def add_heading_styled(text, level=1):
                h = doc.add_heading(text, level=level)
                for run in h.runs:
                    run.font.color.rgb = navy
                return h

            def set_cell_shading(cell, color_hex):
                from docx.oxml.ns import qn
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): color_hex,
                    qn("w:val"): "clear",
                })
                shading.append(shd)

            def add_table_from_data(headers, rows):
                t = doc.add_table(rows=1 + len(rows), cols=len(headers))
                t.style = "Light Grid Accent 1"
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, h in enumerate(headers):
                    cell = t.rows[0].cells[i]
                    cell.text = h
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                for r_idx, row in enumerate(rows):
                    for c_idx, val in enumerate(row):
                        cell = t.rows[r_idx + 1].cells[c_idx]
                        cell.text = str(val)
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)
                return t

            # ── 차트 이미지 삽입 헬퍼 ──
            try:
                _test_fig = go.Figure()
                _test_fig.to_image(format="png", engine="kaleido", width=10, height=10)
                charts_available = True
                del _test_fig
            except Exception:
                charts_available = False

            fig_num = [0]

            def add_chart(fig, caption, w=6.0, fw=1200, fh=600):
                """Plotly Figure → PNG → DOCX 이미지 삽입 + 캡션"""
                fig_num[0] += 1
                png = fig.to_image(format="png", width=fw, height=fh, scale=2, engine="kaleido")
                doc.add_picture(io.BytesIO(png), width=Inches(w))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(f"그림 {fig_num[0]}. {caption}")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                r.italic = True

            topo_kr = "Full Grid (격자형)" if params.get("topology") == "grid" else "Tree (가지형)"
            now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
            tp = mc_results["terminal_pressures"]
            tp_arr = np.array(tp)
            mc_mean = float(np.mean(tp_arr))
            mc_std = float(np.std(tp_arr, ddof=1)) if len(tp_arr) > 1 else 0.0
            mc_min = float(np.min(tp_arr))
            mc_max = float(np.max(tp_arr))
            mc_n = len(tp_arr)
            p_below = mc_results["p_below_threshold"] * 100

            comp_A = check_nfpc_compliance(case_results["system_A"])
            comp_B = check_nfpc_compliance(case_results["system_B"])

            # ═══ 표지 ═══
            title = doc.add_heading("FiPLSim Simulation Analysis Report", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.color.rgb = navy

            sub = doc.add_paragraph("소화배관 시뮬레이션 상세 분석 리포트")
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.runs[0].font.size = Pt(12)
            sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            meta = doc.add_paragraph(f"생성 일시: {now_str}  |  FiPLSim: Advanced Fire Protection Pipe Let Simulator")
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.runs[0].font.size = Pt(8)
            meta.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            doc.add_paragraph()  # 빈 줄

            # ═══ Section 1: 시뮬레이션 개요 ═══
            add_heading_styled("1. 시뮬레이션 개요 (Simulation Overview)", level=1)
            add_heading_styled("1.1 경계 조건 (Boundary Conditions)", level=2)

            bc_rows = [
                ("배관망 토폴로지", topo_kr),
                ("가지배관 수 (n)", f"{params['num_branches']}개"),
                ("가지배관당 헤드 수 (m)", f"{params['heads_per_branch']}개"),
                ("총 헤드 수", f"{case_results['total_heads']}개"),
                ("교차배관 구경", case_results["cross_main_size"]),
                ("가지배관 간격", f"{params['branch_spacing']} m"),
                ("헤드 간격", f"{params['head_spacing']} m"),
                ("입구 압력", f"{params['inlet_pressure']} MPa"),
                ("설계 유량", f"{params['design_flow']} LPM"),
                ("기존 비드 높이 (Case A)", f"{params['bead_height']} mm"),
                ("직관 용접 비드 수", f"가지배관당 {params.get('beads_per_branch', 0)}개"),
                ("펌프 모델", params.get("pump_model", "N/A")),
                ("몬테카를로 반복 횟수", f"{mc_n}회"),
            ]
            sys_A_data = case_results.get("system_A", {})
            if sys_A_data.get("topology") == "grid":
                bc_rows.append(("HC 수렴 반복 횟수", f"{sys_A_data.get('hc_iterations', 'N/A')}회"))
                bc_rows.append(("최종 루프 오차", f"{sys_A_data.get('hc_max_imbalance_m', 0):.6f} m"))
                bc_rows.append(("수렴 여부", "Yes" if sys_A_data.get("hc_converged", False) else "No"))
            add_table_from_data(["항목", "값"], bc_rows)

            doc.add_paragraph()
            add_heading_styled("1.2 해석 조건", level=2)
            add_table_from_data(
                ["항목", "Case A (기존 용접)", "Case B (형상제어 신기술)"],
                [
                    ("이음쇠 비드 높이", f"{params['bead_height']} mm", "0.0 mm"),
                    ("최악 가지배관", f"B#{case_results['worst_branch_A']+1}", f"B#{case_results['worst_branch_B']+1}"),
                    ("최악 말단 수압", f"{case_results['terminal_A_mpa']:.4f} MPa", f"{case_results['terminal_B_mpa']:.4f} MPa"),
                ],
            )

            # ── Section 1 차트 삽입 ──
            if charts_available:
                worst_A_doc = case_results["case_A"]
                worst_B_doc = case_results["case_B"]
                n_h_doc = params["heads_per_branch"]
                n_b_doc = params["num_branches"]

                # 1.3 압력 프로파일 차트
                doc.add_paragraph()
                add_heading_styled("1.3 압력 프로파일 (Pressure Profile)", level=2)
                ps_doc = sens_results.get("pipe_sizes", [])
                labels_doc = ["입구"] + [
                    f"H#{i+1} ({ps_doc[i]})" if i < len(ps_doc) else f"H#{i+1}"
                    for i in range(n_h_doc)
                ]
                fig_p_doc = go.Figure()
                fig_p_doc.add_trace(go.Scatter(
                    x=labels_doc, y=worst_A_doc["pressures_mpa"],
                    name=f"Case A (비드 {params['bead_height']}mm)",
                    mode="lines+markers",
                    line=dict(color="#EF553B", dash="dash", width=2), marker=dict(size=8),
                ))
                fig_p_doc.add_trace(go.Scatter(
                    x=labels_doc, y=worst_B_doc["pressures_mpa"],
                    name="Case B (비드 0mm, 신기술)",
                    mode="lines+markers",
                    line=dict(color="#636EFA", width=3), marker=dict(size=8),
                ))
                fig_p_doc.add_hline(y=MIN_TERMINAL_PRESSURE_MPA, line_dash="dot",
                                    line_color="green", line_width=2,
                                    annotation_text=f"최소 방수압 {MIN_TERMINAL_PRESSURE_MPA} MPa")
                fig_p_doc.update_layout(
                    xaxis_title="위치", yaxis_title="압력 (MPa)",
                    template="plotly_white", height=500,
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                    font=dict(family="Arial", size=13),
                )
                add_chart(fig_p_doc, "최악 가지배관 전 구간 누적 압력 프로파일")

                # 구간별 상세 데이터 테이블
                det_A = worst_A_doc.get("segment_details", [])
                det_B = worst_B_doc.get("segment_details", [])
                if det_A and det_B:
                    seg_rows = []
                    for i in range(len(det_A)):
                        seg_rows.append((
                            str(det_A[i]["head_number"]),
                            det_A[i]["pipe_size"],
                            f"{det_A[i]['flow_lpm']:.1f}",
                            f"{det_A[i]['velocity_ms']:.2f}",
                            f"{det_A[i]['total_seg_loss_mpa']:.4f}",
                            f"{det_B[i]['total_seg_loss_mpa']:.4f}",
                            f"{det_A[i]['pressure_after_mpa']:.4f}",
                            f"{det_B[i]['pressure_after_mpa']:.4f}",
                        ))
                    add_table_from_data(
                        ["헤드#", "관경", "유량(LPM)", "유속(m/s)",
                         "A 손실(MPa)", "B 손실(MPa)", "A 잔여(MPa)", "B 잔여(MPa)"],
                        seg_rows,
                    )

                # 1.4 가지배관별 말단 압력 비교
                doc.add_paragraph()
                add_heading_styled("1.4 가지배관별 말단 압력 비교", level=2)
                tp_A_all = case_results["system_A"]["all_terminal_pressures"]
                tp_B_all = case_results["system_B"]["all_terminal_pressures"]
                fig_br_doc = go.Figure()
                fig_br_doc.add_trace(go.Bar(
                    x=[f"B#{i+1}" for i in range(n_b_doc)], y=tp_A_all,
                    name=f"Case A (비드 {params['bead_height']}mm)",
                    marker_color="#EF553B", opacity=0.7,
                ))
                fig_br_doc.add_trace(go.Bar(
                    x=[f"B#{i+1}" for i in range(n_b_doc)], y=tp_B_all,
                    name="Case B (비드 0mm)",
                    marker_color="#636EFA", opacity=0.7,
                ))
                fig_br_doc.add_hline(y=MIN_TERMINAL_PRESSURE_MPA, line_dash="dot", line_color="green")
                fig_br_doc.update_layout(
                    barmode="group", xaxis_title="가지배관", yaxis_title="말단 압력 (MPa)",
                    template="plotly_white", height=400,
                    font=dict(family="Arial", size=13),
                )
                add_chart(fig_br_doc, "전체 가지배관 말단 압력 비교")

                # 가지배관별 데이터 테이블
                br_rows = []
                for i in range(n_b_doc):
                    br_rows.append((
                        f"B#{i+1}",
                        f"{tp_A_all[i]:.4f}",
                        f"{tp_B_all[i]:.4f}",
                        f"{(tp_B_all[i] - tp_A_all[i])*1000:.2f}",
                    ))
                add_table_from_data(["가지배관", "Case A (MPa)", "Case B (MPa)", "차이 (kPa)"], br_rows)

                # 1.5 Hardy-Cross 수렴 이력 (Grid 전용)
                sys_A_doc = case_results.get("system_A", {})
                if sys_A_doc.get("topology") == "grid" and sys_A_doc.get("imbalance_history"):
                    doc.add_paragraph()
                    add_heading_styled("1.5 Hardy-Cross 수렴 이력", level=2)
                    hist_imb = sys_A_doc["imbalance_history"]
                    hist_dq = sys_A_doc.get("delta_Q_history", [])
                    fig_conv_doc = make_subplots(
                        rows=1, cols=2,
                        subplot_titles=("루프 수두 불균형 수렴", "유량 보정값 수렴"),
                        horizontal_spacing=0.15,
                    )
                    fig_conv_doc.add_trace(go.Scatter(
                        x=list(range(1, len(hist_imb)+1)), y=hist_imb,
                        mode="lines", name="Max Imbalance (m)",
                        line=dict(color="#636EFA", width=2),
                    ), row=1, col=1)
                    fig_conv_doc.add_hline(y=0.001, line_dash="dash", line_color="red", row=1, col=1)
                    if hist_dq:
                        fig_conv_doc.add_trace(go.Scatter(
                            x=list(range(1, len(hist_dq)+1)), y=hist_dq,
                            mode="lines", name="Max dQ (LPM)",
                            line=dict(color="#EF553B", width=2),
                        ), row=1, col=2)
                        fig_conv_doc.add_hline(y=0.0001, line_dash="dash", line_color="red", row=1, col=2)
                    fig_conv_doc.update_yaxes(type="log", row=1, col=1)
                    fig_conv_doc.update_yaxes(type="log", row=1, col=2)
                    fig_conv_doc.update_xaxes(title_text="반복 횟수", row=1, col=1)
                    fig_conv_doc.update_xaxes(title_text="반복 횟수", row=1, col=2)
                    fig_conv_doc.update_layout(
                        template="plotly_white", height=420,
                        font=dict(family="Arial", size=13), showlegend=False,
                    )
                    add_chart(fig_conv_doc,
                              f"Hardy-Cross 수렴 이력 (총 {sys_A_doc.get('hc_iterations', '?')}회)",
                              fw=1400, fh=500)

            # ═══ Section 2: 몬테카를로 통계 분석 ═══
            doc.add_page_break()
            add_heading_styled("2. 몬테카를로 통계 분석 (Statistical Analysis)", level=1)
            doc.add_paragraph(
                f"기존 용접 기술(Case A) 조건에서 용접 결함 위치를 무작위로 변화시킨 "
                f"{mc_n}회 시뮬레이션의 말단 수압 산포도 통계입니다."
            )

            add_table_from_data(
                ["통계 항목", "값"],
                [
                    ("평균 (Mean)", f"{mc_mean:.4f} MPa"),
                    ("표준편차 (Std Dev)", f"{mc_std:.4f} MPa"),
                    ("최솟값 (Min)", f"{mc_min:.4f} MPa"),
                    ("최댓값 (Max)", f"{mc_max:.4f} MPa"),
                    ("분산 (Variance)", f"{mc_std**2:.6f} MPa²"),
                    ("시행 횟수 (N)", f"{mc_n}"),
                ],
            )

            doc.add_paragraph()
            p_crit = doc.add_paragraph()
            run_label = p_crit.add_run("치명적 결함 확률: ")
            run_label.bold = True
            run_val = p_crit.add_run(
                f"시뮬레이션 {mc_n}회 중 최소 방수압(0.1 MPa) 미달 발생 확률: {p_below:.1f}%"
            )
            if p_below > 0:
                run_val.font.color.rgb = red
                p_crit.add_run(" — 규정 미달 위험이 존재합니다.").font.color.rgb = red
            else:
                run_val.font.color.rgb = green
                p_crit.add_run(" — 전 시행에서 규정을 만족합니다.").font.color.rgb = green

            # ── Section 2 차트 삽입: MC 히스토그램 + 박스플롯 ──
            if charts_available:
                doc.add_paragraph()
                add_heading_styled("2.2 말단 압력 분포 및 결함 빈도", level=2)
                mc_tp_doc = mc_results["terminal_pressures"]
                mean_p_doc = float(np.mean(mc_tp_doc))
                n_b_mc = params["num_branches"]

                fig_mc_doc = make_subplots(
                    rows=1, cols=2,
                    subplot_titles=(
                        f"최악 말단 압력 분포 (N={mc_n})",
                        "가지배관별 결함 빈도",
                    ),
                    horizontal_spacing=0.15,
                )
                fig_mc_doc.add_trace(go.Histogram(
                    x=mc_tp_doc, nbinsx=30,
                    marker_color="rgba(99,110,250,0.7)",
                    name="빈도",
                ), row=1, col=1)
                fig_mc_doc.add_vline(
                    x=MIN_TERMINAL_PRESSURE_MPA, line_dash="dash", line_color="red",
                    annotation_text=f"최소 기준 ({MIN_TERMINAL_PRESSURE_MPA} MPa)", row=1, col=1,
                )
                fig_mc_doc.add_vline(
                    x=mean_p_doc, line_dash="dot", line_color="#00CC96",
                    annotation_text=f"μ = {mean_p_doc:.4f}", row=1, col=1,
                )
                fig_mc_doc.add_trace(go.Bar(
                    x=[f"B#{i+1}" for i in range(n_b_mc)],
                    y=list(mc_results["defect_frequency"]),
                    marker_color="rgba(239,85,59,0.7)", name="결함 빈도",
                ), row=1, col=2)
                fig_mc_doc.update_xaxes(title_text="말단 압력 (MPa)", row=1, col=1)
                fig_mc_doc.update_yaxes(title_text="빈도 (Frequency)", row=1, col=1)
                fig_mc_doc.update_xaxes(title_text="가지배관 (Branch)", row=1, col=2)
                fig_mc_doc.update_yaxes(title_text="결함 빈도 (Count)", row=1, col=2)
                fig_mc_doc.update_layout(
                    template="plotly_white", height=500, showlegend=False,
                    font=dict(family="Arial", size=13),
                )
                add_chart(fig_mc_doc, f"몬테카를로 시뮬레이션 — 말단 압력 분포 및 결함 빈도 (N={mc_n})",
                          fw=1400, fh=500)

                # 2.3 박스플롯
                doc.add_paragraph()
                add_heading_styled("2.3 말단 압력 산포도 (Box Plot)", level=2)
                fig_box_doc = go.Figure()
                fig_box_doc.add_trace(go.Box(
                    y=mc_tp_doc, name="말단 압력",
                    boxpoints="all", jitter=0.3, pointpos=-1.5,
                    marker=dict(color="rgba(99,110,250,0.4)", size=4),
                    line=dict(color="#636EFA"),
                ))
                fig_box_doc.add_hline(y=MIN_TERMINAL_PRESSURE_MPA, line_dash="dot",
                                      line_color="red",
                                      annotation_text=f"최소 기준 ({MIN_TERMINAL_PRESSURE_MPA} MPa)")
                fig_box_doc.add_hline(y=MAX_TERMINAL_PRESSURE_MPA, line_dash="dot",
                                      line_color="orange",
                                      annotation_text=f"최대 기준 ({MAX_TERMINAL_PRESSURE_MPA} MPa)")
                fig_box_doc.update_layout(
                    yaxis_title="말단 압력 (MPa)", template="plotly_white", height=400,
                    font=dict(family="Arial", size=13),
                )
                add_chart(fig_box_doc, "몬테카를로 말단 압력 산포도 (Box Plot + Jitter)")

            # ═══ Section 3: 기술 비교 및 경제성 ═══
            add_heading_styled("3. 기술 비교 및 경제성 (Comparative & Economic Analysis)", level=1)
            add_heading_styled("3.1 최악 조건 압력 개선", level=2)

            p_imp = doc.add_paragraph()
            p_imp.add_run("형상제어 신기술(Case B) 적용 시, 최악 조건 말단 수압이 ")
            p_imp.add_run(
                f"{case_results['terminal_A_mpa']:.4f} MPa → {case_results['terminal_B_mpa']:.4f} MPa"
            ).bold = True
            p_imp.add_run("로 ")
            run_pct = p_imp.add_run(f"+{case_results['improvement_pct']:.1f}% 개선")
            run_pct.bold = True
            run_pct.font.color.rgb = green
            p_imp.add_run("되었습니다.")

            add_heading_styled("3.2 펌프 운전점 및 LCC 경제성 분석", level=2)
            if energy:
                e_rows = [
                    ("요구 양정 (m)",
                     f"{op_A['head_m']:.2f}" if op_A else "N/A",
                     f"{op_B['head_m']:.2f}" if op_B else "N/A",
                     f"{energy['delta_head_m']:.2f} m"),
                    ("요구 유량 (LPM)",
                     f"{op_A['flow_lpm']:.1f}" if op_A else "N/A",
                     f"{op_B['flow_lpm']:.1f}" if op_B else "N/A",
                     f"{energy['delta_flow_lpm']:.2f} LPM"),
                    ("소비 동력 (kW)",
                     f"{energy['case_A_power_kw']:.3f}",
                     f"{energy['case_B_power_kw']:.3f}",
                     f"{energy['delta_power_kw']:.3f} kW"),
                    ("연간 에너지 절감 (kWh/yr)", "—", "—",
                     f"{energy['annual_energy_kwh']:.1f}"),
                    ("연간 비용 절감 (KRW/yr)", "—", "—",
                     f"₩{energy['annual_cost_savings_krw']:,.0f}"),
                ]
                add_table_from_data(["항목", "Case A", "Case B", "절감량"], e_rows)
            else:
                doc.add_paragraph("펌프 운전점을 찾을 수 없어 경제성 분석이 생략되었습니다.")

            # ── Section 3 차트: P-Q 곡선 ──
            if charts_available and op_A and op_B:
                doc.add_paragraph()
                add_heading_styled("3.3 펌프 P-Q 곡선 및 운전점", level=2)
                fig_pq_doc = go.Figure()
                Q_pump_d, H_pump_d = pump.get_curve_points(100)
                fig_pq_doc.add_trace(go.Scatter(
                    x=Q_pump_d, y=H_pump_d,
                    name=f"펌프: {pump.name}", line=dict(color="#00CC96", width=3),
                ))
                sys_A_c = res["sys_A"]
                sys_B_c = res["sys_B"]
                Q_sA_d, H_sA_d = sys_A_c.get_curve_points(30, q_max=pump.max_flow)
                Q_sB_d, H_sB_d = sys_B_c.get_curve_points(30, q_max=pump.max_flow)
                fig_pq_doc.add_trace(go.Scatter(
                    x=Q_sA_d, y=H_sA_d,
                    name=f"시스템 A (비드 {params['bead_height']}mm)",
                    line=dict(color="#EF553B", dash="dash", width=2),
                ))
                fig_pq_doc.add_trace(go.Scatter(
                    x=Q_sB_d, y=H_sB_d,
                    name="시스템 B (비드 0mm)",
                    line=dict(color="#636EFA", dash="dash", width=2),
                ))
                fig_pq_doc.add_trace(go.Scatter(
                    x=[op_A["flow_lpm"]], y=[op_A["head_m"]],
                    name=f"운전점 A ({op_A['flow_lpm']:.0f}LPM, {op_A['head_m']:.1f}m)",
                    mode="markers", marker=dict(size=15, color="#EF553B", symbol="circle"),
                ))
                fig_pq_doc.add_trace(go.Scatter(
                    x=[op_B["flow_lpm"]], y=[op_B["head_m"]],
                    name=f"운전점 B ({op_B['flow_lpm']:.0f}LPM, {op_B['head_m']:.1f}m)",
                    mode="markers", marker=dict(size=15, color="#636EFA", symbol="circle"),
                ))
                fig_pq_doc.update_layout(
                    xaxis_title="유량 Q (LPM)", yaxis_title="양정 H (m)",
                    template="plotly_white", height=500,
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                    font=dict(family="Arial", size=13),
                )
                add_chart(fig_pq_doc, "펌프 P-Q 곡선 및 시스템 운전점")

            # ═══ Section 4: 민감도 분석 ═══
            doc.add_page_break()
            add_heading_styled("4. 민감도 분석 (Sensitivity Analysis)", level=1)
            doc.add_paragraph(
                f"가지배관 B#{sens_results['worst_branch']+1}의 각 헤드에 "
                f"비드({params['bead_height']}mm) 단독 배치 시 말단 압력 변화량을 분석합니다."
            )

            if charts_available:
                n_h_sens = params["heads_per_branch"]
                ps_sens = sens_results.get("pipe_sizes", [])
                crit_pt = sens_results["critical_point"]
                colors_s = ["#EF553B" if i == crit_pt else "#636EFA" for i in range(n_h_sens)]
                fig_s_doc = go.Figure()
                fig_s_doc.add_trace(go.Bar(
                    x=[f"H#{i+1} ({ps_sens[i]})" if i < len(ps_sens) else f"H#{i+1}" for i in range(n_h_sens)],
                    y=[d * 1000 for d in sens_results["deltas"]],
                    marker_color=colors_s,
                    text=[f"{d*1000:.2f}" for d in sens_results["deltas"]],
                    textposition="outside",
                ))
                fig_s_doc.update_layout(
                    xaxis_title="헤드 위치", yaxis_title="압력 강하 (kPa)",
                    template="plotly_white", height=450,
                    font=dict(family="Arial", size=13),
                )
                add_chart(fig_s_doc, f"민감도 분석 — 헤드 위치별 압력 강하 (임계점: H#{crit_pt+1})")

            # 민감도 순위 테이블
            add_heading_styled("4.1 민감도 순위", level=2)
            ps_rank = sens_results.get("pipe_sizes", [])
            sens_rows = []
            for rank, idx in enumerate(sens_results["ranking"]):
                sens_rows.append((
                    str(rank + 1),
                    f"Head #{idx+1}",
                    ps_rank[idx] if idx < len(ps_rank) else "N/A",
                    f"{sens_results['single_bead_pressures'][idx]:.4f}",
                    f"{sens_results['deltas'][idx]*1000:.2f}",
                ))
            add_table_from_data(["순위", "위치", "관경", "말단 압력 (MPa)", "강하량 (kPa)"], sens_rows)

            # ═══ Section 5: 변수 스캐닝 (조건부) ═══
            sweep_doc = st.session_state.get("sweep_results")
            if sweep_doc:
                doc.add_page_break()
                _is_mc_doc = sweep_doc.get("sweep_variable") == "mc_iterations"
                _is_bern_doc = sweep_doc.get("sweep_variable") == "bernoulli_p"
                sw_var_names = {
                    "design_flow": "설계 유량 (LPM)",
                    "inlet_pressure": "입구 압력 (MPa)",
                    "bead_height": "비드 높이 (mm)",
                    "heads_per_branch": "가지배관당 헤드 수",
                    "mc_iterations": "몬테카를로 반복 횟수",
                    "bernoulli_p": "비드 존재 확률 (p_b)",
                }
                sw_label = sw_var_names.get(sweep_doc["sweep_variable"], sweep_doc["sweep_variable"])
                sw_vals_doc = sweep_doc["sweep_values"]

                if _is_bern_doc:
                    # ── 베르누이 확률 스캔 DOCX ──
                    add_heading_styled("5. 베르누이 확률 스캔 분석 (Bernoulli p Sweep)", level=1)

                    add_heading_styled("5.1 스캔 설정", level=2)
                    add_table_from_data(["항목", "값"], [
                        ("스캔 변수", sw_label),
                        ("범위", f"{sw_vals_doc[0]:.2f} ~ {sw_vals_doc[-1]:.2f}"),
                        ("총 케이스 수", f"{len(sw_vals_doc)}"),
                    ])

                    doc.add_paragraph()
                    add_heading_styled("5.2 통계 요약", level=2)
                    _bd_mean = sweep_doc["bern_mean"]
                    _bd_std = sweep_doc["bern_std"]
                    _bd_min = sweep_doc["bern_min"]
                    _bd_max = sweep_doc["bern_max"]
                    _bd_pb = sweep_doc["bern_p_below"]
                    _bd_exp = sweep_doc["bern_expected"]
                    _bd_act = sweep_doc["bern_actual"]

                    bd_rows = []
                    for i in range(len(sw_vals_doc)):
                        bd_rows.append((
                            f"{sw_vals_doc[i]:.2f}",
                            f"{_bd_exp[i]:.1f}",
                            f"{_bd_act[i]:.1f}",
                            f"{_bd_mean[i]:.4f}",
                            f"{_bd_std[i]:.6f}",
                            f"{_bd_min[i]:.4f}",
                            f"{_bd_max[i]:.4f}",
                            f"{_bd_pb[i]*100:.2f}",
                        ))
                    add_table_from_data(
                        ["p_b", "기대비드", "실측비드", "평균(MPa)", "표준편차", "최솟값", "최댓값", "Pf(%)"],
                        bd_rows,
                    )

                elif _is_mc_doc:
                    # ── MC 반복 횟수 스캔 DOCX ──
                    add_heading_styled("5. MC 반복 횟수 스캔 분석", level=1)

                    add_heading_styled("5.1 스캔 설정", level=2)
                    add_table_from_data(["항목", "값"], [
                        ("스캔 변수", sw_label),
                        ("범위", f"{int(sw_vals_doc[0])} ~ {int(sw_vals_doc[-1])}"),
                        ("총 케이스 수", f"{len(sw_vals_doc)}"),
                    ])

                    doc.add_paragraph()
                    add_heading_styled("5.2 MC 수렴성 요약", level=2)
                    _mc_mean_doc = sweep_doc["mc_mean"]
                    _mc_std_doc = sweep_doc["mc_std"]
                    _mc_min_doc = sweep_doc["mc_min"]
                    _mc_max_doc = sweep_doc["mc_max"]
                    _mc_pb_doc = sweep_doc["mc_p_below"]
                    p_mc = doc.add_paragraph()
                    p_mc.add_run(f"최종 평균 수압: {_mc_mean_doc[-1]:.4f} MPa  |  "
                                 f"표준편차: {_mc_std_doc[-1]:.6f} MPa  |  "
                                 f"기준 미달 확률: {_mc_pb_doc[-1]*100:.1f}%")

                    # MC 데이터 테이블
                    doc.add_paragraph()
                    add_heading_styled("5.3 스캔 결과 데이터", level=2)
                    mc_doc_rows = []
                    for i in range(len(sw_vals_doc)):
                        mc_doc_rows.append((
                            f"{int(sw_vals_doc[i])}",
                            f"{_mc_mean_doc[i]:.4f}",
                            f"{_mc_std_doc[i]:.6f}",
                            f"{_mc_min_doc[i]:.4f}",
                            f"{_mc_max_doc[i]:.4f}",
                            f"{_mc_pb_doc[i]*100:.1f}",
                        ))
                    add_table_from_data(
                        ["반복 횟수", "평균(MPa)", "표준편차(MPa)", "최솟값(MPa)", "최댓값(MPa)", "미달(%)"],
                        mc_doc_rows,
                    )
                else:
                    # ── 기존 변수 스캔 DOCX ──
                    add_heading_styled("5. 변수 스캐닝 분석 (Variable Sweep)", level=1)

                    add_heading_styled("5.1 스캔 설정", level=2)
                    add_table_from_data(["항목", "값"], [
                        ("스캔 변수", sw_label),
                        ("범위", f"{sw_vals_doc[0]} ~ {sw_vals_doc[-1]}"),
                        ("총 케이스 수", f"{len(sw_vals_doc)}"),
                    ])

                    doc.add_paragraph()
                    add_heading_styled("5.2 임계점 탐지", level=2)
                    ca_str = f"{sweep_doc['critical_A']:.2f}" if sweep_doc.get("critical_A") is not None else "해당 없음 (전 구간 PASS)"
                    cb_str = f"{sweep_doc['critical_B']:.2f}" if sweep_doc.get("critical_B") is not None else "해당 없음 (전 구간 PASS)"
                    p_ca = doc.add_paragraph()
                    p_ca.add_run(f"Case A 임계점: ").bold = True
                    p_ca.add_run(ca_str)
                    p_cb = doc.add_paragraph()
                    p_cb.add_run(f"Case B 임계점: ").bold = True
                    p_cb.add_run(cb_str)

                    # 스캔 그래프
                    if charts_available:
                        doc.add_paragraph()
                        add_heading_styled("5.3 변수-수압 응답 곡선", level=2)
                        fig_sw_doc = go.Figure()
                        fig_sw_doc.add_trace(go.Scatter(
                            x=sw_vals_doc, y=sweep_doc["terminal_A"],
                            name="Case A", mode="lines+markers",
                            line=dict(color="#EF553B", dash="dash", width=2), marker=dict(size=6),
                        ))
                        fig_sw_doc.add_trace(go.Scatter(
                            x=sw_vals_doc, y=sweep_doc["terminal_B"],
                            name="Case B", mode="lines+markers",
                            line=dict(color="#636EFA", width=3), marker=dict(size=6),
                        ))
                        fig_sw_doc.add_hline(y=MIN_TERMINAL_PRESSURE_MPA, line_dash="dot",
                                             line_color="green", line_width=2)
                        if sweep_doc.get("critical_A") is not None:
                            idx_d = sw_vals_doc.index(sweep_doc["critical_A"])
                            fig_sw_doc.add_trace(go.Scatter(
                                x=[sweep_doc["critical_A"]], y=[sweep_doc["terminal_A"][idx_d]],
                                mode="markers", name="A 임계점",
                                marker=dict(size=16, color="#EF553B", symbol="diamond"),
                            ))
                        if sweep_doc.get("critical_B") is not None:
                            idx_d = sw_vals_doc.index(sweep_doc["critical_B"])
                            fig_sw_doc.add_trace(go.Scatter(
                                x=[sweep_doc["critical_B"]], y=[sweep_doc["terminal_B"][idx_d]],
                                mode="markers", name="B 임계점",
                                marker=dict(size=16, color="#636EFA", symbol="diamond"),
                            ))
                        fig_sw_doc.update_layout(
                            xaxis_title=sw_label, yaxis_title="최악 말단 수압 (MPa)",
                            template="plotly_white", height=500,
                            font=dict(family="Arial", size=13),
                        )
                        add_chart(fig_sw_doc, f"{sw_label} 변화에 따른 말단 수압 응답", fw=1200, fh=500)

                    # 전체 데이터 테이블
                    doc.add_paragraph()
                    add_heading_styled("5.4 스캔 결과 데이터", level=2)
                    _int_keys_doc = {"heads_per_branch", "mc_iterations"}
                    sw_data_rows = []
                    for i in range(len(sw_vals_doc)):
                        sw_data_rows.append((
                            f"{sw_vals_doc[i]:.2f}" if sweep_doc["sweep_variable"] not in _int_keys_doc else f"{int(sw_vals_doc[i])}",
                            f"{sweep_doc['terminal_A'][i]:.4f}",
                            f"{sweep_doc['terminal_B'][i]:.4f}",
                            f"{sweep_doc['improvement_pct'][i]:.1f}",
                            "PASS" if sweep_doc["pass_fail_A"][i] else "FAIL",
                            "PASS" if sweep_doc["pass_fail_B"][i] else "FAIL",
                        ))
                    add_table_from_data(
                        [sw_label, "A 수압(MPa)", "B 수압(MPa)", "개선율(%)", "A 판정", "B 판정"],
                        sw_data_rows,
                    )
                nfpc_section_num = "6"
            else:
                nfpc_section_num = "5"

            # ═══ NFPC 규정 준수 판정 ═══
            doc.add_page_break()
            add_heading_styled(f"{nfpc_section_num}. NFPC 규정 준수 판정 (Code Compliance)", level=1)

            def pf(cond):
                return "PASS" if cond else "FAIL"

            nfpc_rows = [
                ("가지배관 유속 제한", "≤ 6.0 m/s",
                 pf(not any(v["pipe_type"] == "branch" for v in comp_A["velocity_violations"])),
                 pf(not any(v["pipe_type"] == "branch" for v in comp_B["velocity_violations"]))),
                ("교차배관 유속 제한", "≤ 10.0 m/s",
                 pf(not any(v["pipe_type"] == "cross_main" for v in comp_A["velocity_violations"])),
                 pf(not any(v["pipe_type"] == "cross_main" for v in comp_B["velocity_violations"]))),
                ("말단 수압 하한", "≥ 0.1 MPa",
                 pf(not any(v["type"] == "under" for v in comp_A["pressure_violations"])),
                 pf(not any(v["type"] == "under" for v in comp_B["pressure_violations"]))),
                ("말단 수압 상한", "≤ 1.2 MPa",
                 pf(not any(v["type"] == "over" for v in comp_A["pressure_violations"])),
                 pf(not any(v["type"] == "over" for v in comp_B["pressure_violations"]))),
                ("종합 판정", "—",
                 pf(comp_A["is_compliant"]),
                 pf(comp_B["is_compliant"])),
            ]
            t_nfpc = add_table_from_data(["규정 항목", "기준", "Case A", "Case B"], nfpc_rows)

            # PASS/FAIL 셀 색상 적용
            for r_idx in range(1, len(t_nfpc.rows)):
                for c_idx in [2, 3]:
                    cell = t_nfpc.rows[r_idx].cells[c_idx]
                    txt = cell.text.strip()
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True
                            if txt == "PASS":
                                run.font.color.rgb = green
                            elif txt == "FAIL":
                                run.font.color.rgb = red

            # 위반 상세 (있을 경우)
            all_violations = comp_A["velocity_violations"] + comp_A["pressure_violations"] \
                           + comp_B["velocity_violations"] + comp_B["pressure_violations"]
            if all_violations:
                doc.add_paragraph()
                add_heading_styled("위반 사항 상세", level=2)
                v_rows = []
                for v in comp_A["velocity_violations"]:
                    loc = f"교차배관 ({v['pipe_size']})" if v["pipe_type"] == "cross_main" \
                        else f"B#{v['branch']+1} H#{v['head']} ({v['pipe_size']})"
                    v_rows.append(("Case A", "유속", loc, f"{v['velocity_ms']:.2f} > {v['limit_ms']} m/s"))
                for v in comp_A["pressure_violations"]:
                    kind = "상한 초과" if v["type"] == "over" else "하한 미달"
                    v_rows.append(("Case A", "수압", f"B#{v['branch']+1}", f"{v['pressure_mpa']:.4f} MPa — {kind}"))
                for v in comp_B["velocity_violations"]:
                    loc = f"교차배관 ({v['pipe_size']})" if v["pipe_type"] == "cross_main" \
                        else f"B#{v['branch']+1} H#{v['head']} ({v['pipe_size']})"
                    v_rows.append(("Case B", "유속", loc, f"{v['velocity_ms']:.2f} > {v['limit_ms']} m/s"))
                for v in comp_B["pressure_violations"]:
                    kind = "상한 초과" if v["type"] == "over" else "하한 미달"
                    v_rows.append(("Case B", "수압", f"B#{v['branch']+1}", f"{v['pressure_mpa']:.4f} MPa — {kind}"))
                add_table_from_data(["Case", "위반 유형", "위치", "상세"], v_rows)

            # ═══ Section 6: 베르누이 MC (조건부) ═══
            bern_doc = st.session_state.get("bernoulli_results")
            if bern_doc:
                doc.add_page_break()
                add_heading_styled("6. 베르누이 MC 분석 (Bernoulli Monte Carlo)", level=1)
                bern_sum = bern_doc["summary"]
                doc.add_paragraph(
                    f"각 접합부에 독립적 확률 p_b로 비드 존재를 설정한 "
                    f"베르누이 MC 시뮬레이션 결과입니다. "
                    f"(MC 반복: {bern_doc['n_iterations']}회, "
                    f"접합부: {bern_doc['total_fittings']}개)"
                )
                bern_headers = ["p_b", "기대 비드", "실측 비드", "평균(MPa)",
                                "표준편차", "최솟값", "최댓값", "Pf(%)"]
                bern_rows = []
                for _bi in range(len(bern_sum["p_values"])):
                    bern_rows.append([
                        f"{bern_sum['p_values'][_bi]:.2f}",
                        f"{bern_sum['expected_bead_counts'][_bi]:.1f}",
                        f"{bern_sum['mean_bead_counts'][_bi]:.1f}",
                        f"{bern_sum['mean_pressures'][_bi]:.4f}",
                        f"{bern_sum['std_pressures'][_bi]:.6f}",
                        f"{bern_sum['min_pressures'][_bi]:.4f}",
                        f"{bern_sum['max_pressures'][_bi]:.4f}",
                        f"{bern_sum['pf_percents'][_bi]:.2f}",
                    ])
                add_table_from_data(bern_headers, bern_rows)

            # ── 푸터 ──
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_f = footer.add_run(
                "본 리포트는 FiPLSim (Fire Protection Pipe Let Simulator)에 의해 자동 생성되었습니다.\n"
                f"동적 배관망 생성 및 몬테카를로 기반 유체역학 해석 엔진 (PLS) | {now_str}"
            )
            run_f.font.size = Pt(8)
            run_f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        c1, c2 = st.columns(2)
        with c1:
            st.download_button("Excel 다운로드", gen_excel(),
                                "FiPLSim_시뮬레이션_결과.xlsx",
                                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                use_container_width=True)
        with c2:
            csv = pd.DataFrame({
                "위치": ["입구"] + [f"Head #{i+1}" for i in range(n_h)],
                "Case A (MPa)": case_results["case_A"]["pressures_mpa"],
                "Case B (MPa)": case_results["case_B"]["pressures_mpa"],
            }).to_csv(index=False).encode("utf-8-sig")
            st.download_button("CSV 다운로드", csv,
                                "FiPLSim_압력_프로파일.csv", "text/csv",
                                use_container_width=True)
        c3, c4 = st.columns(2)
        with c3:
            st.download_button("분석 리포트 (HTML)", gen_report_html(),
                                "FiPLSim_분석_리포트.html", "text/html",
                                use_container_width=True)
        with c4:
            st.download_button("분석 리포트 (DOCX)", gen_report_docx(),
                                "FiPLSim_분석_리포트.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True)

    # ═══════════════════════════════════════════
    #  Tab 6: 변수 스캐닝 (Variable Sweep)
    # ═══════════════════════════════════════════
    with tab6:
        st.header("연속 변수 스캐닝 (Variable Sweep)")
        st.caption("특정 설계 변수를 연속 변화시키며 시스템 임계점(PASS→FAIL)을 자동 탐지합니다.")

        # ── 입력 인터페이스 ──
        sweep_options = {
            "설계 유량 (LPM)": ("design_flow", 100.0, 3000.0, 100.0),
            "입구 압력 (MPa)": ("inlet_pressure", 0.1, 2.0, 0.05),
            "비드 높이 (mm)": ("bead_height", 0.1, 5.0, 0.1),
            "가지배관당 헤드 수": ("heads_per_branch", 1.0, 50.0, 1.0),
            "몬테카를로 반복 횟수": ("mc_iterations", 10.0, 500.0, 10.0),
            "비드 존재 확률 (p_b)": ("bernoulli_p", 0.05, 0.95, 0.05),
        }
        _int_format_keys = {"heads_per_branch", "mc_iterations"}
        col_a, col_b = st.columns([1, 2])
        with col_a:
            sweep_label = st.selectbox("스캔 대상 변수", list(sweep_options.keys()))
        sv_key, sv_start, sv_end, sv_step = sweep_options[sweep_label]
        _is_int_var = sv_key in _int_format_keys
        _fmt = "%.0f" if _is_int_var else "%.2f"

        with col_b:
            sc1, sc2, sc3 = st.columns(3)
            sw_start = sc1.number_input("시작값", value=sv_start, step=sv_step, format=_fmt)
            sw_end = sc2.number_input("종료값", value=sv_end, step=sv_step, format=_fmt)
            sw_step = sc3.number_input("증감 간격", value=sv_step, min_value=sv_step, step=sv_step, format=_fmt)

        n_steps = int((sw_end - sw_start) / sw_step) + 1 if sw_step > 0 else 0
        if sv_key == "mc_iterations":
            st.info(f"총 **{n_steps}개** 반복 횟수 조건으로 몬테카를로 시뮬레이션 수행 예정")
        elif sv_key == "bernoulli_p":
            st.info(f"총 **{n_steps}개** 비드 확률(p) 조건으로 베르누이 MC 수행 예정 (각 {mc_iterations}회 반복)")
        else:
            st.info(f"총 **{n_steps}개** 시뮬레이션 수행 예정 (현재 설정 기준)")

        if st.button(":material/search: 스캔 시작", use_container_width=True):
            with st.spinner(f"변수 스캐닝 중... ({n_steps}개 케이스)"):
                sweep_res = run_variable_sweep(
                    sweep_variable=sv_key,
                    start_val=sw_start, end_val=sw_end, step_val=sw_step,
                    num_branches=num_branches,
                    heads_per_branch=heads_per_branch,
                    branch_spacing_m=branch_spacing,
                    head_spacing_m=head_spacing,
                    inlet_pressure_mpa=inlet_pressure,
                    total_flow_lpm=float(design_flow),
                    bead_height_mm=bead_height,
                    beads_per_branch=beads_per_branch,
                    topology=topology_key,
                    relaxation=hc_relaxation,
                    mc_iterations=mc_iterations,
                    equipment_k_factors=equipment_k_factors,
                    supply_pipe_size=supply_pipe_size,
                )
            st.session_state["sweep_results"] = sweep_res
            st.success(f"스캔 완료! {len(sweep_res['sweep_values'])}개 케이스 분석됨")

        # ── 결과 표시 ──
        if "sweep_results" in st.session_state:
            sw = st.session_state["sweep_results"]
            sv_vals = sw["sweep_values"]
            _is_mc_sweep = sw["sweep_variable"] == "mc_iterations"
            _is_bern_sweep = sw["sweep_variable"] == "bernoulli_p"

            # ────── 베르누이 확률 스캔 전용 결과 ──────
            if _is_bern_sweep:
                _bm = sw["bern_mean"]
                _bs = sw["bern_std"]
                _bmin = sw["bern_min"]
                _bmax = sw["bern_max"]
                _bpb = sw["bern_p_below"]
                _bexp = sw["bern_expected"]
                _bact = sw["bern_actual"]

                st.markdown("#### 베르누이 MC 스캔 결과")
                _bk1, _bk2, _bk3 = st.columns(3)
                _bk1.metric("최종 평균 수압", f"{_bm[-1]:.4f} MPa")
                _bk2.metric("최종 표준편차", f"{_bs[-1]:.6f} MPa")
                _bk3.metric("기준 미달 확률", f"{_bpb[-1]*100:.1f}%")

                # 차트: p별 평균 수압 + 표준편차 밴드
                _bu = [m + s for m, s in zip(_bm, _bs)]
                _bl = [m - s for m, s in zip(_bm, _bs)]
                fig_bs = go.Figure()
                fig_bs.add_trace(go.Scatter(x=sv_vals, y=_bu, mode="lines",
                    line=dict(width=0), showlegend=False, hoverinfo="skip"))
                fig_bs.add_trace(go.Scatter(x=sv_vals, y=_bl, mode="lines",
                    line=dict(width=0), fill="tonexty", fillcolor="rgba(99,110,250,0.15)",
                    name="평균 +/- 1 표준편차"))
                fig_bs.add_trace(go.Scatter(x=sv_vals, y=_bm,
                    name="평균 말단 수압", mode="lines+markers",
                    line=dict(color="#636EFA", width=3), marker=dict(size=7)))
                fig_bs.add_trace(go.Scatter(x=sv_vals, y=_bmin,
                    name="최솟값", mode="lines", line=dict(color="#EF553B", dash="dot")))
                fig_bs.add_trace(go.Scatter(x=sv_vals, y=_bmax,
                    name="최댓값", mode="lines", line=dict(color="#00CC96", dash="dot")))
                fig_bs.add_hline(y=MIN_TERMINAL_PRESSURE_MPA, line_dash="dash",
                    line_color="orange", line_width=2,
                    annotation_text=f"최소 기준 {MIN_TERMINAL_PRESSURE_MPA} MPa")
                fig_bs.update_layout(
                    xaxis_title="비드 존재 확률 (p_b)",
                    yaxis_title="말단 수압 (MPa)",
                    template="plotly_white", height=500,
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                )
                st.plotly_chart(fig_bs, use_container_width=True)

                # Pf 차트
                fig_bpf = go.Figure()
                fig_bpf.add_trace(go.Scatter(x=sv_vals, y=[p*100 for p in _bpb],
                    name="규정 미달 확률 (%)", mode="lines+markers",
                    line=dict(color="#EF553B", width=3), marker=dict(size=7),
                    fill="tozeroy", fillcolor="rgba(239,85,59,0.1)"))
                fig_bpf.update_layout(
                    xaxis_title="비드 존재 확률 (p_b)",
                    yaxis_title="규정 미달 확률 (%)",
                    template="plotly_white", height=350,
                )
                st.plotly_chart(fig_bpf, use_container_width=True)

                # 데이터프레임
                st.dataframe(pd.DataFrame({
                    "p_b": sv_vals,
                    "기대 비드 수": [f"{v:.1f}" for v in _bexp],
                    "실측 비드 수": [f"{v:.1f}" for v in _bact],
                    "평균 (MPa)": [f"{v:.4f}" for v in _bm],
                    "표준편차 (MPa)": [f"{v:.6f}" for v in _bs],
                    "최솟값 (MPa)": [f"{v:.4f}" for v in _bmin],
                    "최댓값 (MPa)": [f"{v:.4f}" for v in _bmax],
                    "Pf (%)": [f"{v*100:.2f}" for v in _bpb],
                }), use_container_width=True, hide_index=True)

                # Excel
                def gen_sweep_excel():
                    df_exp = pd.DataFrame({
                        "p_b": sv_vals,
                        "기대 비드 수": _bexp,
                        "실측 비드 수": _bact,
                        "평균 수압 (MPa)": _bm,
                        "표준편차 (MPa)": _bs,
                        "최솟값 (MPa)": _bmin,
                        "최댓값 (MPa)": _bmax,
                        "규정 미달 Pf (%)": [v * 100 for v in _bpb],
                    })
                    buf = io.BytesIO()
                    with pd.ExcelWriter(buf, engine="openpyxl") as w:
                        df_exp.to_excel(w, sheet_name="Bernoulli p Sweep", index=False)
                    return buf.getvalue()

            # ────── 몬테카를로 반복 횟수 스캔 전용 결과 ──────
            elif _is_mc_sweep:
                mc_mean = sw["mc_mean"]
                mc_std = sw["mc_std"]
                mc_min = sw["mc_min"]
                mc_max = sw["mc_max"]
                mc_pbelow = sw["mc_p_below"]

                # KPI 카드
                st.markdown("#### MC 수렴성 분석 결과")
                _mc_k1, _mc_k2, _mc_k3 = st.columns(3)
                _mc_k1.metric("최종 평균 수압", f"{mc_mean[-1]:.4f} MPa")
                _mc_k2.metric("최종 표준편차", f"{mc_std[-1]:.6f} MPa")
                _mc_k3.metric("기준 미달 확률", f"{mc_pbelow[-1]*100:.1f}%")

                # 그래프 1: 평균 수압 수렴 곡선 + 표준편차 밴드
                st.markdown("#### 반복 횟수별 평균 수압 수렴 곡선")
                fig_mc = go.Figure()
                _upper = [m + s for m, s in zip(mc_mean, mc_std)]
                _lower = [m - s for m, s in zip(mc_mean, mc_std)]
                fig_mc.add_trace(go.Scatter(
                    x=sv_vals, y=_upper, mode="lines", line=dict(width=0),
                    showlegend=False, hoverinfo="skip",
                ))
                fig_mc.add_trace(go.Scatter(
                    x=sv_vals, y=_lower, mode="lines", line=dict(width=0),
                    fill="tonexty", fillcolor="rgba(99,110,250,0.15)",
                    name="평균 +/- 1 표준편차", hoverinfo="skip",
                ))
                fig_mc.add_trace(go.Scatter(
                    x=sv_vals, y=mc_mean, name="평균 말단 수압",
                    mode="lines+markers",
                    line=dict(color="#636EFA", width=3), marker=dict(size=6),
                ))
                fig_mc.add_trace(go.Scatter(
                    x=sv_vals, y=mc_min, name="최솟값",
                    mode="lines", line=dict(color="#EF553B", dash="dot", width=1.5),
                ))
                fig_mc.add_trace(go.Scatter(
                    x=sv_vals, y=mc_max, name="최댓값",
                    mode="lines", line=dict(color="#00CC96", dash="dot", width=1.5),
                ))
                fig_mc.add_hline(y=MIN_TERMINAL_PRESSURE_MPA, line_dash="dash",
                                 line_color="orange", line_width=2,
                                 annotation_text=f"최소 기준 {MIN_TERMINAL_PRESSURE_MPA} MPa")
                fig_mc.update_layout(
                    xaxis_title="몬테카를로 반복 횟수", yaxis_title="말단 수압 (MPa)",
                    template="plotly_white", height=500,
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                    font=dict(family="Arial", size=13),
                )
                st.plotly_chart(fig_mc, use_container_width=True)

                # 그래프 2: 기준 미달 확률 변화
                st.markdown("#### 반복 횟수별 기준 미달 확률 변화")
                fig_pb = go.Figure()
                fig_pb.add_trace(go.Scatter(
                    x=sv_vals, y=[p * 100 for p in mc_pbelow],
                    name="기준 미달 확률",
                    mode="lines+markers",
                    line=dict(color="#EF553B", width=3), marker=dict(size=6),
                    fill="tozeroy", fillcolor="rgba(239,85,59,0.1)",
                ))
                fig_pb.update_layout(
                    xaxis_title="몬테카를로 반복 횟수", yaxis_title="기준 미달 확률 (%)",
                    template="plotly_white", height=400,
                    font=dict(family="Arial", size=13),
                )
                st.plotly_chart(fig_pb, use_container_width=True)

                # 데이터 테이블
                st.markdown("#### 스캔 결과 상세 테이블")
                df_mc = pd.DataFrame({
                    "반복 횟수": [int(v) for v in sv_vals],
                    "평균 수압 (MPa)": [f"{v:.4f}" for v in mc_mean],
                    "표준편차 (MPa)": [f"{v:.6f}" for v in mc_std],
                    "최솟값 (MPa)": [f"{v:.4f}" for v in mc_min],
                    "최댓값 (MPa)": [f"{v:.4f}" for v in mc_max],
                    "기준 미달 (%)": [f"{v*100:.1f}" for v in mc_pbelow],
                })
                st.dataframe(df_mc, use_container_width=True, height=400)

                # Excel 다운로드
                def gen_sweep_excel():
                    df_exp = pd.DataFrame({
                        "반복 횟수": [int(v) for v in sv_vals],
                        "평균 수압 (MPa)": mc_mean,
                        "표준편차 (MPa)": mc_std,
                        "최솟값 (MPa)": mc_min,
                        "최댓값 (MPa)": mc_max,
                        "기준 미달 확률 (%)": [v * 100 for v in mc_pbelow],
                    })
                    buf = io.BytesIO()
                    with pd.ExcelWriter(buf, engine="openpyxl") as w:
                        df_exp.to_excel(w, sheet_name="MC Iterations Sweep", index=False)
                    return buf.getvalue()

            # ────── 기존 변수 스캔 결과 (설계 유량/압력/비드/헤드수) ──────
            else:
                t_A = sw["terminal_A"]
                t_B = sw["terminal_B"]

                # 임계점 KPI
                st.markdown("#### 임계점 탐지 결과 (Critical Point Detection)")
                kc1, kc2 = st.columns(2)
                crit_A_str = f"{sw['critical_A']:.2f}" if sw["critical_A"] is not None else "해당 없음 (전 구간 PASS)"
                crit_B_str = f"{sw['critical_B']:.2f}" if sw["critical_B"] is not None else "해당 없음 (전 구간 PASS)"
                kc1.metric(f"Case A 임계점 ({sweep_label})", crit_A_str)
                kc2.metric(f"Case B 임계점 ({sweep_label})", crit_B_str)

                # 스캔 그래프
                st.markdown("#### 변수-수압 응답 곡선")
                fig_sw = go.Figure()
                fig_sw.add_trace(go.Scatter(
                    x=sv_vals, y=t_A,
                    name=f"Case A (비드 {bead_height}mm)",
                    mode="lines+markers",
                    line=dict(color="#EF553B", dash="dash", width=2), marker=dict(size=6),
                ))
                fig_sw.add_trace(go.Scatter(
                    x=sv_vals, y=t_B,
                    name="Case B (비드 0mm, 신기술)",
                    mode="lines+markers",
                    line=dict(color="#636EFA", width=3), marker=dict(size=6),
                ))
                fig_sw.add_hline(y=MIN_TERMINAL_PRESSURE_MPA, line_dash="dot",
                                 line_color="green", line_width=2,
                                 annotation_text=f"최소 기준 {MIN_TERMINAL_PRESSURE_MPA} MPa")
                # 임계점 마커
                if sw["critical_A"] is not None:
                    idx_ca = sv_vals.index(sw["critical_A"])
                    fig_sw.add_trace(go.Scatter(
                        x=[sw["critical_A"]], y=[t_A[idx_ca]],
                        mode="markers", name=f"A 임계점 ({sw['critical_A']:.2f})",
                        marker=dict(size=16, color="#EF553B", symbol="diamond"),
                        showlegend=True,
                    ))
                if sw["critical_B"] is not None:
                    idx_cb = sv_vals.index(sw["critical_B"])
                    fig_sw.add_trace(go.Scatter(
                        x=[sw["critical_B"]], y=[t_B[idx_cb]],
                        mode="markers", name=f"B 임계점 ({sw['critical_B']:.2f})",
                        marker=dict(size=16, color="#636EFA", symbol="diamond"),
                        showlegend=True,
                    ))
                fig_sw.update_layout(
                    xaxis_title=sweep_label, yaxis_title="최악 말단 수압 (MPa)",
                    template="plotly_white", height=500,
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                    font=dict(family="Arial", size=13),
                )
                st.plotly_chart(fig_sw, use_container_width=True)

                # PASS/FAIL 데이터 테이블
                st.markdown("#### 스캔 결과 상세 테이블")
                df_sw = pd.DataFrame({
                    sweep_label: sv_vals,
                    "Case A 수압 (MPa)": [f"{v:.4f}" for v in t_A],
                    "Case B 수압 (MPa)": [f"{v:.4f}" for v in t_B],
                    "개선율 (%)": [f"{v:.1f}" for v in sw["improvement_pct"]],
                    "Case A": ["PASS" if p else "FAIL" for p in sw["pass_fail_A"]],
                    "Case B": ["PASS" if p else "FAIL" for p in sw["pass_fail_B"]],
                })
                st.dataframe(df_sw, use_container_width=True, height=400)

                # Excel 다운로드
                def gen_sweep_excel():
                    df_exp = pd.DataFrame({
                        sweep_label: sv_vals,
                        "Case A 수압 (MPa)": t_A,
                        "Case B 수압 (MPa)": t_B,
                        "개선율 (%)": sw["improvement_pct"],
                        "Case A 판정": ["PASS" if p else "FAIL" for p in sw["pass_fail_A"]],
                        "Case B 판정": ["PASS" if p else "FAIL" for p in sw["pass_fail_B"]],
                    })
                    buf = io.BytesIO()
                    with pd.ExcelWriter(buf, engine="openpyxl") as w:
                        df_exp.to_excel(w, sheet_name="Variable Sweep", index=False)
                    return buf.getvalue()

            # DOCX 다운로드
            def gen_sweep_docx():
                from datetime import datetime
                from docx import Document
                from docx.shared import Pt, Inches, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.table import WD_TABLE_ALIGNMENT

                doc = Document()
                style = doc.styles["Normal"]
                style.font.name = "맑은 고딕"
                style.font.size = Pt(10)
                navy = RGBColor(0x1A, 0x3C, 0x6E)
                now_str = datetime.now().strftime("%Y-%m-%d %H:%M")

                def heading_s(text, lv=1):
                    h = doc.add_heading(text, level=lv)
                    for r in h.runs:
                        r.font.color.rgb = navy
                    return h

                def tbl(headers, rows):
                    t = doc.add_table(rows=1+len(rows), cols=len(headers))
                    t.style = "Light Grid Accent 1"
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for i, hd in enumerate(headers):
                        c = t.rows[0].cells[i]; c.text = hd
                        for p in c.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in p.runs: r.bold = True; r.font.size = Pt(9)
                    for ri, row in enumerate(rows):
                        for ci, val in enumerate(row):
                            c = t.rows[ri+1].cells[ci]; c.text = str(val)
                            for p in c.paragraphs:
                                for r in p.runs: r.font.size = Pt(9)
                    return t

                # 표지
                if _is_bern_sweep:
                    _doc_title = "FiPLSim Bernoulli p Sweep Report"
                elif _is_mc_sweep:
                    _doc_title = "FiPLSim MC Iterations Sweep Report"
                else:
                    _doc_title = "FiPLSim Variable Sweep Report"
                title = doc.add_heading(_doc_title, level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in title.runs: r.font.color.rgb = navy
                meta = doc.add_paragraph(f"생성 일시: {now_str}  |  FiPLSim: Advanced Fire Protection Pipe Let Simulator")
                meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta.runs[0].font.size = Pt(8)
                meta.runs[0].font.color.rgb = RGBColor(0x99,0x99,0x99)
                doc.add_paragraph()

                # 1. 스캔 설정
                heading_s("1. 스캔 설정 (Sweep Configuration)")
                tbl(["항목", "값"], [
                    ("스캔 변수", sweep_label),
                    ("시작값", f"{sw_start}"),
                    ("종료값", f"{sw_end}"),
                    ("증감 간격", f"{sw_step}"),
                    ("총 케이스 수", f"{len(sv_vals)}"),
                ])

                if _is_bern_sweep:
                    # 베르누이 확률 스캔 DOCX 내용
                    doc.add_paragraph()
                    heading_s("2. 베르누이 MC 요약")
                    doc.add_paragraph(
                        f"최종(p={sv_vals[-1]:.2f}) 평균 수압: {_bm[-1]:.4f} MPa  |  "
                        f"표준편차: {_bs[-1]:.6f} MPa  |  "
                        f"기준 미달 확률: {_bpb[-1]*100:.1f}%"
                    )

                    # 차트
                    try:
                        _png_bs = fig_bs.to_image(format="png", width=1200, height=600, scale=2, engine="kaleido")
                        doc.add_paragraph()
                        heading_s("3. 비드 확률별 평균 수압 곡선")
                        doc.add_picture(io.BytesIO(_png_bs), width=Inches(6.0))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
                    try:
                        _png_bpf = fig_bpf.to_image(format="png", width=1200, height=500, scale=2, engine="kaleido")
                        doc.add_paragraph()
                        heading_s("4. 규정 미달 확률 변화")
                        doc.add_picture(io.BytesIO(_png_bpf), width=Inches(6.0))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass

                    # 데이터 테이블
                    doc.add_page_break()
                    heading_s("5. 스캔 결과 데이터 (Full Data)")
                    bd_rows = []
                    for i in range(len(sv_vals)):
                        bd_rows.append((
                            f"{sv_vals[i]:.2f}",
                            f"{_bexp[i]:.1f}", f"{_bact[i]:.1f}",
                            f"{_bm[i]:.4f}", f"{_bs[i]:.6f}",
                            f"{_bmin[i]:.4f}", f"{_bmax[i]:.4f}",
                            f"{_bpb[i]*100:.2f}",
                        ))
                    tbl(["p_b", "기대비드", "실측비드", "평균(MPa)", "표준편차", "최솟값", "최댓값", "Pf(%)"], bd_rows)

                elif _is_mc_sweep:
                    # MC 전용 DOCX 내용
                    doc.add_paragraph()
                    heading_s("2. MC 수렴성 요약")
                    doc.add_paragraph(
                        f"최종 평균 수압: {mc_mean[-1]:.4f} MPa  |  "
                        f"최종 표준편차: {mc_std[-1]:.6f} MPa  |  "
                        f"기준 미달 확률: {mc_pbelow[-1]*100:.1f}%"
                    )

                    # 그래프
                    try:
                        png1 = fig_mc.to_image(format="png", width=1200, height=600, scale=2, engine="kaleido")
                        doc.add_paragraph()
                        heading_s("3. 반복 횟수별 수렴 곡선")
                        doc.add_picture(io.BytesIO(png1), width=Inches(6.0))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
                    try:
                        png2 = fig_pb.to_image(format="png", width=1200, height=500, scale=2, engine="kaleido")
                        doc.add_paragraph()
                        heading_s("4. 기준 미달 확률 변화")
                        doc.add_picture(io.BytesIO(png2), width=Inches(6.0))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass

                    # 데이터 테이블
                    doc.add_page_break()
                    heading_s("5. 스캔 결과 데이터 (Full Data)")
                    mc_rows = []
                    for i in range(len(sv_vals)):
                        mc_rows.append((
                            f"{int(sv_vals[i])}",
                            f"{mc_mean[i]:.4f}", f"{mc_std[i]:.6f}",
                            f"{mc_min[i]:.4f}", f"{mc_max[i]:.4f}",
                            f"{mc_pbelow[i]*100:.1f}",
                        ))
                    tbl(["반복 횟수", "평균(MPa)", "표준편차(MPa)", "최솟값(MPa)", "최댓값(MPa)", "미달(%)"], mc_rows)

                else:
                    # 기존 변수 스캔 DOCX 내용
                    # 2. 임계점
                    doc.add_paragraph()
                    heading_s("2. 임계점 탐지 (Critical Point)")
                    doc.add_paragraph(
                        f"Case A 임계점: {crit_A_str}  |  Case B 임계점: {crit_B_str}"
                    )

                    # 3. 스캔 그래프
                    try:
                        png = fig_sw.to_image(format="png", width=1200, height=600, scale=2, engine="kaleido")
                        doc.add_paragraph()
                        heading_s("3. 변수-수압 응답 곡선")
                        doc.add_picture(io.BytesIO(png), width=Inches(6.0))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        rc = cap.add_run(f"그림 1. {sweep_label} 변화에 따른 말단 수압 응답")
                        rc.font.size = Pt(9); rc.font.color.rgb = RGBColor(0x66,0x66,0x66); rc.italic = True
                    except Exception:
                        pass

                    # 4. 전체 데이터 테이블
                    doc.add_page_break()
                    heading_s("4. 스캔 결과 데이터 (Full Data)")
                    data_rows = []
                    for i in range(len(sv_vals)):
                        data_rows.append((
                            f"{sv_vals[i]:.2f}" if sv_key not in _int_format_keys else f"{int(sv_vals[i])}",
                            f"{t_A[i]:.4f}", f"{t_B[i]:.4f}",
                            f"{sw['improvement_pct'][i]:.1f}",
                            "PASS" if sw["pass_fail_A"][i] else "FAIL",
                            "PASS" if sw["pass_fail_B"][i] else "FAIL",
                        ))
                    t_data = tbl([sweep_label, "A 수압(MPa)", "B 수압(MPa)", "개선율(%)", "A 판정", "B 판정"], data_rows)

                    # PASS/FAIL 셀 색상
                    g = RGBColor(0x27,0xAE,0x60)
                    rd = RGBColor(0xC0,0x39,0x2B)
                    for ri in range(1, len(t_data.rows)):
                        for ci in [4, 5]:
                            cell = t_data.rows[ri].cells[ci]
                            txt = cell.text.strip()
                            for p in cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for rn in p.runs:
                                    rn.bold = True
                                    rn.font.color.rgb = g if txt == "PASS" else rd

                # 푸터
                doc.add_paragraph()
                ft = doc.add_paragraph()
                ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rf = ft.add_run(f"FiPLSim Variable Sweep Report | {now_str}")
                rf.font.size = Pt(8); rf.font.color.rgb = RGBColor(0x99,0x99,0x99)

                buf = io.BytesIO()
                doc.save(buf)
                return buf.getvalue()

            dc1, dc2 = st.columns(2)
            with dc1:
                st.download_button("스캔 결과 Excel", gen_sweep_excel(),
                                    "FiPLSim_변수스캐닝.xlsx",
                                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True)
            with dc2:
                st.download_button("스캔 리포트 DOCX", gen_sweep_docx(),
                                    "FiPLSim_변수스캐닝_리포트.docx",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True)

    # ═══ Tab 7: 베르누이 MC ═══
    with tab7:
        st.header(":material/science: 베르누이 MC (Bernoulli Monte Carlo)")
        st.caption(
            "각 접합부(이음쇠)에 독립적으로 비드가 존재할 확률 **p_b**를 설정하여 "
            "시공 품질에 따른 말단 수압 산포를 분석합니다. "
            "기존 MC(결함 N개 무작위 선택)와 달리, 각 접합부가 독립적으로 비드 존재 여부를 판정합니다."
        )

        # ── 입력 모드 선택 ──
        bern_mode = st.radio(
            "분석 모드",
            [":material/apps: 이산 프리셋 (Preset Levels)", ":material/timeline: 연속 스캔 (Continuous Sweep)"],
            horizontal=True,
            key="bern_mode_radio",
        )
        _is_preset = "이산" in bern_mode

        if _is_preset:
            preset_options = {
                "우수 시공 (p=0.1)": 0.1,
                "양호 시공 (p=0.3)": 0.3,
                "보통 시공 (p=0.5)": 0.5,
                "열악 시공 (p=0.7)": 0.7,
                "매우 열악 (p=0.9)": 0.9,
            }
            selected_presets = st.multiselect(
                "비드 존재 확률 수준 선택",
                list(preset_options.keys()),
                default=list(preset_options.keys()),
                key="bern_preset_select",
            )
            bern_p_values = [preset_options[k] for k in selected_presets]
            bern_n_iter = st.number_input(
                "MC 반복 횟수 (N)", min_value=10, max_value=50000,
                value=1000, step=100, key="bern_n_iter_preset",
            )
        else:
            bc1, bc2, bc3 = st.columns(3)
            p_start = bc1.number_input("시작 p", value=0.1, min_value=0.01,
                                        max_value=0.99, step=0.05, format="%.2f",
                                        key="bern_p_start")
            p_end = bc2.number_input("종료 p", value=0.9, min_value=0.01,
                                      max_value=0.99, step=0.05, format="%.2f",
                                      key="bern_p_end")
            p_step = bc3.number_input("증감 간격", value=0.1, min_value=0.01,
                                       max_value=0.5, step=0.01, format="%.2f",
                                       key="bern_p_step")
            bern_p_values = np.arange(p_start, p_end + p_step / 2, p_step).tolist()
            bern_n_iter = st.number_input(
                "MC 반복 횟수 (N)", min_value=10, max_value=50000,
                value=1000, step=100, key="bern_n_iter_sweep",
            )

        total_fittings_bern = num_branches * heads_per_branch
        n_p_levels = len(bern_p_values) if bern_p_values else 0
        st.info(
            f"접합부: **{total_fittings_bern}개** ({num_branches} x {heads_per_branch}) | "
            f"비드 높이: **{bead_height} mm** | "
            f"분석 수준: **{n_p_levels}개** | "
            f"MC 반복: **{bern_n_iter}회** | "
            f"총 시뮬레이션: **{n_p_levels * bern_n_iter:,}회**"
        )

        if st.button(":material/science: 베르누이 MC 실행", use_container_width=True,
                      key="bern_run_btn", disabled=(n_p_levels == 0)):
            with st.spinner(f"베르누이 MC 실행 중... ({n_p_levels}개 수준 x {bern_n_iter}회)"):
                bern_results = run_bernoulli_sweep(
                    p_values=bern_p_values,
                    n_iterations=bern_n_iter,
                    bead_height_mm=bead_height,
                    num_branches=num_branches,
                    heads_per_branch=heads_per_branch,
                    branch_spacing_m=branch_spacing,
                    head_spacing_m=head_spacing,
                    inlet_pressure_mpa=inlet_pressure,
                    total_flow_lpm=float(design_flow),
                    beads_per_branch=beads_per_branch,
                    topology=topology_key,
                    relaxation=hc_relaxation,
                    equipment_k_factors=equipment_k_factors,
                    supply_pipe_size=supply_pipe_size,
                )
            st.session_state["bernoulli_results"] = bern_results
            st.success(f"완료! {n_p_levels}개 수준 분석됨")

        # ── 결과 표시 ──
        if "bernoulli_results" in st.session_state:
            br = st.session_state["bernoulli_results"]
            bsm = br["summary"]

            # KPI 카드
            st.markdown("#### 주요 결과 요약")
            bk1, bk2, bk3 = st.columns(3)
            bk1.metric(
                f"p={bsm['p_values'][0]:.1f} 평균 수압",
                f"{bsm['mean_pressures'][0]:.4f} MPa",
            )
            bk2.metric(
                f"p={bsm['p_values'][-1]:.1f} 평균 수압",
                f"{bsm['mean_pressures'][-1]:.4f} MPa",
            )
            fail_p = None
            for _i, _pf in enumerate(bsm["pf_percents"]):
                if _pf > 0:
                    fail_p = bsm["p_values"][_i]
                    break
            bk3.metric(
                "최초 규정 미달 발생 p",
                f"{fail_p:.2f}" if fail_p else "해당 없음 (전 구간 PASS)",
            )

            # ── 차트 1: p별 평균 수압 + 표준편차 밴드 ──
            st.markdown("#### 비드 확률(p)별 평균 말단 수압")
            _upper_b = [m + s for m, s in zip(bsm["mean_pressures"], bsm["std_pressures"])]
            _lower_b = [m - s for m, s in zip(bsm["mean_pressures"], bsm["std_pressures"])]

            fig_bern = go.Figure()
            fig_bern.add_trace(go.Scatter(
                x=bsm["p_values"], y=_upper_b, mode="lines", line=dict(width=0),
                showlegend=False, hoverinfo="skip",
            ))
            fig_bern.add_trace(go.Scatter(
                x=bsm["p_values"], y=_lower_b, mode="lines", line=dict(width=0),
                fill="tonexty", fillcolor="rgba(99,110,250,0.15)",
                name="평균 +/- 1 표준편차",
            ))
            fig_bern.add_trace(go.Scatter(
                x=bsm["p_values"], y=bsm["mean_pressures"],
                name="평균 말단 수압", mode="lines+markers",
                line=dict(color="#636EFA", width=3), marker=dict(size=8),
            ))
            fig_bern.add_trace(go.Scatter(
                x=bsm["p_values"], y=bsm["min_pressures"],
                name="최솟값", mode="lines+markers",
                line=dict(color="#EF553B", dash="dot"), marker=dict(size=5),
            ))
            fig_bern.add_trace(go.Scatter(
                x=bsm["p_values"], y=bsm["max_pressures"],
                name="최댓값", mode="lines+markers",
                line=dict(color="#00CC96", dash="dot"), marker=dict(size=5),
            ))
            fig_bern.add_hline(
                y=MIN_TERMINAL_PRESSURE_MPA, line_dash="dash",
                line_color="orange", line_width=2,
                annotation_text=f"최소 기준 {MIN_TERMINAL_PRESSURE_MPA} MPa",
            )
            fig_bern.update_layout(
                xaxis_title="비드 존재 확률 (p_b)",
                yaxis_title="말단 수압 (MPa)",
                template="plotly_white", height=500,
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                font=dict(family="Arial", size=13),
            )
            st.plotly_chart(fig_bern, use_container_width=True)

            # ── 차트 2: Pf 변화 ──
            st.markdown("#### 비드 확률(p)별 규정 미달 확률 (Pf)")
            fig_pf_b = go.Figure()
            fig_pf_b.add_trace(go.Scatter(
                x=bsm["p_values"], y=bsm["pf_percents"],
                name="규정 미달 확률 (%)", mode="lines+markers",
                line=dict(color="#EF553B", width=3), marker=dict(size=8),
                fill="tozeroy", fillcolor="rgba(239,85,59,0.1)",
            ))
            fig_pf_b.update_layout(
                xaxis_title="비드 존재 확률 (p_b)",
                yaxis_title="규정 미달 확률 (%)",
                template="plotly_white", height=400,
                font=dict(family="Arial", size=13),
            )
            st.plotly_chart(fig_pf_b, use_container_width=True)

            # ── 차트 3: 기대 비드 수 vs 실측 ──
            st.markdown("#### 비드 확률(p)별 기대/실측 비드 개수")
            fig_cnt_b = go.Figure()
            fig_cnt_b.add_trace(go.Bar(
                x=[f"p={p:.2f}" for p in bsm["p_values"]],
                y=bsm["expected_bead_counts"], name="기대 비드 수 (n x p)",
                marker_color="rgba(99,110,250,0.6)",
            ))
            fig_cnt_b.add_trace(go.Bar(
                x=[f"p={p:.2f}" for p in bsm["p_values"]],
                y=bsm["mean_bead_counts"], name="실측 평균 비드 수",
                marker_color="rgba(239,85,59,0.6)",
            ))
            fig_cnt_b.update_layout(
                barmode="group", template="plotly_white", height=400,
                yaxis_title="비드 개수",
                font=dict(family="Arial", size=13),
            )
            st.plotly_chart(fig_cnt_b, use_container_width=True)

            # ── 데이터 테이블 ──
            st.markdown("#### 요약 테이블")
            df_bern = pd.DataFrame({
                "p (비드 확률)": bsm["p_values"],
                "기대 비드 수": [f"{v:.1f}" for v in bsm["expected_bead_counts"]],
                "실측 비드 수": [f"{v:.1f}" for v in bsm["mean_bead_counts"]],
                "평균 수압 (MPa)": [f"{v:.4f}" for v in bsm["mean_pressures"]],
                "표준편차 (MPa)": [f"{v:.6f}" for v in bsm["std_pressures"]],
                "최솟값 (MPa)": [f"{v:.4f}" for v in bsm["min_pressures"]],
                "최댓값 (MPa)": [f"{v:.4f}" for v in bsm["max_pressures"]],
                "규정 미달 Pf (%)": [f"{v:.2f}" for v in bsm["pf_percents"]],
                "판정": ["PASS" if pf == 0 else "FAIL" for pf in bsm["pf_percents"]],
            })
            st.dataframe(df_bern, use_container_width=True, hide_index=True)

            # ── 다운로드: Excel ──
            def gen_bernoulli_excel() -> bytes:
                _buf = io.BytesIO()
                with pd.ExcelWriter(_buf, engine="openpyxl") as _w:
                    # Sheet 1: 요약
                    pd.DataFrame({
                        "p (비드 확률)": bsm["p_values"],
                        "기대 비드 수": bsm["expected_bead_counts"],
                        "실측 비드 수": bsm["mean_bead_counts"],
                        "평균 수압 (MPa)": bsm["mean_pressures"],
                        "표준편차 (MPa)": bsm["std_pressures"],
                        "최솟값 (MPa)": bsm["min_pressures"],
                        "최댓값 (MPa)": bsm["max_pressures"],
                        "규정 미달 Pf (%)": bsm["pf_percents"],
                    }).to_excel(_w, sheet_name="Bernoulli 요약", index=False)

                    # Sheet 2~N: 각 p별 상세 (누적 통계 포함)
                    for _idx, _p_val in enumerate(bsm["p_values"]):
                        _res_i = br["results"][_idx]
                        _tp = np.array(_res_i["terminal_pressures"])
                        _n = len(_tp)
                        _cm = np.cumsum(_tp) / np.arange(1, _n + 1)
                        _cs = np.array([float(np.std(_tp[:j+1], ddof=1)) if j > 0 else 0.0 for j in range(_n)])
                        _cmin = np.minimum.accumulate(_tp)
                        _cmax = np.maximum.accumulate(_tp)
                        _cpf = np.cumsum(_tp < MIN_TERMINAL_PRESSURE_MPA) / np.arange(1, _n + 1) * 100.0

                        pd.DataFrame({
                            "Trial": range(1, _n + 1),
                            "말단 수압 (MPa)": _tp,
                            "비드 개수": _res_i["bead_counts"],
                            "누적 평균 (MPa)": _cm,
                            "누적 표준편차 (MPa)": _cs,
                            "누적 최솟값 (MPa)": _cmin,
                            "누적 최댓값 (MPa)": _cmax,
                            "규정 미달 확률 (%)": _cpf,
                        }).to_excel(_w, sheet_name=f"p={_p_val:.2f}", index=False)

                    # 입력 파라미터 시트
                    pd.DataFrame([{
                        "토폴로지": topology_key,
                        "가지배관 수": num_branches,
                        "가지배관당 헤드 수": heads_per_branch,
                        "비드 높이 (mm)": bead_height,
                        "입구 압력 (MPa)": inlet_pressure,
                        "설계 유량 (LPM)": design_flow,
                        "가지배관당 용접 비드": beads_per_branch,
                        "MC 반복 횟수": br["n_iterations"],
                    }]).to_excel(_w, sheet_name="입력 파라미터", index=False)

                return _buf.getvalue()

            # ── 다운로드: DOCX ──
            def gen_bernoulli_docx() -> bytes:
                from datetime import datetime
                now_str = datetime.now().strftime("%Y-%m-%d %H:%M")

                _doc = Document()
                _style = _doc.styles["Normal"]
                _style.font.size = Pt(10)
                _style.paragraph_format.space_after = Pt(4)
                _style.paragraph_format.line_spacing = 1.3
                _navy = RGBColor(0x1A, 0x3C, 0x6E)

                def _h(text, lv=1):
                    _hd = _doc.add_heading(text, level=lv)
                    for _r in _hd.runs:
                        _r.font.color.rgb = _navy
                    return _hd

                def _tbl(headers, rows):
                    t = _doc.add_table(rows=1 + len(rows), cols=len(headers))
                    t.style = "Table Grid"
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for j, h in enumerate(headers):
                        c = t.rows[0].cells[j]
                        c.text = ""
                        p = c.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run(h)
                        r.bold = True
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        tc = c._tc
                        tcPr = tc.get_or_add_tcPr()
                        shading = tcPr.makeelement(qn("w:shd"), {
                            qn("w:fill"): "1A3C6E", qn("w:val"): "clear"})
                        tcPr.append(shading)
                    for i, row in enumerate(rows):
                        for j, val in enumerate(row):
                            c = t.rows[i + 1].cells[j]
                            c.text = ""
                            p = c.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r = p.add_run(str(val))
                            r.font.size = Pt(9)
                    return t

                # 표지
                _title = _doc.add_heading("FiPLSim Bernoulli MC Analysis Report", level=0)
                _title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for _r in _title.runs:
                    _r.font.color.rgb = _navy
                _doc.add_paragraph()
                _sub = _doc.add_paragraph()
                _sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _sr = _sub.add_run(
                    "베르누이 몬테카를로 시뮬레이션 결과 리포트\n"
                    f"생성: {now_str}"
                )
                _sr.font.size = Pt(11)
                _sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                _doc.add_page_break()

                # 1. 설정
                _h("1. 시뮬레이션 설정", 1)
                _tbl(["항목", "값"], [
                    ("배관망 토폴로지", topology_key),
                    ("가지배관 수 x 헤드 수", f"{num_branches} x {heads_per_branch} = {total_fittings_bern}개"),
                    ("비드 높이", f"{bead_height} mm"),
                    ("입구 압력", f"{inlet_pressure} MPa"),
                    ("설계 유량", f"{design_flow} LPM"),
                    ("가지배관당 용접 비드", f"{beads_per_branch}개"),
                    ("MC 반복 횟수", f"{br['n_iterations']}회"),
                    ("분석 p 수준", f"{len(bsm['p_values'])}개: {bsm['p_values']}"),
                ])

                # 2. 요약 테이블
                _doc.add_paragraph()
                _h("2. 비드 확률별 통계 요약", 1)
                _s_rows = []
                for i in range(len(bsm["p_values"])):
                    _s_rows.append((
                        f"{bsm['p_values'][i]:.2f}",
                        f"{bsm['expected_bead_counts'][i]:.1f}",
                        f"{bsm['mean_bead_counts'][i]:.1f}",
                        f"{bsm['mean_pressures'][i]:.4f}",
                        f"{bsm['std_pressures'][i]:.6f}",
                        f"{bsm['min_pressures'][i]:.4f}",
                        f"{bsm['max_pressures'][i]:.4f}",
                        f"{bsm['pf_percents'][i]:.2f}",
                    ))
                _tbl(["p_b", "기대비드", "실측비드", "평균(MPa)",
                      "표준편차", "최솟값", "최댓값", "Pf(%)"], _s_rows)

                # 3. 그래프 삽입
                try:
                    _png1 = fig_bern.to_image(format="png", width=1200, height=600, scale=2, engine="kaleido")
                    _doc.add_paragraph()
                    _h("3. 비드 확률별 평균 말단 수압 곡선", 1)
                    _doc.add_picture(io.BytesIO(_png1), width=Inches(6.0))
                    _doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    _doc.add_paragraph()
                    _h("3. 비드 확률별 평균 말단 수압 곡선", 1)
                    _doc.add_paragraph("(차트 이미지 생성 실패 — kaleido 패키지 필요)")

                try:
                    _png2 = fig_pf_b.to_image(format="png", width=1200, height=500, scale=2, engine="kaleido")
                    _doc.add_paragraph()
                    _h("4. 규정 미달 확률 변화", 1)
                    _doc.add_picture(io.BytesIO(_png2), width=Inches(6.0))
                    _doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    _doc.add_paragraph()
                    _h("4. 규정 미달 확률 변화", 1)
                    _doc.add_paragraph("(차트 이미지 생성 실패 — kaleido 패키지 필요)")

                # 푸터
                _doc.add_paragraph()
                _ft = _doc.add_paragraph()
                _ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _fr = _ft.add_run(f"FiPLSim Bernoulli MC Report | {now_str}")
                _fr.font.size = Pt(8)
                _fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

                _buf = io.BytesIO()
                _doc.save(_buf)
                return _buf.getvalue()

            st.markdown("---")
            st.markdown("#### 다운로드")
            dc_b1, dc_b2 = st.columns(2)
            with dc_b1:
                st.download_button(
                    ":material/download: 베르누이 MC Excel",
                    gen_bernoulli_excel(),
                    "FiPLSim_Bernoulli_MC.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            with dc_b2:
                st.download_button(
                    ":material/download: 베르누이 MC 리포트 (DOCX)",
                    gen_bernoulli_docx(),
                    "FiPLSim_Bernoulli_MC_리포트.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

else:
    # ── 초기 안내 화면 ──
    st.markdown("---")
    st.markdown("""
    ### 사용 방법

    1. **배관망 구성**: 가지배관 개수(n), 헤드 수(m), 간격을 설정합니다.
    2. **운전 조건**: 입구 압력, 설계 유량을 조정합니다.
    3. **시뮬레이션 실행** 버튼을 클릭하면 동적 배관망이 자동 생성되고 분석됩니다.

    ---
    """)

    if topology_key == "tree":
        st.markdown("#### 동적 배관망 구조 — Tree (가지형)")
        st.code(
            "입구 (Riser)\n"
            "   │\n"
            "═══╤═══╤═══╤═══╤═══  교차배관 (자동 구경: 65A/80A/100A)\n"
            "   │   │   │   │\n"
            "   B1  B2  B3  B4   가지배관 n개 (양방향)\n"
            "   │   │   │   │\n"
            "   H1  H1  H1  H1   각 가지배관에 헤드 m개\n"
            "   H2  H2  H2  H2   (관경 자동 선정: NFSC 103)\n"
            "   ..  ..  ..  ..\n"
            "   Hm  Hm  Hm  Hm   최말단 헤드",
            language=None,
        )
        st.caption("교차배관 1개에서 각 가지배관으로 일방향 분배하는 구조")
    else:
        st.markdown("#### 동적 배관망 구조 — Full Grid (격자형)")
        st.code(
            "입구 1 (Riser 1)          입구 2 (Riser 2)\n"
            "   │                         │\n"
            "═══╤════╤════╤════╤════╤══════╧═══  교차배관 (상단, TOP)\n"
            "   │    │    │    │    │\n"
            "   B1   B2   B3   B4  ...  가지배관 n개 (상/하 연결)\n"
            "   │    │    │    │    │\n"
            "   H1   H1   H1   H1  ...  각 가지배관에 헤드 m개\n"
            "   H2   H2   H2   H2  ...  (관경 자동 선정: NFSC 103)\n"
            "   ..   ..   ..   ..  ..\n"
            "   Hm   Hm   Hm   Hm  ...  최말단 헤드\n"
            "   │    │    │    │    │\n"
            "═══╧════╧════╧════╧════╧═════════  교차배관 (하단, BOTTOM)",
            language=None,
        )
        st.caption(
            "교차배관 2개(TOP/BOT)가 평행 배치되어 각 가지배관의 양 끝이 연결된 격자 구조. "
            "물이 양방향으로 순환하므로 Tree 대비 마찰 손실이 감소하고 말단 압력이 균등해집니다. "
            "Hardy-Cross 반복법으로 각 루프의 유량 균형을 수렴시킵니다."
        )

    st.markdown("""
    #### 핵심 기능
    | 항목 | 설명 |
    |---|---|
    | 동적 생성 | 가지배관 수·헤드 수를 자유롭게 설정 (최대 200×50) |
    | 자동 관경 | 하류 헤드 수 기준 NFSC 103 자동 선정 |
    | 교차배관 | 전체 헤드 수 기준 65A/80A/100A 자동 선정 |
    | 용접 비드 | 가지배관 직관 구간 내 무작위 배치, MC 시 위치 재배치로 산포도 분석 |
    | 방어 프로그래밍 | 0, 음수, 과도한 값 입력 시 에러 메시지 |
    """)
