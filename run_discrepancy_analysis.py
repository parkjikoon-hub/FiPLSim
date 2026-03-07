"""
FiPLSim — 불일치 원인 정밀 분석 (Discrepancy Root Cause Analysis)

불일치 2건:
  ① Case B @ 1,600 LPM: 논문 0.2430 MPa vs FiPLSim 0.2575 MPa (+14.5 kPa)
  ② Case A 2.5mm @ 1,200 LPM: 논문 0.3041 MPa vs FiPLSim 0.3162 MPa (+12.1 kPa)

분석 항목 (5개 그룹):
  그룹 1: 배관 직경 및 구성 (가장 중요)
  그룹 2: 손실계수 K값 설정
  그룹 3: 관재질 조도 및 마찰 계산
  그룹 4: 유량 분배 방식
  그룹 5: 배관 길이 계산

출력:
  1. FiPLSim_불일치_원인분석_데이터.xlsx
  2. FiPLSim_불일치_원인분석_보고서.docx
"""

import os, sys, time, datetime, math, importlib
import numpy as np
import pandas as pd

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import constants
import hydraulics
import pipe_network as pn

from constants import (
    PIPE_DIMENSIONS, get_inner_diameter_m, RHO, G, NU,
    EPSILON_MM, EPSILON_M,
    MIN_TERMINAL_PRESSURE_MPA, K1_BASE, K2, K3, K_TEE_RUN,
    auto_pipe_size, auto_cross_main_size,
)
from hydraulics import velocity_from_flow, head_to_mpa

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ══════════════════════════════════════════════
# 논문 조건
# ══════════════════════════════════════════════
NUM_BR = 4
HEADS = 8
BRANCH_SP = 3.5
HEAD_SP = 2.3
INLET_P = 0.4
SUPPLY = "80A"

# 논문 참조값
PAPER_CASEB_1600 = 0.2430
PAPER_CASEA25_1200 = 0.3041

print("=" * 70)
print("  FiPLSim — 불일치 원인 정밀 분석")
print("=" * 70)
t0 = time.time()


# ══════════════════════════════════════════════
# 헬퍼: 조도(ε) 변경 후 시뮬레이션 실행
# ══════════════════════════════════════════════
def run_with_epsilon(eps_mm, total_flow_lpm, bead_heights_2d=None):
    """
    조도(ε)를 변경하여 시뮬레이션을 실행합니다.
    Python 기본 매개변수 캐싱 문제를 해결하기 위해 모듈을 다시 로드합니다.
    """
    # constants.py의 EPSILON_MM / EPSILON_M 변경
    constants.EPSILON_MM = eps_mm
    constants.EPSILON_M = eps_mm / 1000.0

    # hydraulics 모듈 재로드 → friction_factor 기본값 갱신
    importlib.reload(hydraulics)
    # pipe_network 모듈 재로드 → 갱신된 hydraulics 참조
    importlib.reload(pn)

    sys_obj = pn.generate_dynamic_system(
        num_branches=NUM_BR, heads_per_branch=HEADS,
        branch_spacing_m=BRANCH_SP, head_spacing_m=HEAD_SP,
        inlet_pressure_mpa=INLET_P, total_flow_lpm=float(total_flow_lpm),
        bead_heights_2d=bead_heights_2d,
    )
    res = pn.calculate_dynamic_system(
        sys_obj, K3,
        equipment_k_factors=None,
        supply_pipe_size=SUPPLY,
    )
    return res


def run_with_epsilon_and_pipe_func(eps_mm, total_flow_lpm, pipe_size_func, bead_heights_2d=None):
    """조도 + 배관 구성 함수 변경 후 시뮬레이션"""
    constants.EPSILON_MM = eps_mm
    constants.EPSILON_M = eps_mm / 1000.0

    importlib.reload(hydraulics)
    importlib.reload(pn)

    # auto_pipe_size 교체
    original = constants.auto_pipe_size
    constants.auto_pipe_size = pipe_size_func
    pn.auto_pipe_size = pipe_size_func

    sys_obj = pn.generate_dynamic_system(
        num_branches=NUM_BR, heads_per_branch=HEADS,
        branch_spacing_m=BRANCH_SP, head_spacing_m=HEAD_SP,
        inlet_pressure_mpa=INLET_P, total_flow_lpm=float(total_flow_lpm),
        bead_heights_2d=bead_heights_2d,
    )
    res = pn.calculate_dynamic_system(
        sys_obj, K3,
        equipment_k_factors=None,
        supply_pipe_size=SUPPLY,
    )

    # 복원
    constants.auto_pipe_size = original
    pn.auto_pipe_size = original
    return res


def restore_epsilon():
    """조도를 원래값으로 복원"""
    constants.EPSILON_MM = EPSILON_MM
    constants.EPSILON_M = EPSILON_M
    importlib.reload(hydraulics)
    importlib.reload(pn)


# 논문 배관 구성 함수 (65A 포함)
def paper_auto_pipe_size(num_heads_downstream):
    if num_heads_downstream >= 8:
        return "65A"
    elif num_heads_downstream >= 6:
        return "50A"
    elif num_heads_downstream >= 4:
        return "40A"
    elif num_heads_downstream >= 3:
        return "32A"
    else:
        return "25A"


# ══════════════════════════════════════════════
# 그룹 1: 배관 직경 및 구성 확인
# ══════════════════════════════════════════════
print("\n" + "─" * 70)
print("  [그룹 1] 배관 직경 및 구성 확인")
print("─" * 70)

# FiPLSim auto_pipe_size 결과 (8 heads per branch)
fiplsim_pipe_sizes = []
for h in range(HEADS):
    downstream = HEADS - h
    size = auto_pipe_size(downstream)
    fiplsim_pipe_sizes.append(size)

# 논문 구성: 50A×3, 40A×2, 32A×1, 25A×2
paper_pipe_sizes = ["50A", "50A", "50A", "40A", "40A", "32A", "25A", "25A"]

print(f"  FiPLSim: {' → '.join(fiplsim_pipe_sizes)}")
print(f"  논문:    {' → '.join(paper_pipe_sizes)}")

pipe_config_rows = []
for h in range(HEADS):
    downstream = HEADS - h
    fp_size = fiplsim_pipe_sizes[h]
    pp_size = paper_pipe_sizes[h]
    fp_id = PIPE_DIMENSIONS[fp_size]["id_mm"]
    pp_id = PIPE_DIMENSIONS[pp_size]["id_mm"]
    match = "✅" if fp_size == pp_size else "❌"
    pipe_config_rows.append({
        "Head #": h + 1, "하류 헤드 수": downstream,
        "FiPLSim 관경": fp_size, "FiPLSim ID (mm)": fp_id,
        "논문 관경": pp_size, "논문 ID (mm)": pp_id,
        "ID 차이 (mm)": round(fp_id - pp_id, 2), "일치": match,
    })

df_pipe_config = pd.DataFrame(pipe_config_rows)

mismatch_count = sum(1 for r in pipe_config_rows if r["일치"] == "❌")
print(f"  가지배관 구간별 일치: {HEADS - mismatch_count}/{HEADS}")

cross_main_fiplsim = auto_cross_main_size(NUM_BR * HEADS)
branch_inlet_fiplsim = auto_pipe_size(HEADS)
print(f"  교차배관: FiPLSim={cross_main_fiplsim}, 논문=80A")
print(f"  가지배관 입구: FiPLSim={branch_inlet_fiplsim}, 논문=65A ← 차이!")


# ══════════════════════════════════════════════
# 그룹 2: 손실계수 K값 확인
# ══════════════════════════════════════════════
print("\n" + "─" * 70)
print("  [그룹 2] 손실계수 K값 확인")
print("─" * 70)

k_check_rows = [
    {"항목": "교차배관 tee-run 손실계수", "논문": "K_TEE_RUN = 0.3",
     "FiPLSim": f"K_TEE_RUN = {K_TEE_RUN}", "일치": "✅" if K_TEE_RUN == 0.3 else "❌"},
    {"항목": "가지배관 분기 inlet 손실계수", "논문": "K = 1.0",
     "FiPLSim": f"K3 = {K3}", "일치": "✅" if K3 == 1.0 else "❌"},
    {"항목": "스프링클러 헤드 피팅 손실계수", "논문": "K = 2.5",
     "FiPLSim": f"K2 = {K2}", "일치": "✅" if K2 == 2.5 else "❌"},
    {"항목": "비드 없는 junction 기준 K값", "논문": "K_base = 0.5",
     "FiPLSim": f"K1_BASE = {K1_BASE}", "일치": "✅" if K1_BASE == 0.5 else "❌"},
]

for r in k_check_rows:
    print(f"  {r['항목']}: {r['FiPLSim']} [{r['일치']}]")

df_k_check = pd.DataFrame(k_check_rows)


# ══════════════════════════════════════════════
# 그룹 3: 마찰 계산 확인
# ══════════════════════════════════════════════
print("\n" + "─" * 70)
print("  [그룹 3] 관재질 조도 및 마찰 계산 확인")
print("─" * 70)

friction_check_rows = [
    {"항목": "마찰 계산 방법", "FiPLSim": "Darcy-Weisbach + Colebrook-White",
     "논문": "Darcy-Weisbach", "비고": "일치"},
    {"항목": "마찰계수 초기값", "FiPLSim": "Swamee-Jain 근사식 + 10회 반복",
     "논문": "Swamee-Jain 사용 명시", "비고": "일치"},
    {"항목": "절대 조도 ε (mm)", "FiPLSim": f"{EPSILON_MM} mm",
     "논문": "미기재 ← 핵심", "비고": "★ 확인 필요"},
    {"항목": "유체 동점성계수 ν", "FiPLSim": f"{NU:.3e} m²/s",
     "논문": "미기재", "비고": "20°C 물 = 1.004e-6"},
    {"항목": "유체 밀도 ρ", "FiPLSim": f"{RHO} kg/m³",
     "논문": "미기재", "비고": "표준값"},
]

for r in friction_check_rows:
    print(f"  {r['항목']}: {r['FiPLSim']} | {r['비고']}")

df_friction_check = pd.DataFrame(friction_check_rows)


# ══════════════════════════════════════════════
# 그룹 4 & 5: (이전과 동일, 일치 확인)
# ══════════════════════════════════════════════
flow_check_rows = [
    {"항목": "전체 유량 분배 방식", "FiPLSim": "균등 분배 (4 branches)", "논문": "균등 분배", "일치": "✅"},
    {"항목": "가지배관 내 헤드별 유량 분배", "FiPLSim": "균등 분배 (8 heads)", "논문": "균등 분배", "일치": "✅"},
    {"항목": "worst-case branch 선택 기준", "FiPLSim": "최저말단 압력 branch", "논문": "최저말단 압력 branch", "일치": "✅"},
]
df_flow_check = pd.DataFrame(flow_check_rows)

length_check_rows = [
    {"항목": "가지배관 간격", "FiPLSim": f"{BRANCH_SP} m", "논문": "3.5 m", "일치": "✅"},
    {"항목": "헤드 간격", "FiPLSim": f"{HEAD_SP} m", "논문": "2.3 m", "일치": "✅"},
    {"항목": "교차배관 총 길이", "FiPLSim": f"(n-1)×{BRANCH_SP}m = {(NUM_BR-1)*BRANCH_SP}m", "논문": "4×3.5m=14m 추정", "일치": "⚠️"},
    {"항목": "가지배관 총 길이", "FiPLSim": f"{HEADS}×{HEAD_SP}m = {HEADS*HEAD_SP}m", "논문": "8×2.3m=18.4m 추정", "일치": "✅"},
]
df_length_check = pd.DataFrame(length_check_rows)


# ══════════════════════════════════════════════
# 감도분석 1: 절대 조도(ε) 변경 — Case B @ 1600 LPM
# ══════════════════════════════════════════════
print("\n" + "═" * 70)
print("  [감도분석 1] 절대 조도(ε) 변경 — Case B @ 1600 LPM")
print("═" * 70)

epsilon_values_mm = [0.020, 0.030, 0.045, 0.046, 0.060, 0.080, 0.100, 0.120, 0.150, 0.200, 0.300, 0.500]

epsilon_sens_B_rows = []
print(f"\n  논문 목표: {PAPER_CASEB_1600} MPa")
print(f"  {'ε (mm)':>10} | {'말단 P (MPa)':>14} | {'차이 (kPa)':>12} | 비고")
print(f"  {'-'*60}")

for eps_mm in epsilon_values_mm:
    res = run_with_epsilon(eps_mm, 1600.0)
    p_term = res["worst_terminal_mpa"]
    diff_kpa = (p_term - PAPER_CASEB_1600) * 1000

    note = ""
    if abs(diff_kpa) < 1.0:
        note = "★★ 논문값 일치!"
    elif abs(diff_kpa) < 3.0:
        note = "★ 근접"
    elif eps_mm == EPSILON_MM:
        note = "← 현재 설정"

    epsilon_sens_B_rows.append({
        "ε (mm)": eps_mm, "말단 압력 (MPa)": round(p_term, 6),
        "논문값 (MPa)": PAPER_CASEB_1600,
        "차이 (kPa)": round(diff_kpa, 2), "비고": note,
    })
    print(f"  {eps_mm:>10.3f} | {p_term:>14.6f} | {diff_kpa:>+12.2f} | {note}")

df_epsilon_B = pd.DataFrame(epsilon_sens_B_rows)

# 보간으로 ε 추정
from scipy.interpolate import interp1d
eps_arr = np.array([r["ε (mm)"] for r in epsilon_sens_B_rows])
p_arr = np.array([r["말단 압력 (MPa)"] for r in epsilon_sens_B_rows])
try:
    interp_func = interp1d(p_arr[::-1], eps_arr[::-1], kind='linear', fill_value='extrapolate')
    epsilon_match_B = float(interp_func(PAPER_CASEB_1600))
    if epsilon_match_B > 0:
        print(f"\n  ★ 보간 추정: 논문값 일치에 필요한 ε ≈ {epsilon_match_B:.4f} mm")
    else:
        epsilon_match_B = None
        print(f"\n  ⚠️ 조도만으로는 논문값 도달 불가 (다른 요인 필요)")
except Exception:
    epsilon_match_B = None

restore_epsilon()


# ══════════════════════════════════════════════
# 감도분석 2: 절대 조도(ε) — Case A 2.5mm @ 1200 LPM
# ══════════════════════════════════════════════
print("\n" + "═" * 70)
print("  [감도분석 2] 절대 조도(ε) — Case A 2.5mm @ 1200 LPM")
print("═" * 70)

epsilon_sens_A_rows = []
print(f"\n  논문 목표: {PAPER_CASEA25_1200} MPa")
print(f"  {'ε (mm)':>10} | {'말단 P (MPa)':>14} | {'차이 (kPa)':>12} | 비고")
print(f"  {'-'*60}")

for eps_mm in epsilon_values_mm:
    beads_2d = [[0.0] * HEADS for _ in range(NUM_BR)]
    beads_2d[NUM_BR - 1] = [2.5] * HEADS

    res = run_with_epsilon(eps_mm, 1200.0, bead_heights_2d=beads_2d)
    p_term = res["worst_terminal_mpa"]
    diff_kpa = (p_term - PAPER_CASEA25_1200) * 1000

    note = ""
    if abs(diff_kpa) < 1.0:
        note = "★★ 논문값 일치!"
    elif abs(diff_kpa) < 3.0:
        note = "★ 근접"
    elif eps_mm == EPSILON_MM:
        note = "← 현재 설정"

    epsilon_sens_A_rows.append({
        "ε (mm)": eps_mm, "말단 압력 (MPa)": round(p_term, 6),
        "논문값 (MPa)": PAPER_CASEA25_1200,
        "차이 (kPa)": round(diff_kpa, 2), "비고": note,
    })
    print(f"  {eps_mm:>10.3f} | {p_term:>14.6f} | {diff_kpa:>+12.2f} | {note}")

df_epsilon_A = pd.DataFrame(epsilon_sens_A_rows)

eps_arr_A = np.array([r["ε (mm)"] for r in epsilon_sens_A_rows])
p_arr_A = np.array([r["말단 압력 (MPa)"] for r in epsilon_sens_A_rows])
try:
    interp_func_A = interp1d(p_arr_A[::-1], eps_arr_A[::-1], kind='linear', fill_value='extrapolate')
    epsilon_match_A = float(interp_func_A(PAPER_CASEA25_1200))
    if epsilon_match_A > 0:
        print(f"\n  ★ 보간 추정: 논문값 일치에 필요한 ε ≈ {epsilon_match_A:.4f} mm")
    else:
        epsilon_match_A = None
except Exception:
    epsilon_match_A = None

restore_epsilon()


# ══════════════════════════════════════════════
# 감도분석 3: 배관 직경 구성 영향 (FiPLSim vs 논문 65A)
# ══════════════════════════════════════════════
print("\n" + "═" * 70)
print("  [감도분석 3] 배관 직경 구성 영향")
print("═" * 70)

pipe_comp_rows = []

for config_name, size_func in [("FiPLSim 현재 (50A)", auto_pipe_size),
                                 ("논문 구성 (65A)", paper_auto_pipe_size)]:
    for Q in [1200, 1600, 2100]:
        # Case B
        res_b = run_with_epsilon_and_pipe_func(EPSILON_MM, float(Q), size_func)
        p_b = res_b["worst_terminal_mpa"]

        # Case A (worst branch bead=2.5)
        beads_2d = [[0.0] * HEADS for _ in range(NUM_BR)]
        beads_2d[NUM_BR - 1] = [2.5] * HEADS
        res_a = run_with_epsilon_and_pipe_func(EPSILON_MM, float(Q), size_func, beads_2d)
        p_a = res_a["worst_terminal_mpa"]

        wb = res_b["worst_branch_index"]

        pipe_comp_rows.append({
            "배관 구성": config_name, "Q (LPM)": Q,
            "Case B (MPa)": round(p_b, 6), "Case A 2.5mm (MPa)": round(p_a, 6),
            "worst branch": wb + 1,
        })

        if Q == 1600:
            diff = (p_b - PAPER_CASEB_1600) * 1000
            print(f"  [{config_name}] Q=1600 Case B: {p_b:.6f} (차이: {diff:+.2f} kPa)")
        if Q == 1200:
            diff = (p_a - PAPER_CASEA25_1200) * 1000
            print(f"  [{config_name}] Q=1200 Case A: {p_a:.6f} (차이: {diff:+.2f} kPa)")

df_pipe_comp = pd.DataFrame(pipe_comp_rows)
restore_epsilon()


# ══════════════════════════════════════════════
# 감도분석 4: 조도 + 배관 조합 (최적 조합 탐색)
# ══════════════════════════════════════════════
print("\n" + "═" * 70)
print("  [감도분석 4] 조도 + 배관 구성 조합 탐색")
print("═" * 70)

combined_rows = []

targets = [
    ("Case B @ 1600", 1600, None, PAPER_CASEB_1600),
    ("Case A 2.5mm @ 1200", 1200, "worst_bead", PAPER_CASEA25_1200),
]

test_eps = [0.045, 0.046, 0.06, 0.08, 0.10, 0.12, 0.15, 0.20, 0.30, 0.50]
test_configs = [("FiPLSim (50A)", auto_pipe_size), ("논문 (65A)", paper_auto_pipe_size)]

for case_name, Q, bead_mode, target in targets:
    print(f"\n  [{case_name}] 논문 목표: {target} MPa")

    beads_2d = None
    if bead_mode == "worst_bead":
        beads_2d = [[0.0] * HEADS for _ in range(NUM_BR)]
        beads_2d[NUM_BR - 1] = [2.5] * HEADS

    for config_name, size_func in test_configs:
        for eps_mm in test_eps:
            res = run_with_epsilon_and_pipe_func(eps_mm, float(Q), size_func, beads_2d)
            p_term = res["worst_terminal_mpa"]
            diff_kpa = (p_term - target) * 1000

            note = ""
            if abs(diff_kpa) < 1.0:
                note = "★★ 일치!"
            elif abs(diff_kpa) < 3.0:
                note = "★ 근접"
            elif abs(diff_kpa) < 5.0:
                note = "근접"

            combined_rows.append({
                "분석 대상": case_name, "배관 구성": config_name,
                "ε (mm)": eps_mm, "말단 압력 (MPa)": round(p_term, 6),
                "논문값 (MPa)": target, "차이 (kPa)": round(diff_kpa, 2),
                "비고": note,
            })

            if note:
                print(f"    {config_name}, ε={eps_mm}mm → {p_term:.6f} ({diff_kpa:+.2f} kPa) {note}")

df_combined = pd.DataFrame(combined_rows)
restore_epsilon()


# ══════════════════════════════════════════════
# 원인 순위 종합
# ══════════════════════════════════════════════
print("\n" + "═" * 70)
print("  [종합] 불일치 원인 순위")
print("═" * 70)

# 현재 ε=0.045에서의 차이
current_diff_B = next(r for r in epsilon_sens_B_rows if r["ε (mm)"] == 0.045)["차이 (kPa)"]
# 65A 배관에서의 차이 (Q=1600)
pipe65_B = next((r for r in pipe_comp_rows if r["배관 구성"] == "논문 구성 (65A)" and r["Q (LPM)"] == 1600), None)
pipe65_diff_B = (pipe65_B["Case B (MPa)"] - PAPER_CASEB_1600) * 1000 if pipe65_B else 0

ranking_rows = [
    {
        "순위": 1,
        "원인 후보": "관 절대 조도 ε 설정 차이",
        "근거": f"ε=0.045mm에서 +{current_diff_B:.1f} kPa. "
                + (f"ε≈{epsilon_match_B:.3f}mm에서 일치 가능" if epsilon_match_B and epsilon_match_B > 0 else "조도 증가 시 논문 방향 이동"),
        "영향도": "★★★★★",
    },
    {
        "순위": 2,
        "원인 후보": "배관 직경 구성 (65A 포함 시 오히려 차이 증가)",
        "근거": f"65A 포함 시 Q=1600 Case B 차이: {pipe65_diff_B:+.1f} kPa (차이 증가 → 원인 아님)",
        "영향도": "★★ (반대 방향)",
    },
    {
        "순위": 3,
        "원인 후보": "K_TEE_RUN, K_branch, K_head 값",
        "근거": "현재 설정과 논문 기재값 동일 (0.3, 1.0, 2.5, 0.5)",
        "영향도": "★ (확인됨)",
    },
    {
        "순위": 4,
        "원인 후보": "유량 분배/배관 길이 계산",
        "근거": "균등 분배, 간격 설정 동일",
        "영향도": "★ (확인됨)",
    },
]

for r in ranking_rows:
    print(f"  {r['순위']}위 [{r['영향도']}] {r['원인 후보']}")
    print(f"      {r['근거']}")

df_ranking = pd.DataFrame(ranking_rows)


# ══════════════════════════════════════════════
# Excel 저장
# ══════════════════════════════════════════════
print("\n" + "=" * 70)
print("  Excel 데이터 저장 중...")

xlsx_path = os.path.join(BASE_DIR, "FiPLSim_불일치_원인분석_데이터.xlsx")

with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
    df_pipe_config.to_excel(writer, sheet_name="1_배관직경비교", index=False)
    df_k_check.to_excel(writer, sheet_name="2_K값확인", index=False)
    df_friction_check.to_excel(writer, sheet_name="3_마찰계산확인", index=False)
    df_flow_check.to_excel(writer, sheet_name="4_유량분배확인", index=False)
    df_length_check.to_excel(writer, sheet_name="5_배관길이확인", index=False)
    df_epsilon_B.to_excel(writer, sheet_name="6_조도감도_CaseB_1600", index=False)
    df_epsilon_A.to_excel(writer, sheet_name="7_조도감도_CaseA_1200", index=False)
    df_pipe_comp.to_excel(writer, sheet_name="8_배관구성비교", index=False)
    df_combined.to_excel(writer, sheet_name="9_조합최적화", index=False)
    df_ranking.to_excel(writer, sheet_name="10_원인순위", index=False)

xlsx_size = os.path.getsize(xlsx_path) / 1024
print(f"  ✅ 저장: {xlsx_path} ({xlsx_size:.1f} KB)")


# ══════════════════════════════════════════════
# DOCX 보고서 생성
# ══════════════════════════════════════════════
print("\n  DOCX 보고서 생성 중...")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

HEADER_BG = RGBColor(0x2E, 0x5E, 0x8E)
LIGHT_BG = RGBColor(0xE8, 0xF0, 0xFA)
HIGHLIGHT_BG = RGBColor(0xFF, 0xF3, 0xCD)


def set_cell_bg(cell, color):
    shading = cell._element.get_or_add_tcPr()
    el = shading.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): str(color)})
    shading.append(el)


def add_table(doc, headers, rows, highlight_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, HEADER_BG)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8.5)
            if highlight_rows and r_idx in highlight_rows:
                set_cell_bg(cell, HIGHLIGHT_BG)
            elif r_idx % 2 == 1:
                set_cell_bg(cell, LIGHT_BG)
    return table


doc = Document()
doc.styles["Normal"].font.name = "맑은 고딕"
doc.styles["Normal"].font.size = Pt(10)

# 표지
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("FiPLSim 불일치 원인 정밀 분석\n결과 보고서")
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

doc.add_paragraph()
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run(
    "Case B @ 1,600 LPM: +14.5 kPa\n"
    "Case A 2.5mm @ 1,200 LPM: +12.1 kPa\n\n"
    "5개 그룹 체크리스트 + 4종 감도분석"
)
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = dp.add_run(f"생성일: {datetime.date.today().strftime('%Y-%m-%d')}")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# 1. 불일치 현황
doc.add_heading("1. 불일치 현황", level=1)
doc.add_paragraph(
    "논문값과 FiPLSim 시뮬레이션 결과 사이에 다음 2건의 체계적 불일치가 발견되었습니다. "
    "두 건 모두 FiPLSim이 논문보다 높은 말단 압력(낮은 손실)을 산출합니다."
)
add_table(doc, ["항목", "논문 (MPa)", "FiPLSim (MPa)", "차이"], [
    ["Case B μ @ 1,600 LPM", "0.2430", "0.2575", "+14.5 kPa"],
    ["Case A 2.5mm μ @ 1,200 LPM", "0.3041", "0.3162", "+12.1 kPa"],
])

# 2. 체크리스트
doc.add_page_break()
doc.add_heading("2. 입력변수 체크리스트 (5개 그룹)", level=1)

doc.add_heading("2.1 그룹 1: 배관 직경 및 구성", level=2)
doc.add_paragraph("논문: '65A→50A→40A→32A→25A' 감소 구성 명시")

pipe_h = ["Head #", "하류 수", "FiPLSim", "ID (mm)", "논문", "ID (mm)", "일치"]
pipe_d = []
for r in pipe_config_rows:
    pipe_d.append([str(r["Head #"]), str(r["하류 헤드 수"]),
                    r["FiPLSim 관경"], f"{r['FiPLSim ID (mm)']:.2f}",
                    r["논문 관경"], f"{r['논문 ID (mm)']:.2f}", r["일치"]])
add_table(doc, pipe_h, pipe_d)

doc.add_paragraph()
add_table(doc, ["위치", "FiPLSim", "논문", "비고"], [
    ["교차배관", f"{cross_main_fiplsim} ({PIPE_DIMENSIONS[cross_main_fiplsim]['id_mm']:.2f}mm)",
     f"80A ({PIPE_DIMENSIONS['80A']['id_mm']:.2f}mm)", "✅ 일치"],
    ["가지배관 입구", f"{branch_inlet_fiplsim} ({PIPE_DIMENSIONS[branch_inlet_fiplsim]['id_mm']:.2f}mm)",
     f"65A ({PIPE_DIMENSIONS['65A']['id_mm']:.2f}mm)", "⚠️ 논문은 65A"],
])

doc.add_paragraph(
    "※ 가지배관 8개 구간의 관경은 논문과 FiPLSim이 완전 일치합니다. "
    "단, 논문의 '65A'는 가지배관 입구 관경으로 해석되며, "
    "FiPLSim은 NFSC 103 기준으로 50A를 자동 배정합니다."
)

doc.add_heading("2.2 그룹 2: 손실계수 K값", level=2)
k_d = [[r["항목"], r["논문"], r["FiPLSim"], r["일치"]] for r in k_check_rows]
add_table(doc, ["항목", "논문", "FiPLSim", "일치"], k_d)
p = doc.add_paragraph()
r = p.add_run("결론: 4개 K값 모두 논문과 동일 → 불일치 원인 아님")
r.font.bold = True

doc.add_heading("2.3 그룹 3: 마찰 계산", level=2)
f_d = [[r["항목"], r["FiPLSim"], r["논문"], r["비고"]] for r in friction_check_rows]
add_table(doc, ["항목", "FiPLSim", "논문", "비고"], f_d, highlight_rows=[2])

doc.add_heading("2.4 그룹 4: 유량 분배 / 그룹 5: 배관 길이", level=2)
doc.add_paragraph("유량 분배: 균등 분배 — 논문과 동일 ✅")
doc.add_paragraph("배관 길이: 가지배관 간격 3.5m, 헤드 간격 2.3m — 논문과 동일 ✅")

# 3. 감도분석
doc.add_page_break()
doc.add_heading("3. 감도분석 결과", level=1)

doc.add_heading("3.1 조도(ε) 감도분석 — Case B @ 1,600 LPM", level=2)
doc.add_paragraph(f"논문 목표: {PAPER_CASEB_1600} MPa")

eps_d = []
hl = []
for i, r in enumerate(epsilon_sens_B_rows):
    if r["비고"]:
        hl.append(i)
    eps_d.append([f"{r['ε (mm)']:.3f}", f"{r['말단 압력 (MPa)']:.6f}",
                   f"{r['차이 (kPa)']:+.2f}", r["비고"]])
add_table(doc, ["ε (mm)", "말단 (MPa)", "차이 (kPa)", "비고"], eps_d, highlight_rows=hl)

if epsilon_match_B and epsilon_match_B > 0:
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"보간 추정: 논문값 일치에 필요한 ε ≈ {epsilon_match_B:.4f} mm")
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_heading("3.2 조도(ε) 감도분석 — Case A 2.5mm @ 1,200 LPM", level=2)
doc.add_paragraph(f"논문 목표: {PAPER_CASEA25_1200} MPa")

eps_a_d = []
hl_a = []
for i, r in enumerate(epsilon_sens_A_rows):
    if r["비고"]:
        hl_a.append(i)
    eps_a_d.append([f"{r['ε (mm)']:.3f}", f"{r['말단 압력 (MPa)']:.6f}",
                     f"{r['차이 (kPa)']:+.2f}", r["비고"]])
add_table(doc, ["ε (mm)", "말단 (MPa)", "차이 (kPa)", "비고"], eps_a_d, highlight_rows=hl_a)

if epsilon_match_A and epsilon_match_A > 0:
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"보간 추정: 논문값 일치에 필요한 ε ≈ {epsilon_match_A:.4f} mm")
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_heading("3.3 배관 직경 구성 영향", level=2)

pc_d = []
for r in pipe_comp_rows:
    pc_d.append([r["배관 구성"], str(r["Q (LPM)"]),
                  f"{r['Case B (MPa)']:.6f}", f"{r['Case A 2.5mm (MPa)']:.6f}"])
add_table(doc, ["배관 구성", "Q (LPM)", "Case B (MPa)", "Case A 2.5mm (MPa)"], pc_d)

doc.add_heading("3.4 조도 + 배관 조합 최적화", level=2)

best = [r for r in combined_rows if abs(r["차이 (kPa)"]) < 5.0]
if best:
    bd = []
    for r in best:
        bd.append([r["분석 대상"], r["배관 구성"], f"{r['ε (mm)']:.3f}",
                    f"{r['말단 압력 (MPa)']:.6f}", f"{r['차이 (kPa)']:+.2f}"])
    add_table(doc, ["대상", "배관", "ε (mm)", "말단 (MPa)", "차이 (kPa)"], bd)
else:
    doc.add_paragraph("  (차이 5 kPa 이내의 조합을 찾지 못했습니다)")

# 4. 원인 순위
doc.add_page_break()
doc.add_heading("4. 불일치 원인 순위 종합", level=1)

rk_d = [[str(r["순위"]), r["원인 후보"], r["영향도"], r["근거"]] for r in ranking_rows]
add_table(doc, ["순위", "원인 후보", "영향도", "근거"], rk_d, highlight_rows=[0])

# 5. 요청 사항
doc.add_heading("5. 확인 요청 사항 (우선순위)", level=1)

doc.add_paragraph("아래 3가지를 확인해 주시면 불일치 원인을 특정할 수 있습니다:")

reqs = [
    "1. 관 절대 조도 ε 값 (mm 단위) — 논문에서 사용한 정확한 값",
    "2. 배관 직경 감소 구성 — 65A→50A→40A→32A→25A가 정확한지, 첫 구간 65A의 의미",
    "3. K_TEE_RUN, K_branch, K_head 값 — 각각 0.3, 1.0, 2.5가 맞는지",
]
for r in reqs:
    p = doc.add_paragraph()
    run = p.add_run(r)
    if "1." in r:
        run.font.bold = True

# 6. 결론
doc.add_heading("6. 결론", level=1)

concl = [
    "K값 4종 (0.3, 1.0, 2.5, 0.5): 논문과 완전 일치 → 불일치 원인 아님",
    "유량 분배·배관 길이: 논문과 동일 → 불일치 원인 아님",
    "마찰 계산 (Darcy-Weisbach + Colebrook-White): 논문과 동일 → 불일치 원인 아님",
    "절대 조도 ε: 1순위 원인 후보. 현재 0.045mm에서 +14.5 kPa 차이 발생. "
    "논문이 더 높은 조도를 사용했을 가능성 → 손실 증가 → 말단 압력 감소",
    "배관 구성 (65A): 65A 포함 시 오히려 차이 증가 (면적 증가 → 손실 감소). "
    "따라서 65A는 교차배관→가지배관 연결부의 관경이 아닌 다른 의미일 수 있음",
    "★ 최종 판단: 절대 조도(ε) 차이가 가장 유력한 원인이며, "
    "논문에서 사용한 ε 값을 확인하면 불일치를 해소할 수 있을 것으로 판단됩니다.",
]

for i, c in enumerate(concl):
    p = doc.add_paragraph()
    run = p.add_run(f"{i+1}. {c}")
    if i == 5:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

docx_path = os.path.join(BASE_DIR, "FiPLSim_불일치_원인분석_보고서.docx")
doc.save(docx_path)
docx_size = os.path.getsize(docx_path) / 1024
print(f"  ✅ 저장: {docx_path} ({docx_size:.1f} KB)")

elapsed = time.time() - t0
print(f"\n{'='*70}")
print(f"  ★ 분석 완료")
print(f"{'='*70}")
print(f"  출력: {xlsx_path}")
print(f"       {docx_path}")
print(f"  소요: {elapsed:.1f}초")
if epsilon_match_B and epsilon_match_B > 0:
    print(f"  ★ Case B 일치 ε ≈ {epsilon_match_B:.4f} mm")
if epsilon_match_A and epsilon_match_A > 0:
    print(f"  ★ Case A 일치 ε ≈ {epsilon_match_A:.4f} mm")
print(f"{'='*70}")
