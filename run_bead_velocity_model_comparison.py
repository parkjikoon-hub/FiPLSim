"""
FiPLSim --- 비드 K-factor 속도 기준 모델 비교 (D^4 vs D^8)

첨부 이미지 분석 기반:
  - 현재 FiPLSim: K_eff × V_upstream² → (D/D_eff)^4 모델
  - 가능성 A: K_eff × V_eff² → (D/D_eff)^8 모델

분석 구성:
  Part 1: D^4 vs D^8 결정론적 비교 (ε=0.154mm)
  Part 2: D^4 vs D^8 Scenario 1 MC 비교 (Table 8 핵심)
  Part 3: 논문 대비 최종 비교
  Part 4: (D/D_eff)^n 지수별 감도분석

출력:
  FiPLSim_비드모델비교_데이터.xlsx
  FiPLSim_비드모델비교_보고서.docx
"""

import os, sys, time, datetime, importlib
import numpy as np
import pandas as pd

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import constants
import hydraulics
import pipe_network as pn
from pipe_network import generate_dynamic_system, calculate_dynamic_system
from constants import (
    PIPE_DIMENSIONS, get_inner_diameter_m, RHO, G,
    MIN_TERMINAL_PRESSURE_MPA, K1_BASE, K2, K3,
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# 논문 조건
NUM_BR = 4; HEADS = 8; BRANCH_SP = 3.5; HEAD_SP = 2.3
INLET_P = 0.4; SUPPLY = "80A"; N_ITER = 10000; SC1_P = 0.5

# 논문 목표값
PAPER = {
    (2.0, 2100): {"mu": 0.1199, "sigma": 0.0039, "Pf": 0.37},
    (2.5, 2100): {"mu": 0.1100, "sigma": 0.0048, "Pf": 2.43},
}


def set_epsilon(eps_mm):
    constants.EPSILON_MM = eps_mm
    constants.EPSILON_M = eps_mm / 1000.0
    importlib.reload(hydraulics)
    importlib.reload(pn)
    globals()['generate_dynamic_system'] = pn.generate_dynamic_system
    globals()['calculate_dynamic_system'] = pn.calculate_dynamic_system


def run_det(Q, bead_mm=0.0, model="upstream"):
    beads_2d = [[bead_mm] * HEADS for _ in range(NUM_BR)] if bead_mm > 0 else None
    sys_ = generate_dynamic_system(
        num_branches=NUM_BR, heads_per_branch=HEADS,
        branch_spacing_m=BRANCH_SP, head_spacing_m=HEAD_SP,
        inlet_pressure_mpa=INLET_P, total_flow_lpm=float(Q),
        bead_heights_2d=beads_2d)
    return calculate_dynamic_system(sys_, K3, equipment_k_factors=None,
        supply_pipe_size=SUPPLY, bead_velocity_model=model)


def run_mc(bead_mm, Q, model="upstream"):
    rng = np.random.default_rng(42)
    ps = np.zeros(N_ITER); cs = np.zeros(N_ITER, dtype=int)
    wb = NUM_BR - 1
    for t in range(N_ITER):
        b2d = [[0.0]*HEADS for _ in range(NUM_BR)]
        rv = rng.uniform(0, 1, size=HEADS)
        c = 0
        for h in range(HEADS):
            if rv[h] <= SC1_P:
                b2d[wb][h] = bead_mm; c += 1
        cs[t] = c
        sys_ = generate_dynamic_system(
            num_branches=NUM_BR, heads_per_branch=HEADS,
            branch_spacing_m=BRANCH_SP, head_spacing_m=HEAD_SP,
            inlet_pressure_mpa=INLET_P, total_flow_lpm=float(Q),
            bead_heights_2d=b2d)
        r = calculate_dynamic_system(sys_, K3, equipment_k_factors=None,
            supply_pipe_size=SUPPLY, bead_velocity_model=model)
        ps[t] = r["worst_terminal_mpa"]
    bel = np.sum(ps < MIN_TERMINAL_PRESSURE_MPA)
    return {
        "mu": float(np.mean(ps)), "sigma": float(np.std(ps, ddof=1)),
        "min": float(np.min(ps)), "max": float(np.max(ps)),
        "Pf": float(bel / N_ITER), "mean_beads": float(np.mean(cs)),
        "pct": {k: float(np.percentile(ps, v)) for k, v in
                [("P5",5),("P25",25),("P50",50),("P75",75),("P95",95)]},
        "raw_P": ps, "raw_N": cs,
    }


# ═══════════════════════════════════════════════
print("=" * 70)
print("  FiPLSim --- 비드 모델 비교: D^4 vs D^8")
print("=" * 70)
t0 = time.time()

# epsilon = 0.154mm (이전 분석에서 Case B 일치값)
EPS = 0.154
set_epsilon(EPS)
print(f"  epsilon = {EPS}mm")


# ╔═══════════════════════════════════════════════╗
# ║  Part 1: 결정론적 비교                          ║
# ╚═══════════════════════════════════════════════╝
print("\n" + "=" * 70)
print("  Part 1: D^4 vs D^8 결정론적 비교 (전체 비드)")
print("=" * 70)

det_rows = []
for bead_mm in [0, 1.5, 2.0, 2.5]:
    for Q in [1200, 1600, 2100, 2300]:
        r4 = run_det(Q, bead_mm, "upstream")["worst_terminal_mpa"]
        r8 = run_det(Q, bead_mm, "constriction")["worst_terminal_mpa"]
        # Case B reference (bead=0)
        rB = run_det(Q, 0.0, "upstream")["worst_terminal_mpa"]
        bl4 = (rB - r4) * 1000
        bl8 = (rB - r8) * 1000

        det_rows.append({
            "bead_mm": bead_mm, "Q": Q,
            "CaseB_mpa": round(rB, 6),
            "D4_mpa": round(r4, 6), "D8_mpa": round(r8, 6),
            "D4_loss_kpa": round(bl4, 2), "D8_loss_kpa": round(bl8, 2),
            "D8_D4_ratio": round(bl8 / bl4, 2) if bl4 > 0 else "---",
        })
        if bead_mm > 0:
            print(f"  bead={bead_mm}mm Q={Q:>5}: D^4={r4:.4f} D^8={r8:.4f}  "
                  f"손실 D^4={bl4:.1f} D^8={bl8:.1f} kPa  비율={bl8/bl4:.2f}x" if bl4 > 0 else "")

df_det = pd.DataFrame(det_rows)


# ╔═══════════════════════════════════════════════╗
# ║  Part 2: Scenario 1 MC 비교                     ║
# ╚═══════════════════════════════════════════════╝
print("\n" + "=" * 70)
print("  Part 2: Scenario 1 MC 비교 (p=0.5, N=10000)")
print("=" * 70)

mc_rows = []
mc_raw = {}
mc_configs = [(2.0, [2100]), (2.5, [1200, 1600, 2100, 2300])]

for bead_mm, flows in mc_configs:
    for Q in flows:
        for model, label in [("upstream", "D^4"), ("constriction", "D^8")]:
            print(f"  bead={bead_mm}mm Q={Q} {label}: ", end="", flush=True)
            ts = time.time()
            r = run_mc(bead_mm, Q, model)
            rB = run_det(Q, 0.0, "upstream")["worst_terminal_mpa"]
            dt = time.time() - ts

            mc_rows.append({
                "bead": bead_mm, "Q": Q, "model": label,
                "CaseB": round(rB, 6),
                "mu": round(r["mu"], 6), "sigma": round(r["sigma"], 6),
                "min": round(r["min"], 6), "max": round(r["max"], 6),
                "P5": round(r["pct"]["P5"], 6), "P50": round(r["pct"]["P50"], 6),
                "P95": round(r["pct"]["P95"], 6),
                "delta_kpa": round((rB - r["mu"]) * 1000, 2),
                "Pf": round(r["Pf"] * 100, 2),
                "mean_beads": round(r["mean_beads"], 2),
            })
            mc_raw[f"b{bead_mm}_Q{Q}_{label}_P"] = r["raw_P"]
            mc_raw[f"b{bead_mm}_Q{Q}_{label}_N"] = r["raw_N"]

            print(f"mu={r['mu']:.6f} sigma={r['sigma']:.6f} Pf={r['Pf']*100:.2f}% ({dt:.1f}s)")

df_mc = pd.DataFrame(mc_rows)


# ╔═══════════════════════════════════════════════╗
# ║  Part 3: 논문 대비 최종 비교                     ║
# ╚═══════════════════════════════════════════════╝
print("\n" + "=" * 70)
print("  Part 3: 논문 대비 비교")
print("=" * 70)

comp_rows = []
for bead_mm, Q_val in [(2.0, 2100), (2.5, 2100)]:
    target = PAPER[(bead_mm, Q_val)]
    for model_label in ["D^4", "D^8"]:
        mc_r = next(r for r in mc_rows if r["bead"]==bead_mm and r["Q"]==Q_val and r["model"]==model_label)
        diff = (mc_r["mu"] - target["mu"]) * 1000
        comp_rows.append({
            "bead": bead_mm, "Q": Q_val, "model": model_label,
            "FiPLSim_mu": mc_r["mu"], "논문_mu": target["mu"],
            "diff_kpa": round(diff, 2),
            "FiPLSim_sigma": mc_r["sigma"], "논문_sigma": target["sigma"],
            "FiPLSim_Pf": mc_r["Pf"], "논문_Pf": target["Pf"],
        })
        print(f"  bead={bead_mm}mm Q={Q_val} {model_label}: "
              f"FiPLSim={mc_r['mu']:.4f} vs 논문={target['mu']:.4f} "
              f"({diff:+.2f} kPa)  Pf={mc_r['Pf']:.2f}% vs {target['Pf']:.2f}%")

df_comp = pd.DataFrame(comp_rows)


# ╔═══════════════════════════════════════════════╗
# ║  Part 4: (D/D_eff)^n 이론 비교                  ║
# ╚═══════════════════════════════════════════════╝
print("\n" + "=" * 70)
print("  Part 4: (D/D_eff)^n 이론 계산")
print("=" * 70)

pipe_ids = {"50A": 52.51, "40A": 40.90, "32A": 35.04, "25A": 26.64}
theory_rows = []
for bead_mm in [1.5, 2.0, 2.5]:
    for name, d_mm in pipe_ids.items():
        d_eff = d_mm - 2 * bead_mm
        if d_eff <= 0: continue
        ratio = d_mm / d_eff
        theory_rows.append({
            "bead_mm": bead_mm, "pipe": name, "D_mm": d_mm, "D_eff_mm": round(d_eff, 2),
            "D/D_eff": round(ratio, 4),
            "(D/D_eff)^4": round(ratio**4, 4),
            "(D/D_eff)^8": round(ratio**8, 4),
            "^8/^4 비율": round(ratio**4, 4),  # ^8/^4 = (D/D_eff)^4
        })
        print(f"  bead={bead_mm}mm {name}: D/D_eff={ratio:.3f}  "
              f"^4={ratio**4:.3f}  ^8={ratio**8:.3f}  비율={ratio**4:.3f}x")

df_theory = pd.DataFrame(theory_rows)


# ╔═══════════════════════════════════════════════╗
# ║  조건 시트                                       ║
# ╚═══════════════════════════════════════════════╝
df_cond = pd.DataFrame({
    "항목": [
        "가지배관", "헤드/가지배관", "입구 압력", "교차배관", "밸브",
        "MC 반복", "epsilon", "비드 확률 (p)", "Scenario",
        "D^4 모델", "D^8 모델",
    ],
    "값": [
        f"{NUM_BR}개", f"{HEADS}개", f"{INLET_P} MPa", SUPPLY, "OFF",
        f"{N_ITER:,}회", f"{EPS} mm", SC1_P, "Scenario 1 (결함 집중)",
        "K_eff × V_upstream² / (2g)  — 기존",
        "K_eff × V_eff² / (2g)  — 협착부 속도 기준",
    ],
})


# ═══════════════════════════════════════════════
# Excel 저장
# ═══════════════════════════════════════════════
print("\n" + "=" * 70)
print("  Excel 저장 중...")

xlsx = os.path.join(BASE_DIR, "FiPLSim_비드모델비교_데이터.xlsx")
df_raw = pd.DataFrame({k: np.round(v, 6) for k, v in mc_raw.items()}) if mc_raw else pd.DataFrame()

with pd.ExcelWriter(xlsx, engine="openpyxl") as w:
    df_cond.to_excel(w, sheet_name="1_조건", index=False)
    df_det.to_excel(w, sheet_name="2_결정론적비교", index=False)
    df_mc.to_excel(w, sheet_name="3_MC비교", index=False)
    df_comp.to_excel(w, sheet_name="4_논문비교", index=False)
    df_theory.to_excel(w, sheet_name="5_이론계산", index=False)
    if not df_raw.empty:
        df_raw.to_excel(w, sheet_name="6_MC원시", index=False)

print(f"  => {xlsx} ({os.path.getsize(xlsx)/1024:.1f} KB)")


# ═══════════════════════════════════════════════
# DOCX 보고서
# ═══════════════════════════════════════════════
print("\n  DOCX 보고서 생성 중...")

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

HDR = RGBColor(0x2E, 0x5E, 0x8E)
LT = RGBColor(0xE8, 0xF0, 0xFA)
HL = RGBColor(0xFF, 0xF3, 0xCD)
GR = RGBColor(0xD4, 0xED, 0xDA)
RD = RGBColor(0xF8, 0xD7, 0xDA)


def bg(c, clr):
    s = c._element.get_or_add_tcPr()
    s.append(s.makeelement(qn("w:shd"), {qn("w:val"):"clear",qn("w:color"):"auto",qn("w:fill"):str(clr)}))


def tbl(doc, hdr, rows, hi=None, gr_=None, rd_=None):
    t = doc.add_table(rows=1+len(rows), cols=len(hdr))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    for i, h in enumerate(hdr):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(255,255,255)
        bg(c, HDR)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[1+ri].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size=Pt(9)
            if hi and ri in hi: bg(c, HL)
            elif gr_ and ri in gr_: bg(c, GR)
            elif rd_ and ri in rd_: bg(c, RD)
            elif ri%2==1: bg(c, LT)


doc = Document()
st = doc.styles["Normal"]; st.font.name = "맑은 고딕"; st.font.size = Pt(10)

# 표지
doc.add_paragraph(); doc.add_paragraph()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("FiPLSim 비드 K-factor\n속도 기준 모델 비교 보고서")
r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = HDR

doc.add_paragraph()
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run(
    "(D/D_eff)^4 모델 vs (D/D_eff)^8 모델\n"
    "불일치 ② (비드 손실 크기) 원인 규명\n\n"
    f"epsilon = {EPS} mm")
r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(0x66,0x66,0x66)

doc.add_paragraph()
dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = dp.add_run(f"생성일: {datetime.date.today().strftime('%Y-%m-%d')}")
r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0x88,0x88,0x88)
doc.add_page_break()


# 1. 개요
doc.add_heading("1. 분석 개요", level=1)
doc.add_paragraph(
    "논문과 FiPLSim의 비드 손실 차이(약 2.3배)의 원인을 규명하기 위해, "
    "비드 K-factor 손실 계산 시 사용하는 속도 기준을 비교합니다.\n\n"
    "현재 FiPLSim (D^4 모델):\n"
    "  h = K_eff × V_upstream² / (2g)\n"
    "  K_eff = K_base × (D/D_eff)^4\n"
    "  → 총 손실 ∝ (D/D_eff)^4\n\n"
    "가능성 A (D^8 모델):\n"
    "  h = K_eff × V_eff² / (2g)\n"
    "  V_eff = V_upstream × (D/D_eff)²\n"
    "  → 총 손실 ∝ (D/D_eff)^8\n\n"
    "물리적 의미: D^8 모델은 K_eff가 이미 (D/D_eff)^4를 포함하고 있는 상태에서 "
    "추가로 협착부 속도(V_eff)를 사용하므로, 속도 제곱의 효과가 이중 적용됩니다."
)


# 2. 결정론적 비교
doc.add_page_break()
doc.add_heading("2. 결정론적 비교 (전체 비드)", level=1)
doc.add_paragraph(f"epsilon = {EPS}mm, 전체 junction에 비드 적용 (결정론적)")

# bead 2.5mm 데이터만 추출
d2 = []
for r in det_rows:
    if r["bead_mm"] == 2.5:
        d2.append([f"{r['Q']:,}", f"{r['CaseB_mpa']:.4f}",
            f"{r['D4_mpa']:.4f}", f"{r['D8_mpa']:.4f}",
            f"{r['D4_loss_kpa']:.1f}", f"{r['D8_loss_kpa']:.1f}",
            f"{r['D8_D4_ratio']}" if isinstance(r['D8_D4_ratio'], str) else f"{r['D8_D4_ratio']:.2f}"])
tbl(doc, ["Q (LPM)", "Case B", "D^4 (MPa)", "D^8 (MPa)",
          "D^4 손실(kPa)", "D^8 손실(kPa)", "D^8/D^4"], d2)

doc.add_paragraph()
doc.add_heading("2.2 비드 2.0mm", level=2)
d2b = []
for r in det_rows:
    if r["bead_mm"] == 2.0:
        d2b.append([f"{r['Q']:,}", f"{r['CaseB_mpa']:.4f}",
            f"{r['D4_mpa']:.4f}", f"{r['D8_mpa']:.4f}",
            f"{r['D4_loss_kpa']:.1f}", f"{r['D8_loss_kpa']:.1f}",
            f"{r['D8_D4_ratio']}" if isinstance(r['D8_D4_ratio'], str) else f"{r['D8_D4_ratio']:.2f}"])
tbl(doc, ["Q (LPM)", "Case B", "D^4 (MPa)", "D^8 (MPa)",
          "D^4 손실(kPa)", "D^8 손실(kPa)", "D^8/D^4"], d2b)


# 3. MC 비교
doc.add_page_break()
doc.add_heading("3. Scenario 1 MC 비교", level=1)
doc.add_paragraph(f"Scenario 1: worst branch {HEADS} junctions, p={SC1_P}, N={N_ITER:,}")

d3 = []; h3 = []
for i, r in enumerate(mc_rows):
    d3.append([f"{r['bead']:.1f}mm", f"{r['Q']:,}", r["model"],
        f"{r['mu']:.6f}", f"{r['sigma']:.6f}", f"{r['Pf']:.2f}",
        f"{r['delta_kpa']:.2f}"])
    if r['bead']==2.5 and r['Q']==2100: h3.append(i)
tbl(doc, ["Bead", "Q", "모델", "mu (MPa)", "sigma", "Pf (%)", "비드손실 (kPa)"],
    d3, hi=h3)


# 4. 논문 대비 비교
doc.add_heading("4. 논문 대비 최종 비교", level=1)
doc.add_paragraph("Table 8 핵심 값(bead 2.0mm/2.5mm, Q=2100)과 비교합니다.")

d4 = []; g4 = []; r4_ = []
for i, r in enumerate(comp_rows):
    d4.append([f"{r['bead']:.1f}mm", r["model"],
        f"{r['FiPLSim_mu']:.4f}", f"{r['논문_mu']:.4f}", f"{r['diff_kpa']:+.2f}",
        f"{r['FiPLSim_Pf']:.2f}%", f"{r['논문_Pf']:.2f}%"])
    if abs(r['diff_kpa']) < 5: g4.append(i)
    elif abs(r['diff_kpa']) < 15: pass
    else: r4_.append(i)
tbl(doc, ["Bead", "모델", "FiPLSim mu", "논문 mu", "차이(kPa)", "FiPLSim Pf", "논문 Pf"],
    d4, gr_=g4, rd_=r4_)


# 5. 이론 계산
doc.add_page_break()
doc.add_heading("5. (D/D_eff)^n 이론 계산", level=1)
doc.add_paragraph(
    "각 배관 구경별로 D^4와 D^8 모델의 K-factor 증폭 비율을 계산합니다.\n"
    "^8/^4 비율 = (D/D_eff)^4 — 이것이 D^8 모델에서 추가되는 증폭입니다.")

d5 = []
for r in theory_rows:
    if r["bead_mm"] == 2.5:
        d5.append([r["pipe"], f"{r['D_mm']:.2f}", f"{r['D_eff_mm']:.2f}",
            f"{r['D/D_eff']:.3f}", f"{r['(D/D_eff)^4']:.3f}",
            f"{r['(D/D_eff)^8']:.3f}", f"{r['^8/^4 비율']:.3f}x"])
tbl(doc, ["배관", "D (mm)", "D_eff (mm)", "D/D_eff", "^4", "^8", "^8/^4"],
    d5)


# 6. 결론
doc.add_page_break()
doc.add_heading("6. 결론", level=1)

# 핵심 비교값 추출
comp_25_d4 = next(r for r in comp_rows if r["bead"]==2.5 and r["model"]=="D^4")
comp_25_d8 = next(r for r in comp_rows if r["bead"]==2.5 and r["model"]=="D^8")
comp_20_d4 = next(r for r in comp_rows if r["bead"]==2.0 and r["model"]=="D^4")
comp_20_d8 = next(r for r in comp_rows if r["bead"]==2.0 and r["model"]=="D^8")

doc.add_heading("6.1 핵심 발견", level=2)
findings = [
    f"D^4 모델 (현재): bead 2.5mm Q=2100 → mu={comp_25_d4['FiPLSim_mu']:.4f} MPa "
    f"(논문 대비 {comp_25_d4['diff_kpa']:+.2f} kPa)",
    f"D^8 모델 (협착부 속도): bead 2.5mm Q=2100 → mu={comp_25_d8['FiPLSim_mu']:.4f} MPa "
    f"(논문 대비 {comp_25_d8['diff_kpa']:+.2f} kPa)",
    f"D^4 모델: bead 2.0mm Q=2100 → mu={comp_20_d4['FiPLSim_mu']:.4f} (논문 {comp_20_d4['diff_kpa']:+.2f} kPa)",
    f"D^8 모델: bead 2.0mm Q=2100 → mu={comp_20_d8['FiPLSim_mu']:.4f} (논문 {comp_20_d8['diff_kpa']:+.2f} kPa)",
]
for f_ in findings:
    doc.add_paragraph(f_, style="List Bullet")

doc.add_heading("6.2 모델 선택 판단", level=2)

# 어느 모델이 더 가까운지 판단
d4_err = abs(comp_25_d4['diff_kpa'])
d8_err = abs(comp_25_d8['diff_kpa'])
better = "D^8" if d8_err < d4_err else "D^4"
doc.add_paragraph(
    f"bead 2.5mm Q=2100 기준: D^4 오차 {d4_err:.1f} kPa vs D^8 오차 {d8_err:.1f} kPa\n"
    f"→ {better} 모델이 논문에 더 가깝습니다.\n\n"
    f"Pf 비교: D^4={comp_25_d4['FiPLSim_Pf']:.2f}% vs D^8={comp_25_d8['FiPLSim_Pf']:.2f}% "
    f"vs 논문={comp_25_d4['논문_Pf']:.2f}%"
)

doc.add_heading("6.3 물리적 해석", level=2)
doc.add_paragraph(
    "D^8 모델은 K_eff에 이미 포함된 (D/D_eff)^4 효과 위에 "
    "협착부 속도의 (D/D_eff)^4 효과를 추가로 적용합니다.\n\n"
    "이것이 물리적으로 타당한지는 논문의 비드 손실 정의를 확인해야 합니다:\n"
    "  - K_eff가 '상류 속도 기준 등가 K'라면 → D^4 모델이 올바름\n"
    "  - K_eff가 '국부 형상 K'로서 협착부 속도와 함께 사용해야 한다면 → D^8 모델이 올바름"
)

doc.add_heading("6.4 향후 과제", level=2)
for t in [
    "논문 Eq.(1)의 K_eff 정의 정밀 확인: 어느 속도 기준의 K인지",
    "논문 저자에게 비드 손실 계산 코드 또는 상세 수식 확인 요청",
    f"확정 시 constants.py EPSILON_MM = {EPS} 영구 반영",
    f"확정 시 bead_velocity_model 기본값 변경 (현재: upstream)",
]:
    doc.add_paragraph(t, style="List Bullet")


docx_path = os.path.join(BASE_DIR, "FiPLSim_비드모델비교_보고서.docx")
doc.save(docx_path)
print(f"  => {docx_path} ({os.path.getsize(docx_path)/1024:.1f} KB)")


# epsilon 원복
set_epsilon(0.045)


# ═══════════════════════════════════════════════
elapsed = time.time() - t0
print(f"\n{'='*70}")
print(f"  === 최종 결과 ===")
print(f"{'='*70}")

print(f"\n  [핵심 비교: bead 2.5mm Q=2100 Scenario 1 MC]")
print(f"    D^4: mu={comp_25_d4['FiPLSim_mu']:.4f} (논문 {comp_25_d4['diff_kpa']:+.2f} kPa) Pf={comp_25_d4['FiPLSim_Pf']:.2f}%")
print(f"    D^8: mu={comp_25_d8['FiPLSim_mu']:.4f} (논문 {comp_25_d8['diff_kpa']:+.2f} kPa) Pf={comp_25_d8['FiPLSim_Pf']:.2f}%")
print(f"    논문: mu=0.1100                    Pf=2.43%")
print(f"    → {better} 모델이 논문에 더 가까움")

print(f"\n  출력:")
print(f"    1. {xlsx}")
print(f"    2. {docx_path}")
print(f"  소요: {elapsed:.1f}초 ({elapsed/60:.1f}분)")
print(f"{'='*70}")
