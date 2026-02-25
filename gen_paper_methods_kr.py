"""
FiPLSim 논문용 '실험방법(Methodology)' 섹션 DOCX 생성 스크립트 — 한국어 버전
SCI 저널 투고용 — 수치해석 이론, 수식, 검증 로직 포함
"""

import io, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def build_methods_docx_kr() -> bytes:
    doc = Document()

    # ── 스타일 ──
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    # 한글 폰트 설정 (East Asian)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    navy = RGBColor(0x00, 0x00, 0x00)  # 논문은 흑색 기본

    def _set_font(run, name="맑은 고딕", size=Pt(11)):
        run.font.name = name
        run.font.size = size
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    def h1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            _set_font(r, size=Pt(14))
            r.font.color.rgb = navy
        return h

    def h2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            _set_font(r, size=Pt(12))
            r.font.color.rgb = navy
        return h

    def h3(text):
        h = doc.add_heading(text, level=3)
        for r in h.runs:
            _set_font(r, size=Pt(11))
            r.font.color.rgb = navy
        return h

    def para(text, bold=False, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        _set_font(r)
        r.bold = bold
        r.italic = italic
        return p

    def eq(text, label=""):
        """수식 삽입 (가운데 정렬)"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Cambria Math"
        r.font.size = Pt(11)
        r.italic = True
        if label:
            r2 = p.add_run(f"    ({label})")
            _set_font(r2)
        return p

    def tbl(headers, rows, caption=""):
        if caption:
            pc = doc.add_paragraph()
            rc = pc.add_run(caption)
            _set_font(rc, size=Pt(10))
            rc.bold = True
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, hd in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = hd
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
                    _set_font(r, size=Pt(9))
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                c = t.rows[ri + 1].cells[ci]
                c.text = str(val)
                for p in c.paragraphs:
                    for r in p.runs:
                        _set_font(r, size=Pt(9))
        doc.add_paragraph()  # spacing
        return t

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        _set_font(r)
        return p

    def numbered(text):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(text)
        _set_font(r)
        return p

    # ═══════════════════════════════════════════
    #  메인 제목
    # ═══════════════════════════════════════════
    title = doc.add_heading(
        "실험방법: 소방배관 시스템 수치 시뮬레이션", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        _set_font(r, size=Pt(16))
        r.font.color.rgb = navy

    doc.add_paragraph()

    # ═══════════════════════════════════════════
    #  1. 시뮬레이션 개발
    # ═══════════════════════════════════════════
    h1("1. 시뮬레이션 개발")

    # ── 1.1 개발 환경 및 도구 ──
    h2("1.1 개발 환경 및 도구")

    para(
        "수치 시뮬레이션 플랫폼 FiPLSim(Fire Protection Pipe Let Simulator)은 "
        "Python 프로그래밍 언어(버전 3.12 이상)로 개발되었으며, 다음의 수치 연산 "
        "라이브러리를 활용하였다: NumPy(≥1.24, 배열 연산 및 난수 생성), "
        "SciPy(≥1.11, 3차 스플라인 보간 및 근 탐색 알고리즘), "
        "Pandas(≥2.0, 표 형태 데이터 처리), Plotly(≥5.18, 대화형 시각화). "
        "그래픽 사용자 인터페이스는 웹 기반 대시보드 프레임워크인 "
        "Streamlit(≥1.30)으로 구현하여, 실시간 매개변수 조정 및 결과 확인이 "
        "가능하도록 하였다."
    )

    para(
        "모든 수력 계산은 SI 단위계로 수행하며, 압력은 MPa, 유량은 L/min(LPM), "
        "배관 치수는 m, 유속은 m/s로 표현한다. 계산 경계에서 단위 변환을 "
        "적용하여 연산 파이프라인 전체의 일관성을 유지한다."
    )

    # ── 1.2 수치해석 이론 ──
    h2("1.2 지배 방정식 및 수치해석 이론")

    # 1.2.1 Darcy-Weisbach
    h3("1.2.1 Darcy–Weisbach 방정식 (주요 마찰 손실)")

    para(
        "직선 배관 구간에서의 벽면 마찰에 의한 압력 손실은 Darcy–Weisbach "
        "방정식으로 산출한다. 이 방정식은 모든 유동 영역과 배관 재질에 적용 "
        "가능하며, Hazen–Williams 또는 Manning 공식의 경험적 한계가 없다:"
    )

    eq("h_f = f · (L / D) · (V² / 2g)", "1")

    para(
        "여기서 h_f는 마찰 수두 손실(m), f는 Darcy 마찰계수(무차원), "
        "L은 배관 구간 길이(m), D는 내경(m), V는 평균 유속(m/s), "
        "g는 중력가속도(9.81 m/s²)이다."
    )

    # 1.2.2 Colebrook-White
    h3("1.2.2 Colebrook–White 방정식 (마찰계수 산정)")

    para(
        "Darcy 마찰계수 f는 상용 배관의 전체 난류 유동 영역에서 유효한 "
        "Colebrook–White 음함수 방정식으로 결정한다:"
    )

    eq("1/√f = −2.0 · log₁₀ [ (ε/D)/3.7 + 2.51/(Re·√f) ]", "2")

    para(
        "여기서 ε는 배관 내벽의 절대 조도(m), Re는 레이놀즈 수이다. "
        "소방설비에 사용되는 탄소강 배관의 경우, ε = 0.045 mm "
        "(4.5 × 10⁻⁵ m)를 적용하였으며, 이는 표준 공학 참고문헌 "
        "(Moody, 1944; Colebrook, 1939)에 근거한 값이다."
    )

    para(
        "식 (2)는 f에 대해 음함수이므로, 고정점 반복법(fixed-point iteration)으로 "
        "풀이한다. 초기 추정값은 Swamee–Jain 양함수 근사식(Swamee and Jain, "
        "1976)으로부터 구한다:"
    )

    eq("f₀ = 0.25 / { log₁₀ [ (ε/D)/3.7 + 2.51/(Re·√0.02) ] }²", "3")

    para(
        "상대 변화량 |f_{n+1} − f_n| / f_n < 10⁻⁸일 때 수렴으로 판정하며, "
        "일반적으로 3~5회 반복 이내에 수렴한다. 층류 유동(Re < 2300)의 경우, "
        "Hagen–Poiseuille 해석해 f = 64/Re를 반복 없이 직접 적용한다."
    )

    # 1.2.3 Reynolds Number
    h3("1.2.3 레이놀즈 수 및 유동 영역 분류")

    eq("Re = V · D / ν", "4")

    para(
        "여기서 ν는 20°C 물의 동점성계수(1.004 × 10⁻⁶ m²/s)이다. "
        "유동 영역은 다음과 같이 분류한다: Re < 2300(층류), "
        "2300 ≤ Re ≤ 4000(천이 영역), Re > 4000(완전 난류). "
        "소방 배관 시스템은 통상 난류 영역(Re ≈ 10⁴~10⁵)에서 운전된다."
    )

    # 1.2.4 Minor Losses
    h3("1.2.4 부차적(국부) 손실 모델")

    para(
        "이음쇠, 분기점 및 부속품에서의 국부 압력 손실은 K-factor 방법으로 "
        "산출한다:"
    )

    eq("h_m = K · (V² / 2g)", "5")

    para(
        "여기서 K는 무차원 손실 계수이다. 본 시뮬레이터는 가지배관의 각 구간에서 "
        "다음과 같은 손실 성분을 고려한다:"
    )

    tbl(
        ["손실 성분", "기호", "값", "설명"],
        [
            ("용접 이음(접합부)", "K₁", "f(h_bead)", "용접 비드 돌출 높이에 따라 변동"),
            ("스프링클러 헤드 오리피스", "K₂", "2.5", "헤드 피팅 고유 저항"),
            ("가지배관 입구(Tee)", "K₃", "1.0", "교차배관-가지배관 분기점"),
            ("교차배관 Tee-직통", "K_tee", "0.3", "교차배관 직통 방향 유동"),
        ],
        caption="표 1. 각 수력 구성 요소에 대한 손실 계수(K-factor) 설정값.",
    )

    # 1.2.5 Weld Bead K-factor model
    h3("1.2.5 용접 비드 돌출부 손실 모델")

    para(
        "본 연구의 핵심 혁신은 내부 용접 비드 돌출부가 수력 저항에 미치는 영향을 "
        "정량적으로 모델링한 것이다. 일반적인 용접으로 배관을 접합할 때, 내부 비드가 "
        "유로 단면 내부로 돌출되어 유효 직경을 감소시킨다. 이로 인한 국부 손실 계수의 "
        "증가는 다음과 같이 모델링하였다:"
    )

    eq("K₁ = K₁,base · (D / D_eff)⁴", "6")

    eq("D_eff = D − 2h", "7")

    para(
        "여기서 K₁,base는 이상적 접합부의 기준 손실 계수(0.5), D는 공칭 내경(mm), "
        "h는 비드 돌출 높이(mm), D_eff는 양측 비드 돌출을 고려한 유효 직경이다."
    )

    para(
        "식 (6)에서 직경비의 4승 의존성은 Hagen–Poiseuille 법칙에서 유도된다. "
        "점성 유동이 축소부를 통과할 때 유량은 D⁴에 비례한다. 난류 유동이 "
        "오리피스형 축소부를 통과할 때, 압력 손실은 단면적비의 제곱, "
        "즉 (A/A_eff)² = (D/D_eff)⁴에 비례하며, 이는 급축소부에 적용한 "
        "에너지 방정식과 일치한다(Idelchik, 1986). 이 모델을 통해 "
        "일반 용접(h > 0, Case A)과 형상 제어 용접 기술(h ≈ 0, Case B)의 "
        "직접 비교가 가능하다."
    )

    # 1.2.6 Branch Pressure Profile
    h3("1.2.6 가지배관 압력 프로파일 계산")

    para(
        "각 가지배관을 따른 누적 압력 강하는 교차배관 접합부에서 말단 "
        "스프링클러 헤드까지 순차적으로 계산한다. 각 구간 j(j = 1, 2, ..., m)에서 "
        "상류측 헤드의 순차적 방류로 인해 국부 유량이 선형적으로 감소한다:"
    )

    eq("Q_j = Q_total − (j − 1) · (Q_total / m)", "8")

    para(
        "여기서 Q_total은 가지배관 총유량, m은 스프링클러 헤드 수이다. "
        "위치 j에서의 누적 압력은 다음과 같다:"
    )

    eq("P_j = P_{j-1} − (ρg)⁻¹ · 10⁶ · [ h_{f,j} + h_{K1,j} + h_{K2,j} + Σ h_{weld,j} ]", "9")

    para(
        "여기서 P₀는 가지배관 입구 압력(교차배관 압력에서 K₃ 분기 입구 손실을 "
        "차감한 값)이고, 합산항 Σh_weld,j는 구간 j 내의 모든 용접 비드 손실을 "
        "포함한다. 마지막 헤드에서의 말단 압력 P_m이 기준 적합성 검증의 핵심 값이다."
    )

    # 1.2.7 Pipe sizing
    h3("1.2.7 자동 배관 구경 산정 알고리즘")

    para(
        "배관 구경은 NFSC 103(국가화재안전기준)의 배관 구경 산정표에 따라 "
        "하류측 스프링클러 헤드 수를 기준으로 자동 선정한다. 이 산정 규칙은 "
        "적절한 유량 수용 능력을 확보하면서 경제적으로 효율적인 배관 구경을 "
        "보장한다:"
    )

    tbl(
        ["하류 헤드 수", "가지배관 구경", "사양"],
        [
            ("1~2", "25A", "외경 33.40 mm, 내경 26.64 mm"),
            ("3", "32A", "외경 42.16 mm, 내경 35.04 mm"),
            ("4~5", "40A", "외경 48.26 mm, 내경 40.90 mm"),
            ("6~11", "50A", "외경 60.33 mm, 내경 52.51 mm"),
            ("12 이상", "65A", "외경 73.03 mm, 내경 62.71 mm"),
        ],
        caption="표 2. 자동 배관 구경 산정 규칙 (JIS/KS Schedule 40 탄소강).",
    )

    tbl(
        ["총 시스템 헤드 수", "교차배관 구경", "내경 (mm)"],
        [
            ("20 미만", "65A", "62.71"),
            ("20~39", "80A", "77.92"),
            ("40 이상", "100A", "102.26"),
        ],
        caption="표 3. 총 시스템 헤드 수에 따른 교차배관 구경 산정.",
    )

    # 1.3 Hardy-Cross
    h2("1.3 그리드 배관망 해석: Hardy–Cross 방법")

    para(
        "상하부 교차배관을 모두 갖는 그리드(루프형) 배관 구성에서는 "
        "단순 순차 계산으로 유량 분배를 결정할 수 없다. Hardy–Cross 방법"
        "(Cross, 1936)을 적용하여 각 폐루프의 유량을 에너지 보존이 만족될 "
        "때까지 반복적으로 균형시킨다."
    )

    h3("1.3.1 정식화")

    para(
        "n개의 가지배관이 n개의 직사각형 루프를 형성하는 배관망을 고려한다. "
        "각 루프 L에서 수두 손실의 대수적 합은 0이어야 한다"
        "(키르히호프 전압 법칙의 유사 원리):"
    )

    eq("Σ h_{f,i} · sign(Q_i) = 0    (각 루프 L에 대해)", "10")

    para(
        "각 반복에서 루프 L에 대한 유량 보정값은 다음과 같이 산출한다:"
    )

    eq("ΔQ_L = − Σ h_{f,i} / Σ (∂h_{f,i}/∂Q_i)", "11")

    para(
        "난류 유동에서 수두 손실-유량 관계는 근사적으로 2차 관계(h_f ∝ Q²)이므로, "
        "미분 근사는 다음과 같다:"
    )

    eq("∂h_f/∂Q ≈ 2 · h_f / |Q|", "12")

    para(
        "각 배관의 유량은 다음과 같이 갱신한다:"
    )

    eq("Q_i^(k+1) = Q_i^(k) + ΔQ_L · d_i · ω", "13")

    para(
        "여기서 d_i는 루프 내에서 해당 배관의 방향을 나타내는 방향 계수(+1 또는 −1), "
        "ω는 감쇠 계수(기본값 0.5, 조정 범위 0.1~1.0), k는 반복 인덱스이다."
    )

    h3("1.3.2 수렴 기준")

    para(
        "다음의 이중 수렴 기준이 동시에 만족될 때 반복 과정을 종료한다:"
    )

    tbl(
        ["기준", "허용 오차", "물리적 의미"],
        [
            ("모든 루프에 대한 max |Σh_f|", "< 0.001 m (≈ 0.01 kPa)", "루프 에너지 불균형"),
            ("모든 루프에 대한 max |ΔQ|", "< 0.0001 LPM", "유량 보정량 크기"),
        ],
        caption="표 4. Hardy–Cross 이중 수렴 기준.",
    )

    para(
        "최대 반복 횟수는 1,000회로 설정한다. 발산 감지를 위해 알고리즘은 "
        "루프 불균형이 3회 연속 증가하는지를 감시하며, 발산이 감지되면 "
        "계산을 조기 종료하고 경고 메시지를 발생시켜 신뢰할 수 없는 결과의 "
        "산출을 방지한다. 이 이중 수렴 기준은 에너지 균형(압력)과 "
        "질량 보존(유량) 모두가 공학적 정밀도 수준에서 만족됨을 보장한다."
    )

    h3("1.3.3 BFS 기반 절점 압력 계산")

    para(
        "유량 수렴 후, 절점 압력은 입구 절점(기지 압력 경계조건)에서 출발하는 "
        "너비 우선 탐색(BFS)을 통해 계산한다. 기지 압력 P_i를 갖는 절점 i와 "
        "절점 j를 연결하는 각 배관에 대해, 하류측 압력은 다음과 같다:"
    )

    eq("P_j = P_i − Δp_{i→j}(Q_{pipe})", "14")

    para(
        "여기서 Δp는 수렴된 유량에 대해 Darcy–Weisbach 방정식과 모든 해당 "
        "부차적 손실 계수를 적용하여 계산한 압력 강하이다. 이를 통해 그리드 "
        "배관망 전체에 걸친 일관된 압력 분포를 보장한다."
    )

    # 1.4 Pump analysis
    h2("1.4 펌프 성능 해석")

    h3("1.4.1 펌프 곡선 보간")

    para(
        "제조사 제공 펌프 성능 데이터(Q–H 점)를 3차 스플라인"
        "(SciPy interp1d, kind='cubic')으로 보간하여 연속적인 펌프 특성 곡선 "
        "H_pump(Q)를 생성한다. 이 방법은 가정된 다항식 차수 없이 "
        "폐쇄 수두 및 곡선 말단 거동을 포함한 펌프 곡선의 물리적 형상을 "
        "보존한다."
    )

    h3("1.4.2 배관 시스템 저항 곡선")

    para(
        "시스템 저항 곡선 H_sys(Q)는 여러 유량에서 배관망 전체의 총 수두 손실을 "
        "계산하여 구성한다. 각 시험 유량 Q_test에 대해 전체 배관망 시뮬레이션을 "
        "수행하고, 총 수두 손실을 다음과 같이 결정한다:"
    )

    eq("H_sys(Q) = [P_inlet − P_terminal(Q)] · 10⁶ / (ρg) + H_min", "15")

    para(
        "여기서 H_min은 최소 말단 수두 요구량(설계 기준 0.1 MPa에 해당)이다."
    )

    h3("1.4.3 운전점 결정")

    para(
        "운전점은 펌프 곡선과 시스템 저항 곡선의 교차점, 즉 다음을 만족하는 "
        "유량 Q*에서 결정된다:"
    )

    eq("H_pump(Q*) − H_sys(Q*) = 0", "16")

    para(
        "이 근 탐색 문제는 Brent 방법(scipy.optimize.brentq)으로 풀이하며, "
        "허용 오차는 ±0.1 LPM이다. 이 방법은 구간 내에서 연속적이고 "
        "단조적으로 분기하는 펌프 곡선과 시스템 곡선에 대해 수렴을 보장한다. "
        "해당 운전점에서의 펌프 소비 전력은 다음과 같이 산출한다:"
    )

    eq("P_pump = ρ · g · Q* · H* / η", "17")

    para(
        "여기서 η는 펌프 효율(소수 분율), Q*는 m³/s 단위로 환산한 값이다."
    )

    # ── 1.5 신뢰성 로직 ──
    h2("1.5 신뢰성 및 정확도 보증")

    h3("1.5.1 입력값 검증")

    para(
        "모든 사용자 입력 매개변수는 계산 전에 경계값 검사를 거친다. "
        "유효하지 않은 구성(예: 직경 0, 음의 압력, 물리적 한계를 초과하는 "
        "헤드 수 등)은 ValidationError 예외를 발생시켜 무의미한 결과의 "
        "전파를 방지한다."
    )

    h3("1.5.2 마찰계수 수렴 제어")

    para(
        "Colebrook–White 반복 연산은 상대 수렴 기준 10⁻⁸과 최대 반복 횟수 "
        "10회를 적용한다. Colebrook–White 함수는 난류 영역에서 축약 사상"
        "(contractive mapping)이므로 이론적으로 수렴이 보장된다(Brkić, 2011). "
        "Swamee–Jain 초기 추정값이 우수한 시작점을 제공하여, 일반적으로 "
        "필요한 반복 횟수를 3~5회로 줄인다."
    )

    h3("1.5.3 Hardy–Cross 안정성 대책")

    para(
        "Hardy–Cross 해석기의 안정적 수렴을 보장하기 위해 세 가지 메커니즘을 "
        "구현하였다:"
    )

    bullet(
        "감쇠(ω = 0.5 기본값): 강하게 결합된 루프에서 유량 보정의 "
        "진동성 발산을 방지한다."
    )

    bullet(
        "발산 감지: 최대 루프 불균형이 3회 연속 증가하면 해석기를 조기 "
        "종료하고 발산 경고를 보고하여 신뢰할 수 없는 결과의 산출을 방지한다."
    )

    bullet(
        "이중 수렴 기준: 에너지 균형(수두)과 질량 보존(유량)이 동시에 "
        "만족되어야 하며, 하나의 기준만 충족된 상태에서의 조기 종료를 방지한다."
    )

    h3("1.5.4 NFPC 법규 적합성 검증")

    para(
        "시뮬레이션 완료 후, 자동 적합성 검사를 통해 국가화재예방법(NFPC) "
        "요구 사항에 대한 결과를 검증한다:"
    )

    tbl(
        ["검증 항목", "요구 기준", "검증 방법"],
        [
            ("가지배관 유속", "≤ 6.0 m/s", "모든 구간 유속 확인"),
            ("교차배관 유속", "≤ 10.0 m/s", "교차배관에 대해 V = Q/(πD²/4) 산출"),
            ("말단 압력 (최소)", "≥ 0.1 MPa", "모든 가지배관 말단 압력 확인"),
            ("말단 압력 (최대)", "≤ 1.2 MPa", "모든 가지배관 말단 압력 확인"),
        ],
        caption="표 5. NFPC 적합성 검증 기준.",
    )

    h3("1.5.5 회귀 테스트 스위트")

    para(
        "시뮬레이션 엔진은 총 130개의 단위 및 통합 테스트로 구성된 포괄적 "
        "회귀 테스트 스위트로 검증된다: 그리드 배관망 계산 46개 테스트"
        "(Hardy–Cross 수렴, 절점 압력 일관성, 키르히호프 법칙 검증), "
        "통합 시스템 거동 46개 테스트(Case A/B 비교, 개선율, 배관 구경 산정), "
        "용접 비드 손실 모델 정확도 38개 테스트(K-factor 스케일링, 무작위 비드 "
        "배치, 구간 매핑). 모든 130개 테스트를 통과해야 코드 수정이 승인되며, "
        "회귀 없는 개발을 보장한다."
    )

    # ═══════════════════════════════════════════
    #  2. 시뮬레이션 방법
    # ═══════════════════════════════════════════
    doc.add_page_break()
    h1("2. 시뮬레이션 방법")

    h2("2.1 배관망 토폴로지 구성")

    para(
        "실제 소방설비에서 사용되는 구성의 범위를 표현하기 위해, "
        "두 가지 배관망 토폴로지를 지원한다:"
    )

    h3("2.1.1 트리(가지형) 토폴로지")

    para(
        "단일 교차배관이 n개의 가지배관을 공급하며, 각 가지배관은 m개의 "
        "스프링클러 헤드를 담당한다. 유동은 교차배관 입구에서 각 가지배관을 거쳐 "
        "말단 헤드까지 단방향으로 진행된다. 압력은 식 (1)~(9)를 사용하여 "
        "각 가지배관을 따라 순차적으로 계산한다. 이 토폴로지는 중소규모 "
        "건물에서 가장 일반적인 설치 형태를 나타낸다."
    )

    h3("2.1.2 전체 그리드(루프형) 토폴로지")

    para(
        "가지배관의 상단과 하단 모두가 교차배관에 연결되어 n개의 폐 직사각형 "
        "루프를 형성한다. 유동은 양방향으로, 양 끝에서 각 가지배관으로 유입된다. "
        "평형 유량 분배를 결정하기 위해 Hardy–Cross 반복법(1.3절)을 적용한다. "
        "이 토폴로지는 수력적 중복성을 제공하며, 대규모 또는 고위험 "
        "용도의 건축물에서 의무적으로 적용된다."
    )

    h2("2.2 비교 사례 분석 (Case A 대 Case B)")

    para(
        "각 시뮬레이션 실행은 동일한 경계 조건 하에서 두 가지 용접 기술 "
        "시나리오 간의 쌍대 비교를 수행한다:"
    )

    tbl(
        ["매개변수", "Case A (일반 용접)", "Case B (형상 제어 용접)"],
        [
            ("용접 비드 높이 h", "사용자 정의 (예: 1.5 mm)", "0 mm (이상적 접합)"),
            ("이음 손실 K₁", "K₁,base · (D/D_eff)⁴", "K₁,base = 0.5"),
            ("물리적 의미", "기존 용접 기술", "첨단 형상 제어 기술"),
        ],
        caption="표 6. 비교 사례 정의.",
    )

    para(
        "개선 지표는 최악 조건 말단 압력의 백분율 변화로 산출한다:"
    )

    eq("개선율 (%) = (P_B − P_A) / |P_A| × 100", "18")

    h2("2.3 몬테카를로 확률론적 분석")

    h3("2.3.1 확률적 결함 모델링")

    para(
        "용접 결함 위치 변동성의 확률적 영향을 평가하기 위해 몬테카를로 "
        "시뮬레이션 프레임워크를 구현하였다. 각 시행 i(i = 1, 2, ..., N)에서:"
    )

    numbered(
        "결함 이음쇠 수 n_d를 [n_min, n_max] 범위에서 균일 분포로 추출한다."
    )

    numbered(
        "n_d개의 결함 위치를 전체 이음쇠 위치 집합 {0, 1, ..., n×m − 1}에서 "
        "비복원 추출한다(NumPy의 default_rng 난수 생성기, PCG64 알고리즘 사용)."
    )

    numbered(
        "직선 배관 용접 비드가 활성화된 경우(beads_per_branch > 0), "
        "비드 위치를 각 가지배관 내에서 가지배관 길이에 대한 균일 분포로 "
        "재무작위화하여 시공 변동성을 모의한다."
    )

    numbered(
        "전체 시스템 수력 계산을 수행하고, 최악 조건 말단 압력 "
        "P_terminal,i를 기록한다."
    )

    h3("2.3.2 통계적 출력 지표")

    para(
        "N개의 말단 압력 표본 {P₁, P₂, ..., P_N}으로부터 다음의 통계량을 산출한다:"
    )

    eq("μ = (1/N) · Σ P_i", "19")

    eq("σ = √[ Σ(P_i − μ)² / (N − 1) ]", "20")

    para(
        "표본 표준편차(Bessel 보정, N − 1 분모)를 모표준편차 대신 사용하여 "
        "불편 추정량을 제공한다. 부적합 확률 P_fail은 말단 압력이 "
        "법규 최소 기준 이하인 시행의 경험적 비율로 산출한다:"
    )

    eq("P_fail = (1/N) · Σ I(P_i < 0.1 MPa)", "21")

    para(
        "여기서 I(·)는 지시 함수이다. N = 1,000 반복의 경우, 이 추정량의 "
        "표준 오차는 최대 √[0.5 × 0.5 / 1000] = 0.016으로, 공학적 "
        "의사 결정에 충분한 정밀도를 제공한다. "
        "더 높은 정밀도가 필요한 경우 N을 최대 10,000까지 증가시킬 수 있다."
    )

    h2("2.4 결정론적 민감도 분석")

    para(
        "배관망 내에서 수력적으로 가장 민감한 위치를 식별하기 위해 "
        "단일 비드 섭동 분석을 수행한다. 절차는 다음과 같다:"
    )

    numbered(
        "기준 계산: 이음쇠 비드 돌출이 없는 상태(모든 이음쇠에서 "
        "K₁ = K₁,base)로 시스템을 시뮬레이션하여 기준 말단 압력 P_base를 구한다."
    )

    numbered(
        "최악 조건 가지배관의 각 헤드 위치 j(j = 1, ..., m)에 대해, "
        "위치 j에만 높이 h의 단일 비드를 배치하고 나머지 모든 위치는 h = 0으로 유지한다."
    )

    numbered(
        "단일 비드로 인한 압력 강하를 Δp_j = P_base − P_j로 산출한다."
    )

    numbered(
        "위치를 Δp_j의 내림차순으로 순위를 매긴다. 가장 큰 Δp를 갖는 위치를 "
        "임계점으로 지정한다."
    )

    para(
        "이 분석을 통해 시스템 성능에 가장 큰 영향력을 갖는 접합 위치가 "
        "드러나며, 시공 중 품질 관리 우선순위 결정에 실행 가능한 "
        "지침을 제공한다."
    )

    h2("2.5 변수 스캐닝 (매개변수 연구)")

    para(
        "시스템의 운전 영역 및 부적합 경계를 파악하기 위해, 변수 스캐닝 "
        "분석은 다른 모든 매개변수를 일정하게 유지한 채 단일 설계 매개변수를 "
        "체계적으로 변화시킨다. 스캔 절차는 각 매개변수 값에서 전체 "
        "Case A/B 비교를 수행하고 다음을 기록한다:"
    )

    bullet("Case A 및 Case B의 최악 조건 말단 압력.")

    bullet("개선율(%).")

    bullet("0.1 MPa 최소 기준에 대한 PASS/FAIL 상태.")

    para(
        "임계점 Q_crit(또는 P_crit, h_crit 등)는 시스템이 최초로 PASS에서 "
        "FAIL로 전환되는 매개변수 값으로 정의한다. 스캔 분석에 사용 가능한 "
        "설계 매개변수는 다음과 같다: 설계 유량(LPM), 입구 압력(MPa), "
        "용접 비드 높이(mm), 가지배관당 헤드 수."
    )

    h2("2.6 경제성 분석 (생애주기 비용)")

    para(
        "수력 저항 감소(Case B 대 Case A)로 인한 에너지 절감량을 "
        "생애주기 비용(LCC) 분석을 통해 정량화한다. 각 운전점에서 "
        "펌프 소비 전력을 식 (17)로 산출하고, 차등 전력은 다음과 같다:"
    )

    eq("ΔP = P_pump,A − P_pump,B", "22")

    eq("연간 에너지 절감량 = ΔP × t_op  (kWh/년)", "23")

    eq("연간 비용 절감액 = ΔP × t_op × C_e  (원/년)", "24")

    para(
        "여기서 t_op는 연간 펌프 운전 시간(기본값 2,000시간/년), "
        "C_e는 단위 전기 요금(기본값 120원/kWh)이다. "
        "이 기본값은 현장 특성 조건을 반영하도록 사용자가 조정할 수 있다."
    )

    # ── 물성치 요약 ──
    doc.add_page_break()
    h2("2.7 물성치 및 상수")

    tbl(
        ["물성치", "기호", "값", "단위", "출처"],
        [
            ("물 밀도 (20°C)", "ρ", "998.0", "kg/m³", "공학 데이터"),
            ("동점성계수 (20°C)", "ν", "1.004 × 10⁻⁶", "m²/s", "공학 데이터"),
            ("동적 점성계수 (20°C)", "μ", "1.002 × 10⁻³", "Pa·s", "공학 데이터"),
            ("중력가속도", "g", "9.81", "m/s²", "—"),
            ("배관 조도 (탄소강)", "ε", "0.045", "mm", "Moody (1944)"),
            ("기준 이음 손실 계수", "K₁,base", "0.5", "—", "Idelchik (1986)"),
            ("스프링클러 헤드 손실 계수", "K₂", "2.5", "—", "NFSC 103"),
            ("가지배관 입구 손실 계수", "K₃", "1.0", "—", "NFSC 103"),
            ("최소 말단 압력", "P_min", "0.1", "MPa", "NFPC"),
            ("최대 말단 압력", "P_max", "1.2", "MPa", "NFPC"),
            ("최대 가지배관 유속", "V_max,br", "6.0", "m/s", "NFPC"),
            ("최대 교차배관 유속", "V_max,cm", "10.0", "m/s", "NFPC"),
        ],
        caption="표 7. 시뮬레이션에 사용된 물성치 및 설계 상수.",
    )

    # ── 참고문헌 ──
    doc.add_page_break()
    h1("참고문헌")

    refs = [
        "Brkić, D. (2011). Review of explicit approximations to the Colebrook relation for flow friction. "
        "Journal of Petroleum Science and Engineering, 77(1), 34–48.",

        "Colebrook, C. F. (1939). Turbulent flow in pipes, with particular reference to the transition "
        "region between the smooth and rough pipe laws. Journal of the Institution of Civil Engineers, "
        "11(4), 133–156.",

        "Cross, H. (1936). Analysis of flow in networks of conduits or conductors. "
        "Bulletin No. 286, University of Illinois Engineering Experiment Station.",

        "Idelchik, I. E. (1986). Handbook of Hydraulic Resistance, 2nd ed. Hemisphere Publishing.",

        "Moody, L. F. (1944). Friction factors for pipe flow. Transactions of the ASME, 66(8), 671–684.",

        "Swamee, P. K., & Jain, A. K. (1976). Explicit equations for pipe-flow problems. "
        "Journal of the Hydraulics Division, 102(5), 657–664.",

        "국가화재예방법(NFPC), 대한민국.",

        "국가화재안전기준 103 (NFSC 103): 스프링클러설비의 화재안전기준, 대한민국.",
    ]

    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(ref)
        _set_font(r, size=Pt(10))

    # ── 저장 ──
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == "__main__":
    data = build_methods_docx_kr()
    out_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "FiPLSim_Paper_Methodology_KR.docx",
    )
    with open(out_path, "wb") as f:
        f.write(data)
    print(f"한국어 문서 생성 완료: {out_path}")
    print(f"파일 크기: {len(data):,} bytes")
