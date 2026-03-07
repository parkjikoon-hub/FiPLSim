"""
FiPLSim — Table 8 재현 시뮬레이션 (6종 실행)

논문: "Stochastic assessment of weld bead effects on fire sprinkler piping networks"

시뮬레이션 조건:
  - 4 branches × 8 heads = 32 junctions, tree topology
  - 입구 압력: 0.4 MPa
  - 밸브: OFF (equipment_k_factors=None)
  - 라이저 관경: 80A (교차배관)
  - MC 반복 횟수: 10,000

실행 계획 (6종):
  ① Case B (0mm), Q=1,200 LPM — 기준값 확인
  ② Case B (0mm), Q=1,600 LPM — 기준값 확인
  ③ Case B (0mm), Q=2,100 / 2,300 LPM — 기준값 확인
  ④ Case A (2.0mm), Q=1,200/1,600/2,100 LPM — 핵심 재현
  ⑤ Case A (2.5mm), Q=1,200/1,600 LPM — 기준값 확인
  ⑥ Case A (2.5mm), Q=2,100 & 2,300 LPM — Table 8 핵심

출력:
  1. FiPLSim_Table8_재현_데이터.xlsx  (Excel 데이터)
  2. FiPLSim_Table8_재현_결과보고서.docx  (DOCX 보고서)
"""

import os, sys, time, datetime, math
import numpy as np
import pandas as pd

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from constants import (
    PIPE_DIMENSIONS, get_inner_diameter_m, RHO, G,
    MIN_TERMINAL_PRESSURE_MPA, K1_BASE, K2, K3,
    DEFAULT_EQUIPMENT_K_FACTORS,
)
from hydraulics import velocity_from_flow
from pipe_network import generate_dynamic_system, calculate_dynamic_system

# ══════════════════════════════════════════════
# 공통 설정 (논문 조건)
# ══════════════════════════════════════════════
NUM_BR = 4
HEADS = 8
BRANCH_SP = 3.5      # m
HEAD_SP = 2.3         # m
INLET_P = 0.4         # MPa
SUPPLY = "80A"        # 라이저 관경 (논문 조건)
N_ITER = 10000        # MC 반복 횟수
SCENARIO1_P = 0.5     # Scenario 1 비드 확률
BEADS_PER_BRANCH = 8  # 가지배관당 용접 비드 개수

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


# ══════════════════════════════════════════════
# Scenario 1 MC 함수 (worst branch에만 비드 배치)
# ══════════════════════════════════════════════
def run_scenario1_mc(
    p_bead, n_iterations, bead_height_mm, total_flow_lpm,
    beads_per_branch=0,
):
    """
    Scenario 1: worst branch(마지막 가지배관)의 8개 junction에만
    독립 Bernoulli(p_bead) 비드 배치. 나머지 branch는 bead=0.
    """
    rng = np.random.default_rng()
    worst_pressures = np.zeros(n_iterations)
    bead_counts = np.zeros(n_iterations, dtype=int)
    worst_branch_idx = NUM_BR - 1

    for trial in range(n_iterations):
        beads_2d = [[0.0] * HEADS for _ in range(NUM_BR)]

        # worst branch에만 Bernoulli 비드 배치
        rand_vals = rng.uniform(0, 1, size=HEADS)
        count = 0
        for h in range(HEADS):
            if rand_vals[h] <= p_bead:
                beads_2d[worst_branch_idx][h] = bead_height_mm
                count += 1
        bead_counts[trial] = count

        system = generate_dynamic_system(
            num_branches=NUM_BR, heads_per_branch=HEADS,
            branch_spacing_m=BRANCH_SP, head_spacing_m=HEAD_SP,
            inlet_pressure_mpa=INLET_P, total_flow_lpm=float(total_flow_lpm),
            bead_heights_2d=beads_2d,
            beads_per_branch=beads_per_branch,
            bead_height_for_weld_mm=bead_height_mm,
            rng=rng,
        )
        result = calculate_dynamic_system(
            system, K3,
            equipment_k_factors=None,
            supply_pipe_size=SUPPLY,
        )
        worst_pressures[trial] = result["worst_terminal_mpa"]

    below = np.sum(worst_pressures < MIN_TERMINAL_PRESSURE_MPA)

    return {
        "terminal_pressures": worst_pressures,
        "bead_counts": bead_counts,
        "mean_pressure": float(np.mean(worst_pressures)),
        "std_pressure": float(np.std(worst_pressures, ddof=1)) if n_iterations > 1 else 0.0,
        "min_pressure": float(np.min(worst_pressures)),
        "max_pressure": float(np.max(worst_pressures)),
        "p_below_threshold": float(below / n_iterations),
        "mean_bead_count": float(np.mean(bead_counts)),
        "expected_bead_count": float(HEADS * p_bead),
        "percentiles": {
            "P5": float(np.percentile(worst_pressures, 5)),
            "P25": float(np.percentile(worst_pressures, 25)),
            "P50": float(np.percentile(worst_pressures, 50)),
            "P75": float(np.percentile(worst_pressures, 75)),
            "P95": float(np.percentile(worst_pressures, 95)),
        },
    }


def run_deterministic(total_flow_lpm, bead_height_mm=0.0, worst_branch_only=False):
    """결정론적 계산 (단일 실행)"""
    if worst_branch_only and bead_height_mm > 0:
        beads_2d = [[0.0] * HEADS for _ in range(NUM_BR)]
        beads_2d[NUM_BR - 1] = [bead_height_mm] * HEADS
    elif bead_height_mm > 0:
        beads_2d = [[bead_height_mm] * HEADS for _ in range(NUM_BR)]
    else:
        beads_2d = None

    system = generate_dynamic_system(
        num_branches=NUM_BR, heads_per_branch=HEADS,
        branch_spacing_m=BRANCH_SP, head_spacing_m=HEAD_SP,
        inlet_pressure_mpa=INLET_P, total_flow_lpm=float(total_flow_lpm),
        bead_heights_2d=beads_2d,
    )
    result = calculate_dynamic_system(
        system, K3,
        equipment_k_factors=None,
        supply_pipe_size=SUPPLY,
    )
    return result


# ══════════════════════════════════════════════
# 메인 실행
# ══════════════════════════════════════════════
print("=" * 70)
print("  FiPLSim — Table 8 재현 시뮬레이션 (6종)")
print("=" * 70)
print(f"  {NUM_BR} branches × {HEADS} heads = {NUM_BR * HEADS} junctions")
print(f"  입구 압력: {INLET_P} MPa | 밸브: OFF | MC: {N_ITER:,}회")
print(f"  라이저 관경: {SUPPLY} | Scenario 1: p={SCENARIO1_P}")
t0 = time.time()


# ──────────────────────────────────────────────
# Sheet 1: 시뮬레이션 조건
# ──────────────────────────────────────────────
df_cond = pd.DataFrame({
    "항목": [
        "가지배관 수 (n)", "가지배관당 헤드 (m)", "전체 junction 수",
        "가지배관 간격 (m)", "헤드 간격 (m)", "입구 압력 (MPa)",
        "라이저 관경", "밸브 상태", "MC 반복 횟수",
        "말단 최소 기준 (MPa)", "Scenario 유형",
        "Scenario 1 비드 확률 (p)", "비드 적용 범위",
    ],
    "값": [
        NUM_BR, HEADS, NUM_BR * HEADS,
        BRANCH_SP, HEAD_SP, INLET_P,
        SUPPLY, "OFF (밸브 미포함)", f"{N_ITER:,}",
        MIN_TERMINAL_PRESSURE_MPA, "Scenario 1 (결함 집중)",
        SCENARIO1_P, f"worst branch {HEADS} junctions만",
    ],
})


# ──────────────────────────────────────────────
# ① ② ③ Case B 기준선 (bead=0mm, 결정론적)
# ──────────────────────────────────────────────
print("\n" + "─" * 70)
print("  [①②③] Case B 기준선 (bead=0mm) — 결정론적 계산")
print("─" * 70)

caseB_flows = [1200, 1600, 2100, 2300]
caseB_rows = []

for Q in caseB_flows:
    res = run_deterministic(Q, bead_height_mm=0.0)
    p_term = res["worst_terminal_mpa"]
    status = "PASS" if p_term >= MIN_TERMINAL_PRESSURE_MPA else "FAIL"
    margin = (p_term - 0.1) / 0.1 * 100 if p_term > 0 else -100

    # 교차배관 손실
    cm_loss = res["cross_main_cumulative"]

    # worst branch 상세
    wb = res["worst_branch_index"]
    wp = res["branch_profiles"][wb]

    caseB_rows.append({
        "Q (LPM)": Q,
        "말단 압력 (MPa)": round(p_term, 6),
        "말단 압력 (kPa)": round(p_term * 1000, 2),
        "교차배관 손실 (kPa)": round(cm_loss * 1000, 2),
        "K3 분기 손실 (kPa)": round(wp["K3_loss_mpa"] * 1000, 2),
        "가지배관 총 손실 (kPa)": round((INLET_P - p_term) * 1000, 2),
        "최악 가지배관": wb + 1,
        "0.1 MPa 여유 (%)": round(margin, 1),
        "판정": status,
    })
    print(f"  Q={Q:>5} LPM: 말단 {p_term:.6f} MPa ({p_term*1000:.2f} kPa) [{status}]")

df_caseB = pd.DataFrame(caseB_rows)


# ──────────────────────────────────────────────
# ④ Case A (2.0mm) — Scenario 1 MC, p=0.5
# ──────────────────────────────────────────────
print("\n" + "─" * 70)
print("  [④] Case A (bead=2.0mm) — Scenario 1 MC, p=0.5")
print("─" * 70)

caseA_20_flows = [1200, 1600, 2100]
caseA_20_rows = []
caseA_20_raw = {}

for Q in caseA_20_flows:
    print(f"  bead=2.0mm, Q={Q}: ", end="", flush=True)
    t_s = time.time()

    res_mc = run_scenario1_mc(
        p_bead=SCENARIO1_P, n_iterations=N_ITER,
        bead_height_mm=2.0, total_flow_lpm=float(Q),
    )

    # 동일 유량의 Case B 기준선
    caseB_ref = next((r for r in caseB_rows if r["Q (LPM)"] == Q), None)
    caseB_val = caseB_ref["말단 압력 (MPa)"] if caseB_ref else 0

    pct = res_mc["percentiles"]

    caseA_20_rows.append({
        "Bead (mm)": 2.0, "Q (LPM)": Q,
        "Case B (MPa)": round(caseB_val, 6),
        "μ (MPa)": round(res_mc["mean_pressure"], 6),
        "σ (MPa)": round(res_mc["std_pressure"], 6),
        "Min (MPa)": round(res_mc["min_pressure"], 6),
        "Max (MPa)": round(res_mc["max_pressure"], 6),
        "P5% (MPa)": round(pct["P5"], 6),
        "P50% (MPa)": round(pct["P50"], 6),
        "P95% (MPa)": round(pct["P95"], 6),
        "Δμ vs CaseB (kPa)": round((caseB_val - res_mc["mean_pressure"]) * 1000, 2),
        "E[beads]": round(res_mc["expected_bead_count"], 1),
        "실제 평균 beads": round(res_mc["mean_bead_count"], 2),
        "Pf (%)": round(res_mc["p_below_threshold"] * 100, 2),
    })

    caseA_20_raw[f"Q{Q}_pressure"] = res_mc["terminal_pressures"]
    caseA_20_raw[f"Q{Q}_beads"] = res_mc["bead_counts"]

    dt = time.time() - t_s
    print(f"μ={res_mc['mean_pressure']:.6f}, σ={res_mc['std_pressure']:.6f}, "
          f"Pf={res_mc['p_below_threshold']*100:.2f}% ({dt:.1f}s)")

df_caseA_20 = pd.DataFrame(caseA_20_rows)


# ──────────────────────────────────────────────
# ⑤ Case A (2.5mm) — Scenario 1 MC, p=0.5, Q=1200/1600
# ──────────────────────────────────────────────
print("\n" + "─" * 70)
print("  [⑤] Case A (bead=2.5mm) — Scenario 1 MC, Q=1200/1600")
print("─" * 70)

caseA_25_base_flows = [1200, 1600]
caseA_25_base_rows = []
caseA_25_base_raw = {}

for Q in caseA_25_base_flows:
    print(f"  bead=2.5mm, Q={Q}: ", end="", flush=True)
    t_s = time.time()

    res_mc = run_scenario1_mc(
        p_bead=SCENARIO1_P, n_iterations=N_ITER,
        bead_height_mm=2.5, total_flow_lpm=float(Q),
    )

    caseB_ref = next((r for r in caseB_rows if r["Q (LPM)"] == Q), None)
    caseB_val = caseB_ref["말단 압력 (MPa)"] if caseB_ref else 0

    pct = res_mc["percentiles"]

    caseA_25_base_rows.append({
        "Bead (mm)": 2.5, "Q (LPM)": Q,
        "Case B (MPa)": round(caseB_val, 6),
        "μ (MPa)": round(res_mc["mean_pressure"], 6),
        "σ (MPa)": round(res_mc["std_pressure"], 6),
        "Min (MPa)": round(res_mc["min_pressure"], 6),
        "Max (MPa)": round(res_mc["max_pressure"], 6),
        "P5% (MPa)": round(pct["P5"], 6),
        "P50% (MPa)": round(pct["P50"], 6),
        "P95% (MPa)": round(pct["P95"], 6),
        "Δμ vs CaseB (kPa)": round((caseB_val - res_mc["mean_pressure"]) * 1000, 2),
        "E[beads]": round(res_mc["expected_bead_count"], 1),
        "실제 평균 beads": round(res_mc["mean_bead_count"], 2),
        "Pf (%)": round(res_mc["p_below_threshold"] * 100, 2),
    })

    caseA_25_base_raw[f"Q{Q}_pressure"] = res_mc["terminal_pressures"]
    caseA_25_base_raw[f"Q{Q}_beads"] = res_mc["bead_counts"]

    dt = time.time() - t_s
    print(f"μ={res_mc['mean_pressure']:.6f}, σ={res_mc['std_pressure']:.6f}, "
          f"Pf={res_mc['p_below_threshold']*100:.2f}% ({dt:.1f}s)")

df_caseA_25_base = pd.DataFrame(caseA_25_base_rows)


# ──────────────────────────────────────────────
# ⑥ Case A (2.5mm) — Scenario 1 MC, p=0.5, Q=2100/2300
#    ★ Table 8 핵심 — 오류 원인 규명
# ──────────────────────────────────────────────
print("\n" + "─" * 70)
print("  [⑥] ★ Case A (bead=2.5mm) — Table 8 핵심, Q=2100/2300")
print("─" * 70)

caseA_25_core_flows = [2100, 2300]
caseA_25_core_rows = []
caseA_25_core_raw = {}

for Q in caseA_25_core_flows:
    print(f"  bead=2.5mm, Q={Q}: ", end="", flush=True)
    t_s = time.time()

    res_mc = run_scenario1_mc(
        p_bead=SCENARIO1_P, n_iterations=N_ITER,
        bead_height_mm=2.5, total_flow_lpm=float(Q),
    )

    caseB_ref = next((r for r in caseB_rows if r["Q (LPM)"] == Q), None)
    caseB_val = caseB_ref["말단 압력 (MPa)"] if caseB_ref else 0

    pct = res_mc["percentiles"]

    caseA_25_core_rows.append({
        "Bead (mm)": 2.5, "Q (LPM)": Q,
        "Case B (MPa)": round(caseB_val, 6),
        "μ (MPa)": round(res_mc["mean_pressure"], 6),
        "σ (MPa)": round(res_mc["std_pressure"], 6),
        "Min (MPa)": round(res_mc["min_pressure"], 6),
        "Max (MPa)": round(res_mc["max_pressure"], 6),
        "P5% (MPa)": round(pct["P5"], 6),
        "P25% (MPa)": round(pct["P25"], 6),
        "P50% (MPa)": round(pct["P50"], 6),
        "P75% (MPa)": round(pct["P75"], 6),
        "P95% (MPa)": round(pct["P95"], 6),
        "Δμ vs CaseB (kPa)": round((caseB_val - res_mc["mean_pressure"]) * 1000, 2),
        "E[beads]": round(res_mc["expected_bead_count"], 1),
        "실제 평균 beads": round(res_mc["mean_bead_count"], 2),
        "Pf (%)": round(res_mc["p_below_threshold"] * 100, 2),
    })

    caseA_25_core_raw[f"Q{Q}_pressure"] = res_mc["terminal_pressures"]
    caseA_25_core_raw[f"Q{Q}_beads"] = res_mc["bead_counts"]

    dt = time.time() - t_s
    print(f"μ={res_mc['mean_pressure']:.6f}, σ={res_mc['std_pressure']:.6f}, "
          f"Pf={res_mc['p_below_threshold']*100:.2f}% ({dt:.1f}s)")

df_caseA_25_core = pd.DataFrame(caseA_25_core_rows)


# ──────────────────────────────────────────────
# Sheet: Table 8 통합 요약 (전체 결과 비교표)
# ──────────────────────────────────────────────
print("\n" + "─" * 70)
print("  통합 요약표 구성 중...")
print("─" * 70)

summary_rows = []

# Case B (0mm) — 결정론적
for r in caseB_rows:
    summary_rows.append({
        "Case": "Case B (0mm)",
        "Bead (mm)": 0.0,
        "Q (LPM)": r["Q (LPM)"],
        "μ (MPa)": r["말단 압력 (MPa)"],
        "σ (MPa)": 0.0,
        "Min (MPa)": r["말단 압력 (MPa)"],
        "Max (MPa)": r["말단 압력 (MPa)"],
        "Pf (%)": 0.0 if r["판정"] == "PASS" else 100.0,
        "E[beads]": 0,
        "비고": "결정론적 (①②③)",
    })

# Case A (2.0mm)
for r in caseA_20_rows:
    summary_rows.append({
        "Case": "Case A (2.0mm)",
        "Bead (mm)": 2.0,
        "Q (LPM)": r["Q (LPM)"],
        "μ (MPa)": r["μ (MPa)"],
        "σ (MPa)": r["σ (MPa)"],
        "Min (MPa)": r["Min (MPa)"],
        "Max (MPa)": r["Max (MPa)"],
        "Pf (%)": r["Pf (%)"],
        "E[beads]": r["E[beads]"],
        "비고": "MC Scenario 1 (④)",
    })

# Case A (2.5mm) — 기준
for r in caseA_25_base_rows:
    summary_rows.append({
        "Case": "Case A (2.5mm)",
        "Bead (mm)": 2.5,
        "Q (LPM)": r["Q (LPM)"],
        "μ (MPa)": r["μ (MPa)"],
        "σ (MPa)": r["σ (MPa)"],
        "Min (MPa)": r["Min (MPa)"],
        "Max (MPa)": r["Max (MPa)"],
        "Pf (%)": r["Pf (%)"],
        "E[beads]": r["E[beads]"],
        "비고": "MC Scenario 1 (⑤)",
    })

# Case A (2.5mm) — Table 8 핵심
for r in caseA_25_core_rows:
    summary_rows.append({
        "Case": "Case A (2.5mm) ★",
        "Bead (mm)": 2.5,
        "Q (LPM)": r["Q (LPM)"],
        "μ (MPa)": r["μ (MPa)"],
        "σ (MPa)": r["σ (MPa)"],
        "Min (MPa)": r["Min (MPa)"],
        "Max (MPa)": r["Max (MPa)"],
        "Pf (%)": r["Pf (%)"],
        "E[beads]": r["E[beads]"],
        "비고": "★ Table 8 핵심 (⑥)",
    })

df_summary = pd.DataFrame(summary_rows)


# ──────────────────────────────────────────────
# Sheet: 논문 대비 비교 (Table 8 format)
# ──────────────────────────────────────────────
# 논문 Table 8 참고값 (있는 경우)
paper_ref = {
    (2.5, 2100): {"μ": 0.1100, "σ": 0.0048, "Pf": 2.43},
    (2.5, 2300): {"μ": None, "σ": None, "Pf": None},  # 논문에 없을 수 있음
}

comparison_rows = []
for r in caseA_25_core_rows:
    Q = r["Q (LPM)"]
    bh = r["Bead (mm)"]
    ref = paper_ref.get((bh, Q), {})

    row = {
        "Q (LPM)": Q,
        "Bead (mm)": bh,
        "FiPLSim μ (MPa)": r["μ (MPa)"],
        "FiPLSim σ (MPa)": r["σ (MPa)"],
        "FiPLSim Pf (%)": r["Pf (%)"],
    }
    if ref.get("μ") is not None:
        row["논문 μ (MPa)"] = ref["μ"]
        row["논문 σ (MPa)"] = ref["σ"]
        row["논문 Pf (%)"] = ref["Pf"]
        row["μ 차이 (kPa)"] = round((r["μ (MPa)"] - ref["μ"]) * 1000, 2)
        row["μ 비율 (%)"] = round(r["μ (MPa)"] / ref["μ"] * 100, 1) if ref["μ"] else "—"
    else:
        row["논문 μ (MPa)"] = "—"
        row["논문 σ (MPa)"] = "—"
        row["논문 Pf (%)"] = "—"
        row["μ 차이 (kPa)"] = "—"
        row["μ 비율 (%)"] = "—"

    comparison_rows.append(row)

df_comparison = pd.DataFrame(comparison_rows)


# ──────────────────────────────────────────────
# Sheet: 비드 높이별 영향 분석 (2.0mm vs 2.5mm)
# ──────────────────────────────────────────────
bead_effect_rows = []
for Q in [1200, 1600]:
    r_20 = next(r for r in caseA_20_rows if r["Q (LPM)"] == Q)
    r_25 = next(r for r in caseA_25_base_rows if r["Q (LPM)"] == Q)
    caseB_ref = next(r for r in caseB_rows if r["Q (LPM)"] == Q)

    bead_effect_rows.append({
        "Q (LPM)": Q,
        "Case B (MPa)": caseB_ref["말단 압력 (MPa)"],
        "2.0mm μ (MPa)": r_20["μ (MPa)"],
        "2.5mm μ (MPa)": r_25["μ (MPa)"],
        "2.0mm 손실 (kPa)": r_20["Δμ vs CaseB (kPa)"],
        "2.5mm 손실 (kPa)": r_25["Δμ vs CaseB (kPa)"],
        "2.5/2.0 비율": round(r_25["Δμ vs CaseB (kPa)"] / r_20["Δμ vs CaseB (kPa)"], 2) if r_20["Δμ vs CaseB (kPa)"] != 0 else "—",
        "2.0mm Pf (%)": r_20["Pf (%)"],
        "2.5mm Pf (%)": r_25["Pf (%)"],
    })

# Q=2100 (20mm와 25mm 모두 있음)
r_20_2100 = next(r for r in caseA_20_rows if r["Q (LPM)"] == 2100)
r_25_2100 = next(r for r in caseA_25_core_rows if r["Q (LPM)"] == 2100)
caseB_2100 = next(r for r in caseB_rows if r["Q (LPM)"] == 2100)

bead_effect_rows.append({
    "Q (LPM)": 2100,
    "Case B (MPa)": caseB_2100["말단 압력 (MPa)"],
    "2.0mm μ (MPa)": r_20_2100["μ (MPa)"],
    "2.5mm μ (MPa)": r_25_2100["μ (MPa)"],
    "2.0mm 손실 (kPa)": r_20_2100["Δμ vs CaseB (kPa)"],
    "2.5mm 손실 (kPa)": r_25_2100["Δμ vs CaseB (kPa)"],
    "2.5/2.0 비율": round(r_25_2100["Δμ vs CaseB (kPa)"] / r_20_2100["Δμ vs CaseB (kPa)"], 2) if r_20_2100["Δμ vs CaseB (kPa)"] != 0 else "—",
    "2.0mm Pf (%)": r_20_2100["Pf (%)"],
    "2.5mm Pf (%)": r_25_2100["Pf (%)"],
})

df_bead_effect = pd.DataFrame(bead_effect_rows)


# ──────────────────────────────────────────────
# Sheet: MC 원시 분포 데이터 (④⑤⑥)
# ──────────────────────────────────────────────
raw_dist = {}

# ④ bead 2.0mm
for Q in caseA_20_flows:
    key_p = f"Q{Q}_pressure"
    key_b = f"Q{Q}_beads"
    if key_p in caseA_20_raw:
        raw_dist[f"A2.0_Q{Q}_P(MPa)"] = np.round(caseA_20_raw[key_p], 6)
        raw_dist[f"A2.0_Q{Q}_beads"] = caseA_20_raw[key_b]

# ⑤ bead 2.5mm base
for Q in caseA_25_base_flows:
    key_p = f"Q{Q}_pressure"
    key_b = f"Q{Q}_beads"
    if key_p in caseA_25_base_raw:
        raw_dist[f"A2.5_Q{Q}_P(MPa)"] = np.round(caseA_25_base_raw[key_p], 6)
        raw_dist[f"A2.5_Q{Q}_beads"] = caseA_25_base_raw[key_b]

# ⑥ bead 2.5mm core
for Q in caseA_25_core_flows:
    key_p = f"Q{Q}_pressure"
    key_b = f"Q{Q}_beads"
    if key_p in caseA_25_core_raw:
        raw_dist[f"A2.5_Q{Q}_P(MPa)"] = np.round(caseA_25_core_raw[key_p], 6)
        raw_dist[f"A2.5_Q{Q}_beads"] = caseA_25_core_raw[key_b]

df_raw = pd.DataFrame(raw_dist) if raw_dist else pd.DataFrame()


# ══════════════════════════════════════════════
# Excel 저장
# ══════════════════════════════════════════════
print("\n" + "=" * 70)
print("  Excel 데이터 저장 중...")
print("=" * 70)

xlsx_path = os.path.join(BASE_DIR, "FiPLSim_Table8_재현_데이터.xlsx")

with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
    df_cond.to_excel(writer, sheet_name="1_시뮬레이션조건", index=False)
    df_caseB.to_excel(writer, sheet_name="2_CaseB_기준선", index=False)
    df_caseA_20.to_excel(writer, sheet_name="3_CaseA_2.0mm_MC", index=False)
    df_caseA_25_base.to_excel(writer, sheet_name="4_CaseA_2.5mm_기준", index=False)
    df_caseA_25_core.to_excel(writer, sheet_name="5_CaseA_2.5mm_핵심", index=False)
    df_summary.to_excel(writer, sheet_name="6_통합요약_Table8", index=False)
    df_comparison.to_excel(writer, sheet_name="7_논문대비비교", index=False)
    df_bead_effect.to_excel(writer, sheet_name="8_비드높이별효과", index=False)
    if not df_raw.empty:
        df_raw.to_excel(writer, sheet_name="9_MC원시분포", index=False)

xlsx_size = os.path.getsize(xlsx_path) / 1024
print(f"  ✅ Excel 저장 완료: {xlsx_path}")
print(f"     파일 크기: {xlsx_size:.1f} KB")


# ══════════════════════════════════════════════
# DOCX 보고서 생성
# ══════════════════════════════════════════════
print("\n" + "=" * 70)
print("  DOCX 결과보고서 생성 중...")
print("=" * 70)

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

HEADER_BG = RGBColor(0x2E, 0x5E, 0x8E)
LIGHT_BG = RGBColor(0xE8, 0xF0, 0xFA)
HIGHLIGHT_BG = RGBColor(0xFF, 0xF3, 0xCD)  # 강조 (노란색)


def set_cell_bg(cell, color: RGBColor):
    shading = cell._element.get_or_add_tcPr()
    shading_el = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): str(color),
    })
    shading.append(shading_el)


def add_styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    """헤더 + 데이터행 테이블 생성"""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, HEADER_BG)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
            if highlight_rows and r_idx in highlight_rows:
                set_cell_bg(cell, HIGHLIGHT_BG)
            elif r_idx % 2 == 1:
                set_cell_bg(cell, LIGHT_BG)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


doc = Document()

# 기본 스타일
style = doc.styles["Normal"]
font = style.font
font.name = "맑은 고딕"
font.size = Pt(10)

# ── 표지 ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("FiPLSim Table 8 재현\n시뮬레이션 결과보고서")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run(
    "논문: Stochastic assessment of weld bead effects\n"
    "on fire sprinkler piping networks\n\n"
    "Scenario 1 (결함 집중 모델) — 6종 시뮬레이션"
)
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_p.add_run(f"생성일: {datetime.date.today().strftime('%Y-%m-%d')}")
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()


# ══════════════════════════════════════════════
# 1. 시뮬레이션 개요
# ══════════════════════════════════════════════
doc.add_heading("1. 시뮬레이션 개요", level=1)

doc.add_paragraph(
    "본 보고서는 논문 'Stochastic assessment of weld bead effects on fire sprinkler "
    "piping networks' (Fire Safety Journal)의 Table 8을 재현하기 위한 "
    "6종 시뮬레이션 결과를 정리한 문서입니다."
)

doc.add_heading("1.1 시뮬레이션 조건", level=2)

cond_table_rows = [
    ["가지배관 수 (n)", f"{NUM_BR}개"],
    ["가지배관당 헤드 (m)", f"{HEADS}개"],
    ["전체 junction 수", f"{NUM_BR * HEADS}개"],
    ["가지배관 간격", f"{BRANCH_SP} m"],
    ["헤드 간격", f"{HEAD_SP} m"],
    ["입구 압력", f"{INLET_P} MPa"],
    ["라이저 관경", SUPPLY],
    ["밸브 상태", "OFF (미포함)"],
    ["MC 반복 횟수", f"{N_ITER:,}회"],
    ["Scenario 유형", "Scenario 1 (결함 집중)"],
    ["비드 확률 (p)", f"{SCENARIO1_P}"],
    ["비드 적용 범위", f"worst branch {HEADS} junctions만"],
]
add_styled_table(doc, ["항목", "설정값"], cond_table_rows)

doc.add_heading("1.2 실행 계획 (6종)", level=2)

plan_rows = [
    ["①", "0 mm (Case B)", "1,200", "기준값 확인"],
    ["②", "0 mm (Case B)", "1,600", "기준값 확인"],
    ["③", "0 mm (Case B)", "2,100 / 2,300", "기준값 확인"],
    ["④", "2.0 mm (Case A)", "1,200 / 1,600 / 2,100", "핵심 재현"],
    ["⑤", "2.5 mm (Case A)", "1,200 / 1,600", "기준값 확인"],
    ["⑥", "2.5 mm (Case A)", "2,100 & 2,300", "★ Table 8 핵심"],
]
add_styled_table(
    doc, ["실행 번호", "비드 높이", "유량 (LPM)", "목적"],
    plan_rows, highlight_rows=[5],
)


# ══════════════════════════════════════════════
# 2. Case B 기준선 결과 (①②③)
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("2. Case B 기준선 결과 (①②③)", level=1)

doc.add_paragraph(
    "비드가 없는 Case B (bead=0mm)에서의 결정론적 계산 결과입니다. "
    "모든 유량에서 단일 값이 산출됩니다 (MC 불필요)."
)

cb_headers = ["유량 (LPM)", "말단 압력 (MPa)", "말단 압력 (kPa)",
              "총 손실 (kPa)", "여유 (%)", "판정"]
cb_data = []
for r in caseB_rows:
    cb_data.append([
        f"{r['Q (LPM)']:,}", f"{r['말단 압력 (MPa)']:.6f}",
        f"{r['말단 압력 (kPa)']:.2f}",
        f"{r['가지배관 총 손실 (kPa)']:.2f}",
        f"{r['0.1 MPa 여유 (%)']:.1f}",
        r["판정"],
    ])
add_styled_table(doc, cb_headers, cb_data)

doc.add_paragraph()
doc.add_paragraph(
    "※ '여유'는 말단 최소 기준 0.1 MPa 대비 여유율입니다. "
    "음수이면 기준 미달을 의미합니다."
)


# ══════════════════════════════════════════════
# 3. Case A (2.0mm) 결과 (④)
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("3. Case A (bead=2.0mm) MC 결과 (④)", level=1)

doc.add_paragraph(
    f"Scenario 1 MC 시뮬레이션 (p={SCENARIO1_P}, N={N_ITER:,})에서 "
    "비드 높이 2.0mm의 결과입니다. worst branch의 8개 junction에만 "
    "독립 Bernoulli 비드를 배치했습니다."
)

a20_headers = ["유량 (LPM)", "Case B (MPa)", "μ (MPa)", "σ (MPa)",
               "Δμ (kPa)", "Pf (%)", "E[beads]"]
a20_data = []
for r in caseA_20_rows:
    a20_data.append([
        f"{r['Q (LPM)']:,}", f"{r['Case B (MPa)']:.6f}",
        f"{r['μ (MPa)']:.6f}", f"{r['σ (MPa)']:.6f}",
        f"{r['Δμ vs CaseB (kPa)']:.2f}",
        f"{r['Pf (%)']:.2f}", f"{r['E[beads]']:.1f}",
    ])
add_styled_table(doc, a20_headers, a20_data)

doc.add_paragraph()
doc.add_paragraph(
    "※ Δμ = Case B 기준선 대비 평균 압력 강하량 (양수 = 비드로 인한 손실)"
)


# ══════════════════════════════════════════════
# 4. Case A (2.5mm) 기준 결과 (⑤)
# ══════════════════════════════════════════════
doc.add_heading("4. Case A (bead=2.5mm) 기준 결과 (⑤)", level=1)

doc.add_paragraph(
    "비드 높이 2.5mm, Q=1200/1600 LPM에서의 MC 결과입니다."
)

a25b_headers = ["유량 (LPM)", "Case B (MPa)", "μ (MPa)", "σ (MPa)",
                "Δμ (kPa)", "Pf (%)", "E[beads]"]
a25b_data = []
for r in caseA_25_base_rows:
    a25b_data.append([
        f"{r['Q (LPM)']:,}", f"{r['Case B (MPa)']:.6f}",
        f"{r['μ (MPa)']:.6f}", f"{r['σ (MPa)']:.6f}",
        f"{r['Δμ vs CaseB (kPa)']:.2f}",
        f"{r['Pf (%)']:.2f}", f"{r['E[beads]']:.1f}",
    ])
add_styled_table(doc, a25b_headers, a25b_data)


# ══════════════════════════════════════════════
# 5. ★ Case A (2.5mm) Table 8 핵심 결과 (⑥)
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("5. ★ Case A (bead=2.5mm) Table 8 핵심 결과 (⑥)", level=1)

doc.add_paragraph(
    "Table 8 재현의 핵심 결과입니다. bead=2.5mm, Q=2100/2300 LPM에서의 "
    f"Scenario 1 MC ({N_ITER:,}회) 상세 통계입니다."
)

a25c_headers = ["항목", "Q=2,100 LPM", "Q=2,300 LPM"]
a25c_data = []

r_2100 = caseA_25_core_rows[0]
r_2300 = caseA_25_core_rows[1]

stat_items = [
    ("Case B 기준 (MPa)", "Case B (MPa)"),
    ("평균 μ (MPa)", "μ (MPa)"),
    ("표준편차 σ (MPa)", "σ (MPa)"),
    ("최솟값 (MPa)", "Min (MPa)"),
    ("P5% (MPa)", "P5% (MPa)"),
    ("P25% (MPa)", "P25% (MPa)"),
    ("P50% 중앙값 (MPa)", "P50% (MPa)"),
    ("P75% (MPa)", "P75% (MPa)"),
    ("P95% (MPa)", "P95% (MPa)"),
    ("최댓값 (MPa)", "Max (MPa)"),
    ("비드 평균 기대 개수", "E[beads]"),
    ("비드 실제 평균 개수", "실제 평균 beads"),
    ("Case B 대비 손실 (kPa)", "Δμ vs CaseB (kPa)"),
    ("기준 미달 확률 Pf (%)", "Pf (%)"),
]

for label, key in stat_items:
    v1 = r_2100.get(key, "—")
    v2 = r_2300.get(key, "—")
    if isinstance(v1, float):
        v1 = f"{v1:.6f}" if v1 < 1 else f"{v1:.2f}"
    if isinstance(v2, float):
        v2 = f"{v2:.6f}" if v2 < 1 else f"{v2:.2f}"
    a25c_data.append([label, str(v1), str(v2)])

highlight = [13]  # Pf row
add_styled_table(doc, a25c_headers, a25c_data, highlight_rows=highlight)


# ══════════════════════════════════════════════
# 6. 논문 대비 비교
# ══════════════════════════════════════════════
doc.add_heading("6. 논문 대비 비교", level=1)

doc.add_paragraph(
    "논문 Table 8의 참고값과 FiPLSim 시뮬레이션 결과를 비교합니다."
)

comp_headers = ["항목", "FiPLSim (Q=2100)", "논문 Table 8"]
comp_data = []

ref_2100 = paper_ref.get((2.5, 2100), {})
sim_2100 = r_2100

comp_items = [
    ("평균 μ (MPa)", sim_2100["μ (MPa)"], ref_2100.get("μ", "—")),
    ("표준편차 σ (MPa)", sim_2100["σ (MPa)"], ref_2100.get("σ", "—")),
    ("기준 미달 확률 Pf (%)", sim_2100["Pf (%)"], ref_2100.get("Pf", "—")),
]

for label, sim_val, paper_val in comp_items:
    sv = f"{sim_val:.6f}" if isinstance(sim_val, float) and sim_val < 1 else f"{sim_val:.2f}" if isinstance(sim_val, float) else str(sim_val)
    pv = f"{paper_val:.4f}" if isinstance(paper_val, float) and paper_val < 1 else f"{paper_val:.2f}" if isinstance(paper_val, float) else str(paper_val)
    comp_data.append([label, sv, pv])

add_styled_table(doc, comp_headers, comp_data)


# ══════════════════════════════════════════════
# 7. 비드 높이별 영향 분석
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("7. 비드 높이별 영향 분석 (2.0mm vs 2.5mm)", level=1)

doc.add_paragraph(
    "동일한 Scenario 1 조건에서 비드 높이에 따른 압력 손실 차이를 비교합니다."
)

be_headers = ["유량 (LPM)", "Case B (MPa)", "2.0mm μ (MPa)", "2.5mm μ (MPa)",
              "2.0mm 손실 (kPa)", "2.5mm 손실 (kPa)", "2.5/2.0 비율"]
be_data = []
for r in bead_effect_rows:
    be_data.append([
        f"{r['Q (LPM)']:,}", f"{r['Case B (MPa)']:.6f}",
        f"{r['2.0mm μ (MPa)']:.6f}", f"{r['2.5mm μ (MPa)']:.6f}",
        f"{r['2.0mm 손실 (kPa)']:.2f}", f"{r['2.5mm 손실 (kPa)']:.2f}",
        f"{r['2.5/2.0 비율']}" if isinstance(r['2.5/2.0 비율'], str) else f"{r['2.5/2.0 비율']:.2f}",
    ])
add_styled_table(doc, be_headers, be_data)

doc.add_paragraph()
doc.add_paragraph(
    "※ 2.5/2.0 비율이 1보다 큰 것은 비드 높이 증가에 따른 손실 증가를 의미합니다."
)


# ══════════════════════════════════════════════
# 8. 통합 요약표 (Table 8 형식)
# ══════════════════════════════════════════════
doc.add_heading("8. 통합 요약표", level=1)

doc.add_paragraph(
    "모든 실행 결과를 하나의 표로 정리한 통합 요약입니다."
)

sum_headers = ["Case", "Bead", "Q (LPM)", "μ (MPa)", "σ (MPa)",
               "Pf (%)", "비고"]
sum_data = []
highlight_sum = []
for i, r in enumerate(summary_rows):
    if "★" in str(r.get("비고", "")):
        highlight_sum.append(i)
    mu_val = r["μ (MPa)"]
    sigma_val = r["σ (MPa)"]
    sum_data.append([
        r["Case"], f"{r['Bead (mm)']:.1f}mm", f"{r['Q (LPM)']:,}",
        f"{mu_val:.6f}" if isinstance(mu_val, float) else str(mu_val),
        f"{sigma_val:.6f}" if isinstance(sigma_val, float) else str(sigma_val),
        f"{r['Pf (%)']:.2f}",
        r["비고"],
    ])

add_styled_table(doc, sum_headers, sum_data, highlight_rows=highlight_sum)


# ══════════════════════════════════════════════
# 9. 분석 및 결론
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("9. 분석 및 결론", level=1)

doc.add_heading("9.1 Case B 기준선 분석", level=2)

caseB_2100_val = next(r for r in caseB_rows if r["Q (LPM)"] == 2100)["말단 압력 (MPa)"]
caseB_2300_val = next(r for r in caseB_rows if r["Q (LPM)"] == 2300)["말단 압력 (MPa)"]

doc.add_paragraph(
    f"Case B (bead=0mm) 기준선에서 Q=2100 LPM일 때 말단 압력은 "
    f"{caseB_2100_val:.6f} MPa, Q=2300 LPM일 때 {caseB_2300_val:.6f} MPa입니다."
)

doc.add_heading("9.2 비드 효과 분석", level=2)

delta_20_2100 = r_20_2100["Δμ vs CaseB (kPa)"]
delta_25_2100 = r_25_2100["Δμ vs CaseB (kPa)"]

doc.add_paragraph(
    f"Q=2100 LPM 기준, Scenario 1 (p=0.5)에서:\n"
    f"  - 비드 2.0mm: Case B 대비 평균 {delta_20_2100:.2f} kPa 추가 손실\n"
    f"  - 비드 2.5mm: Case B 대비 평균 {delta_25_2100:.2f} kPa 추가 손실\n\n"
    f"비드 높이 증가(2.0→2.5mm)에 따른 손실 비율: "
    f"{delta_25_2100/delta_20_2100:.2f}배" if delta_20_2100 != 0 else ""
)

doc.add_heading("9.3 기준 미달 확률 (Pf) 분석", level=2)

pf_20_2100 = r_20_2100["Pf (%)"]
pf_25_2100 = r_25_2100["Pf (%)"]
pf_25_2300 = r_2300["Pf (%)"]

doc.add_paragraph(
    f"Q=2100 LPM에서의 기준 미달 확률:\n"
    f"  - bead=2.0mm: Pf = {pf_20_2100:.2f}%\n"
    f"  - bead=2.5mm: Pf = {pf_25_2100:.2f}%\n\n"
    f"Q=2300 LPM에서의 기준 미달 확률:\n"
    f"  - bead=2.5mm: Pf = {pf_25_2300:.2f}%"
)

doc.add_heading("9.4 논문 Table 8 대비 결과", level=2)

if ref_2100.get("μ") is not None:
    diff_mu = (sim_2100["μ (MPa)"] - ref_2100["μ"]) * 1000
    doc.add_paragraph(
        f"논문 Table 8 (bead=2.5mm, Q=2100, p=0.5):\n"
        f"  - 논문: μ = {ref_2100['μ']:.4f} MPa, Pf = {ref_2100['Pf']:.2f}%\n"
        f"  - FiPLSim: μ = {sim_2100['μ (MPa)']:.6f} MPa, Pf = {sim_2100['Pf (%)']:.2f}%\n"
        f"  - 차이: Δμ = {diff_mu:+.2f} kPa\n\n"
        f"{'FiPLSim 결과가 논문보다 높은 압력(보수적)으로 나옵니다.' if diff_mu > 0 else 'FiPLSim 결과가 논문보다 낮은 압력으로 나옵니다.'}"
    )
else:
    doc.add_paragraph(
        f"FiPLSim (bead=2.5mm, Q=2100, p=0.5):\n"
        f"  - μ = {sim_2100['μ (MPa)']:.6f} MPa, Pf = {sim_2100['Pf (%)']:.2f}%"
    )

doc.add_heading("9.5 향후 과제", level=2)

tasks = [
    "D_eff 모델 적용: 비드 있는 junction의 유속·마찰손실에 유효 내경(D_eff) 반영 "
    "→ 비드 효과 증폭 기대",
    "논문 미기재 변수 확인: 가지배관 사이 간격, 헤드 간격, 관재질 조도, 배관 직경 감소 구성 등",
    "Scenario 2 (시공 품질 모델) 추가 비교: 전체 32 junction에 Bernoulli 비드 배치",
    "Table 6과의 교차 검증: 결정론적 Case A/B 값 비교",
]

for t in tasks:
    doc.add_paragraph(t, style="List Bullet")


# ── 저장 ──
docx_path = os.path.join(BASE_DIR, "FiPLSim_Table8_재현_결과보고서.docx")
doc.save(docx_path)
docx_size = os.path.getsize(docx_path) / 1024

print(f"  ✅ DOCX 저장 완료: {docx_path}")
print(f"     파일 크기: {docx_size:.1f} KB")


# ══════════════════════════════════════════════
# 최종 요약
# ══════════════════════════════════════════════
elapsed = time.time() - t0

print(f"\n{'='*70}")
print(f"  ★ 최종 결과 요약")
print(f"{'='*70}")

print(f"\n  [Case B 기준선]")
for r in caseB_rows:
    print(f"    Q={r['Q (LPM)']:>5}: {r['말단 압력 (MPa)']:.6f} MPa [{r['판정']}]")

print(f"\n  [Case A, bead=2.0mm, Scenario 1, p=0.5]")
for r in caseA_20_rows:
    print(f"    Q={r['Q (LPM)']:>5}: μ={r['μ (MPa)']:.6f}, σ={r['σ (MPa)']:.6f}, Pf={r['Pf (%)']:.2f}%")

print(f"\n  [Case A, bead=2.5mm, Scenario 1, p=0.5]")
for r in caseA_25_base_rows + caseA_25_core_rows:
    print(f"    Q={r['Q (LPM)']:>5}: μ={r['μ (MPa)']:.6f}, σ={r['σ (MPa)']:.6f}, Pf={r['Pf (%)']:.2f}%")

print(f"\n{'='*70}")
print(f"  출력 파일:")
print(f"    1. {xlsx_path}")
print(f"    2. {docx_path}")
print(f"  총 소요 시간: {elapsed:.1f}초 ({elapsed/60:.1f}분)")
print(f"{'='*70}")
